#!/usr/bin/env python3
"""
Agente PDF -> DOCX para sintese de objetos alteradores com contexto de parecer.

Fluxo:
1. Coleta PDFs por CLI/GUI e aplica fallback para pasta do script quando nao houver selecao.
2. Detecta um unico parecer no padrao inteiroTeor-<id>.pdf (ou usa --parecer-file).
3. Extrai texto dos PDFs (fitz -> pypdf) e higieniza conteudo.
4. Resume o parecer e depois cada objeto alterador em 2 paragrafos via OpenAI.
5. Processa arquivos em paralelo (workers + lotes + pacing RPM) para maior velocidade.
6. Mantem checkpoint incremental e backup para retomar do ponto interrompido.
7. Gera DOCX em lista numerada por objeto alterador e salva report tecnico em JSON.
"""

from __future__ import annotations

import argparse
import json
import logging
import os
import re
import threading
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

from gui_intuitiva import dedupe_files, list_files_in_directory, open_file_panel
from openai_log_utils import configure_standard_logging, install_print_logger_bridge
from openai_progress_utils import (
    build_file_signature,
    make_backup,
    read_json_dict,
    utc_now_iso,
    write_json_atomic,
)


SCRIPT_DIR = Path(__file__).resolve().parent
REPORT_FILE = SCRIPT_DIR / ".relatorio_objetos_alteradores.report.json"
CHECKPOINT_VERSION = 1
PROMPT_VERSION = 3

PARECER_FILENAME_RE = re.compile(r"^inteiroTeor-\d+\.pdf$", re.IGNORECASE)
OBJETO_RE = re.compile(r"([A-Z]{2,6})[-_ ]*(\d+)", re.IGNORECASE)
PROPOSICAO_RE = re.compile(r"\b(PLP?|PEC|MPV?|PDC|PRC|PLC)[-_ ]*(\d{1,5})[-_ ]*(\d{4})\b", re.IGNORECASE)
EMP_REF_RE = re.compile(r"\bEMP[-_ ]*(\d+)\b", re.IGNORECASE)
EMENDA_REF_RE = re.compile(r"\bEMENDA\s*(?:N[.oº]*\s*)?(\d+)\b", re.IGNORECASE)
URL_RE = re.compile(r"https?://[^\s)>\"]+")

DISPENSABLE_BRACKET_MARKER_RE = re.compile(
    r"\[(?:suprime|substitui|condiciona|acresce|altera|retira|inclui|exclui)\]",
    re.IGNORECASE,
)
DISPENSABLE_LEADING_RULES: List[Tuple[str, re.Pattern[str], str]] = [
    (
        "intro_trata_votacao_separado",
        re.compile(r"^trata\s+da\s+vota[cç][aã]o\s+em\s+separado\s+do\s+", re.IGNORECASE),
        "destaca para votacao em separado o ",
    ),
    (
        "intro_emenda_ao_pl",
        re.compile(
            r"^(?:A\s+)?Emenda(?:\s+de\s+Plen[aá]rio)?\s*(?:N[.oº°]*\s*)?\d+\s+ao\s+Projeto\s+de\s+Lei\s+n[.oº°]*\s*\d+,\s*de\s*\d{4},?\s*",
            re.IGNORECASE,
        ),
        "",
    ),
    (
        "intro_fragmento_plenario_ao_pl",
        re.compile(
            r"^de\s+Plen[aá]rio\s+\d+\s+ao\s+Projeto\s+de\s+Lei\s+n[.oº°]*\s*\d+,\s*de\s*\d{4},?\s*",
            re.IGNORECASE,
        ),
        "",
    ),
    (
        "intro_emenda_numerada",
        re.compile(
            r"^(?:A\s+)?Emenda(?:\s+de\s+Plen[aá]rio)?\s*(?:N[.oº°]*\s*)?\d+\s+",
            re.IGNORECASE,
        ),
        "",
    ),
]
DISPENSABLE_INLINE_RULES: List[Tuple[str, re.Pattern[str], str]] = [
    (
        "inline_na_forma_proposta_pl",
        re.compile(
            r",?\s*na\s+forma\s+proposta\s+pelo\s+Projeto\s+de\s+Lei\s+n[.oº°]*\s*\d+,\s*de\s*\d{4},?\s*",
            re.IGNORECASE,
        ),
        " ",
    ),
    (
        "inline_inserido_pelo_pl",
        re.compile(
            r",?\s*inserid[oa]\s+pelo\s+Projeto\s+de\s+Lei\s+n[.oº°]*\s*\d+(?:/\d{4})?,?\s*",
            re.IGNORECASE,
        ),
        " ",
    ),
]

DEFAULT_OPENAI_MODEL = "gpt-5.1"
DEFAULT_FALLBACK_MODEL = "gpt-5-mini"
DEFAULT_OPENAI_TIMEOUT_S = 120
DEFAULT_OPENAI_RETRIES = 3
DEFAULT_OPENAI_MAX_COMPLETION_TOKENS = 420
DEFAULT_OPENAI_MAX_WORKERS = 10
DEFAULT_OPENAI_BATCH_SIZE = 30
DEFAULT_OPENAI_TARGET_RPM = 480
DEFAULT_OPENAI_DELAY_S = 0.0
DEFAULT_INTEGRAL_LENGTH_RETRIES = 2
OPENAI_WAIT_HEARTBEAT_INTERVAL_S = 10.0
BATCH_HEARTBEAT_INTERVAL_S = 8.0

# Fallback por chunks e excecao: so entra aqui apos tentativas integrais por length.
CHUNK_SIZE_CHARS = 120000
MAX_FINAL_CHUNK_SUMMARY_CHARS = 180000

LOGGER = logging.getLogger("AGENTE_objetos_alteradores_pdf_para_docx")


PARECER_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "additionalProperties": False,
    "required": ["parecer_resumo"],
    "properties": {
        "parecer_resumo": {"type": "string"},
    },
}

ITEM_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "additionalProperties": False,
    "required": ["paragrafo_1", "paragrafo_2", "objeto_alterador", "tipo_objeto"],
    "properties": {
        "paragrafo_1": {"type": "string"},
        "paragrafo_2": {"type": "string"},
        "objeto_alterador": {"type": "string"},
        "tipo_objeto": {"type": "string"},
    },
}

CHUNK_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "additionalProperties": False,
    "required": ["resumo_chunk"],
    "properties": {
        "resumo_chunk": {"type": "string"},
    },
}

BLOCK_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "additionalProperties": False,
    "required": ["resumo_bloco"],
    "properties": {
        "resumo_bloco": {"type": "string"},
    },
}


@dataclass
class InputDoc:
    path: str
    name: str
    object_type: str
    object_number: int
    object_label: str
    kind: str = "objeto"


@dataclass
class SummaryItem:
    file_path: str
    file_name: str
    objeto_alterador: str
    tipo_objeto: str
    numero_objeto: int
    paragrafo_1: str
    paragrafo_2: str
    model_used: str
    status: str
    used_chunk_fallback: bool = False
    error: str = ""


@dataclass
class RunReport:
    started_at: str
    finished_at: str = ""
    status: str = "started"
    model_primary: str = DEFAULT_OPENAI_MODEL
    model_fallback: str = DEFAULT_FALLBACK_MODEL
    input_files: List[str] = field(default_factory=list)
    parecer_file: str = ""
    output_docx: str = ""
    checkpoint_file: str = ""
    total_targets: int = 0
    success_count: int = 0
    failure_count: int = 0
    resumed_from_checkpoint: bool = False
    items: List[Dict[str, Any]] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    notes: List[str] = field(default_factory=list)


@dataclass
class OpenAIErrorInfo:
    kind: str
    message: str
    status_code: int = 0


class AppError(RuntimeError):
    """Erro de aplicacao para retorno controlado de codigo 1."""


class OpenAIRequestFailed(RuntimeError):
    def __init__(self, info: OpenAIErrorInfo, model: str) -> None:
        super().__init__(f"[{info.kind}] {info.message}")
        self.info = info
        self.model = model


class RequestPacer:
    def __init__(self, *, target_rpm: int = 0, min_interval_s: float = 0.0) -> None:
        if target_rpm > 0:
            self._min_interval_s = 60.0 / float(target_rpm)
        else:
            self._min_interval_s = max(0.0, float(min_interval_s))
        self._lock = threading.Lock()
        self._next_at = 0.0

    def wait_turn(self) -> None:
        if self._min_interval_s <= 0.0:
            return
        while True:
            with self._lock:
                now = time.monotonic()
                if now >= self._next_at:
                    self._next_at = now + self._min_interval_s
                    return
                sleep_for = max(0.0, self._next_at - now)
            if sleep_for > 0:
                time.sleep(min(0.2, sleep_for))


class BatchTelemetry:
    def __init__(self, *, label: str, total_docs: int) -> None:
        self.label = label
        self.total_docs = max(0, int(total_docs))
        self._lock = threading.Lock()
        self.started_at = time.monotonic()
        self.last_event_at = self.started_at
        self.docs_inflight = 0
        self.docs_done = 0
        self.docs_ok = 0
        self.docs_err = 0
        self.docs_elapsed_total_s = 0.0
        self.api_started = 0
        self.api_success = 0
        self.api_failed = 0
        self.api_retries = 0
        self.api_elapsed_total_s = 0.0

    def _touch(self) -> None:
        self.last_event_at = time.monotonic()

    def on_doc_start(self) -> None:
        with self._lock:
            self.docs_inflight += 1
            self._touch()

    def on_doc_finish(self, *, ok: bool, elapsed_s: float) -> None:
        with self._lock:
            self.docs_inflight = max(0, self.docs_inflight - 1)
            self.docs_done += 1
            if ok:
                self.docs_ok += 1
            else:
                self.docs_err += 1
            self.docs_elapsed_total_s += max(0.0, float(elapsed_s))
            self._touch()

    def on_api_start(self) -> None:
        with self._lock:
            self.api_started += 1
            self._touch()

    def on_api_success(self, elapsed_s: float) -> None:
        with self._lock:
            self.api_success += 1
            self.api_elapsed_total_s += max(0.0, float(elapsed_s))
            self._touch()

    def on_api_retry(self) -> None:
        with self._lock:
            self.api_retries += 1
            self._touch()

    def on_api_fail(self) -> None:
        with self._lock:
            self.api_failed += 1
            self._touch()

    def snapshot(self) -> Dict[str, Any]:
        with self._lock:
            now = time.monotonic()
            elapsed_total_s = max(0.0, now - self.started_at)
            idle_s = max(0.0, now - self.last_event_at)
            avg_api_s = (self.api_elapsed_total_s / self.api_success) if self.api_success > 0 else 0.0
            avg_doc_s = (self.docs_elapsed_total_s / self.docs_done) if self.docs_done > 0 else 0.0
            completed_api_calls = self.api_success + self.api_failed
            api_error_rate_pct = (self.api_failed / completed_api_calls * 100.0) if completed_api_calls > 0 else 0.0
            qps_success = (self.api_success / elapsed_total_s) if elapsed_total_s > 0 else 0.0
            return {
                "label": self.label,
                "total_docs": self.total_docs,
                "docs_inflight": self.docs_inflight,
                "docs_done": self.docs_done,
                "docs_ok": self.docs_ok,
                "docs_err": self.docs_err,
                "api_started": self.api_started,
                "api_success": self.api_success,
                "api_failed": self.api_failed,
                "api_retries": self.api_retries,
                "elapsed_total_s": elapsed_total_s,
                "idle_s": idle_s,
                "avg_api_s": avg_api_s,
                "avg_doc_s": avg_doc_s,
                "api_error_rate_pct": api_error_rate_pct,
                "qps_success": qps_success,
            }


def start_batch_progress_heartbeat(
    telemetry: BatchTelemetry,
    logger: logging.Logger,
    *,
    interval_s: float = BATCH_HEARTBEAT_INTERVAL_S,
) -> Tuple[threading.Event, Optional[threading.Thread]]:
    stop_event = threading.Event()
    if interval_s <= 0.0:
        return stop_event, None

    def _heartbeat() -> None:
        while not stop_event.wait(interval_s):
            snap = telemetry.snapshot()
            logger.info(
                "BATCH_HEARTBEAT | lote=%s | docs_done=%d/%d | inflight=%d | ok=%d | err=%d | "
                "api_started=%d | api_ok=%d | api_err=%d | retries=%d | avg_api_s=%.2f | idle=%.1fs",
                snap["label"],
                snap["docs_done"],
                snap["total_docs"],
                snap["docs_inflight"],
                snap["docs_ok"],
                snap["docs_err"],
                snap["api_started"],
                snap["api_success"],
                snap["api_failed"],
                snap["api_retries"],
                snap["avg_api_s"],
                snap["idle_s"],
            )

    thread = threading.Thread(
        target=_heartbeat,
        name=f"batch-heartbeat-{telemetry.label}",
        daemon=True,
    )
    thread.start()
    return stop_event, thread


def normalize_ws(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def normalize_multiline(value: Any) -> str:
    raw = str(value or "")
    if not raw:
        return ""
    lines = []
    blank_streak = 0
    for line in raw.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        cleaned = re.sub(r"[ \t]+", " ", line).strip()
        if not cleaned:
            blank_streak += 1
            if blank_streak <= 1:
                lines.append("")
            continue
        blank_streak = 0
        lines.append(cleaned)
    return "\n".join(lines).strip()


def resolve_project_path(path_str: str) -> Path:
    path = Path(path_str).expanduser()
    if not path.is_absolute():
        path = SCRIPT_DIR / path
    return path.resolve()


def read_secret_from_file(path: Path) -> str:
    if not path.exists() or not path.is_file():
        return ""
    try:
        raw = path.read_text(encoding="utf-8").strip()
    except Exception:
        return ""
    if not raw:
        return ""
    first_line = raw.splitlines()[0].strip()
    if "=" in first_line:
        first_line = first_line.split("=", 1)[1].strip()
    return first_line.strip("\"' ")


def resolve_openai_key() -> str:
    env_key = os.getenv("OPENAI_API_KEY", "").strip()
    if env_key:
        return env_key

    candidates = [
        SCRIPT_DIR / "CHAVE_SECRETA_API_Mauricio_local.txt",
        SCRIPT_DIR / "Chave Secreta API_Mauricio_local.txt",
        Path.cwd() / "CHAVE_SECRETA_API_Mauricio_local.txt",
        Path.cwd() / "Chave Secreta API_Mauricio_local.txt",
    ]
    for candidate in candidates:
        value = read_secret_from_file(candidate)
        if value:
            return value
    return ""


def parse_object_from_filename(path: Path) -> Tuple[str, int, str]:
    stem_upper = path.stem.upper()
    match = OBJETO_RE.search(stem_upper)
    if not match:
        return "SEM_TIPO", -1, stem_upper

    object_type = normalize_ws(match.group(1).upper()) or "SEM_TIPO"
    try:
        object_number = int(match.group(2))
    except Exception:
        object_number = -1

    if object_number >= 0:
        return object_type, object_number, f"{object_type}-{object_number}"
    return object_type, object_number, object_type


def build_input_doc(path: Path, kind: str = "objeto") -> InputDoc:
    obj_type, obj_num, obj_label = parse_object_from_filename(path)
    return InputDoc(
        path=str(path.resolve()),
        name=path.name,
        object_type=obj_type,
        object_number=obj_num,
        object_label=obj_label,
        kind=kind,
    )


def choose_recursive_for_gui(default_recursive: bool, logger: logging.Logger) -> bool:
    try:
        import tkinter as tk
        from tkinter import messagebox
    except Exception as exc:
        logger.debug("GUI: tkinter indisponivel para pergunta de recursao: %s", exc)
        return bool(default_recursive)

    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    try:
        answer = messagebox.askyesno(
            "Selecao de pastas",
            "Incluir subpastas ao adicionar pasta?",
        )
        return bool(answer)
    finally:
        root.destroy()


def maybe_collect_gui_files(args: argparse.Namespace, logger: logging.Logger) -> argparse.Namespace:
    if bool(args.no_gui) or list(args.input_files or []) or list(args.input_dirs or []):
        return args

    recursive_choice = choose_recursive_for_gui(bool(args.recursive), logger)
    gui = open_file_panel(
        title="Agente PDF -> DOCX (Objetos Alteradores)",
        subtitle="Selecione PDFs por arquivo ou pasta.",
        filetypes=[("PDF", "*.pdf"), ("Todos os arquivos", "*.*")],
        extensions=[".pdf"],
        initial_files=[],
        allow_add_dir=True,
        recursive_dir=recursive_choice,
        min_files=1,
        output_label="Pasta de saida",
        initial_output=str(resolve_project_path(args.output_dir)),
        extra_texts=[
            ("openai_model", "Modelo principal OpenAI", str(args.openai_model)),
            ("fallback_model", "Modelo fallback OpenAI", str(args.fallback_model)),
            ("openai_max_workers", "OpenAI workers", str(args.openai_max_workers)),
            ("openai_batch_size", "OpenAI batch size", str(args.openai_batch_size)),
            ("openai_target_rpm", "OpenAI target RPM", str(args.openai_target_rpm)),
        ],
    )

    if not gui or not gui.get("confirmed"):
        return args

    files = dedupe_files(gui.get("files") or [], [".pdf"])
    if files:
        args.input_files = files

    output_dir = str(gui.get("output") or "").strip()
    if output_dir:
        args.output_dir = output_dir

    texts = gui.get("texts") or {}
    args.openai_model = str(texts.get("openai_model", args.openai_model)).strip() or args.openai_model
    args.fallback_model = str(texts.get("fallback_model", args.fallback_model)).strip() or args.fallback_model
    try:
        args.openai_max_workers = max(1, int(str(texts.get("openai_max_workers", args.openai_max_workers)).strip()))
    except Exception:
        pass
    try:
        args.openai_batch_size = max(1, int(str(texts.get("openai_batch_size", args.openai_batch_size)).strip()))
    except Exception:
        pass
    try:
        args.openai_target_rpm = max(0, int(str(texts.get("openai_target_rpm", args.openai_target_rpm)).strip()))
    except Exception:
        pass

    return args


def discover_pdf_inputs(input_files: Sequence[str], input_dirs: Sequence[str], recursive: bool) -> List[Path]:
    found: List[str] = []
    for input_dir in input_dirs or []:
        found.extend(list_files_in_directory(input_dir, [".pdf"], recursive=bool(recursive)))
    found.extend(list(input_files or []))
    deduped = dedupe_files(found, [".pdf"])
    return [Path(p).resolve() for p in deduped]


def fallback_script_dir_pdfs() -> List[Path]:
    found = list_files_in_directory(str(SCRIPT_DIR), [".pdf"], recursive=False)
    deduped = dedupe_files(found, [".pdf"])
    return [Path(p).resolve() for p in deduped]


def detect_parecer_path(all_pdfs: Sequence[Path], parecer_cli: str) -> Path:
    if normalize_ws(parecer_cli):
        parecer_path = resolve_project_path(parecer_cli)
        if not parecer_path.exists() or not parecer_path.is_file():
            raise AppError(f"Parecer informado em --parecer-file nao existe: {parecer_path}")
        if parecer_path.suffix.lower() != ".pdf":
            raise AppError(f"Parecer informado nao e PDF: {parecer_path}")
        return parecer_path

    candidates = [path for path in all_pdfs if PARECER_FILENAME_RE.match(path.name)]
    if len(candidates) == 1:
        return candidates[0]

    if not candidates:
        raise AppError(
            "Nao foi encontrado parecer no padrao 'inteiroTeor-<id>.pdf'. "
            "Use --parecer-file para indicar explicitamente o arquivo de referencia."
        )

    candidate_names = ", ".join(path.name for path in candidates)
    raise AppError(
        "Mais de um parecer candidato foi encontrado. "
        f"Candidatos: {candidate_names}. Use --parecer-file para escolher um unico parecer."
    )


def parse_openai_message_content(content_obj: Any) -> str:
    if isinstance(content_obj, str):
        return content_obj.strip()
    if isinstance(content_obj, list):
        chunks: List[str] = []
        for item in content_obj:
            if isinstance(item, dict):
                text = item.get("text")
                if isinstance(text, str):
                    chunks.append(text)
            else:
                text = getattr(item, "text", None)
                if isinstance(text, str):
                    chunks.append(text)
        return "".join(chunks).strip()
    return ""


def status_code_from_exception(exc: Exception) -> int:
    for attr in ("status_code", "status"):
        value = getattr(exc, attr, None)
        if isinstance(value, int):
            return value
    response = getattr(exc, "response", None)
    if response is not None:
        status_code = getattr(response, "status_code", None)
        if isinstance(status_code, int):
            return status_code
    return 0


def classify_openai_exception(exc: Exception) -> OpenAIErrorInfo:
    message = normalize_ws(str(exc))
    low = message.lower()
    status_code = status_code_from_exception(exc)

    if any(
        tok in low
        for tok in (
            "context length",
            "maximum context",
            "too many tokens",
            "max token",
            "prompt is too long",
            "token limit",
            "empty_json_due_length",
        )
    ):
        return OpenAIErrorInfo(kind="length", message=message, status_code=status_code)

    if status_code in (401, 403) or any(tok in low for tok in ("authentication", "permission", "forbidden")):
        return OpenAIErrorInfo(kind="permission", message=message, status_code=status_code)

    if status_code == 404 and "model" in low:
        return OpenAIErrorInfo(kind="model", message=message, status_code=status_code)

    if status_code == 429 or "rate limit" in low:
        return OpenAIErrorInfo(kind="rate_limit", message=message, status_code=status_code)

    cls_name = exc.__class__.__name__.lower()
    if status_code >= 500 or any(tok in cls_name for tok in ("timeout", "connection")):
        return OpenAIErrorInfo(kind="unavailable", message=message, status_code=status_code)

    if any(tok in low for tok in ("timeout", "timed out", "connection", "temporarily unavailable", "service unavailable")):
        return OpenAIErrorInfo(kind="unavailable", message=message, status_code=status_code)

    return OpenAIErrorInfo(kind="error", message=message, status_code=status_code)


def retry_delay_s(attempt: int) -> float:
    return min(20.0, 0.8 * (2 ** max(0, attempt - 1)))


def estimate_prompt_chars(messages: Sequence[Dict[str, str]]) -> int:
    total = 0
    for message in messages or []:
        if not isinstance(message, dict):
            continue
        total += len(str(message.get("content", "") or ""))
    return total


def short_error_message(text: Any, limit: int = 280) -> str:
    cleaned = normalize_ws(text)
    if len(cleaned) <= limit:
        return cleaned
    return f"{cleaned[: max(0, limit - 3)]}..."


def start_openai_wait_heartbeat(
    logger: logging.Logger,
    *,
    request_label: str,
    model: str,
    attempt: int,
    max_attempts: int,
    interval_s: float = OPENAI_WAIT_HEARTBEAT_INTERVAL_S,
) -> Tuple[threading.Event, Optional[threading.Thread], float]:
    started_at = time.monotonic()
    stop_event = threading.Event()

    if interval_s <= 0:
        return stop_event, None, started_at

    def _heartbeat() -> None:
        while not stop_event.wait(interval_s):
            elapsed = time.monotonic() - started_at
            logger.info(
                "OPENAI_WAIT | %s | model=%s | tentativa=%d/%d | aguardando resposta ha %.1fs",
                request_label,
                model,
                attempt,
                max_attempts,
                elapsed,
            )

    thread = threading.Thread(
        target=_heartbeat,
        name=f"openai-heartbeat-{attempt}",
        daemon=True,
    )
    thread.start()
    return stop_event, thread, started_at


def openai_json_request_with_model(
    client: Any,
    *,
    model: str,
    messages: Sequence[Dict[str, str]],
    schema: Dict[str, Any],
    schema_name: str,
    timeout_s: int,
    retries: int,
    max_completion_tokens: int,
    logger: logging.Logger,
    request_label: str,
    pacer: Optional[RequestPacer] = None,
    telemetry: Optional[BatchTelemetry] = None,
) -> Dict[str, Any]:
    max_attempts = max(1, int(retries))
    last_error: Optional[OpenAIErrorInfo] = None
    prompt_chars = estimate_prompt_chars(messages)
    token_limit = max(0, int(max_completion_tokens))

    for attempt in range(1, max_attempts + 1):
        try:
            if pacer is not None:
                pacer.wait_turn()
            if telemetry is not None:
                telemetry.on_api_start()
            logger.info(
                "OPENAI_START | %s | model=%s | tentativa=%d/%d | timeout=%ss | max_completion_tokens=%d | prompt_chars=%d",
                request_label,
                model,
                attempt,
                max_attempts,
                max(5, int(timeout_s)),
                token_limit,
                prompt_chars,
            )

            payload: Dict[str, Any] = {
                "model": model,
                "messages": list(messages),
                "response_format": {
                    "type": "json_schema",
                    "json_schema": {
                        "name": schema_name,
                        "strict": True,
                        "schema": schema,
                    },
                },
                "timeout": max(5, int(timeout_s)),
            }
            if int(max_completion_tokens) > 0:
                payload["max_completion_tokens"] = int(max_completion_tokens)

            stop_event, hb_thread, started_at = start_openai_wait_heartbeat(
                logger,
                request_label=request_label,
                model=model,
                attempt=attempt,
                max_attempts=max_attempts,
            )
            try:
                response = client.chat.completions.create(**payload)
            finally:
                stop_event.set()
                if hb_thread is not None:
                    hb_thread.join(timeout=0.2)

            elapsed = time.monotonic() - started_at
            choice = response.choices[0]
            raw = parse_openai_message_content(choice.message.content)
            finish_reason = str(getattr(choice, "finish_reason", "") or "")
            if not raw:
                if finish_reason == "length":
                    raise RuntimeError("empty_json_due_length")
                refusal = getattr(choice.message, "refusal", None)
                if refusal:
                    raise RuntimeError(f"refusal: {refusal}")
                raise RuntimeError("empty_json")

            parsed = json.loads(raw)
            if not isinstance(parsed, dict):
                raise RuntimeError("A resposta JSON da OpenAI nao e um objeto.")
            logger.info(
                "OPENAI_SUCCESS | %s | model=%s | tentativa=%d/%d | elapsed=%.2fs | finish_reason=%s",
                request_label,
                model,
                attempt,
                max_attempts,
                elapsed,
                finish_reason or "<vazio>",
            )
            if telemetry is not None:
                telemetry.on_api_success(elapsed)
            return parsed
        except Exception as exc:  # noqa: BLE001
            info = classify_openai_exception(exc)
            last_error = info
            retryable = info.kind in {"rate_limit", "unavailable"}
            if retryable and attempt < max_attempts:
                delay = retry_delay_s(attempt)
                if telemetry is not None:
                    telemetry.on_api_retry()
                logger.warning(
                    "OPENAI_RETRY | %s | model=%s | tentativa=%d/%d | kind=%s | status=%s | msg=%s | retry_em=%.2fs",
                    request_label,
                    model,
                    attempt,
                    max_attempts,
                    info.kind,
                    info.status_code,
                    short_error_message(info.message),
                    delay,
                )
                time.sleep(delay)
                continue
            if telemetry is not None:
                telemetry.on_api_fail()
            logger.error(
                "OPENAI_FAIL | %s | model=%s | tentativa=%d/%d | kind=%s | status=%s | msg=%s",
                request_label,
                model,
                attempt,
                max_attempts,
                info.kind,
                info.status_code,
                short_error_message(info.message),
            )
            raise OpenAIRequestFailed(info=info, model=model) from exc

    if last_error is None:
        last_error = OpenAIErrorInfo(kind="error", message="Falha desconhecida na OpenAI")
    raise OpenAIRequestFailed(info=last_error, model=model)


def openai_json_request(
    client: Any,
    *,
    messages: Sequence[Dict[str, str]],
    schema: Dict[str, Any],
    schema_name: str,
    primary_model: str,
    fallback_model: str,
    timeout_s: int,
    retries: int,
    max_completion_tokens: int,
    logger: logging.Logger,
    request_label: str,
    pacer: Optional[RequestPacer] = None,
    telemetry: Optional[BatchTelemetry] = None,
) -> Tuple[Dict[str, Any], str, bool]:
    primary = normalize_ws(primary_model)
    fallback = normalize_ws(fallback_model)

    try:
        payload = openai_json_request_with_model(
            client,
            model=primary,
            messages=messages,
            schema=schema,
            schema_name=schema_name,
            timeout_s=timeout_s,
            retries=retries,
            max_completion_tokens=max_completion_tokens,
            logger=logger,
            request_label=request_label,
            pacer=pacer,
            telemetry=telemetry,
        )
        return payload, primary, False
    except OpenAIRequestFailed as primary_error:
        allow_fallback = (
            bool(fallback)
            and fallback != primary
            and primary_error.info.kind in {"permission", "model", "unavailable"}
        )
        if not allow_fallback:
            raise

        logger.warning(
            "%s | modelo principal '%s' falhou (%s). Tentando fallback '%s'.",
            request_label,
            primary,
            primary_error.info.kind,
            fallback,
        )

        payload = openai_json_request_with_model(
            client,
            model=fallback,
            messages=messages,
            schema=schema,
            schema_name=schema_name,
            timeout_s=timeout_s,
            retries=retries,
            max_completion_tokens=max_completion_tokens,
            logger=logger,
            request_label=f"{request_label} [fallback]",
            pacer=pacer,
            telemetry=telemetry,
        )
        logger.info(
            "OPENAI_FALLBACK_SUCCESS | %s | model=%s",
            request_label,
            fallback,
        )
        return payload, fallback, True


def extract_pdf_text(path: Path) -> str:
    fitz_error = ""
    pypdf_error = ""

    try:
        import fitz  # type: ignore

        with fitz.open(str(path)) as doc:
            texts = [page.get_text("text") or "" for page in doc]
        text = "\n".join(texts)
        if normalize_ws(text):
            return text
    except Exception as exc:  # noqa: BLE001
        fitz_error = normalize_ws(str(exc))

    try:
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception:
            from PyPDF2 import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        texts = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            texts.append(page_text)
        text = "\n".join(texts)
        if normalize_ws(text):
            return text
    except Exception as exc:  # noqa: BLE001
        pypdf_error = normalize_ws(str(exc))

    if "No module named 'fitz'" in fitz_error and (
        "No module named 'pypdf'" in pypdf_error or "No module named 'PyPDF2'" in pypdf_error
    ):
        raise AppError(
            "Bibliotecas de PDF ausentes. Instale uma opcao: "
            "python -m pip install pymupdf  OU  python -m pip install pypdf"
        )

    detail = "; ".join([part for part in [fitz_error, pypdf_error] if part])
    if not detail:
        detail = "conteudo vazio"
    raise AppError(f"Falha ao extrair texto do PDF '{path.name}': {detail}")


def split_text_chunks(text: str, max_chars: int = CHUNK_SIZE_CHARS) -> List[str]:
    normalized = normalize_multiline(text)
    if len(normalized) <= max_chars:
        return [normalized]

    paragraphs = [part.strip() for part in re.split(r"\n\n+", normalized) if part.strip()]
    chunks: List[str] = []
    current: List[str] = []
    current_len = 0

    def flush_current() -> None:
        nonlocal current, current_len
        if current:
            chunks.append("\n\n".join(current).strip())
        current = []
        current_len = 0

    for paragraph in paragraphs:
        if len(paragraph) > max_chars:
            flush_current()
            start = 0
            while start < len(paragraph):
                end = min(start + max_chars, len(paragraph))
                piece = paragraph[start:end].strip()
                if piece:
                    chunks.append(piece)
                start = end
            continue

        projected = current_len + len(paragraph) + (2 if current else 0)
        if projected > max_chars:
            flush_current()
        current.append(paragraph)
        current_len += len(paragraph) + (2 if current_len > 0 else 0)

    flush_current()
    return [chunk for chunk in chunks if chunk]


def is_destaque_doc(doc: InputDoc) -> bool:
    doc_type = normalize_ws(doc.object_type).upper()
    return doc_type in {"DTQ", "DVS"}


def infer_proposicao_ref(*texts: str) -> str:
    for raw in texts:
        content = str(raw or "")
        if not content:
            continue
        match = PROPOSICAO_RE.search(content)
        if not match:
            continue
        sigla = normalize_ws(match.group(1)).upper()
        numero = normalize_ws(match.group(2))
        ano = normalize_ws(match.group(3))
        if sigla == "MP":
            sigla = "MPV"
        return f"{sigla} {numero}/{ano}"
    return ""


def infer_emp_ref(text: str) -> str:
    match = EMP_REF_RE.search(str(text or ""))
    if match:
        return normalize_ws(match.group(1))
    emenda_match = EMENDA_REF_RE.search(str(text or ""))
    if emenda_match:
        return normalize_ws(emenda_match.group(1))
    return ""


def classify_destaque_mode(doc: InputDoc, doc_text: str) -> str:
    doc_type = normalize_ws(doc.object_type).upper()
    if doc_type == "DVS":
        return "dvs"
    if doc_type != "DTQ":
        return "nao_destaque"

    emp_ref = infer_emp_ref(doc_text)
    if emp_ref:
        return "destaque_emenda"

    low = str(doc_text or "").lower()
    if "votacao em separado" in low or "votação em separado" in low or re.search(r"\bdvs\b", low):
        return "dvs"

    # Default prudente para DTQ sem referencia explicita de EMP.
    return "dvs"


def infer_tramitacao_link(*texts: str) -> str:
    urls: List[str] = []
    for raw in texts:
        content = str(raw or "")
        if not content:
            continue
        urls.extend(URL_RE.findall(content))
    if not urls:
        return ""
    for url in urls:
        if "camara.leg.br" in url.lower():
            return url
    return urls[0]


def build_identificacao_header(doc: InputDoc, doc_text: str = "", destaque_mode: str = "") -> str:
    proposicao = infer_proposicao_ref(doc.name, doc_text)
    if not proposicao:
        proposicao = "Proposicao nao identificada"

    doc_type = normalize_ws(doc.object_type).upper()
    if doc_type in {"DTQ", "DVS"}:
        dtq_label = f"DTQ {doc.object_number}" if doc.object_number >= 0 else normalize_ws(doc.object_label or "DTQ")
        mode = normalize_ws(destaque_mode).lower()
        emp_ref = infer_emp_ref(doc_text)
        if mode == "destaque_emenda" and emp_ref:
            return f"{proposicao} - {dtq_label} - EMP {emp_ref}"
        return f"{proposicao} - {dtq_label}"

    if doc_type == "EMP":
        emp_label = f"EMP {doc.object_number}" if doc.object_number >= 0 else normalize_ws(doc.object_label or "EMP")
        return f"{proposicao} - {emp_label}"

    object_label = normalize_ws(doc.object_label or doc.name)
    return f"{proposicao} - {object_label}"


def build_parecer_messages(parecer_text: str) -> List[Dict[str, str]]:
    prompt = f"""
Resumo-base do parecer da comissao:
- Produza um unico paragrafo curto e neutro (4 a 6 linhas).
- Foque no contexto geral, criterio de analise e eixo normativo principal.
- Nao use citacoes literais, aspas, bullet points ou numeracao.
- Retorne apenas JSON valido conforme schema.

PARECER (texto integral):
---
{parecer_text}
---
""".strip()
    return [
        {
            "role": "system",
            "content": (
                "Voce e um analista legislativo. Responda somente com JSON valido. "
                "Nao inclua explicacoes fora do JSON."
            ),
        },
        {"role": "user", "content": prompt},
    ]


def build_item_messages(
    doc: InputDoc,
    doc_text: str,
    parecer_summary: str,
    destaque_mode: str = "",
) -> List[Dict[str, str]]:
    mode = normalize_ws(destaque_mode).lower()
    if mode == "destaque_emenda":
        identificacao = build_identificacao_header(doc, doc_text, destaque_mode=mode)
        tramitacao_link = infer_tramitacao_link(doc_text, parecer_summary)
        tramitacao_txt = tramitacao_link or "Nao informado."
        proposicao_context = normalize_multiline(parecer_summary)
        if not proposicao_context:
            proposicao_context = "Nao fornecido separadamente nesta chamada."

        prompt = f"""
PAPEL
Voce e um Consultor Legislativo da Camara dos Deputados, especialista em Direito Constitucional,
Processo Legislativo e na area tematica especifica da emenda em analise.

TAREFA
Elabore uma explicacao objetiva do Destaque (DTQ) anexo, cujo objeto e a Emenda (EMP) vinculada.

INSUMOS DISPONIVEIS NESTA CHAMADA
1. Texto integral do destaque.
2. Texto da emenda, quando indicado no proprio destaque.
3. Contexto da proposicao principal: fornecido abaixo.
4. Link de tramitacao da proposicao principal: {tramitacao_txt}

INSTRUCOES
1. Foco exclusivo no texto normativo. Explique o que a emenda pretende alterar ou incluir
na proposicao principal. Nao explique justificacao da emenda.
2. Resolva as remissoes. Se houver remissao a leis/dispositivos, localize o teor no
material fornecido e, quando necessario, use referencia legislativa consolidada em
planalto.gov.br/legislacao para tornar o texto autocontido.
3. Formato: no maximo dois paragrafos, em prosa corrida. O primeiro paragrafo DEVE iniciar
exatamente com "{identificacao}".
4. Tom profissional e direto, acessivel a parlamentares.
5. Pense passo a passo internamente antes de redigir (sem expor raciocinio).
6. Evite estruturas introdutorias dispensaveis apos a identificacao (ex.: "A Emenda ... ao Projeto de Lei ...").
7. Nao use marcadores editoriais entre colchetes, como [suprime], [substitui], [condiciona].

Contexto da proposicao principal:
---
{proposicao_context}
---

Texto integral do destaque:
---
{doc_text}
---

Retorne apenas JSON valido no schema.
""".strip()
        return [
            {
                "role": "system",
                "content": (
                    "Voce e Consultor Legislativo da Camara dos Deputados. "
                    "Responda apenas em JSON valido e nao exponha cadeia de raciocinio."
                ),
            },
            {"role": "user", "content": prompt},
        ]

    if mode == "dvs":
        identificacao = build_identificacao_header(doc, doc_text, destaque_mode=mode)
        tramitacao_link = infer_tramitacao_link(doc_text, parecer_summary)
        tramitacao_txt = tramitacao_link or "Nao informado."
        proposicao_context = normalize_multiline(parecer_summary)
        if not proposicao_context:
            proposicao_context = "Nao fornecido separadamente nesta chamada."

        prompt = f"""
PAPEL
Voce e um Consultor Legislativo da Camara dos Deputados, especialista em Direito Constitucional,
Processo Legislativo e na area tematica especifica do destaque em analise.

TAREFA
Elabore uma explicacao objetiva do Destaque para Votacao em Separado (DVS) anexo.

INSUMOS DISPONIVEIS NESTA CHAMADA
1. Texto integral do DVS: fornecido abaixo.
2. Contexto da proposicao principal: fornecido abaixo.
3. Link da ficha de tramitacao: {tramitacao_txt}

INSTRUCOES
1. Foco exclusivo no texto normativo. Explique o que o DVS pretende suprimir ou destacar
na proposicao principal. Nao explique justificacao ou objetivo politico.
2. Resolva as remissoes. Se o dispositivo mencionar outras leis ou artigos do proprio projeto,
localize o teor no material fornecido e, quando necessario, use referencia legislativa consolidada
no portal planalto.gov.br/legislacao. Torne a explicacao autocontida.
3. Formato: no maximo dois paragrafos, em prosa corrida. O primeiro paragrafo DEVE iniciar
exatamente com "{identificacao}".
4. Tom: profissional e direto, acessivel a parlamentares, sem jargoes desnecessarios.
5. Pense passo a passo internamente antes de redigir (sem expor raciocinio).
6. Evite estruturas introdutorias dispensaveis apos a identificacao (ex.: "trata da votacao em separado do...").
7. Nao use marcadores editoriais entre colchetes, como [suprime], [substitui], [condiciona].

Contexto da proposicao principal:
---
{proposicao_context}
---

Texto integral do DVS em analise:
---
{doc_text}
---

Retorne apenas JSON valido no schema.
""".strip()
        return [
            {
                "role": "system",
                "content": (
                    "Voce e Consultor Legislativo da Camara dos Deputados. "
                    "Responda apenas em JSON valido e nao exponha cadeia de raciocinio."
                ),
            },
            {"role": "user", "content": prompt},
        ]

    prompt = f"""
Voce ira sintetizar um objeto alterador com base no texto integral do documento e no parecer de referencia.

Regras obrigatorias:
- Produza exatamente 2 paragrafos curtos e objetivos em portugues juridico claro.
- Nao reproduza perguntas no texto final.
- Nao use citacao literal, aspas, bullets ou numeracao.
- O primeiro paragrafo deve explicar o que e o destaque/emenda e qual alteracao proposicional ele faz.
- O segundo paragrafo deve explicar o efeito sobre os dispositivos ou sobre o conteudo normativo afetado.
- Mantenha estilo sintetico, no nivel do exemplo fornecido pelo usuario.
- Evite estruturas introdutorias dispensaveis e nao use marcadores editoriais entre colchetes.

Metadados do documento:
- arquivo: {doc.name}
- objeto alterador esperado: {doc.object_label}
- tipo esperado: {doc.object_type}

Contexto do parecer da comissao:
{parecer_summary}

Documento alvo (texto integral):
---
{doc_text}
---

Retorne apenas JSON valido no schema.
""".strip()
    return [
        {
            "role": "system",
            "content": (
                "Voce e especialista em tecnica legislativa e redacao parlamentar. "
                "Responda exclusivamente em JSON valido."
            ),
        },
        {"role": "user", "content": prompt},
    ]


def build_item_messages_compact(
    doc: InputDoc,
    doc_text: str,
    parecer_summary: str,
    destaque_mode: str = "",
) -> List[Dict[str, str]]:
    mode = normalize_ws(destaque_mode).lower()
    if mode == "destaque_emenda":
        identificacao = build_identificacao_header(doc, doc_text, destaque_mode=mode)
        tramitacao_link = infer_tramitacao_link(doc_text, parecer_summary)
        tramitacao_txt = tramitacao_link or "Nao informado."
        context_short = normalize_ws(parecer_summary) or "Nao fornecido separadamente nesta chamada."

        prompt = f"""
Explique o destaque de emenda em no maximo dois paragrafos, com foco exclusivo no texto normativo.

Regras obrigatorias:
- Nao explique justificacao da emenda.
- Explique o que a emenda vinculada pretende alterar/incluir na proposicao principal.
- Resolva remissoes a leis/dispositivos usando o material disponivel e, se preciso,
  referencia consolidada em planalto.gov.br/legislacao.
- O primeiro paragrafo DEVE iniciar com "{identificacao}".
- Pense passo a passo internamente (sem expor o raciocinio).
- Evite estruturas introdutorias dispensaveis apos a identificacao.
- Nao use marcadores editoriais entre colchetes, como [suprime], [substitui], [condiciona].

Contexto da proposicao principal (resumo):
{context_short}

Link de tramitacao:
{tramitacao_txt}

Texto integral do destaque:
{doc_text}

Retorne somente JSON valido no schema.
""".strip()

        return [
            {
                "role": "system",
                "content": "Responda apenas em JSON valido, em tom tecnico-legislativo, sem raciocinio explicito.",
            },
            {"role": "user", "content": prompt},
        ]

    if mode == "dvs":
        identificacao = build_identificacao_header(doc, doc_text, destaque_mode=mode)
        tramitacao_link = infer_tramitacao_link(doc_text, parecer_summary)
        tramitacao_txt = tramitacao_link or "Nao informado."
        context_short = normalize_ws(parecer_summary) or "Nao fornecido separadamente nesta chamada."

        prompt = f"""
Explique o DVS em no maximo dois paragrafos, com foco exclusivo no texto normativo.

Regras obrigatorias:
- Nao explique justificacao nem objetivo politico.
- Resolva remissoes a leis/dispositivos usando o material disponivel e, se preciso,
  referencia consolidada em planalto.gov.br/legislacao.
- O primeiro paragrafo DEVE iniciar com "{identificacao}".
- O texto deve descrever o objeto do DVS e o efeito normativo da supressao/votacao em separado.
- Pense passo a passo internamente (sem expor o raciocinio).
- Evite estruturas introdutorias dispensaveis apos a identificacao.
- Nao use marcadores editoriais entre colchetes, como [suprime], [substitui], [condiciona].

Contexto da proposicao principal (resumo):
{context_short}

Link de tramitacao:
{tramitacao_txt}

Texto integral do DVS:
{doc_text}

Retorne somente JSON valido no schema.
""".strip()

        return [
            {
                "role": "system",
                "content": "Responda apenas em JSON valido, em tom tecnico-legislativo, sem raciocinio explicito.",
            },
            {"role": "user", "content": prompt},
        ]

    compact_summary = normalize_ws(parecer_summary)
    prompt = f"""
Sintetize o documento integral abaixo em exatamente 2 paragrafos curtos.

Paragrafo 1:
- Identifique o que e o destaque/emenda e qual alteracao ele pretende fazer.

Paragrafo 2:
- Explique o efeito da alteracao sobre os dispositivos/conteudo normativo atingido.

Restricoes:
- Sem citacao literal.
- Sem perguntas explicitas, sem bullets, sem numeracao.
- Linguagem juridica objetiva.
- Evite estruturas introdutorias dispensaveis e nao use marcadores editoriais entre colchetes.

Metadados:
arquivo={doc.name}; objeto_esperado={doc.object_label}; tipo_esperado={doc.object_type}
contexto_parecer={compact_summary}

Documento integral:
{doc_text}

Retorne apenas JSON valido no schema.
""".strip()

    return [
        {
            "role": "system",
            "content": "Responda apenas JSON valido e sem texto fora do schema.",
        },
        {"role": "user", "content": prompt},
    ]


def build_chunk_messages(
    doc: InputDoc,
    chunk_text: str,
    parecer_summary: str,
    chunk_index: int,
    total_chunks: int,
    destaque_mode: str = "",
) -> List[Dict[str, str]]:
    mode = normalize_ws(destaque_mode).lower()
    if mode == "destaque_emenda":
        identificacao = build_identificacao_header(doc, chunk_text, destaque_mode=mode)
        context_short = normalize_ws(parecer_summary) or "Nao fornecido separadamente nesta chamada."
        prompt = f"""
Documento em analise: Destaque de Emenda ({doc.name})
Identificacao de referencia: {identificacao}
Trecho {chunk_index}/{total_chunks}

Tarefa deste passo:
- Resumir apenas este trecho em ate 6 frases curtas.
- Focar no texto normativo da emenda vinculada e no que sera alterado/incluido na proposicao.
- Resolver remissoes quando identificaveis no proprio trecho/contexto.
- Nao tratar de justificacao.
- Nao usar marcadores editoriais entre colchetes.

Contexto da proposicao principal:
{context_short}

Trecho:
---
{chunk_text}
---

Retorne somente JSON valido no schema.
""".strip()
        return [
            {
                "role": "system",
                "content": "Voce resume destaque de emenda por trecho e responde apenas em JSON valido.",
            },
            {"role": "user", "content": prompt},
        ]

    if mode == "dvs":
        identificacao = build_identificacao_header(doc, chunk_text, destaque_mode=mode)
        context_short = normalize_ws(parecer_summary) or "Nao fornecido separadamente nesta chamada."
        prompt = f"""
Documento em analise: DVS ({doc.name})
Identificacao de referencia: {identificacao}
Trecho {chunk_index}/{total_chunks}

Tarefa deste passo:
- Resumir apenas este trecho em ate 6 frases curtas.
- Focar no texto normativo do que sera suprimido/destacado para votacao em separado.
- Resolver remissoes quando identificaveis no proprio trecho/contexto.
- Nao tratar de justificacao politica.
- Nao usar marcadores editoriais entre colchetes.

Contexto da proposicao principal:
{context_short}

Trecho do DVS:
---
{chunk_text}
---

Retorne somente JSON valido no schema.
""".strip()
        return [
            {
                "role": "system",
                "content": "Voce resume DVS por trecho e responde apenas em JSON valido.",
            },
            {"role": "user", "content": prompt},
        ]

    prompt = f"""
Documento: {doc.name}
Objeto esperado: {doc.object_label}
Trecho {chunk_index}/{total_chunks}

Contexto do parecer:
{parecer_summary}

Instrucao:
- Resuma apenas este trecho em ate 6 frases curtas.
- Explique pontos materiais sobre o que o objeto alterador pretende mudar e o efeito normativo associado.
- Nao use citacoes literais.
- Retorne apenas JSON valido no schema.

Trecho:
---
{chunk_text}
---
""".strip()
    return [
        {
            "role": "system",
            "content": "Voce resume documentos legislativos e responde somente com JSON valido.",
        },
        {"role": "user", "content": prompt},
    ]


def build_condense_messages(doc: InputDoc, block_text: str, parecer_summary: str) -> List[Dict[str, str]]:
    prompt = f"""
Documento: {doc.name}
Objeto esperado: {doc.object_label}

Contexto do parecer:
{parecer_summary}

Ha multiplos resumos parciais abaixo. Condense-os em um unico resumo tecnico curto,
sem citacoes literais, focando no nucleo da alteracao e no efeito sobre dispositivos.
Retorne apenas JSON valido no schema.

Resumos parciais:
---
{block_text}
---
""".strip()
    return [
        {
            "role": "system",
            "content": "Voce consolida resumos legislativos e responde somente com JSON valido.",
        },
        {"role": "user", "content": prompt},
    ]


def build_item_from_chunk_messages(
    doc: InputDoc,
    chunk_summary_text: str,
    parecer_summary: str,
    destaque_mode: str = "",
) -> List[Dict[str, str]]:
    mode = normalize_ws(destaque_mode).lower()
    if mode == "destaque_emenda":
        identificacao = build_identificacao_header(doc, chunk_summary_text, destaque_mode=mode)
        tramitacao_link = infer_tramitacao_link(chunk_summary_text, parecer_summary)
        tramitacao_txt = tramitacao_link or "Nao informado."
        context_short = normalize_ws(parecer_summary) or "Nao fornecido separadamente nesta chamada."
        prompt = f"""
Consolide os resumos de chunks em explicacao final de destaque de emenda.

Instrucoes obrigatorias:
- No maximo dois paragrafos em prosa corrida.
- O primeiro paragrafo DEVE iniciar com "{identificacao}".
- Explicar somente o texto normativo: o que a emenda vinculada altera/inclui e o efeito normativo.
- Nao abordar justificacao.
- Resolver remissoes para tornar a explicacao autocontida.
- Pense passo a passo internamente (sem expor raciocinio).
- Evitar estruturas introdutorias dispensaveis apos a identificacao.
- Nao usar marcadores editoriais entre colchetes.

Contexto da proposicao principal:
{context_short}

Link de tramitacao:
{tramitacao_txt}

Resumos consolidados:
---
{chunk_summary_text}
---

Retorne apenas JSON valido no schema.
""".strip()
        return [
            {
                "role": "system",
                "content": "Voce consolida explicacoes de destaque de emenda e responde apenas com JSON valido.",
            },
            {"role": "user", "content": prompt},
        ]

    if mode == "dvs":
        identificacao = build_identificacao_header(doc, chunk_summary_text, destaque_mode=mode)
        tramitacao_link = infer_tramitacao_link(chunk_summary_text, parecer_summary)
        tramitacao_txt = tramitacao_link or "Nao informado."
        context_short = normalize_ws(parecer_summary) or "Nao fornecido separadamente nesta chamada."
        prompt = f"""
Consolide os resumos de chunks em explicacao final de DVS.

Instrucoes obrigatorias:
- No maximo dois paragrafos em prosa corrida.
- O primeiro paragrafo DEVE iniciar com "{identificacao}".
- Explicar somente o texto normativo: objeto do DVS e efeito normativo da supressao/votacao em separado.
- Nao abordar justificacao nem objetivo politico.
- Resolver remissoes para tornar a explicacao autocontida.
- Pense passo a passo internamente (sem expor raciocinio).
- Evitar estruturas introdutorias dispensaveis apos a identificacao.
- Nao usar marcadores editoriais entre colchetes.

Contexto da proposicao principal:
{context_short}

Link de tramitacao:
{tramitacao_txt}

Resumos consolidados:
---
{chunk_summary_text}
---

Retorne apenas JSON valido no schema.
""".strip()
        return [
            {
                "role": "system",
                "content": "Voce consolida explicacoes de DVS e responde apenas com JSON valido.",
            },
            {"role": "user", "content": prompt},
        ]

    prompt = f"""
Voce ira produzir a sintese final de um objeto alterador com base em resumos de chunks.

Regras obrigatorias:
- Exatamente 2 paragrafos curtos.
- Sem perguntas explicitas, sem citacao literal, sem bullets.
- Paragrafo 1: o que e o destaque/emenda e qual alteracao pretende realizar.
- Paragrafo 2: efeito sobre os dispositivos/conteudo normativo afetado.
- Texto sintetico e objetivo.
- Evite estruturas introdutorias dispensaveis e nao use marcadores editoriais entre colchetes.

Metadados:
- arquivo: {doc.name}
- objeto alterador esperado: {doc.object_label}
- tipo esperado: {doc.object_type}

Contexto do parecer:
{parecer_summary}

Resumos parciais consolidados:
---
{chunk_summary_text}
---

Retorne apenas JSON valido no schema.
""".strip()
    return [
        {
            "role": "system",
            "content": "Voce e analista legislativo e responde apenas em JSON valido.",
        },
        {"role": "user", "content": prompt},
    ]


def sanitize_paragraph_output(text: Any) -> str:
    value = normalize_ws(text)
    value = re.sub(r"^[-*\d\.)\s]+", "", value).strip()
    return value


def cleanup_dispensable_structures(
    text: str,
    *,
    protected_prefix: str = "",
) -> Tuple[str, List[str]]:
    source = normalize_ws(text)
    if not source:
        return source, []

    warnings: List[str] = []
    prefix = normalize_ws(protected_prefix)
    prefix_part = ""
    body = source

    if prefix and body.casefold().startswith(prefix.casefold()):
        prefix_part = body[: len(prefix)]
        body = body[len(prefix) :]
        body = re.sub(r"^[\s\-:;,.]+", "", body)

    cleaned_body = body
    marker_match_count = len(DISPENSABLE_BRACKET_MARKER_RE.findall(cleaned_body))
    if marker_match_count > 0:
        cleaned_body = DISPENSABLE_BRACKET_MARKER_RE.sub("", cleaned_body)
        warnings.append(f"editorial_bracket_marker:{marker_match_count}")

    for tag, pattern, repl in DISPENSABLE_LEADING_RULES:
        updated, count = pattern.subn(repl, cleaned_body, count=1)
        if count > 0:
            cleaned_body = updated
            warnings.append(f"{tag}:{count}")

    for tag, pattern, repl in DISPENSABLE_INLINE_RULES:
        updated, count = pattern.subn(repl, cleaned_body)
        if count > 0:
            cleaned_body = updated
            warnings.append(f"{tag}:{count}")

    cleaned_body = re.sub(r"\s+", " ", cleaned_body)
    cleaned_body = re.sub(r"\s+([,.;:])", r"\1", cleaned_body)
    cleaned_body = re.sub(r"([,;:])(?=\S)", r"\1 ", cleaned_body)
    cleaned_body = re.sub(r"\s{2,}", " ", cleaned_body).strip(" ,;")
    if cleaned_body and cleaned_body[0].islower():
        cleaned_body = cleaned_body[0].upper() + cleaned_body[1:]

    if prefix_part:
        cleaned = prefix_part if not cleaned_body else f"{prefix_part} - {cleaned_body}"
    else:
        cleaned = cleaned_body

    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
    return cleaned, warnings


def coerce_item_payload(
    payload: Dict[str, Any],
    doc: InputDoc,
    *,
    paragraph_1_prefix: str = "",
) -> Tuple[str, str, str, str, List[str]]:
    p1 = sanitize_paragraph_output(payload.get("paragrafo_1", ""))
    p2 = sanitize_paragraph_output(payload.get("paragrafo_2", ""))
    if not p1 or not p2:
        raise AppError("Resposta da IA sem os dois paragrafos obrigatorios.")

    prefix = normalize_ws(paragraph_1_prefix)
    p1, p1_warnings = cleanup_dispensable_structures(p1, protected_prefix=prefix)
    p2, p2_warnings = cleanup_dispensable_structures(p2)
    style_warnings = p1_warnings + p2_warnings

    if prefix and not p1.casefold().startswith(prefix.casefold()):
        p1 = f"{prefix} - {p1}"

    objeto_resp = normalize_ws(payload.get("objeto_alterador", ""))
    tipo_resp = normalize_ws(payload.get("tipo_objeto", "")).upper()

    objeto = doc.object_label or objeto_resp or doc.name
    tipo = doc.object_type or tipo_resp or "SEM_TIPO"
    return p1, p2, objeto, tipo, style_warnings


def condense_chunk_summaries(
    client: Any,
    *,
    doc: InputDoc,
    chunk_summaries: List[str],
    parecer_summary: str,
    args: argparse.Namespace,
    logger: logging.Logger,
    pacer: Optional[RequestPacer],
    telemetry: Optional[BatchTelemetry],
) -> str:
    current = [normalize_ws(item) for item in chunk_summaries if normalize_ws(item)]
    if not current:
        raise AppError("Falha ao consolidar chunks: resumos vazios.")

    while len("\n".join(current)) > MAX_FINAL_CHUNK_SUMMARY_CHARS and len(current) > 1:
        next_stage: List[str] = []
        block_size = 8
        for i in range(0, len(current), block_size):
            block = current[i : i + block_size]
            block_text = "\n".join(f"- {line}" for line in block)
            messages = build_condense_messages(doc, block_text, parecer_summary)
            payload, used_model, _ = openai_json_request(
                client,
                messages=messages,
                schema=BLOCK_SCHEMA,
                schema_name="resumo_bloco_objeto_alterador",
                primary_model=args.openai_model,
                fallback_model=args.fallback_model,
                timeout_s=args.openai_timeout,
                retries=args.openai_retries,
                max_completion_tokens=300,
                logger=logger,
                request_label=f"OpenAI condensacao de blocos ({doc.name})",
                pacer=pacer,
                telemetry=telemetry,
            )
            summary = normalize_ws(payload.get("resumo_bloco", ""))
            if not summary:
                raise AppError(f"Condensacao de bloco retornou vazio para {doc.name}.")
            logger.debug("Condensacao de bloco concluida para %s com modelo %s", doc.name, used_model)
            next_stage.append(summary)
        current = next_stage

    return "\n".join(current)


def summarize_parecer(
    client: Any,
    *,
    parecer_doc: InputDoc,
    parecer_text: str,
    args: argparse.Namespace,
    logger: logging.Logger,
    pacer: Optional[RequestPacer],
    telemetry: Optional[BatchTelemetry],
) -> Tuple[str, str]:
    messages = build_parecer_messages(parecer_text)
    payload, used_model, _ = openai_json_request(
        client,
        messages=messages,
        schema=PARECER_SCHEMA,
        schema_name="parecer_contexto",
        primary_model=args.openai_model,
        fallback_model=args.fallback_model,
        timeout_s=args.openai_timeout,
        retries=args.openai_retries,
        max_completion_tokens=260,
        logger=logger,
        request_label=f"OpenAI resumo do parecer ({parecer_doc.name})",
        pacer=pacer,
        telemetry=telemetry,
    )

    summary = normalize_ws(payload.get("parecer_resumo", ""))
    if not summary:
        raise AppError("Resumo do parecer retornou vazio.")
    return summary, used_model


def summarize_single_doc(
    client: Any,
    *,
    doc: InputDoc,
    doc_text: str,
    parecer_summary: str,
    args: argparse.Namespace,
    logger: logging.Logger,
    pacer: Optional[RequestPacer],
    telemetry: Optional[BatchTelemetry],
) -> SummaryItem:
    used_chunk_fallback = False
    destaque_mode = classify_destaque_mode(doc, doc_text) if is_destaque_doc(doc) else "nao_destaque"
    expected_prefix = (
        build_identificacao_header(doc, doc_text, destaque_mode=destaque_mode)
        if destaque_mode in {"dvs", "destaque_emenda"}
        else ""
    )
    if destaque_mode in {"dvs", "destaque_emenda"}:
        logger.info(
            "Modo de analise detectado para %s: %s (prefixo esperado: %s)",
            doc.name,
            destaque_mode,
            expected_prefix,
        )

    # Regra principal: insistir no envio integral. Chunk fallback e excepcional.
    integral_attempts = max(1, int(args.integral_length_retries))
    model_used = ""
    last_length_error: Optional[OpenAIRequestFailed] = None

    for integral_attempt in range(1, integral_attempts + 1):
        if integral_attempt == 1:
            full_messages = build_item_messages(
                doc,
                doc_text,
                parecer_summary,
                destaque_mode=destaque_mode,
            )
        else:
            full_messages = build_item_messages_compact(
                doc,
                doc_text,
                parecer_summary,
                destaque_mode=destaque_mode,
            )

        try:
            payload, model_used, _ = openai_json_request(
                client,
                messages=full_messages,
                schema=ITEM_SCHEMA,
                schema_name="sintese_objeto_alterador",
                primary_model=args.openai_model,
                fallback_model=args.fallback_model,
                timeout_s=args.openai_timeout,
                retries=args.openai_retries,
                max_completion_tokens=args.openai_max_completion_tokens,
                logger=logger,
                request_label=f"OpenAI sintese integral ({doc.name}) tentativa {integral_attempt}/{integral_attempts}",
                pacer=pacer,
                telemetry=telemetry,
            )
            p1, p2, objeto, tipo, style_warnings = coerce_item_payload(
                payload,
                doc,
                paragraph_1_prefix=expected_prefix,
            )
            if style_warnings:
                logger.warning(
                    "ADVERTENCIA_ESTILO | arquivo=%s | estruturas_dispensaveis_removidas=%s",
                    doc.name,
                    ",".join(sorted(set(style_warnings))),
                )
            return SummaryItem(
                file_path=doc.path,
                file_name=doc.name,
                objeto_alterador=objeto,
                tipo_objeto=tipo,
                numero_objeto=doc.object_number,
                paragrafo_1=p1,
                paragrafo_2=p2,
                model_used=model_used,
                status="ok",
                used_chunk_fallback=used_chunk_fallback,
            )
        except OpenAIRequestFailed as exc:
            if exc.info.kind != "length":
                raise
            last_length_error = exc
            if integral_attempt < integral_attempts:
                logger.warning(
                    "Length em %s (tentativa integral %d/%d). Repetindo em modo integral compacto.",
                    doc.name,
                    integral_attempt,
                    integral_attempts,
                )
                continue

    if not bool(args.enable_chunk_fallback):
        if last_length_error is not None:
            raise AppError(
                "Documento excedeu limite de contexto em todas as tentativas integrais e o fallback por chunks esta desativado."
            ) from last_length_error
        raise AppError("Falha inesperada antes do fallback por chunks.")

    logger.warning(
        "Documento %s excedeu limite apos tentativas integrais. Ativando fallback excepcional por chunks.",
        doc.name,
    )
    used_chunk_fallback = True

    chunks = split_text_chunks(doc_text, CHUNK_SIZE_CHARS)
    if not chunks:
        raise AppError(f"Nao foi possivel montar chunks para {doc.name}.")

    chunk_summaries: List[str] = []
    for idx, chunk in enumerate(chunks, start=1):
        messages = build_chunk_messages(
            doc,
            chunk,
            parecer_summary,
            idx,
            len(chunks),
            destaque_mode=destaque_mode,
        )
        payload, model_used, _ = openai_json_request(
            client,
            messages=messages,
            schema=CHUNK_SCHEMA,
            schema_name="resumo_chunk_objeto_alterador",
            primary_model=args.openai_model,
            fallback_model=args.fallback_model,
            timeout_s=args.openai_timeout,
            retries=args.openai_retries,
            max_completion_tokens=260,
            logger=logger,
            request_label=f"OpenAI resumo chunk {idx}/{len(chunks)} ({doc.name})",
            pacer=pacer,
            telemetry=telemetry,
        )
        chunk_summary = normalize_ws(payload.get("resumo_chunk", ""))
        if not chunk_summary:
            raise AppError(f"Chunk {idx}/{len(chunks)} retornou resumo vazio para {doc.name}.")
        chunk_summaries.append(f"[{idx}/{len(chunks)}] {chunk_summary}")

    summary_text = "\n".join(chunk_summaries)
    if len(summary_text) > MAX_FINAL_CHUNK_SUMMARY_CHARS:
        logger.info("Consolidando resumos de chunks para %s antes da sintese final.", doc.name)
        summary_text = condense_chunk_summaries(
            client,
            doc=doc,
            chunk_summaries=chunk_summaries,
            parecer_summary=parecer_summary,
            args=args,
            logger=logger,
            pacer=pacer,
            telemetry=telemetry,
        )

    final_messages = build_item_from_chunk_messages(
        doc,
        summary_text,
        parecer_summary,
        destaque_mode=destaque_mode,
    )
    payload, model_used, _ = openai_json_request(
        client,
        messages=final_messages,
        schema=ITEM_SCHEMA,
        schema_name="sintese_objeto_alterador_final",
        primary_model=args.openai_model,
        fallback_model=args.fallback_model,
        timeout_s=args.openai_timeout,
        retries=args.openai_retries,
        max_completion_tokens=args.openai_max_completion_tokens,
        logger=logger,
        request_label=f"OpenAI sintese final por chunks ({doc.name})",
        pacer=pacer,
        telemetry=telemetry,
    )

    p1, p2, objeto, tipo, style_warnings = coerce_item_payload(
        payload,
        doc,
        paragraph_1_prefix=expected_prefix,
    )
    if style_warnings:
        logger.warning(
            "ADVERTENCIA_ESTILO | arquivo=%s | estruturas_dispensaveis_removidas=%s",
            doc.name,
            ",".join(sorted(set(style_warnings))),
        )
    return SummaryItem(
        file_path=doc.path,
        file_name=doc.name,
        objeto_alterador=objeto,
        tipo_objeto=tipo,
        numero_objeto=doc.object_number,
        paragrafo_1=p1,
        paragrafo_2=p2,
        model_used=model_used,
        status="ok",
        used_chunk_fallback=used_chunk_fallback,
    )


def sort_items(items: Sequence[SummaryItem]) -> List[SummaryItem]:
    def key(item: SummaryItem) -> Tuple[str, int, str]:
        type_key = (item.tipo_objeto or "SEM_TIPO").upper()
        num_key = item.numero_objeto if item.numero_objeto >= 0 else 10**9
        name_key = item.file_name.lower()
        return (type_key, num_key, name_key)

    return sorted(items, key=key)


def resolve_output_docx_path(args: argparse.Namespace) -> Path:
    output_dir = resolve_project_path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    raw_output_docx = normalize_ws(args.output_docx)
    if raw_output_docx:
        out_path = Path(raw_output_docx).expanduser()
        if not out_path.is_absolute():
            out_path = output_dir / out_path
        out_path.parent.mkdir(parents=True, exist_ok=True)
        return out_path.resolve()

    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"relatorio_objetos_alteradores_{stamp}.docx"
    return (output_dir / filename).resolve()


def checkpoint_path_for_output(output_docx_path: Path) -> Path:
    return output_docx_path.parent / ".relatorio_objetos_alteradores.openai.checkpoint.json"


def write_docx_report(
    output_path: Path,
    *,
    parecer_doc: InputDoc,
    items: Sequence[SummaryItem],
) -> None:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise AppError(
            "python-docx nao encontrado. Instale com: python -m pip install python-docx"
        ) from exc

    doc = Document()
    doc.add_heading("Relatorio de Objetos Alteradores", level=1)
    doc.add_paragraph(f"Gerado em {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Parecer de referencia: {parecer_doc.name}")

    for item in items:
        header = doc.add_paragraph(style="List Number")
        title = f"{item.objeto_alterador} - {item.file_name}"
        run = header.add_run(title)
        run.bold = True

        doc.add_paragraph(item.paragrafo_1)
        doc.add_paragraph(item.paragrafo_2)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def summary_item_from_dict(raw: Dict[str, Any], fallback_doc: Optional[InputDoc] = None) -> SummaryItem:
    if fallback_doc is None:
        fallback_doc = InputDoc(path="", name="", object_type="SEM_TIPO", object_number=-1, object_label="SEM_TIPO")
    return SummaryItem(
        file_path=str(raw.get("file_path", fallback_doc.path)),
        file_name=str(raw.get("file_name", fallback_doc.name)),
        objeto_alterador=str(raw.get("objeto_alterador", fallback_doc.object_label)),
        tipo_objeto=str(raw.get("tipo_objeto", fallback_doc.object_type)),
        numero_objeto=int(raw.get("numero_objeto", fallback_doc.object_number) or fallback_doc.object_number),
        paragrafo_1=str(raw.get("paragrafo_1", "")),
        paragrafo_2=str(raw.get("paragrafo_2", "")),
        model_used=str(raw.get("model_used", "")),
        status=str(raw.get("status", "error")),
        used_chunk_fallback=bool(raw.get("used_chunk_fallback", False)),
        error=str(raw.get("error", "")),
    )


def build_manifest(paths: Sequence[Path]) -> List[Dict[str, Any]]:
    entries: List[Dict[str, Any]] = []
    for path in sorted({str(p.resolve()) for p in paths}):
        entries.append(build_file_signature(Path(path)))
    return entries


def build_checkpoint_payload(
    *,
    status: str,
    checkpoint_created_at: str,
    output_docx_path: Path,
    manifest: Sequence[Dict[str, Any]],
    parecer_path: Path,
    target_docs: Sequence[InputDoc],
    args: argparse.Namespace,
    parecer_summary: str,
    parecer_model_used: str,
    results_by_path: Dict[str, SummaryItem],
) -> Dict[str, Any]:
    target_paths = [doc.path for doc in target_docs]
    ok_count = sum(1 for doc in target_docs if (results_by_path.get(doc.path) and results_by_path[doc.path].status == "ok"))
    pending_count = sum(1 for doc in target_docs if not (results_by_path.get(doc.path) and results_by_path[doc.path].status == "ok"))

    return {
        "version": CHECKPOINT_VERSION,
        "prompt_version": PROMPT_VERSION,
        "status": status,
        "created_at": checkpoint_created_at,
        "updated_at": utc_now_iso(),
        "output_docx": str(output_docx_path.resolve()),
        "openai_model": str(args.openai_model),
        "fallback_model": str(args.fallback_model),
        "openai_timeout": int(args.openai_timeout),
        "openai_retries": int(args.openai_retries),
        "integral_length_retries": int(args.integral_length_retries),
        "enable_chunk_fallback": bool(args.enable_chunk_fallback),
        "parecer_file": str(parecer_path.resolve()),
        "parecer_summary": str(parecer_summary or ""),
        "parecer_model_used": str(parecer_model_used or ""),
        "input_manifest": list(manifest),
        "target_paths": list(target_paths),
        "total_targets": len(target_paths),
        "ok_count": ok_count,
        "pending_count": pending_count,
        "results_by_path": {path: asdict(item) for path, item in results_by_path.items()},
    }


def checkpoint_is_compatible(
    checkpoint: Dict[str, Any],
    *,
    manifest: Sequence[Dict[str, Any]],
    parecer_path: Path,
    target_docs: Sequence[InputDoc],
    args: argparse.Namespace,
) -> bool:
    if int(checkpoint.get("version", 0) or 0) != CHECKPOINT_VERSION:
        return False
    if int(checkpoint.get("prompt_version", 0) or 0) != PROMPT_VERSION:
        return False
    if str(checkpoint.get("parecer_file", "")) != str(parecer_path.resolve()):
        return False
    if list(checkpoint.get("input_manifest", [])) != list(manifest):
        return False
    if str(checkpoint.get("openai_model", "")).strip() != str(args.openai_model).strip():
        return False
    if str(checkpoint.get("fallback_model", "")).strip() != str(args.fallback_model).strip():
        return False

    cp_targets = [str(x) for x in (checkpoint.get("target_paths", []) or [])]
    current_targets = [doc.path for doc in target_docs]
    return cp_targets == current_targets


def load_checkpoint_results(
    checkpoint: Dict[str, Any],
    *,
    target_docs: Sequence[InputDoc],
) -> Dict[str, SummaryItem]:
    out: Dict[str, SummaryItem] = {}
    by_path = checkpoint.get("results_by_path", {})
    if not isinstance(by_path, dict):
        return out

    doc_by_path = {doc.path: doc for doc in target_docs}
    for path, raw in by_path.items():
        if path not in doc_by_path:
            continue
        if not isinstance(raw, dict):
            continue
        try:
            out[path] = summary_item_from_dict(raw, fallback_doc=doc_by_path[path])
        except Exception:
            continue
    return out


def write_checkpoint(path: Path, payload: Dict[str, Any]) -> None:
    write_json_atomic(path, payload, pretty=True)


def write_report(report: RunReport) -> None:
    write_json_atomic(REPORT_FILE, asdict(report), pretty=True)


def sanitize_runtime_args(args: argparse.Namespace) -> argparse.Namespace:
    args.openai_timeout = max(5, int(args.openai_timeout))
    args.openai_retries = max(1, int(args.openai_retries))
    args.openai_max_completion_tokens = max(64, int(args.openai_max_completion_tokens))
    args.openai_max_workers = max(1, int(args.openai_max_workers))
    args.openai_batch_size = max(1, int(args.openai_batch_size))
    args.openai_target_rpm = max(0, int(args.openai_target_rpm))
    args.openai_delay = max(0.0, float(args.openai_delay))
    args.integral_length_retries = max(1, int(args.integral_length_retries))
    return args


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Agente de sintese de objetos alteradores: PDF -> DOCX com contexto de parecer.",
    )
    parser.add_argument("--input-files", nargs="*", default=[], help="Arquivos PDF especificos.")
    parser.add_argument("--input-dirs", nargs="*", default=[], help="Pastas para buscar PDFs.")
    parser.add_argument("--parecer-file", default="", help="PDF de parecer de referencia.")
    parser.add_argument("--output-dir", default=str(SCRIPT_DIR), help="Pasta de saida para o DOCX.")
    parser.add_argument("--output-docx", default="", help="Nome/caminho do DOCX de saida.")

    parser.add_argument("--openai-model", default=DEFAULT_OPENAI_MODEL, help="Modelo principal da OpenAI.")
    parser.add_argument("--fallback-model", default=DEFAULT_FALLBACK_MODEL, help="Modelo fallback da OpenAI.")
    parser.add_argument("--openai-timeout", type=int, default=DEFAULT_OPENAI_TIMEOUT_S, help="Timeout por chamada OpenAI (s).")
    parser.add_argument("--openai-retries", type=int, default=DEFAULT_OPENAI_RETRIES, help="Tentativas por chamada OpenAI.")
    parser.add_argument(
        "--openai-max-completion-tokens",
        type=int,
        default=DEFAULT_OPENAI_MAX_COMPLETION_TOKENS,
        help="Max completion tokens da chamada de sintese.",
    )

    parser.add_argument(
        "--openai-max-workers",
        type=int,
        default=DEFAULT_OPENAI_MAX_WORKERS,
        help="Numero maximo de workers paralelos para documentos alvo.",
    )
    parser.add_argument(
        "--openai-batch-size",
        type=int,
        default=DEFAULT_OPENAI_BATCH_SIZE,
        help="Quantidade de documentos por lote de processamento paralelo.",
    )
    parser.add_argument(
        "--openai-target-rpm",
        type=int,
        default=DEFAULT_OPENAI_TARGET_RPM,
        help="Limite alvo global de requests por minuto para pacing.",
    )
    parser.add_argument(
        "--openai-delay",
        type=float,
        default=DEFAULT_OPENAI_DELAY_S,
        help="Pausa entre lotes (segundos).",
    )

    parser.add_argument(
        "--integral-length-retries",
        type=int,
        default=DEFAULT_INTEGRAL_LENGTH_RETRIES,
        help="Quantas tentativas integrais por arquivo antes de chunk fallback.",
    )
    parser.add_argument(
        "--disable-chunk-fallback",
        action="store_true",
        help="Desativa fallback por chunks (falha apos tentativas integrais).",
    )

    parser.add_argument("--no-gui", action="store_true", help="Desativa selecao grafica.")
    parser.add_argument("--recursive", action="store_true", help="Busca recursiva em --input-dirs.")
    parser.add_argument("--verbose", action="store_true", help="Compatibilidade: modo verbose ja e forcado.")
    parser.add_argument("--quiet", action="store_true", help="Compatibilidade: ignorado (verbose forcado).")
    parser.add_argument("--debug", action="store_true", help="Ativa ainda mais detalhes tecnicos no modo verbose.")
    parser.add_argument("--log-file", default="", help="Arquivo opcional para salvar logs.")
    return parser


def run(args: argparse.Namespace, logger: logging.Logger) -> int:
    args.enable_chunk_fallback = not bool(args.disable_chunk_fallback)
    args = sanitize_runtime_args(args)

    report = RunReport(
        started_at=utc_now_iso(),
        model_primary=str(args.openai_model or DEFAULT_OPENAI_MODEL).strip() or DEFAULT_OPENAI_MODEL,
        model_fallback=str(args.fallback_model or DEFAULT_FALLBACK_MODEL).strip() or DEFAULT_FALLBACK_MODEL,
    )

    checkpoint_created_at = utc_now_iso()

    try:
        args = maybe_collect_gui_files(args, logger)
        args = sanitize_runtime_args(args)

        discovered = discover_pdf_inputs(args.input_files or [], args.input_dirs or [], bool(args.recursive))
        if not discovered:
            if list(args.input_files or []) or list(args.input_dirs or []):
                raise AppError("Nenhum PDF encontrado nas entradas fornecidas.")

            logger.info("Sem selecao explicita. Aplicando fallback para PDFs da pasta do script (nao recursivo).")
            discovered = fallback_script_dir_pdfs()

        if not discovered:
            raise AppError("Nenhum PDF encontrado para processamento.")

        report.input_files = [str(path.resolve()) for path in discovered]
        logger.info("PDFs detectados: %d", len(discovered))

        parecer_path = detect_parecer_path(discovered, args.parecer_file)
        parecer_doc = build_input_doc(parecer_path, kind="parecer")
        report.parecer_file = parecer_doc.path

        targets: List[Path] = [path for path in discovered if path.resolve() != parecer_path.resolve()]
        if not targets:
            raise AppError("Nenhum PDF de objeto alterador foi encontrado apos excluir o parecer.")

        target_docs = [build_input_doc(path, kind="objeto") for path in targets]
        report.total_targets = len(target_docs)

        output_docx_path = resolve_output_docx_path(args)
        checkpoint_path = checkpoint_path_for_output(output_docx_path)

        report.notes.append(
            "config: workers={} batch={} rpm={} delay={}s integral_retries={} chunk_fallback={}".format(
                args.openai_max_workers,
                args.openai_batch_size,
                args.openai_target_rpm,
                args.openai_delay,
                args.integral_length_retries,
                args.enable_chunk_fallback,
            )
        )
        report.notes.append("style_guard: filtros de estruturas dispensaveis e marcadores editoriais ativados")
        logger.info("Guardrail de estilo ativo: estruturas dispensaveis serao removidas e sinalizadas com ADVERTENCIA_ESTILO.")

        manifest_paths = [parecer_path] + targets
        manifest = build_manifest(manifest_paths)

        checkpoint_raw = read_json_dict(checkpoint_path)
        checkpoint_compatible = False
        if checkpoint_raw:
            checkpoint_compatible = checkpoint_is_compatible(
                checkpoint_raw,
                manifest=manifest,
                parecer_path=parecer_path,
                target_docs=target_docs,
                args=args,
            )
            if checkpoint_compatible:
                cp_output = normalize_ws(checkpoint_raw.get("output_docx", ""))
                if cp_output:
                    cp_output_path = Path(cp_output).expanduser()
                    if not cp_output_path.is_absolute():
                        cp_output_path = resolve_project_path(cp_output)
                    output_docx_path = cp_output_path.resolve()

        report.output_docx = str(output_docx_path)
        report.checkpoint_file = str(checkpoint_path)

        if checkpoint_raw and not checkpoint_compatible:
            backup = make_backup(checkpoint_path, label="stale_checkpoint")
            if backup is not None:
                logger.info("Checkpoint antigo incompativel: backup criado em %s", backup)

        if output_docx_path.exists() and not checkpoint_compatible:
            backup = make_backup(output_docx_path, label="startup_backup")
            if backup is not None:
                logger.info("Backup do DOCX existente criado em %s", backup)

        results_by_path: Dict[str, SummaryItem] = {}
        parecer_summary = ""
        parecer_model_used = ""

        if checkpoint_raw and checkpoint_compatible:
            report.resumed_from_checkpoint = True
            checkpoint_created_at = str(checkpoint_raw.get("created_at", checkpoint_created_at) or checkpoint_created_at)
            results_by_path = load_checkpoint_results(checkpoint_raw, target_docs=target_docs)
            parecer_summary = str(checkpoint_raw.get("parecer_summary", "") or "")
            parecer_model_used = str(checkpoint_raw.get("parecer_model_used", "") or "")
            done_ok = sum(1 for doc in target_docs if (results_by_path.get(doc.path) and results_by_path[doc.path].status == "ok"))
            logger.info("[resume] checkpoint carregado: %d/%d itens ja concluidos.", done_ok, len(target_docs))

        checkpoint_payload = build_checkpoint_payload(
            status="running",
            checkpoint_created_at=checkpoint_created_at,
            output_docx_path=output_docx_path,
            manifest=manifest,
            parecer_path=parecer_path,
            target_docs=target_docs,
            args=args,
            parecer_summary=parecer_summary,
            parecer_model_used=parecer_model_used,
            results_by_path=results_by_path,
        )
        write_checkpoint(checkpoint_path, checkpoint_payload)

        openai_key = resolve_openai_key()
        if not openai_key:
            raise AppError(
                "OPENAI_API_KEY ausente. Defina no ambiente/.env ou em arquivo de chave suportado."
            )

        try:
            from openai import OpenAI  # type: ignore
        except Exception as exc:
            raise AppError("Pacote openai nao encontrado. Instale com: python -m pip install openai") from exc

        pacer = RequestPacer(target_rpm=args.openai_target_rpm)

        if not normalize_ws(parecer_summary):
            logger.info("Extraindo texto do parecer: %s", parecer_doc.name)
            parecer_text_raw = extract_pdf_text(Path(parecer_doc.path))
            parecer_text = normalize_multiline(parecer_text_raw)
            if not normalize_ws(parecer_text):
                raise AppError(f"Texto do parecer esta vazio/ilegivel: {parecer_doc.name}")

            logger.info("Gerando resumo-base do parecer via OpenAI...")
            parecer_client = OpenAI(api_key=openai_key, max_retries=0)
            parecer_summary, parecer_model_used = summarize_parecer(
                parecer_client,
                parecer_doc=parecer_doc,
                parecer_text=parecer_text,
                args=args,
                logger=logger,
                pacer=pacer,
                telemetry=None,
            )
            report.notes.append(f"Modelo usado no resumo do parecer: {parecer_model_used}")

            checkpoint_payload = build_checkpoint_payload(
                status="running",
                checkpoint_created_at=checkpoint_created_at,
                output_docx_path=output_docx_path,
                manifest=manifest,
                parecer_path=parecer_path,
                target_docs=target_docs,
                args=args,
                parecer_summary=parecer_summary,
                parecer_model_used=parecer_model_used,
                results_by_path=results_by_path,
            )
            write_checkpoint(checkpoint_path, checkpoint_payload)

        pending_docs = [
            doc
            for doc in target_docs
            if not (results_by_path.get(doc.path) and results_by_path[doc.path].status == "ok")
        ]

        if pending_docs:
            logger.info(
                "Processamento paralelo: pendentes=%d | workers=%d | batch=%d | rpm=%d",
                len(pending_docs),
                args.openai_max_workers,
                args.openai_batch_size,
                args.openai_target_rpm,
            )

        def process_doc_worker(doc: InputDoc, telemetry: BatchTelemetry) -> SummaryItem:
            doc_started = time.monotonic()
            telemetry.on_doc_start()
            logger.info("DOC_START | arquivo=%s", doc.name)
            client = OpenAI(api_key=openai_key, max_retries=0)
            try:
                doc_text_raw = extract_pdf_text(Path(doc.path))
                doc_text = normalize_multiline(doc_text_raw)
                if not normalize_ws(doc_text):
                    raise AppError(f"Texto vazio/ilegivel no arquivo {doc.name}")
                item = summarize_single_doc(
                    client,
                    doc=doc,
                    doc_text=doc_text,
                    parecer_summary=parecer_summary,
                    args=args,
                    logger=logger,
                    pacer=pacer,
                    telemetry=telemetry,
                )
                telemetry.on_doc_finish(ok=True, elapsed_s=(time.monotonic() - doc_started))
                return item
            except Exception:
                telemetry.on_doc_finish(ok=False, elapsed_s=(time.monotonic() - doc_started))
                raise

        total_pending = len(pending_docs)
        completed_pending = 0

        for start in range(0, total_pending, max(1, int(args.openai_batch_size))):
            end = min(start + max(1, int(args.openai_batch_size)), total_pending)
            batch_docs = pending_docs[start:end]
            if not batch_docs:
                continue

            batch_workers = min(max(1, int(args.openai_max_workers)), len(batch_docs))
            logger.info(
                "Lote pendente %d-%d/%d | workers=%d",
                start + 1,
                end,
                total_pending,
                batch_workers,
            )
            batch_label = f"{start + 1}-{end}/{total_pending}"
            batch_telemetry = BatchTelemetry(label=batch_label, total_docs=len(batch_docs))
            batch_hb_stop, batch_hb_thread = start_batch_progress_heartbeat(
                batch_telemetry,
                logger,
                interval_s=BATCH_HEARTBEAT_INTERVAL_S,
            )
            batch_started_at = time.monotonic()

            try:
                with ThreadPoolExecutor(max_workers=batch_workers) as executor:
                    future_map = {
                        executor.submit(process_doc_worker, doc, batch_telemetry): doc
                        for doc in batch_docs
                    }
                    for future in as_completed(future_map):
                        doc = future_map[future]
                        try:
                            item = future.result()
                            logger.info(
                                "Concluido %s | modelo=%s | chunk_fallback=%s",
                                doc.name,
                                item.model_used,
                                item.used_chunk_fallback,
                            )
                        except Exception as exc:  # noqa: BLE001
                            err_msg = normalize_ws(str(exc)) or "erro desconhecido"
                            logger.error("Falha em %s: %s", doc.name, err_msg)
                            item = SummaryItem(
                                file_path=doc.path,
                                file_name=doc.name,
                                objeto_alterador=doc.object_label,
                                tipo_objeto=doc.object_type,
                                numero_objeto=doc.object_number,
                                paragrafo_1="",
                                paragrafo_2="",
                                model_used="",
                                status="error",
                                used_chunk_fallback=False,
                                error=err_msg,
                            )

                        results_by_path[doc.path] = item
                        completed_pending += 1

                        checkpoint_payload = build_checkpoint_payload(
                            status="running",
                            checkpoint_created_at=checkpoint_created_at,
                            output_docx_path=output_docx_path,
                            manifest=manifest,
                            parecer_path=parecer_path,
                            target_docs=target_docs,
                            args=args,
                            parecer_summary=parecer_summary,
                            parecer_model_used=parecer_model_used,
                            results_by_path=results_by_path,
                        )
                        write_checkpoint(checkpoint_path, checkpoint_payload)
                        logger.info(
                            "CHECKPOINT_SAVED | progresso_pendentes=%d/%d | arquivo=%s | status=%s",
                            completed_pending,
                            total_pending,
                            doc.name,
                            item.status,
                        )
            except KeyboardInterrupt as exc:
                checkpoint_payload = build_checkpoint_payload(
                    status="interrupted",
                    checkpoint_created_at=checkpoint_created_at,
                    output_docx_path=output_docx_path,
                    manifest=manifest,
                    parecer_path=parecer_path,
                    target_docs=target_docs,
                    args=args,
                    parecer_summary=parecer_summary,
                    parecer_model_used=parecer_model_used,
                    results_by_path=results_by_path,
                )
                write_checkpoint(checkpoint_path, checkpoint_payload)
                raise AppError(
                    "Execucao interrompida pelo usuario. Checkpoint mantido para retomada."
                ) from exc
            finally:
                batch_hb_stop.set()
                if batch_hb_thread is not None:
                    batch_hb_thread.join(timeout=0.2)
                snap = batch_telemetry.snapshot()
                batch_elapsed = max(0.0, time.monotonic() - batch_started_at)
                logger.info(
                    "BATCH_SUMMARY | lote=%s | elapsed=%.2fs | docs_ok=%d | docs_err=%d | docs_done=%d/%d | "
                    "api_started=%d | api_ok=%d | api_err=%d | retries=%d | qps=%.3f | avg_api_s=%.2f | avg_doc_s=%.2f | api_error_rate=%.2f%%",
                    snap["label"],
                    batch_elapsed,
                    snap["docs_ok"],
                    snap["docs_err"],
                    snap["docs_done"],
                    snap["total_docs"],
                    snap["api_started"],
                    snap["api_success"],
                    snap["api_failed"],
                    snap["api_retries"],
                    snap["qps_success"],
                    snap["avg_api_s"],
                    snap["avg_doc_s"],
                    snap["api_error_rate_pct"],
                )

            if end < total_pending and float(args.openai_delay) > 0:
                time.sleep(float(args.openai_delay))

        result_items = [
            results_by_path.get(doc.path)
            or SummaryItem(
                file_path=doc.path,
                file_name=doc.name,
                objeto_alterador=doc.object_label,
                tipo_objeto=doc.object_type,
                numero_objeto=doc.object_number,
                paragrafo_1="",
                paragrafo_2="",
                model_used="",
                status="error",
                used_chunk_fallback=False,
                error="item ausente no checkpoint interno",
            )
            for doc in target_docs
        ]

        report.items = [asdict(item) for item in result_items]

        successful_items = [item for item in result_items if item.status == "ok"]
        failed_items = [item for item in result_items if item.status != "ok"]

        report.success_count = len(successful_items)
        report.failure_count = len(failed_items)

        if failed_items:
            report.errors.extend([f"{item.file_name}: {item.error}" for item in failed_items if item.error])

        if not successful_items:
            raise AppError("Nenhuma sintese valida foi gerada para os objetos alteradores.")

        ordered_items = sort_items(successful_items)
        write_docx_report(output_docx_path, parecer_doc=parecer_doc, items=ordered_items)

        final_status = "completed" if not failed_items else "partial"
        report.status = final_status
        report.finished_at = utc_now_iso()
        write_report(report)

        checkpoint_payload = build_checkpoint_payload(
            status=final_status,
            checkpoint_created_at=checkpoint_created_at,
            output_docx_path=output_docx_path,
            manifest=manifest,
            parecer_path=parecer_path,
            target_docs=target_docs,
            args=args,
            parecer_summary=parecer_summary,
            parecer_model_used=parecer_model_used,
            results_by_path=results_by_path,
        )
        write_checkpoint(checkpoint_path, checkpoint_payload)

        logger.info("DOCX gerado: %s", output_docx_path)

        if final_status == "completed" and checkpoint_path.exists():
            checkpoint_path.unlink()
            logger.info("Checkpoint removido apos conclusao total: %s", checkpoint_path)

        if failed_items:
            logger.warning(
                "Execucao parcial: %d sucesso(s), %d falha(s). Consulte %s",
                len(successful_items),
                len(failed_items),
                REPORT_FILE,
            )
            return 1

        logger.info("Execucao concluida com sucesso total. Itens no DOCX: %d", len(successful_items))
        return 0

    except Exception as exc:  # noqa: BLE001
        msg = normalize_ws(str(exc)) or "erro desconhecido"
        report.status = "failed"
        report.finished_at = utc_now_iso()
        report.errors.append(msg)
        write_report(report)
        logger.error("Falha de execucao: %s", msg)
        return 1


def main() -> int:
    parser = build_arg_parser()
    args = parser.parse_args()

    quiet_requested = bool(args.quiet)
    forced_verbose = True
    if quiet_requested:
        args.quiet = False

    logger = configure_standard_logging(
        "AGENTE_objetos_alteradores_pdf_para_docx",
        verbose=forced_verbose,
        quiet=False,
        debug=(bool(args.debug) or forced_verbose),
        log_file=str(args.log_file or ""),
    )
    install_print_logger_bridge(globals(), logger)
    logger.info("Logger em modo verbose continuo para acompanhamento em tempo real.")
    logger.info(
        "Config runtime log: verbose_forcado=%s | debug=%s | quiet_solicitado=%s",
        forced_verbose,
        (bool(args.debug) or forced_verbose),
        quiet_requested,
    )

    return run(args, logger)


if __name__ == "__main__":
    raise SystemExit(main())
