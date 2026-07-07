# -*- coding: utf-8 -*-
"""Conversão de rich text (Notion) em runs do Word."""
from __future__ import annotations

import re
import unicodedata
from dataclasses import replace
from typing import List, Optional

from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from . import config
from .notion_parser import RichText

# Cor azul padrão de hyperlink do Word (igual ao estilo "Hyperlink").
_COR_LINK = "0563C1"

# Domínios internos do Notion: nunca viram hyperlink no .docx (o público externo
# não acessa as bases). notion.site (páginas publicadas) fica de fora de propósito.
_DOMINIOS_NOTION = re.compile(
    r"^(notion://|https?://(www\.)?notion\.so\b|https?://app\.notion\.com\b)", re.I
)


def _link_interno_notion(href: str) -> bool:
    return bool(href) and (href.startswith("/") or _DOMINIOS_NOTION.match(href) is not None)


def resolver_fonte_publica(texto: str) -> Optional[str]:
    """Mapeia o texto de uma mention interna ("RICD - Art. 54") para a URL da
    fonte pública oficial, via config.FONTES_OFICIAIS. None se não mapeada.
    Callable que devolve falsy (ex. resolução TSE fora do mapa) conta como
    não mapeada — segue-se para o próximo padrão."""
    for padrao, template, _aliases in config.FONTES_OFICIAIS:
        m = re.search(padrao, texto or "")
        if m:
            url = template(m) if callable(template) else m.expand(template)
            if url:
                return url
    return None


# "art. 7º", "arts. 14 a 17" (âncora do primeiro), "art. 10-B" (letra entra
# na âncora do Planalto: #art10b)
_RE_ARTIGO = re.compile(r"\bArts?\.?\s*(\d+)(?:[ºo°])?(?:\s*[-–]\s*([A-Za-z])\b)?", re.I)

# Normas hospedadas no Planalto (ccivil_03): todas têm âncora nativa #artN.
_RE_PLANALTO = re.compile(r"^https?://(?:www\.)?planalto\.gov\.br/ccivil_03/", re.I)

# Citação textual de norma digitada no corpo (fora de mention), ex. "art. 3º,
# incisos II e III, da CF", "art. 22, XVI, da LC nº 64/1990", "art. 57-C, §
# 1º, I, da Lei nº 9.504/1997", "art. 299 do Código Eleitoral". Um padrão
# compilado POR NORMA de config.NORMAS_OFICIAIS: "art(s). N[-letra]" + miolo
# (incisos/§§/alíneas) + "da/do <designação da norma>". O miolo não cruza
# ; ( ) nem pode citar OUTRA norma no caminho ("art. 7º da LC 95/1998 ... da
# CF" não pode linkar a CF) — daí a lista negra de nomes de norma dentro do
# miolo; a designação da própria norma casa em grupo próprio, DEPOIS do miolo,
# e por isso não é bloqueada. Lista de artigos ("arts. 36-A, 37 e 57-D da Lei
# nº 9.504/1997") vira um único link com a âncora do PRIMEIRO artigo — mesma
# convenção de _RE_ARTIGO.
_MIOLO_CITACAO = (
    r"(?:(?!\b(?:lei|decreto|c[óo]digo|resolu[çc][ãa]o|res|emenda|regimento|"
    r"ricd|adct|lc|ec)\b)[^;()]){0,60}?"
)


def _padrao_citacao_norma(citacao_re: str) -> "re.Pattern":
    return re.compile(
        rf"\b(?:arts?\.|artigos?)\s*(?P<num>\d+)(?:[ºo°])?"
        rf"(?:\s*[-–‑]\s*(?P<letra>[A-Za-z])\b)?"
        rf"{_MIOLO_CITACAO}"
        rf"\bd[aeo]s?\s+(?:{citacao_re})",
        re.I,
    )


# Ações de controle concentrado no STF, por extenso ou sigla ("Ação Direta de
# Inconstitucionalidade número 3.685", "ADPF nº 492") -> página do processo no
# portal do STF (mesmo formato de URL dos documentos-modelo da casa).
_RE_ACAO_STF = re.compile(
    r"\b(?:(?P<ado>A[çc][ãa]o\s+Direta\s+de\s+Inconstitucionalidade\s+por\s+Omiss[ãa]o)"
    r"|(?P<adi>A[çc][ãa]o\s+Direta\s+de\s+Inconstitucionalidade)"
    r"|(?P<adpf>Argui[çc][ãa]o\s+de\s+Descumprimento\s+de\s+Preceito\s+Fundamental)"
    r"|(?P<adc>A[çc][ãa]o\s+Declarat[óo]ria\s+de\s+Constitucionalidade)"
    r"|(?P<sigla>ADI|ADO|ADC|ADPF))"
    r"(?:\s*\((?:ADI|ADO|ADC|ADPF)\))?"
    r"\s*(?:n[ºo°]\.?\s*|n[uú]mero\s+)?(?P<num>\d[\d.]*\d|\d)",
    re.I,
)


def _url_acao_stf(m: "re.Match") -> str:
    if m.group("sigla"):
        classe = m.group("sigla").upper()
    else:
        classe = next(c.upper() for c in ("ado", "adi", "adpf", "adc") if m.group(c))
    num = m.group("num").replace(".", "")
    return ("https://portal.stf.jus.br/processos/listarProcessos.asp"
            f"?classe={classe}&numeroProcesso={num}")


# "Súmula nº 73 do TSE" / "Súmula número 73 do Tribunal Superior Eleitoral"
# -> página da súmula no site do TSE.
_RE_SUMULA_TSE = re.compile(
    r"\bS[úu]mula\s*(?:n[ºo°]\.?\s*|n[uú]mero\s+)?(?P<num>\d+)"
    r"\s+do\s+(?:Tribunal\s+Superior\s+Eleitoral|TSE)\b(?:\s*\(TSE\))?",
    re.I,
)


def _url_sumula_tse(m: "re.Match") -> str:
    return ("https://www.tse.jus.br/legislacao/codigo-eleitoral/sumulas/"
            f"sumulas-do-tse/sumula-tse-n-{m.group('num')}")


def _montar_padroes_citacao() -> list:
    """(regex, callable(Match) -> URL ou None): citações com artigo de cada
    norma de config.NORMAS_OFICIAIS, ações do STF, súmulas do TSE e, por fim,
    as designações NUMERADAS de norma sem artigo ("a LC nº 64/1990 já
    oferece…" -> link sem âncora). A ordem não arbitra sobreposição — isso é
    papel do sort por (início, -fim) em linkificar_citacoes: "art. 37 da Lei
    nº 9.504/1997" começa antes e engole o match isolado de "Lei nº
    9.504/1997" contido nele."""
    pads = [
        (_padrao_citacao_norma(n.citacao_re), n.montar_url)
        for n in config.NORMAS_OFICIAIS
    ]
    pads += [(_RE_ACAO_STF, _url_acao_stf), (_RE_SUMULA_TSE, _url_sumula_tse)]
    pads += [
        (re.compile(rf"\b(?:{n.citacao_isolada_re})", re.I), n.montar_url)
        for n in config.NORMAS_OFICIAIS
        if n.citacao_isolada_re
    ]
    return pads


_PADROES_CITACAO = _montar_padroes_citacao()


def resolver_links_notion_texto(rich_list: List[RichText]) -> List[RichText]:
    """Hyperlink de TEXTO (não mention) apontando para página interna do Notion
    — o autor linkou a citação a uma base interna (ex. Vademecum), que o leitor
    do .docx não acessa. Texto que casa uma norma de config.FONTES_OFICIAIS
    ganha a URL oficial; senão o href é ZERADO ainda no início da cadeia, para
    a citação textual poder ser linkada por linkificar_citacoes (um href
    presente a faria ser pulada — caso real: "art. 54 do RICD" linkado à página
    do Vademecum saía como texto morto no parecer do PL 3.263/2023). Mentions
    não passam por aqui: as de norma são resolvidas por texto em add_runs e as
    demais já chegam com URL pública via notion_api.resolver_mentions_publicas."""
    out: List[RichText] = []
    for r in rich_list or []:
        if r.kind == "text" and _link_interno_notion((r.href or "").strip()):
            out.append(replace(r, href=resolver_fonte_publica(r.text)))
        else:
            out.append(r)
    return out


def canonicalizar_hrefs(rich_list: List[RichText]) -> List[RichText]:
    """URL "equivalente" que o autor linkou no Notion (página institucional do
    RICD, "compilado" do Planalto sem âncoras) vira a URL oficial canônica do
    padrão da casa (config.URLS_CANONICAS, mapa fechado e verificado). Sem a
    canonização, os passos seguintes não conseguem anexar âncora de artigo —
    achado do usuário (06/07/2026): "arts. 139, I, e 142 do RICD" abria a
    página institucional genérica."""
    out: List[RichText] = []
    for r in rich_list or []:
        href = (r.href or "").strip()
        canon = config.url_canonica(href) if href.startswith("http") else None
        if canon and canon != href.split("#")[0]:
            frag = href.split("#", 1)[1] if "#" in href else ""
            out.append(replace(r, href=canon + (f"#{frag}" if frag else "")))
        else:
            out.append(r)
    return out


def ancorar_links_por_texto(rich_list: List[RichText]) -> List[RichText]:
    """Antecipação da âncora que add_runs aplicaria no final: link http de
    norma SEM "#" cujo TEXTO cita o dispositivo ("art. 51, III, da
    Constituição Federal" linkado pelo autor no topo da norma) ganha a âncora
    JÁ NA CADEIA — os passos seguintes (supressão de parêntese remissivo,
    anáfora) enxergam a URL efetiva e comparam igualdade com as mentions
    (caso real da PRC Bancada Negra, 06/07/2026). Links cujo texto não cita
    artigo (designações, julgados) ficam intactos — a extensão de link e a
    retro-extensão continuam vendo-os sem âncora, como exigem."""
    out: List[RichText] = []
    for r in rich_list or []:
        href = (r.href or "").strip()
        if href.startswith("http") and "#" not in href and r.kind == "text":
            novo = _ancorar_dispositivo(href, r.text)
            out.append(replace(r, href=novo) if novo != href else r)
        else:
            out.append(r)
    return out


def fundir_runs_lisos(rich_list: List[RichText]) -> List[RichText]:
    """Une runs adjacentes SEM formatação e sem href (kind text). Um link
    interno zerado por resolver_links_notion_texto deixa a citação espalhada
    em runs contíguos ("…alínea 'a', do " + "Regimento Interno…" + ", compete…")
    que linkificar_citacoes — limitado a um run — não casaria."""
    out: List[RichText] = []
    for r in rich_list or []:
        liso = not (r.bold or r.italic or r.underline or r.href) and r.kind == "text"
        if liso and out:
            p = out[-1]
            if not (p.bold or p.italic or p.underline or p.href) and p.kind == "text":
                out[-1] = replace(p, text=p.text + r.text)
                continue
        out.append(r)
    return out


def linkificar_citacoes(rich_list: List[RichText]) -> List[RichText]:
    """Citações textuais em texto puro (sem href) ganham o hyperlink da fonte
    oficial — o trecho citado inteiro vira o texto do link, como nos
    documentos-modelo da casa: artigos de CF/RICD, ações de controle
    concentrado no STF (ADI/ADO/ADC/ADPF) e súmulas do TSE. Limitação
    assumida: só citações contidas num único run (o caso real; formatação no
    meio da citação a dividiria em runs)."""
    out: List[RichText] = []
    for r in rich_list or []:
        if r.href or r.kind != "text" or not r.text:
            out.append(r)
            continue
        matches = []
        for padrao, montar_url in _PADROES_CITACAO:
            matches += [(m.start(), m.end(), m, montar_url) for m in padrao.finditer(r.text)]
        matches.sort(key=lambda t: (t[0], -t[1]))
        pos = 0
        for ini, fim, m, montar_url in matches:
            if ini < pos:
                continue  # sobreposto a um match anterior
            url = montar_url(m)
            if not url:
                continue  # fonte não mapeada (ex. resolução TSE fora do mapa)
            if ini > pos:
                out.append(replace(r, text=r.text[pos:ini]))
            out.append(replace(r, text=m.group(0), href=url))
            pos = fim
        if pos == 0:
            out.append(r)
        elif pos < len(r.text):
            out.append(replace(r, text=r.text[pos:]))
    return out


def _eh_citacao_completa(texto: str) -> bool:
    """True se `texto` é, por inteiro, uma citação reconhecida por algum padrão
    de _PADROES_CITACAO com fonte mapeada (usado para decidir se colchetes ao
    redor podem ser removidos)."""
    s = (texto or "").strip()
    if not s:
        return False
    for padrao, montar_url in _PADROES_CITACAO:
        m = padrao.fullmatch(s)
        if m and montar_url(m):
            return True
    return False


# Conteúdo entre colchetes candidato à limpeza do passe B (texto puro).
_RE_COLCHETES = re.compile(r"\[([^\[\]]{2,80})\]")


def limpar_colchetes_citacoes(rich_list: List[RichText]) -> List[RichText]:
    """Remove colchetes editoriais colados em citações, nos dois formatos vindos
    do Notion: (A) o autor linkou o trecho e o "[" ficou DENTRO do texto do
    link, com o "]" abrindo o run seguinte (caso real: '[ADI nº 4.650' com href
    do STF + '], o Tribunal…') — remove o par SEM tocar o href; (B) citação
    textual entre colchetes sem link ('[ADI nº 4.650]') — remove os colchetes
    apenas se o conteúdo interno for uma citação completa reconhecida, expondo-a
    para linkificar_citacoes linkar por inteiro. Colchetes legítimos ('[sic]',
    '[Art. 12] § 4º' do bloco fixo da Resolução) não casam e ficam intactos."""
    out: List[RichText] = list(rich_list or [])
    # Passe A: par de colchetes abraçando um run já linkado pelo autor.
    for i, r in enumerate(out):
        if not (r.href or "").startswith("http") or not r.text:
            continue
        abre_no_link = r.text.startswith("[")
        abre_no_anterior = (
            not abre_no_link and i > 0 and not out[i - 1].href
            and (out[i - 1].text or "").endswith("[")
        )
        if not (abre_no_link or abre_no_anterior):
            continue
        fecha_no_link = r.text.endswith("]") and len(r.text) > 1
        j = next((k for k in range(i + 1, len(out)) if out[k].text), None)
        fecha_no_seguinte = (
            not fecha_no_link and j is not None and not out[j].href
            and out[j].text.startswith("]")
        )
        if not (fecha_no_link or fecha_no_seguinte):
            continue  # colchete sem par -> não mexe
        novo = r.text
        if abre_no_link:
            novo = novo[1:]
        else:
            out[i - 1] = replace(out[i - 1], text=out[i - 1].text[:-1])
        if fecha_no_link:
            novo = novo[:-1]
        else:
            out[j] = replace(out[j], text=out[j].text[1:])
        out[i] = replace(r, text=novo)
    # Passe B: citação textual entre colchetes, em texto puro.
    res: List[RichText] = []
    for r in out:
        if r.href or r.kind != "text" or "[" not in (r.text or ""):
            res.append(r)
            continue
        novo = _RE_COLCHETES.sub(
            lambda m: m.group(1) if _eh_citacao_completa(m.group(1)) else m.group(0),
            r.text,
        )
        res.append(replace(r, text=novo) if novo != r.text else r)
    return res


# Dispositivo citado LOGO APÓS um link de norma, na ordem invertida "Resolução
# TSE nº 23.610/2019, no art. 28, reproduz…" (o autor linkou só o nome da
# norma; o artigo ficou fora do link e sem âncora).
# O conector precisa ser IMEDIATO e curto ("…, no art. 28" / "…, cujo art. 6º"
# — o relativo amarra o dispositivo à norma do link); um advérbio no meio
# (", especialmente nos arts. 1º, V, …") indica enumeração ampla, que fica
# para a transferência do "(cf.)" — não para a extensão do link.
_RE_DISP_APOS_LINK = re.compile(
    r"^(?P<pre>\s*,?\s*(?:n[oa]s?|em|cuj[oa]s?)\s+)"
    r"(?P<disp>arts?\.\s*(?P<num>\d+)(?:[ºo°])?(?:\s*[-–‑]\s*(?P<letra>[A-Za-z])\b)?)",
    re.I,
)


def estender_link_dispositivo(rich_list: List[RichText]) -> List[RichText]:
    """Link de norma SEM âncora seguido, no texto, de ", no art. N": o trecho
    do dispositivo entra no link e a âncora correspondente é anexada — a
    citação abre a fonte já no artigo, como exige a política de links (achado
    do usuário, 05/07/2026: "Resolução TSE nº 23.610…, no art. 28" abria no
    topo e o art. 28 ficava fora do link). Vale tanto para o link cru do autor
    quanto para a designação isolada linkada por linkificar_citacoes (por isso
    roda DEPOIS dela). URL sem regra de âncora conhecida fica como está."""
    out: List[RichText] = list(rich_list or [])
    for i in range(len(out) - 1):
        r = out[i]
        href = (r.href or "").strip()
        if not href.startswith("http") or "#" in href:
            continue
        seg = out[i + 1]
        if seg.href or not seg.text:
            continue
        m = _RE_DISP_APOS_LINK.match(seg.text)
        if not m:
            continue
        novo_href = _ancorar_dispositivo(href, m.group("disp"))
        if novo_href == href:
            continue
        out[i] = replace(r, text=r.text + seg.text[: m.end()], href=novo_href)
        out[i + 1] = replace(seg, text=seg.text[m.end():])
    return [x for x in out if x.text]


# ---------------------------------------------------------------------------
# Lista de dispositivos ("arts. 1º, parágrafo único, 23, IX e XVIII, e 23-A"):
# consumo procedural com tolerância a complementos (caput, §§, incisos,
# alíneas, parágrafo único) entre os itens numéricos. Cada item vira UM link
# com âncora própria — decisão do usuário (06/07/2026), a mesma convenção da
# transferência do "(cf.)" em lista.
# ---------------------------------------------------------------------------
# Primeiro item, com o prefixo "art(s)./artigo(s)" incluído no span do link.
_RE_ARTS_INICIO = re.compile(
    r"\b(?:arts?\.|artigos?)\s*(?P<num>\d+)(?!\.?\d)(?:[ºo°])?"
    r"(?:\s*[-–‑]\s*(?P<letra>[A-Za-z])(?![A-Za-zº°]))?",
    re.I,
)
# Separador entre itens: vírgula e/ou conector ("e", "a", "até") ou traço de
# intervalo ("arts. 173–174" -> itens 173 e 174, cada um com sua âncora).
_RE_SEP_ITEM = re.compile(r"(?:\s*,\s*(?:e\s+|a\s+|at[ée]\s+)?|\s+(?:e|a|at[ée])\s+|\s*[-–‑]\s*)", re.I)
# Item numérico subsequente da lista (sem o prefixo "arts?.").
_RE_ITEM = re.compile(
    r"(?P<num>\d+)(?!\.?\d)(?:[ºo°])?(?:\s*[-–‑]\s*(?P<letra>[A-Za-z])(?![A-Za-zº°]))?"
)
# Complemento de dispositivo que NÃO é item ("caput", "parágrafo único",
# "IX e XVIII", "§§ 1º a 7º", 'alínea "a"'): consumido sem gerar link, mantém
# a lista viva. Os dígitos de §§ ficam dentro do complemento (não viram item).
_RE_COMPLEMENTO = re.compile(
    r"(?:\s*,\s*|\s+)(?:e\s+)?(?:caput\b|par[áa]grafos?\s+[úu]nicos?\b"
    r"|§§?\s*\d+[ºo°]?(?:\s*(?:a|e)\s*\d+[ºo°]?)*"
    r"|incisos?\s+[IVXLCDM]+\b(?:\s+(?:a|e)\s+[IVXLCDM]+\b)*"
    r"|al[íi]neas?\s+[\"'“”‘’]?[a-z][\"'“”‘’]?"
    r"|[IVXLCDM]+\b(?:\s+(?:a|e)\s+[IVXLCDM]+\b)*)",
    re.I,
)


def _consumir_lista_dispositivos(texto: str, m_inicio: "re.Match") -> tuple:
    """A partir do match de _RE_ARTS_INICIO, consome a lista de dispositivos.
    Retorna (itens, fim): itens = [(ini, fim, num, letra)] de cada trecho
    linkável (o 1º inclui o "art(s).") e fim = posição após o último elemento
    consumido (item ou complemento)."""
    itens = [(m_inicio.start(), m_inicio.end(),
              int(m_inicio.group("num")), m_inicio.group("letra") or "")]
    pos = m_inicio.end()
    while True:
        mc = _RE_COMPLEMENTO.match(texto, pos)
        if mc:
            pos = mc.end()
            continue
        ms = _RE_SEP_ITEM.match(texto, pos)
        if not ms:
            break
        mi = _RE_ITEM.match(texto, ms.end())
        if not mi:
            break
        itens.append((ms.end(), mi.end(), int(mi.group("num")), mi.group("letra") or ""))
        pos = mi.end()
    return itens, pos


def _dividir_por_spans(r: RichText, spans) -> List[RichText]:
    """Divide o run `r` conforme spans [(ini, fim, url)], linkando cada trecho."""
    pedacos: List[RichText] = []
    pos = 0
    for ini, fim, url in spans:
        if ini > pos:
            pedacos.append(replace(r, text=r.text[pos:ini]))
        pedacos.append(replace(r, text=r.text[ini:fim], href=url))
        pos = fim
    if pos < len(r.text):
        pedacos.append(replace(r, text=r.text[pos:]))
    return pedacos


def _spans_itens(itens, base: str) -> Optional[list]:
    """Spans (ini, fim, url) dos itens de lista, cada um com a âncora do seu
    dispositivo em `base`. None se a base não tem regra de âncora."""
    if _ancora_para(base, itens[0][2], itens[0][3]) == base:
        return None
    return [(ini, fim, _ancora_para(base, num, letra)) for ini, fim, num, letra in itens]


def _linkar_itens(r: RichText, itens, base: str) -> Optional[List[RichText]]:
    """Divide o run `r` linkando cada item da lista à âncora correspondente em
    `base`. None se a base não tem regra de âncora (nada é linkado)."""
    spans = _spans_itens(itens, base)
    return _dividir_por_spans(r, spans) if spans else None


# Fecho do gatilho da lista ANTES do link: "...arts. 139, I, e 142 do [RICD]"
# — entre o fim da lista e o fim do run só pode haver ", do "/" da ".
_RE_FECHO_DO = re.compile(r"\s*,?\s*d[aeo]s?\s+$", re.I)
# Gatilho da lista DEPOIS do link: "[Código Eleitoral], arts. 1º, ..." —
# vírgula DIRETA (com conector "no/na/em/cujo" quem age é a extensão do link;
# advérbio no meio, ex. ", especialmente nos arts.", continua de fora).
_RE_VIRGULA_ARTS = re.compile(r"^\s*,\s*(?=arts?\.|artigos?\b)", re.I)


def linkificar_dispositivos_junto_a_link(rich_list: List[RichText]) -> List[RichText]:
    """Dispositivos citados em texto puro COLADOS a um link de norma (achados
    do usuário, 06/07/2026). Dois gatilhos, sempre com a norma linkada pelo
    autor (ou pela linkificação) SEM âncora de artigo:

    - ANTES do link: "…arts. 139, I, e 142 do [RICD]" / "O art. 49, V, da
      [Constituição Federal]" / "arts. 66 e 68 da [Lei nº 9.504/1997]" — o run
      de texto anterior termina com a lista + "do/da";
    - DEPOIS do link: "[Código Eleitoral], arts. 1º, parágrafo único, 23, IX
      e XVIII, e 23-A" — vírgula direta (sem conector).

    Cada item da lista vira um link individual com a âncora do dispositivo na
    URL do link vizinho (decisão do usuário: âncora por dispositivo, não um
    link único). O link da norma em si não muda. Roda ANTES do anafórico —
    sem isso o anafórico herdaria outra base do parágrafo e linkaria errado
    (caso real: "arts. 66 e 68 da [Lei 9.504]" recebia âncora da CF)."""
    out: List[RichText] = list(rich_list or [])
    i = 0
    while i < len(out):
        r = out[i]
        href = (r.href or "").strip().split("#")[0]
        if not href.startswith("http") or "#" in (r.href or ""):
            i += 1
            continue
        # -- lista ANTES do link ------------------------------------------
        # A citação pode atravessar runs sem href ("O art. 66, " + "caput"
        # em itálico + " e §§1º a 7º, da " + [Lei…]): concatena os runs de
        # texto anteriores contíguos e redistribui os itens por run.
        if i > 0:
            j0 = i
            while j0 > 0 and not out[j0 - 1].href and out[j0 - 1].kind == "text" \
                    and out[j0 - 1].text and i - j0 < 4:
                j0 -= 1
            if j0 < i:
                offsets = []  # (idx_run, ini_concat)
                concat = ""
                for j in range(j0, i):
                    offsets.append((j, len(concat)))
                    concat += out[j].text
                melhor = None
                for m in _RE_ARTS_INICIO.finditer(concat):
                    itens, fim = _consumir_lista_dispositivos(concat, m)
                    if _RE_FECHO_DO.match(concat, fim):
                        melhor = itens  # última lista que fecha no fim do trecho
                spans = _spans_itens(melhor, href) if melhor else None
                if spans:
                    novos: List[RichText] = []
                    for k, (j, ini_run) in enumerate(offsets):
                        fim_run = ini_run + len(out[j].text)
                        do_run = [(a - ini_run, b - ini_run, u) for a, b, u in spans
                                  if a >= ini_run and b <= fim_run]
                        novos.extend(_dividir_por_spans(out[j], do_run) if do_run else [out[j]])
                    out[j0:i] = novos
                    i = j0 + len(novos)
                    r = out[i]
        # -- lista DEPOIS do link -----------------------------------------
        if i + 1 < len(out):
            seg = out[i + 1]
            if not seg.href and seg.kind == "text" and seg.text:
                mv = _RE_VIRGULA_ARTS.match(seg.text)
                if mv:
                    m = _RE_ARTS_INICIO.search(seg.text, mv.end())
                    if m and m.start() == mv.end():
                        itens, _fim = _consumir_lista_dispositivos(seg.text, m)
                        pedacos = _linkar_itens(seg, itens, href)
                        if pedacos:
                            out[i + 1 : i + 2] = pedacos
        i += 1
    return [x for x in out if x.text]


# Designações textuais das normas mapeadas (citacao_re por norma), para o
# padrão "norma ANTES do dispositivo": "CF, art. 121" / "Constituição (art.
# 121, caput)". O gatilho é imediato: vírgula ou parêntese + "art(s).".
# O \b à esquerda evita casar sufixos de palavra ("esclarece" -> "CE") — nos
# padrões de citação o contexto "d[aeo] " dava esse boundary de graça.
_DESIGNACOES_NORMAS = [
    (re.compile(rf"\b(?:{n.citacao_re})", re.I), n) for n in config.NORMAS_OFICIAIS
]
_RE_GATILHO_PARENS = re.compile(r"\s*\(\s*(?=arts?\.|artigos?\b)", re.I)
_RE_GATILHO_VIRGULA = re.compile(r"\s*,\s*(?=arts?\.|artigos?\b)", re.I)
# Outra norma nomeada logo após a lista ("art. 5º da Lei X"): os itens não são
# da norma-contexto — mesma lista-negra do miolo de citação.
_RE_NORMA_LOGO_APOS = re.compile(
    r"^\s*,?\s*d[aeo]s?\s+(?:lei|decreto|c[óo]digo|resolu[çc]|res\b|emenda|regimento|"
    r"ricd\b|cf\b|constitui[çc]|adct\b|lc\b|ec\b|nr\b|norma\s+regulamentadora)",
    re.I,
)


def _designacao_com_gatilho(texto: str, pos: int):
    """Próxima designação de norma (a partir de pos) seguida IMEDIATAMENTE de
    "(art…" ou ", art…". Retorna (m_designacao, norma, m_gatilho, parens) da
    ocorrência mais à esquerda (mais longa no desempate), ou None."""
    melhor = None
    for pat, norma in _DESIGNACOES_NORMAS:
        for m in pat.finditer(texto, pos):
            gat_p = _RE_GATILHO_PARENS.match(texto, m.end())
            gat_v = None if gat_p else _RE_GATILHO_VIRGULA.match(texto, m.end())
            gat = gat_p or gat_v
            if not gat:
                continue
            chave = (m.start(), -m.end())
            if melhor is None or chave < melhor[0]:
                melhor = (chave, m, norma, gat, bool(gat_p))
            break  # das ocorrências desta norma, só a primeira com gatilho
    return melhor[1:] if melhor else None


def linkificar_norma_antes_dispositivo(rich_list: List[RichText]) -> List[RichText]:
    """Citação com a norma nomeada ANTES dos dispositivos, em texto puro
    (achados do usuário, 06/07/2026): "CF, art. 121, caput" e "pela
    Constituição (art. 121, caput) e pelo Código Eleitoral (art. 1º,
    parágrafo único; art. 23, IX e XVIII; art. 23-A)". Cada dispositivo ganha
    link individual com âncora própria; a designação da norma NÃO é linkada
    solta (anti-overlinking, coerente com citacao_isolada_re=None de CF/RICD).
    O contexto de parêntese persiste entre runs — o "caput" em run itálico
    separado não interrompe a região (limitação de 1 run do linkificar
    clássico). Norma fora do mapa (ex. resolução TSE desconhecida) e item
    seguido da designação de OUTRA norma ficam sem link."""
    out: List[RichText] = []
    contexto: Optional[str] = None  # URL-base da norma do parêntese aberto
    for r in rich_list or []:
        if r.href or r.kind != "text" or not r.text:
            contexto = None
            out.append(r)
            continue
        texto = r.text
        spans: list = []
        pos = 0
        while pos < len(texto):
            if contexto:
                m_arts = _RE_ARTS_INICIO.search(texto, pos)
                i_fecha = texto.find(")", pos)
                if i_fecha != -1 and (m_arts is None or i_fecha < m_arts.start()):
                    contexto = None
                    pos = i_fecha + 1
                    continue
                if m_arts is None:
                    break  # região continua no próximo run
                itens, fim = _consumir_lista_dispositivos(texto, m_arts)
                if not _RE_NORMA_LOGO_APOS.match(texto[fim:]):
                    spans.extend(_spans_itens(itens, contexto) or [])
                pos = fim
                continue
            achado = _designacao_com_gatilho(texto, pos)
            if not achado:
                break
            m_desig, norma, gat, parens = achado
            base = norma.montar_url(m_desig)
            if not base:
                pos = m_desig.end()
                continue
            if parens:
                contexto = base
                pos = gat.end()
                continue
            m_arts = _RE_ARTS_INICIO.match(texto, gat.end())
            if not m_arts:
                pos = m_desig.end()
                continue
            itens, fim = _consumir_lista_dispositivos(texto, m_arts)
            if not _RE_NORMA_LOGO_APOS.match(texto[fim:]):
                spans.extend(_spans_itens(itens, base) or [])
            pos = fim
        out.extend(_dividir_por_spans(r, spans) if spans else [r])
    return out


# URLs-base (sem âncora) das normas oficiais conhecidas — antecedentes válidos
# para a resolução de anáfora ("No art. 36-A, …" depois de um link da Lei
# 9.504). Links de julgados/proposições NÃO entram.
_BASES_NORMAS = (
    {u for u in (n.montar_url(None) for n in config.NORMAS_OFICIAIS) if u}
    | set(config.RESOLUCOES_TSE.values())
)

# Outra norma nomeada logo adiante ("art. 7º da LC nº 95/1998…"): o artigo NÃO
# é anáfora da norma anterior — mesmo critério da lista-negra do miolo.
_RE_NORMA_ADIANTE = re.compile(
    r"^[^;()]{0,60}?\bd[aeo]s?\s+(?:lei|decreto|c[óo]digo|resolu[çc][ãa]o|res\b|emenda|"
    r"regimento|ricd\b|cf\b|constitui[çc][ãa]o|adct\b|lc\b|ec\b)",
    re.I,
)
# Dispositivo precedido de possessivo ("Por meio de seu art. 1º, a proposição
# acrescenta…"): pertence ao sujeito da frase (a proposição em exame), não à
# norma linkada antes — anáfora NÃO se aplica.
_RE_POSSESSIVO_ANTES = re.compile(r"\b(?:seus?|suas?|dest[ea]s?)\s+$", re.I)
# Cláusula de vigência logo após ("O art. 2º dispõe sobre a vigência"): é o
# fecho padrão da PRÓPRIA proposição — nunca ancorável na norma alterada.
_RE_VIGENCIA_ADIANTE = re.compile(r"^[^.;]{0,40}?\b(?:vig[êe]ncia|vigor)\b", re.I)
# Regência de ACRÉSCIMO imediatamente antes ("o acréscimo do art. 9º-A",
# "acrescenta o art. 9º-A à referida lei"): o dispositivo é o próprio objeto
# da deliberação — ainda NÃO existe no texto oficial da norma alterada, então
# fica SEM link (decisão do usuário, 05/07/2026: "abrir a lei não resultaria
# em nada, visto que o próprio artigo está em deliberação").
_RE_ACRESCIMO_ANTES = re.compile(
    r"(?:\b(?:acr[ée]scimo|inclus[ãa]o|inser[çc][ãa]o|cria[çc][ãa]o)\s+d[oa]s?"
    r"|\b(?:acrescent\w+|inclui\w*|insere\w*|inserind\w+|introduz\w*|cria\w*)\s+[oa]s?"
    r"|\b(?:acrescid|inclu[ií]d|inserid|introduzid|criad)\w+\s+[oa]s?)"
    r"\s+(?:novo\s+)?$",
    re.I,
)
# Dispositivo PROPOSTO — "os novos arts. 66-A e 67-A" / "arts. 66-A e 67-A
# propostos": é texto da minuta em deliberação, não existe na norma oficial —
# fica sem link (achado do usuário, 06/07/2026; mesma decisão do acréscimo).
_RE_NOVO_ANTES = re.compile(r"\bnov[oa]s?\s+$", re.I)
_RE_PROPOSTO_APOS = re.compile(r"^\s*(?:ora\s+)?propost[oa]s?\b", re.I)
# Anáfora com a CLASSE da norma ("art. 216 da mesma Resolução", "da referida
# lei", "do mesmo diploma"): designação genérica SEM número — resolve para a
# única base daquela classe já linkada no parágrafo. O lookahead barra as
# designações completas ("da Lei nº 9.504", "do Código Eleitoral"), que são
# assunto da linkificação textual, não desta anáfora.
_RE_ANAFORA_CLASSE = re.compile(
    r"^\s*,?\s*d[aeo]s?\s+(?:mesm[oa]s?\s+|referid[oa]s?\s+|pr[óo]pri[oa]s?\s+|citad[oa]s?\s+)?"
    r"(?P<classe>resolu[çc][ãa]o|lei|c[óo]digo|regimento|constitui[çc][ãa]o|diploma|norma)"
    r"(?!\s*(?:n[ºo°.]?\s*\d|\d|complementar\b|federal\b|eleitoral\b|interno\b|"
    r"regulamentadora\b|d[aeo]s?\b))",
    re.I,
)


def _norm_ascii(s: str) -> str:
    return unicodedata.normalize("NFKD", s or "").encode("ascii", "ignore").decode().lower()


def _classes_da_base(base: str) -> tuple:
    """Classes de designação genérica que podem retomar a norma `base` — o
    Código Eleitoral atende por "código" e por "lei" (é a Lei nº 4.737)."""
    if base in config.RESOLUCOES_TSE.values():
        return ("resolucao",)
    if base == config.RICD_LEGIN_URL:
        return ("regimento",)
    if "constituicao" in base:
        return ("constituicao",)
    if base == config.L4737_PLANALTO_URL:
        return ("codigo", "lei")
    return ("lei",)


def linkificar_dispositivos_anaforicos(rich_list: List[RichText]) -> List[RichText]:
    """Anáfora de dispositivo (achado do usuário, 05/07/2026): "…harmoniza a
    Lei nº 9.504, de 1997… No art. 36-A, esclarece-se… No art. 37,
    preserva-se…" — o artigo citado SEM nomear a norma ganha o link da norma
    oficial linkada antes dele no MESMO parágrafo, com a âncora do
    dispositivo. Só resolve quando o parágrafo, até aquele ponto, linkou UMA
    ÚNICA norma (com duas ou mais, a retomada é ambígua — ex. "…do art. 74 da
    Lei nº 9.504/1997. …critérios ao próprio art. 22…", em que o art. 22 é da
    LC 64: fica sem link em vez de ganhar o errado). Também não linka sem
    antecedente, quando outra norma é nomeada logo após o artigo (mesmo se o
    nome estiver no run linkado seguinte), nem em runs já linkados. Lista de
    dispositivos ("arts. 173–174", "arts. 57-B e 57-C") linka POR ITEM, cada
    um com sua âncora (decisão do usuário, 06/07/2026). Refinos do parecer
    (05/07/2026, PL 3.263/2023): dispositivo com possessivo antes ("seu art.
    1º"), cláusula de vigência depois ("O art. 2º dispõe sobre a vigência")
    ou regido por acréscimo/inclusão ("o acréscimo do art. 9º-A") pertence à
    PRÓPRIA proposição em deliberação — nunca herda link; idem "novos arts.
    66-A e 67-A propostos" (06/07/2026)."""
    out: List[RichText] = []
    bases_vistas: set = set()
    por_classe: dict = {}
    rich = list(rich_list or [])
    for idx, r in enumerate(rich):
        href = (r.href or "").strip()
        if href:
            b = href.split("#")[0]
            if b in _BASES_NORMAS:
                bases_vistas.add(b)
                for cl in _classes_da_base(b):
                    por_classe.setdefault(cl, set()).add(b)
            out.append(r)
            continue
        if r.kind != "text" or not r.text:
            out.append(r)
            continue
        # Designações TEXTUAIS de norma também contam para a ambiguidade, mas
        # só DALI PARA A FRENTE (caso real 06/07/2026: "No Código Eleitoral, o
        # novo art. 165-A cria… evita sobrecarregar o art. 165" — só a 9.504
        # estava LINKADA e o art. 165, do CE, herdava a base errada; já em
        # "…o art. 66, que já assegura… O Código Eleitoral permanece…" a
        # designação vem DEPOIS e não pode ofuscar o art. 66 da 9.504).
        desigs: list = []
        for pat, norma in _DESIGNACOES_NORMAS:
            for m_d in pat.finditer(r.text):
                b_d = norma.montar_url(m_d)
                if b_d:
                    desigs.append((m_d.start(), b_d))
        if not bases_vistas and not desigs:
            out.append(r)
            continue

        def _ate(pos):
            bs = set(bases_vistas)
            pc = {k: set(v) for k, v in por_classe.items()}
            for s, b in desigs:
                if s < pos:
                    bs.add(b)
                    for cl in _classes_da_base(b):
                        pc.setdefault(cl, set()).add(b)
            return bs, pc
        # texto "adiante" p/ os bloqueios inclui o começo do run seguinte — a
        # designação da norma pode estar num run linkado ("arts. 66 e 68 da
        # [Lei nº 9.504/1997]": o run atual acaba em "da " e o nome vem depois)
        prox = rich[idx + 1] if idx + 1 < len(rich) else None
        prox_txt = (prox.text or "")[:80] if prox is not None else ""
        spans: list = []
        pos = 0
        for m in _RE_ARTS_INICIO.finditer(r.text):
            if m.start() < pos:
                continue
            itens, fim = _consumir_lista_dispositivos(r.text, m)
            adiante = r.text[fim:] + prox_txt
            pos = fim
            bases_m, por_classe_m = _ate(m.start())
            m_cl = _RE_ANAFORA_CLASSE.match(adiante)
            if m_cl:
                # "art. 216 da mesma Resolução": resolve pela classe citada
                classe = _norm_ascii(m_cl.group("classe"))
                if classe in ("diploma", "norma"):
                    candidatas = bases_m
                else:
                    candidatas = por_classe_m.get(classe, set())
                if len(candidatas) != 1:
                    continue  # nenhuma ou mais de uma norma da classe: ambíguo
                base = next(iter(candidatas))
            else:
                if len(bases_m) != 1:
                    continue  # 2+ normas no parágrafo: retomada ambígua
                base = next(iter(bases_m))
                if _RE_NORMA_ADIANTE.match(adiante):
                    continue  # "art. 7º da LC 95/1998" — não é anáfora
                if prox is not None and prox.href and _RE_FECHO_DO.match(r.text, fim):
                    continue  # "arts. 66 e 68 da [<norma linkada>]" — idem
            prefixo = r.text[:m.start()]
            if _RE_POSSESSIVO_ANTES.search(prefixo):
                continue  # "seu art. 1º" — dispositivo da proposição
            if _RE_VIGENCIA_ADIANTE.match(r.text[fim:]):
                continue  # "O art. 2º dispõe sobre a vigência" — idem
            if _RE_ACRESCIMO_ANTES.search(prefixo):
                continue  # "o acréscimo do art. 9º-A" — em deliberação, sem link
            if _RE_NOVO_ANTES.search(prefixo) or _RE_PROPOSTO_APOS.match(r.text[fim:]):
                continue  # "novos arts. 66-A e 67-A propostos" — idem
            spans.extend(_spans_itens(itens, base) or [])
        # as designações do run valem integralmente para os runs seguintes
        for _s, b_d in desigs:
            bases_vistas.add(b_d)
            for cl in _classes_da_base(b_d):
                por_classe.setdefault(cl, set()).add(b_d)
        out.extend(_dividir_por_spans(r, spans) if spans else [r])
    return out


# Número de dispositivo solto (texto de um link por item de lista: "142", "23-A").
_RE_NUM_SOLTO = re.compile(
    r"^\s*(?:arts?\.\s*|artigos?\s+)?(\d+)(?:[ºo°])?(?:\s*[-–‑]\s*([A-Za-z])\b)?\s*$", re.I
)


def _zerar_links_proibidos(rich_list: List[RichText], proibidos) -> List[RichText]:
    """Remove o hyperlink dos runs que citam dispositivo PROIBIDO — artigos
    criados pela própria minuta (conjunto de (num, LETRA) vindo de
    splitter.dispositivos_criados). Vale para qualquer origem do link
    (citação textual, anafórico, item de lista, mention): dispositivo em
    deliberação não existe na norma oficial e sai sem link (decisões do
    usuário de 05-06/07/2026). Roda como último passo de linkificação da
    cadeia preparar_rich — depois de TODOS os linkificadores."""
    out: List[RichText] = []
    for r in rich_list or []:
        if not (r.href or "").startswith("http"):
            out.append(r)
            continue
        candidatos = [(int(m.group(1)), (m.group(2) or "").upper())
                      for m in _RE_ARTIGO.finditer(r.text or "")]
        for m in _RE_ARTS_INICIO.finditer(r.text or ""):
            itens, _fim = _consumir_lista_dispositivos(r.text, m)
            candidatos += [(num, (letra or "").upper()) for _a, _b, num, letra in itens]
        m = _RE_NUM_SOLTO.match(r.text or "")
        if m:
            candidatos.append((int(m.group(1)), (m.group(2) or "").upper()))
        if any(c in proibidos for c in candidatos):
            out.append(replace(r, href=None))
        else:
            out.append(r)
    return out


# Abertura de parêntese remissivo no FIM de um run de texto — gatilho da
# supressão do "(cf. <mention>)"/"(<mention>)" redundante. O "cf." é opcional
# (página PRC Bancada Negra, 06/07/2026: o Gerador emite "([RICD - Art.
# 13-A], campo texto_em_vigor)" sem o "cf."); a semântica de "parêntese
# puramente remissivo" é garantida pelo laço, que aborta com conteúdo extra.
_RE_ABRE_CF = re.compile(r"\(\s*(?:cf\.?\s*)?$", re.I)
# Separador aceitável entre duas mentions dentro do parêntese (", " / " e " /
# "; incisos em " — este último também artefato do Gerador via RAG).
_RE_SEP_MENTION = re.compile(r"^\s*(?:[,;]\s*)?(?:e|incisos?\s+em)?\s*$", re.I)
# Sufixo aceitável entre a última mention e o ")": vazio ou rótulo técnico de
# RAG (", campo texto_em_vigor" — pode vir fatiado em vários runs do Notion),
# descartado junto com o parêntese.
_RE_SUFIXO_MENTION = re.compile(r"\s*(?:,\s*campo\s+\w+\s*)?", re.I)
# Janela (chars p/ trás) em que se procura o link equivalente à mention do "(cf.)".
_JANELA_CF = 300


def _span_dispositivo(prefixo: str, num: str, letra: str, em_lista: bool) -> Optional[tuple]:
    """Localiza no texto anterior ao "(cf." o dispositivo citado pela mention.
    Fora de lista, exige a forma completa "art(s). N" (última ocorrência do
    parágrafo — ex. "O art. 22 autoriza … (cf. …)", achado do usuário de que a
    janela curta deixava o parêntese sobrar). Em lista ("arts. 1º, V, 5º, …,
    17 e 220"), o número aparece SOLTO após o "arts." — busca-se o token na
    região depois do último "art(s).", desde que ali não haja "§" (número de
    parágrafo confundiria a remissão)."""
    if letra:
        nucleo = rf"{num}(?:[ºo°])?\s*[-–‑]\s*{letra}\b"
    else:
        nucleo = rf"{num}(?:[ºo°])?\b(?!\s*[-–‑]\s*[A-Za-z])"
    if not em_lista:
        ms = list(re.finditer(rf"\barts?\.?\s*{nucleo}", prefixo, re.I))
        if ms:
            return (ms[-1].start(), ms[-1].end())
    marcas = list(re.finditer(r"\barts?\.", prefixo, re.I))
    if not marcas:
        return None
    base = marcas[-1].end()
    regiao = prefixo[base:]
    if "§" in regiao:
        return None
    ms = list(re.finditer(rf"\b{nucleo}", regiao, re.I))
    if not ms:
        return None
    return (base + ms[-1].start(), base + ms[-1].end())


def _localizar_transferencias(
    prefixo: str, textos_mentions: List[str], urls: List[str]
) -> Optional[List[tuple]]:
    """Spans (ini, fim, url) do prefixo que receberão os links das mentions do
    parêntese — caso (b) da supressão, inclusive com VÁRIAS mentions ("arts.
    1º, V, 5º, … e 220 (cf. CF - Art. 1º, CF - Art. 5º, …)"): cada link vai
    para o número correspondente da lista. None se algum dispositivo não for
    localizado com segurança (aí o parêntese fica como está). Várias mentions
    só transferem quando todas são da MESMA norma — números soltos da lista
    não identificam a norma por si."""
    disps = [_dispositivo_da_mention(t) for t in textos_mentions]
    if any(d is None for d in disps):
        return None
    if len(disps) > 1 and len({u.split("#")[0] for u in urls}) > 1:
        return None
    spans = []
    for (num, letra), url in zip(disps, urls):
        span = _span_dispositivo(prefixo, num, letra, em_lista=len(disps) > 1)
        if span is None:
            return None
        spans.append((span[0], span[1], url))
    spans.sort()
    for a, b in zip(spans, spans[1:]):
        if b[0] < a[1]:
            return None  # sobreposição: remissão ambígua
    return spans


def _url_do_run(r: RichText) -> Optional[str]:
    """URL EFETIVA com que o run sairá no .docx: href externo como está;
    mention interna do Notion pela fonte oficial (None se não mapeada)."""
    h = (r.href or "").strip()
    if not h:
        return None
    if _link_interno_notion(h):
        return resolver_fonte_publica(r.text)
    return h if h.startswith("http") else None


def _dispositivo_da_mention(texto: str) -> Optional[tuple]:
    """(num, letra) do dispositivo citado numa mention de norma reconhecida
    ("Lei das Eleições - Art. 36-A" -> ("36", "A")); None sem artigo/não
    reconhecida."""
    for padrao, _t, _a in config.FONTES_OFICIAIS:
        m = re.search(padrao, texto or "")
        if m:
            d = m.groupdict()
            if d.get("num"):
                return d["num"], (d.get("letra") or "")
            return None
    return None


def suprimir_cf_redundante(rich_list: List[RichText]) -> List[RichText]:
    """Suprime o parêntese remissivo "(cf. <mention>)" quando ele é REDUNDANTE
    com a citação do próprio texto — integração orgânica escolhida pelo usuário
    (jul/2026). Roda DEPOIS de linkificar_citacoes (precisa dos hrefs já
    atribuídos). Dois casos: (a) o texto imediatamente anterior já traz link
    com a MESMA URL da mention (ex. "O art. 37 da Lei nº 9.504/1997 (cf. Lei
    das Eleições - Art. 37)…", ou o acórdão do TSE já linkado pelo autor) ->
    o parêntese inteiro sai; (b) o dispositivo aparece no texto anterior SEM
    link ("O art. 36-A (cf. …)") -> o link é transferido para a citação
    textual e o parêntese sai. Parêntese com conteúdo além de mentions e
    separadores, mention sem fonte pública ou sem equivalente anterior ->
    fica tudo como está (a mention sai linkada dentro do "(cf. …)")."""
    out: List[RichText] = list(rich_list or [])
    i = 0
    while i < len(out):
        r = out[i]
        if r.href or not r.text:
            i += 1
            continue
        m_abre = _RE_ABRE_CF.search(r.text)
        if not m_abre:
            i += 1
            continue
        com_cf = "cf" in m_abre.group(0).lower()
        # Conteúdo do parêntese: só runs com href (mentions/links), separadores
        # e um eventual sufixo técnico antes do ")". O texto entre as refs é
        # ACUMULADO (o Notion fatia livremente: ", campo " + "texto_em_vigor").
        refs: List[int] = []
        j = i + 1
        fecha = None
        fim_fecha = 0
        interm = ""  # texto acumulado desde a última ref (ou desde o "(")
        valido = True
        while j < len(out):
            rj = out[j]
            if (rj.href or "").strip():
                if not _RE_SEP_MENTION.fullmatch(interm):
                    valido = False  # conteúdo extra antes da ref -> não mexe
                    break
                refs.append(j)
                interm = ""
                j += 1
                continue
            txt = rj.text or ""
            pos_f = txt.find(")")
            if pos_f != -1:
                if refs and _RE_SUFIXO_MENTION.fullmatch(interm + txt[:pos_f]):
                    fecha = j
                    fim_fecha = pos_f + 1
                else:
                    valido = False
                break
            interm += txt
            if len(interm) > 60:
                valido = False
                break
            j += 1
        if not valido:
            i += 1
            continue
        if fecha is None or not refs:
            i += 1
            continue
        if not com_cf and not all(
            out[k].kind == "mention" or resolver_fonte_publica(out[k].text or "")
            for k in refs
        ):
            # sem o "cf.", só o parêntese de refs do Gerador é remissivo por
            # definição: @mention OU link de texto cujo conteúdo tem o formato
            # de mention de norma ("Constituicao Federal - Art. 51, III").
            # Um "(art. 151)" que a própria cadeia linkou não casa o formato e
            # não pode ser suprimido na passada seguinte (idempotência).
            i += 1
            continue
        urls = [_url_do_run(out[k]) for k in refs]
        if not all(urls):
            if not com_cf and not any(urls):
                # parêntese de mentions SEM fonte pública (ex. "(RICD -
                # CAPÍTULO VI)", página interna não mapeada): no .docx viraria
                # texto técnico morto — sai inteiro (o "(cf. …)" explícito
                # continua preservado, como aprovado em jul/2026)
                seg_fecha = out[fecha]
                texto_fecha = seg_fecha.text[fim_fecha:]
                prefixo_txt = r.text[: m_abre.start()]
                if texto_fecha[:1] in ",.;:)" and prefixo_txt.endswith(" "):
                    prefixo_txt = prefixo_txt.rstrip()
                novos = [replace(r, text=prefixo_txt)] if prefixo_txt else []
                out[i : fecha + 1] = novos + [replace(seg_fecha, text=texto_fecha)]
                i += len(novos)
                continue
            i += 1
            continue
        # URLs efetivas já presentes antes do "(cf." (janela p/ trás). Um link
        # de LISTA ("arts. 3º, …, 4º, …, e 5º, … da Constituição", âncora do
        # 1º) cobre TODOS os dispositivos citados no seu texto — as âncoras
        # individuais entram no conjunto para o teste de redundância.
        anteriores = set()
        recuo = 0
        for k in range(i - 1, -1, -1):
            u = _url_do_run(out[k])
            if u:
                anteriores.add(u)
                base_k = u.split("#")[0]
                for m_l in _RE_ARTS_INICIO.finditer(out[k].text or ""):
                    itens_k, _f = _consumir_lista_dispositivos(out[k].text, m_l)
                    for _a, _b2, num_k, letra_k in itens_k:
                        anteriores.add(_ancora_para(base_k, num_k, letra_k))
            recuo += len(out[k].text or "")
            if recuo > _JANELA_CF:
                break
        bases_anteriores = {u.split("#")[0] for u in anteriores}
        prefixo = r.text[: m_abre.start()]
        disps = [_dispositivo_da_mention(out[k].text) for k in refs]

        # caso (a): mention já coberta por link anterior — igualdade estrita
        # (mesmo dispositivo); mention SEM dispositivo (norma sem artigo, ex.
        # "Res.-TSE 23.610", ou julgado) compara pela URL sem âncora.
        def _presente(u, d):
            return u in anteriores or (d is None and u.split("#")[0] in bases_anteriores)

        transferencias = None  # caso (b): [(ini, fim, url)] no prefixo
        if not all(_presente(u, d) for u, d in zip(urls, disps)):
            transferencias = _localizar_transferencias(
                prefixo, [out[k].text for k in refs], urls
            )
            if transferencias is None:
                i += 1
                continue
        # Suprime o parêntese (e, no caso (b), transfere os links p/ o texto).
        seg_fecha = out[fecha]
        texto_fecha = seg_fecha.text[fim_fecha:]
        novos: List[RichText] = []
        if transferencias:
            pos = 0
            for ini, fim, url in transferencias:
                if ini > pos:
                    novos.append(replace(r, text=prefixo[pos:ini]))
                novos.append(replace(r, text=prefixo[ini:fim], href=url))
                pos = fim
            if prefixo[pos:]:
                novos.append(replace(r, text=prefixo[pos:]))
        elif prefixo:
            novos.append(replace(r, text=prefixo))
        # Costura: espaço sobrando antes de pontuação/espaço do run de fechamento.
        if novos and not novos[-1].href:
            ultimo = novos[-1].text
            if texto_fecha[:1] in ",.;:)" and ultimo.endswith(" "):
                novos[-1] = replace(novos[-1], text=ultimo.rstrip())
            elif texto_fecha[:1] == " " and ultimo.endswith(" "):
                texto_fecha = texto_fecha.lstrip()
        out[i : fecha + 1] = novos + [replace(seg_fecha, text=texto_fecha)]
        i += len(novos)
    return [x for x in out if x.text]


# Latinismo "caput" SEMPRE em itálico (norma tipográfica; regra do usuário,
# 04/07/26) — vale em corpo, transcrições e articulado.
_RE_CAPUT = re.compile(r"\bcaput\b", re.I)


def italicizar_caput(rich_list: List[RichText]) -> List[RichText]:
    """Divide os runs para que toda ocorrência de "caput" saia em itálico,
    preservando href/negrito do trecho (um link que contém "caput" vira
    hyperlinks adjacentes com o mesmo destino — visualmente contínuo)."""
    out: List[RichText] = []
    for r in rich_list or []:
        if not r.text or r.italic or r.kind == "equation":
            out.append(r)
            continue
        pos = 0
        for m in _RE_CAPUT.finditer(r.text):
            if m.start() > pos:
                out.append(replace(r, text=r.text[pos:m.start()]))
            out.append(replace(r, text=m.group(0), italic=True))
            pos = m.end()
        if pos == 0:
            out.append(r)
        elif pos < len(r.text):
            out.append(replace(r, text=r.text[pos:]))
    return out


def _ancora_para(base: str, num: int, letra: str = "") -> str:
    """URL da norma `base` (sem âncora) apontando para o artigo num[-letra],
    conforme a regra de âncora do host; sem regra conhecida, retorna a própria
    base (o chamador decide se ainda vale linkar)."""
    if base.startswith(config.RICD_LEGIN_URL):
        return config.RICD_LEGIN_URL + config.ancora_artigo_ricd(num, letra or "")
    if _RE_PLANALTO.match(base):
        return f"{base}#art{num}{(letra or '').lower()}"
    if any(base.startswith(u) for u in config.RESOLUCOES_TSE.values()):
        return base + config.ancora_artigo_tse(num, letra or "")
    return base


def _ancorar_dispositivo(href: str, texto: str) -> str:
    """Link http externo para norma SEM âncora: se o texto do trecho cita um
    artigo ("art. 7º da LC 95/1998"), anexa a âncora que abre a página já no
    dispositivo — o "deslocamento" que o usuário exige em toda citação. No
    Planalto a âncora nativa é #artN (letra de artigo acrescido entra: art.
    10-B -> #art10b); no LEGIN do RICD é o text fragment de ancora_artigo_ricd.
    Âncora inexistente degrada para abrir no topo (comportamento antigo)."""
    if not href or "#" in href:
        return href
    m = _RE_ARTIGO.search(texto or "")
    if not m:
        return href
    return _ancora_para(href, int(m.group(1)), m.group(2) or "")


def _aliases_da_norma(texto: str) -> List[str]:
    """Aliases (nome por extenso/sigla) da norma citada em `texto`, se
    reconhecida em config.FONTES_OFICIAIS — usados para podar o eco redundante
    que às vezes segue a mention (ex. "RICD" repetido depois de "RICD - Art. 54")."""
    for padrao, _template, aliases in config.FONTES_OFICIAIS:
        if re.search(padrao, texto or ""):
            return list(aliases)
    return []


# Janela (em caracteres) em que se procura o eco logo após a mention — curta o
# bastante para não alcançar uma citação legítima e distante da MESMA norma
# mais adiante (ex. outro dispositivo) nem uma definição por extenso da sigla
# (ex. "do Regimento Interno da Câmara dos Deputados (RICD)", ~47 chars);
# os ecos redundantes observados ficam a até ~26 chars do início do trecho.
_JANELA_ECO = 32


_RE_ESPACOS_DUPLOS = re.compile(r"[ \t]{2,}")
_RE_ESPACO_ANTES_PONTUACAO = re.compile(r"[ \t]+([,.;:)])")


def _encontrar_eco(texto: str, aliases: List[str]) -> Optional[tuple]:
    """Localiza, dentro de uma janela curta a partir do início de `texto`, o eco
    redundante do nome/sigla de uma norma já citada pela mention anterior.
    Dois padrões: (1) fecha parênteses logo em seguida — o eco vai até (mas sem
    incluir) o ")" (ex. ", RICD)" -> ")"); (2) "do/da <norma>" no meio da frase
    (ex. ", caput e inciso VI, da CF permanecem" -> ", caput e inciso VI, permanecem").
    Retorna (início, fim) do trecho a remover, ou None."""
    if not aliases or not texto:
        return None
    alt = "|".join(re.escape(a) for a in sorted(aliases, key=len, reverse=True))
    padrao = re.compile(
        rf"[,;]?\s*(?:\bdo\b\s+|\bda\b\s+)?(?:{alt})\b\s*(?=\))"
        rf"|\b(?:do|da)\b\s+(?:{alt})\b"
    )
    m = padrao.search(texto[:_JANELA_ECO])
    return (m.start(), m.end()) if m else None


def _remover_eco(texto: str, aliases: List[str]) -> str:
    """Variante de conveniência de _encontrar_eco que já retorna o texto podado
    e com espaços/pontuação normalizados (usada quando o eco cabe num único
    run; limpar_ecos_redundantes trata o caso geral, de vários runs)."""
    span = _encontrar_eco(texto, aliases)
    if not span:
        return texto
    novo = texto[: span[0]] + texto[span[1] :]
    novo = _RE_ESPACOS_DUPLOS.sub(" ", novo)
    novo = _RE_ESPACO_ANTES_PONTUACAO.sub(r"\1", novo)
    return novo


def limpar_ecos_redundantes(rich_list: List[RichText]) -> List[RichText]:
    """Varre a lista de rich text de um parágrafo e, após cada trecho que cita
    uma norma reconhecida (mention já resolvida por extenso), poda o eco
    redundante que segue. O eco pode cair num run diferente do imediatamente
    seguinte (ex.: um trecho em itálico separa a mention do eco em runs
    distintos: ", " + "caput" + " e inciso VI, da CF...") — por isso os runs
    seguintes são concatenados (até a janela) para a busca, e a remoção é
    redistribuída para os runs originais conforme a sobreposição de índices."""
    out: List[RichText] = list(rich_list or [])
    i = 0
    while i < len(out):
        aliases = _aliases_da_norma(out[i].text)
        if not aliases:
            i += 1
            continue
        limites = []  # (índice do run, início, fim) no texto concatenado
        acumulado = ""
        j = i + 1
        while j < len(out) and len(acumulado) < _JANELA_ECO:
            limites.append((j, len(acumulado), len(acumulado) + len(out[j].text)))
            acumulado += out[j].text
            j += 1
        span = _encontrar_eco(acumulado, aliases)
        if span:
            ini, fim = span
            for idx, r_ini, r_fim in limites:
                corte_ini, corte_fim = max(ini, r_ini), min(fim, r_fim)
                if corte_ini < corte_fim:
                    texto_r = out[idx].text
                    novo = texto_r[: corte_ini - r_ini] + texto_r[corte_fim - r_ini :]
                    novo = _RE_ESPACOS_DUPLOS.sub(" ", novo)
                    novo = _RE_ESPACO_ANTES_PONTUACAO.sub(r"\1", novo)
                    out[idx] = replace(out[idx], text=novo)
        i += 1
    return out


def _add_hyperlink(paragraph, text: str, url: str, *, bold: bool = False, italic: bool = False):
    """Insere um hyperlink REAL (clicável) no parágrafo. O python-docx não expõe API
    para isso, então monta-se o elemento w:hyperlink no XML, com a relationship externa
    apontando para a URL. Mantém o link do Notion preservado no .docx (azul + sublinhado)."""
    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _COR_LINK)
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    if bold:
        rpr.append(OxmlElement("w:b"))
    if italic:
        rpr.append(OxmlElement("w:i"))
    run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def preparar_rich(rich_list: List[RichText], *, linkificar: bool = True,
                  dispositivos_proibidos=frozenset()) -> List[RichText]:
    """Cadeia (pura e idempotente) de harmonização das citações de um parágrafo,
    na ordem que os passos exigem: (1) poda de ecos ANTES de qualquer split de
    runs (a janela de busca depende dos offsets originais); (2) resolução de
    links internos do Notion + canonização de URLs "equivalentes" do autor
    (RICD institucional → LEGIN etc.; um href residual do Notion ou fora do
    padrão blindaria a citação contra os passos seguintes); (3) colchetes
    antes de linkificar (expõe "[ADI nº 4.650]" para o padrão casar por
    inteiro); (4) linkificação das citações textuais + padrão "norma ANTES do
    dispositivo" ("CF, art. 121"; "Constituição (art. 121, caput)");
    (5) supressão do "(cf. mention)" redundante (precisa dos hrefs do passo
    4); (6) extensão do link ao dispositivo citado logo após ("Resolução…, no
    art. 28") — depois da supressão, porque o "(cf. mention)" pode estar
    ENTRE o link e o ", no art. N" — e listas de dispositivos coladas a um
    link de norma ("arts. 139, I, e 142 do [RICD]"; "[Código Eleitoral],
    arts. 1º, …"), item a item; (7) anáfora de dispositivo por último entre
    os linkificadores (só herda base quando os passos anteriores não
    resolveram); (8) "caput" em itálico por último (divide runs e quebraria o
    casamento dos padrões)."""
    rich_list = limpar_ecos_redundantes(rich_list)
    if linkificar:
        rich_list = resolver_links_notion_texto(rich_list)
        rich_list = canonicalizar_hrefs(rich_list)
        rich_list = ancorar_links_por_texto(rich_list)
        rich_list = fundir_runs_lisos(rich_list)
        rich_list = limpar_colchetes_citacoes(rich_list)
        rich_list = linkificar_citacoes(rich_list)
        rich_list = linkificar_norma_antes_dispositivo(rich_list)
        rich_list = suprimir_cf_redundante(rich_list)
        rich_list = estender_link_dispositivo(rich_list)
        rich_list = linkificar_dispositivos_junto_a_link(rich_list)
        rich_list = linkificar_dispositivos_anaforicos(rich_list)
        # 2ª passada da supressão remissiva: o dispositivo equivalente à
        # mention pode estar em OUTRO run e só ganhar link nos passos acima
        # (anafórico) — ex. "O art. 13-A disciplina… ([RICD - Art. 13-A],
        # campo texto_em_vigor)" da página PRC: na 1ª passada o art. 13-A do
        # texto ainda não tinha href e o parêntese ficava; sem esta chamada a
        # cadeia nem seria idempotente (a supressão ocorreria na conversão
        # seguinte). Só age nos gatilhos; custo desprezível.
        rich_list = suprimir_cf_redundante(rich_list)
        if dispositivos_proibidos:
            # dispositivos criados pela PRÓPRIA minuta: qualquer link que os
            # cite é desfeito, venha do autor ou dos passos acima
            rich_list = _zerar_links_proibidos(rich_list, dispositivos_proibidos)
    # canonicaliza no fim: os passos acima podem deixar runs lisos adjacentes
    # (ex. supressão do "(cf. …)"), e sem a fusão final uma 2ª passada os
    # uniria — quebrando a idempotência da cadeia
    return fundir_runs_lisos(italicizar_caput(rich_list))


def add_runs(paragraph, rich_list: List[RichText], *, force_bold: Optional[bool] = None,
             linkificar: bool = True, dispositivos_proibidos=frozenset()):
    """Adiciona runs ao parágrafo preservando bold/italic/underline e os hyperlinks
    EXTERNOS (viram links clicáveis no .docx). Links/mentions internos do Notion
    nunca viram hyperlink: quando o texto casa com config.FONTES_OFICIAIS, ganham
    link para a fonte pública oficial; senão, saem como texto puro. Antes de
    renderizar, o parágrafo passa por preparar_rich: poda de ecos redundantes e,
    com `linkificar` (desligado nas transcrições de lei e no articulado), limpeza
    de colchetes editoriais, link às citações textuais de normas/precedentes e
    supressão do "(cf. mention)" redundante. '\n' dentro de um run vira quebra
    de linha (soft break)."""
    rich_list = preparar_rich(rich_list, linkificar=linkificar,
                              dispositivos_proibidos=dispositivos_proibidos)
    for r in rich_list or []:
        if not r.text:
            continue
        if not linkificar:
            # Texto normativo (transcrições, articulado, ementa de proposição)
            # sai LIMPO: nenhum hyperlink, nem mesmo os colocados pelo autor no
            # Notion — o documento oficial protocolado não tem links.
            href = ""
        else:
            href = (r.href or "").strip()
            if _link_interno_notion(href):
                # o match usa o texto completo da mention, não cada parte pós-split
                href = resolver_fonte_publica(r.text) or ""
            else:
                href = _ancorar_dispositivo(href, r.text)
        partes = r.text.split("\n")
        for i, parte in enumerate(partes):
            if i > 0:
                paragraph.add_run().add_break()
            if parte == "":
                continue
            b = r.bold if force_bold is None else force_bold
            if href.startswith("http"):
                _add_hyperlink(paragraph, parte, href, bold=bool(b), italic=bool(r.italic))
                continue
            run = paragraph.add_run(parte)
            if b:
                run.bold = True
            if r.italic:
                run.italic = True
            if r.underline:
                run.underline = True
    return paragraph


# ---------------------------------------------------------------------------
# Detector de normas citadas sem fonte mapeada (vira aviso na conversão)
# ---------------------------------------------------------------------------
# Designação genérica de norma numerada no texto corrido. O lookbehind
# descarta "Projeto de Lei nº 3.263" (a proposição em exame não é lacuna).
_RE_NORMA_GENERICA = re.compile(
    r"(?<!Projeto de )\b"
    r"(?:Lei\s+Complementar|Lei\s+Delegada|Decreto-Lei|Decreto\s+Legislativo|"
    r"Decreto|Emenda\s+Constitucional|Medida\s+Provis[óo]ria|Lei|LC|EC|MPV?)"
    r"\s*n[ºo°.]*\s*\d{1,3}(?:\.\d{3})*"
    r"(?:\s*/\s*\d{2,4}|,?\s+de(?:\s+\d{1,2}[ºo°]?\s+de\s+[a-zA-Zçãéêô]+\s+de)?\s+\d{4})?"
    r"|\bNorma\s+Regulamentadora\s*n[ºo°.]*\s*\d{1,2}(?:\s*\(NR[-\s]?\d{1,2}\))?"
    r"|\bNR[-\s]\d{1,2}\b"
    r"|\bResolu[çc][ãa]o\s+(?:d[oa]\s+)?(?:TSE|CNJ|CNMP|Senado\s+Federal|"
    r"C[âa]mara\s+dos\s+Deputados)\s*n[ºo°.]*\s*[\d.]+(?:\s*/\s*\d{4})?",
    re.I,
)


def _chave_norma(cand: str) -> str:
    """Chave de dedupe: classe + dígitos ("Lei nº 5.889, de 8 de junho de
    1973" e "Lei nº 5.889/1973" são a MESMA lacuna)."""
    s = unicodedata.normalize("NFKD", cand).encode("ascii", "ignore").decode().upper()
    m = re.search(r"[\d.]+", s)
    num = re.sub(r"\D", "", m.group(0)) if m else ""
    if "COMPLEMENTAR" in s or re.match(r"\s*LC\b", s):
        cls = "LC"
    elif "REGULAMENTADORA" in s or re.match(r"\s*NR\b", s):
        cls = "NR"
    elif "EMENDA" in s or re.match(r"\s*EC\b", s):
        cls = "EC"
    elif "PROVIS" in s or re.match(r"\s*MPV?\b", s):
        cls = "MPV"
    elif "DECRETO-LEI" in s:
        cls = "DL"
    elif "DECRETO" in s:
        cls = "DEC"
    elif "RESOLU" in s:
        cls = "RES"
    else:
        cls = "LEI"
    return f"{cls}:{num}"


def detectar_normas_sem_fonte(rich_lists: List[List[RichText]]) -> List[str]:
    """Designações de norma citadas no texto que NÃO sairão com hyperlink no
    .docx: não estão em config.NORMAS_OFICIAIS (a citação não linkifica) nem
    foram linkadas pelo autor. Passa cada parágrafo pela cadeia preparar_rich
    (pura/idempotente — a mesma que o build usa) e só varre os runs que
    restaram SEM href. Retorna a 1ª ocorrência de cada norma, na ordem do
    texto — o pipeline transforma em aviso, para o usuário decidir se mapeia
    a norma na tabela (passo manual: verificar URL/âncoras da fonte oficial)."""
    achados: dict = {}
    for rich in rich_lists or []:
        prep = preparar_rich(list(rich or []))
        # \x00 impede que um candidato atravesse a fronteira entre runs
        solto = "\x00".join(r.text for r in prep if not r.href and r.kind == "text")
        for m in _RE_NORMA_GENERICA.finditer(solto):
            cand = m.group(0).strip()
            chave = _chave_norma(cand)
            if chave in achados:
                continue
            if any(r2.href for r2 in linkificar_citacoes([RichText(cand)])):
                continue  # designação mapeada — não é lacuna
            achados[chave] = cand
    return list(achados.values())


def split_rich_lines(rich_list: List[RichText]) -> List[List[RichText]]:
    """Quebra uma lista de rich text em linhas (separadas por '\n'),
    preservando a formatação de cada trecho."""
    linhas: List[List[RichText]] = [[]]
    for r in rich_list or []:
        partes = r.text.split("\n")
        for i, parte in enumerate(partes):
            if i > 0:
                linhas.append([])
            if parte:
                linhas[-1].append(RichText(parte, r.bold, r.italic, r.underline, r.href))
    return [ln for ln in linhas if any(x.text.strip() for x in ln)]
