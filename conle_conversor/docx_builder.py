# -*- coding: utf-8 -*-
"""Constrói os .docx (Informação Técnica e minuta de proposição) a partir dos
templates, aplicando os estilos nomeados da casa."""
from __future__ import annotations

import re
import unicodedata
from typing import List, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from . import config
from .config import S
from .classifier import TipoProposicao
from .meta import MetaDocumento
from .notion_parser import Block, RichText, plain
from .richtext import add_runs, split_rich_lines
from .splitter import PaginaSeparada, ParecerSeparado, dispositivos_criados
from .harmonizer import Abertura


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------
def _style_name(doc, name: str) -> str:
    try:
        _ = doc.styles[name]
        return name
    except KeyError:
        return "Normal"


# Nível de estrutura de tópicos (painel de Navegação do Word) por estilo de
# título. Nos modelos da casa só o TÍTULO SUBITEM (1.1) traz outlineLvl — os
# itens "1." e os subitens "1.1.1" ficavam FORA da navegação (reclamação do
# usuário, 04/07/26); o builder força o nível correto em cada parágrafo.
_OUTLINE_POR_ESTILO = {S.TITULO_ITEM: 0, S.TITULO_SUB: 1, S.TITULO_SUB2: 2}


def _nivel_topico(para, nivel: int):
    pPr = para._p.get_or_add_pPr()
    for velho in pPr.findall(qn("w:outlineLvl")):
        pPr.remove(velho)
    el = OxmlElement("w:outlineLvl")
    el.set(qn("w:val"), str(nivel))
    pPr.append(el)


def _recuo_alinhado_corpo(doc):
    """Recuo esquerdo que alinha um elemento ao início da 1ª linha dos
    parágrafos do corpo (firstLine do CORPO PADRÃO; 1701 twips nos modelos).
    Padrão do usuário (04/07/26) p/ transcrições e bullets — o left=2268 do
    estilo TRANSCRIÇÃO LEI ficava "mais para dentro" que o corpo."""
    try:
        fl = doc.styles[_style_name(doc, S.CORPO)].paragraph_format.first_line_indent
    except KeyError:
        fl = None
    return fl if fl and int(fl) > 0 else Twips(1701)


def _p(doc, style: str, *, text: Optional[str] = None, rich: Optional[List[RichText]] = None,
       force_bold: Optional[bool] = None, linkificar: bool = True, proibidos=frozenset()):
    para = doc.add_paragraph(style=_style_name(doc, style))
    if style in _OUTLINE_POR_ESTILO:
        _nivel_topico(para, _OUTLINE_POR_ESTILO[style])
    elif style == S.TRANSCRICAO:
        para.paragraph_format.left_indent = _recuo_alinhado_corpo(doc)
    if text is not None:
        run = para.add_run(text)
        if force_bold:
            run.bold = True
        if "  " in text:
            # lacunas de espaços (fecho, epígrafe "Nº     ,") sobrevivem à
            # normalização de whitespace do OOXML só com xml:space=preserve
            for t in run._r.findall(qn("w:t")):
                t.set(qn("xml:space"), "preserve")
    elif rich is not None:
        add_runs(para, rich, force_bold=force_bold, linkificar=linkificar,
                 dispositivos_proibidos=proibidos)
    return para


def _add_bullet(doc, rich: List[RichText], proibidos=frozenset()):
    """Bullet no padrão fixado pelo usuário (04/07/26, ajuste manual replicado):
    recuo esquerdo no início da 1ª linha do corpo (texto alinhado aos demais
    parágrafos) e hanging de 1 twip só para anular o firstLine herdado do
    estilo CORPO PADRÃO."""
    para = doc.add_paragraph(style=_style_name(doc, S.CORPO))
    pf = para.paragraph_format
    pf.left_indent = _recuo_alinhado_corpo(doc)
    pf.first_line_indent = Twips(-1)
    para.add_run("•")
    para.add_run().add_tab()
    add_runs(para, rich, dispositivos_proibidos=proibidos)
    return para


# Linha de supressão do texto legislativo ("Art. 29. ........."): sequência
# longa de pontos. O Notion traz um nº arbitrário de pontos (ex. 125), que em
# fonte proporcional estoura ou não alcança a margem — vira tab com leader.
_RE_PONTILHADO = re.compile(r"\.{8,}")


def _tem_pontilhado(linha: List[RichText]) -> bool:
    return any(_RE_PONTILHADO.search(r.text) for r in linha)


def _tab_margem_direita(doc, style: str):
    """Posição (da margem esquerda do texto) da margem direita efetiva de um
    parágrafo do estilo dado — onde o tab-leader do pontilhado deve terminar."""
    sec = doc.sections[-1]
    largura = sec.page_width - sec.left_margin - sec.right_margin
    try:
        ri = doc.styles[_style_name(doc, style)].paragraph_format.right_indent
    except KeyError:
        ri = None
    return largura - (ri or 0)


def _p_pontilhado(doc, linha: List[RichText], style: str):
    """Parágrafo com linha de supressão: cada sequência longa de pontos vira
    um TAB até tab-stop na margem direita com leader de pontos — preenche a
    linha exata, qualquer que seja o nº de pontos vindo do Notion."""
    para = doc.add_paragraph(style=_style_name(doc, style))
    if style == S.TRANSCRICAO:
        para.paragraph_format.left_indent = _recuo_alinhado_corpo(doc)
    para.paragraph_format.tab_stops.add_tab_stop(
        _tab_margem_direita(doc, style), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    for r in linha:
        partes = _RE_PONTILHADO.split(r.text)
        for i, parte in enumerate(partes):
            if i:
                para.add_run().add_tab()
            if parte:
                run = para.add_run(parte)
                if r.bold:
                    run.bold = True
                if r.italic:
                    run.italic = True
    return para


def _add_quote_as_transcricao(doc, rich: List[RichText]):
    # transcrição de texto de lei: sem linkificação de citações (fidelidade)
    for linha in split_rich_lines(rich):
        if _tem_pontilhado(linha):
            _p_pontilhado(doc, linha, S.TRANSCRICAO)
        else:
            _p(doc, S.TRANSCRICAO, rich=linha, linkificar=False)


def _add_equation(doc, block: Block):
    """Fórmula em bloco (Notion equation). Sem suporte a OMML, renderiza a
    expressão LaTeX como texto centralizado — coerente com o LaTeX que já
    aparece inline no corpo destes documentos."""
    expr = (block.extra.get("expression") or "").strip()
    if not expr:
        return
    para = _p(doc, S.CORPO, text=expr)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Pt(0)
    return para


def _add_table(doc, block: Block, proibidos=frozenset()):
    rows = block.extra.get("rows", [])
    if not rows:
        return
    ncols = block.extra.get("table_width") or max((len(r) for r in rows), default=1)
    has_header = block.extra.get("has_column_header", False) or _looks_like_header(rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    # primeiro estilo de TABELA disponível no template (o fallback "Normal" de
    # _style_name é de parágrafo e não se aplica a tabelas)
    for nome in ("Plain Table 1", "TABELA CONLE", "Table Grid"):
        try:
            table.style = doc.styles[nome]
            break
        except (KeyError, ValueError):
            continue
    table.autofit = True
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            para = cell.paragraphs[0]
            for r in list(para.runs):
                r.text = ""
            cell_rich = row[ci] if ci < len(row) else []
            add_runs(para, cell_rich, force_bold=True if (ri == 0 and has_header) else None,
                     dispositivos_proibidos=proibidos)


def _looks_like_header(rows) -> bool:
    if not rows:
        return False
    first = rows[0]
    return all(any(rt.bold for rt in cell) for cell in first if cell)


def _limpar_hyperlinks_orfaos(doc) -> int:
    """Remove relationships de hyperlink NÃO referenciadas por nenhum w:hyperlink do corpo
    (sujeira herdada do template .docx). Preserva os links reais e as rels de outros tipos."""
    usados = {hl.get(qn("r:id")) for hl in doc.element.iter(qn("w:hyperlink"))}
    usados.discard(None)
    rels = doc.part.rels
    orfas = [rId for rId, rel in list(rels.items())
             if rel.reltype == RELATIONSHIP_TYPE.HYPERLINK and rId not in usados]
    for rId in orfas:
        del rels[rId]
    return len(orfas)


_RE_PROP_CAMARA = re.compile(r"\b(?:PL|PLP|PEC|PRC|PDL|PLV|MPV)\s*n?[ºo°.]*\s*\d", re.IGNORECASE)


def _plain_blocks(blocks: List[Block]) -> str:
    partes: List[str] = []
    for b in blocks:
        if b.rich:
            partes.append(plain(b.rich))
        if b.type == "table":
            for row in b.extra.get("rows", []):
                for cell in row:
                    partes.append(plain(cell))
        if b.children:
            partes.append(_plain_blocks(b.children))
    return "\n".join(partes)


def _tem_risco_apensacao(blocks: List[Block]) -> bool:
    """Há risco de apensação quando a IT cita proposições correlatas concretas em tramitação
    na Câmara (a Seção 6 lista projetos com número: PL/PLP/PEC… nº/ano). Sem correlatas, o
    § 4º do art. 12 da Resolução (projeto coincidente com outro em tramitação) não se aplica
    e é omitido da abertura."""
    return bool(_RE_PROP_CAMARA.search(_plain_blocks(blocks)))


# ---------------------------------------------------------------------------
# corpo da IT
# ---------------------------------------------------------------------------
_FOLD_SUMARIO = ("SUMARIO", "INDICE")


def _substituir_sumario(blocks: List[Block]) -> List[Block]:
    """Página com seção "SUMÁRIO" (heading + lista manual de itens do Notion):
    a lista vira um CAMPO TOC nativo do Word — níveis pelos outline levels dos
    títulos e entradas com hyperlink para o trecho (pedido do usuário,
    06/07/2026; a lista manual saía como texto desconfigurado e sem links).
    O heading vira o tipo sintético "heading_sumario" (fica FORA do próprio
    sumário) e os itens manuais até o próximo heading são descartados."""
    out: List[Block] = []
    i = 0
    while i < len(blocks):
        b = blocks[i]
        if b.type in ("heading_1", "heading_2", "heading_3") and \
                _fold(plain(b.rich)).strip() in _FOLD_SUMARIO:
            out.append(Block(type="heading_sumario", rich=b.rich))
            out.append(Block(type="toc"))
            i += 1
            while i < len(blocks) and not blocks[i].type.startswith("heading"):
                i += 1
            continue
        out.append(b)
        i += 1
    return out


def _add_campo_toc(doc):
    """Campo TOC do Word ({ TOC \\h \\z \\u }): monta o sumário pelos outline
    levels que o builder força nos títulos, com hyperlink em cada entrada.
    w:dirty faz o Word recalcular o campo ao abrir o documento."""
    para = doc.add_paragraph()

    def _fld(tipo, dirty=False):
        r = OxmlElement("w:r")
        f = OxmlElement("w:fldChar")
        f.set(qn("w:fldCharType"), tipo)
        if dirty:
            f.set(qn("w:dirty"), "true")
        r.append(f)
        return r

    para._p.append(_fld("begin", dirty=True))
    r = OxmlElement("w:r")
    t = OxmlElement("w:instrText")
    t.set(qn("xml:space"), "preserve")
    t.text = r" TOC \h \z \u "
    r.append(t)
    para._p.append(r)
    para._p.append(_fld("separate"))
    para.add_run("Sumário gerado pelo Word — se não aparecer, selecione tudo e pressione F9.")
    para._p.append(_fld("end"))
    return para


def _mapa_estilo_headings(blocks: List[Block]) -> dict:
    """Nível dos títulos RELATIVO à página: nas ITs clássicas o topo é o
    heading_2 ("1. INTRODUÇÃO"); em páginas-estudo o topo é o heading_1 e o
    heading_2 é SUBNÍVEL — sem o deslocamento, "5.1" saía achatado no mesmo
    nível do "5" (sumário sem hierarquia, achado do usuário 06/07/2026)."""
    if any(b.type == "heading_1" for b in blocks):
        return {"heading_1": S.TITULO_ITEM, "heading_2": S.TITULO_SUB,
                "heading_3": S.TITULO_SUB2, "heading_4": S.TITULO_SUB2}
    return {"heading_1": S.TITULO_ITEM, "heading_2": S.TITULO_ITEM,
            "heading_3": S.TITULO_SUB, "heading_4": S.TITULO_SUB2}


def _render_it_blocks(doc, blocks: List[Block], proibidos=frozenset()):
    estilo_h = _mapa_estilo_headings(blocks)
    for b in blocks:
        if b.type == "heading_sumario":
            para = _p(doc, S.TITULO_ITEM, rich=b.rich, proibidos=proibidos)
            _nivel_topico(para, 9)  # 9 = sem nível: o título não entra no TOC
        elif b.type == "toc":
            _add_campo_toc(doc)
        elif b.type in estilo_h:
            _p(doc, estilo_h[b.type], rich=b.rich, proibidos=proibidos)
        elif b.type == "paragraph":
            if plain(b.rich).strip():
                _p(doc, S.CORPO, rich=b.rich, proibidos=proibidos)
        elif b.type == "bulleted_list_item":
            _add_bullet(doc, b.rich, proibidos)
        elif b.type == "numbered_list_item":
            _add_bullet(doc, b.rich, proibidos)
        elif b.type == "quote":
            _add_quote_as_transcricao(doc, b.rich)
        elif b.type == "callout":
            if plain(b.rich).strip():
                _p(doc, S.CORPO, rich=b.rich, proibidos=proibidos)
        elif b.type == "table":
            _add_table(doc, b, proibidos)
        elif b.type == "equation":
            _add_equation(doc, b)
        elif b.type == "divider":
            continue
        else:
            if plain(b.rich).strip():
                _p(doc, S.CORPO, rich=b.rich, proibidos=proibidos)


def _fold(s: str) -> str:
    return unicodedata.normalize("NFKD", s or "").encode("ascii", "ignore").decode("ascii").upper()


PRETO = RGBColor(0x00, 0x00, 0x00)


def _run_label(para, texto: str):
    """Rótulo da capa: herda o verde/negrito do estilo CAPA ESTUDO - TÍTULO."""
    r = para.add_run(texto)
    r.font.size = Pt(12)
    r.font.name = "Arial"
    return r


def _run_valor(para, texto: str, *, bold=False):
    """Valor da capa: preto, 12pt, Arial. bold=False reproduz o `b=0` do modelo;
    bold=None herda o negrito do estilo (caso do valor do TEOR)."""
    r = para.add_run(texto)
    r.font.size = Pt(12)
    r.font.name = "Arial"
    r.font.color.rgb = PRETO
    if bold is not None:
        r.bold = bold
    return r


def _espacar(p, before: float, after: float):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.15


def _preencher_capa(doc, teor: str, solicitante: str, autor_linhas):
    """Reconstrói o campo da capa (TEOR / SOLICITANTE / AUTOR) reproduzindo o
    padrão do modelo: rótulo verde-negrito, valor preto; valor do TEOR em
    negrito e os demais sem negrito; espaçamentos próprios por linha."""
    if not doc.tables:
        return
    table = doc.tables[0]
    alvo = None
    for row in table.rows:
        for cell in row.cells:
            if "TEOR DA SOLICITACAO" in _fold(cell.text):
                alvo = cell
                break
        if alvo is not None:
            break
    if alvo is None:
        return

    estilo = _style_name(doc, S.CAPA_TITULO)
    for p in list(alvo.paragraphs):
        p._p.getparent().remove(p._p)

    # TEOR DA SOLICITAÇÃO — valor em negrito (herda), justificado
    p = alvo.add_paragraph(style=estilo)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _espacar(p, 36, 10)
    _run_label(p, "TEOR DA SOLICITAÇÃO: ")
    _run_valor(p, teor, bold=None)

    # SOLICITANTE — valor sem negrito
    p = alvo.add_paragraph(style=estilo)
    _espacar(p, 24, 12)
    _run_label(p, "SOLICITANTE: ")
    _run_valor(p, solicitante, bold=False)

    # AUTOR — valor (consultor + cargo + área) sem negrito, em três linhas
    p = alvo.add_paragraph(style=estilo)
    _espacar(p, 24, 36)
    _run_label(p, "AUTOR: ")
    for i, linha in enumerate(autor_linhas):
        if i > 0:
            _run_valor(p, "", bold=False).add_break()
        _run_valor(p, linha, bold=False)


# ---------------------------------------------------------------------------
# build IT
# ---------------------------------------------------------------------------
def build_it(sep: PaginaSeparada, abertura: Abertura, meta: MetaDocumento) -> Document:
    doc = Document(str(config.TEMPLATE_IT))
    _preencher_capa(doc, abertura.teor, meta.solicitante_capa, meta.autor_linhas)
    # dispositivos criados pelas minutas anexas: citados na IT, saem sem link
    minutas = sep.minutas or [sep]
    proibidos = frozenset().union(
        *(dispositivos_criados(m.articulado_blocks) for m in minutas)
    )

    _p(doc, S.CORPO, text=meta.vocativo)
    _p(doc, S.CORPO, text=abertura.encaminhamento)
    # o § 4º do art. 12 (projeto coincidente com outro em tramitação) só entra se houver
    # risco de apensação — i.e., se a IT cita proposições correlatas concretas
    if _tem_risco_apensacao(sep.it_blocks):
        _p(doc, S.CORPO, text=config.BLOCO_RESOLUCAO_INTRO)
        transcricao = config.BLOCO_RESOLUCAO_TRANSCRICAO
    else:
        _p(doc, S.CORPO, text=config.BLOCO_RESOLUCAO_INTRO_SEM_APENSACAO)
        transcricao = config.BLOCO_RESOLUCAO_TRANSCRICAO_SEM_APENSACAO
    for linha in transcricao:
        if set(linha) == {"."}:
            # linha PONTILHADO do bloco fixo: tab-leader (o nº fixo de pontos
            # era calibrado p/ o recuo antigo; o leader preenche em qualquer)
            _p_pontilhado(doc, [RichText(linha, False, False, False, None)], S.TRANSCRICAO)
        else:
            _p(doc, S.TRANSCRICAO, text=linha)
    _p(doc, S.CORPO, text=abertura.transicao)

    _render_it_blocks(doc, _substituir_sumario(sep.it_blocks), proibidos)

    _p(doc, S.FECHO, text=meta.fecho_it_txt)
    _p(doc, S.ASSINATURA, text=meta.consultor)
    _p(doc, S.ASSINATURA, text=meta.consultor_cargo)
    _sisconle_rodape(doc, meta)
    _limpar_hyperlinks_orfaos(doc)
    return doc


# ---------------------------------------------------------------------------
# build proposição
# ---------------------------------------------------------------------------
_PREAMBULO_RE = re.compile(r"\b(decreta|resolve|promulga)", re.IGNORECASE)


def _epigrafe_text(sep: PaginaSeparada, meta: MetaDocumento) -> str:
    """Normaliza a epígrafe para o padrão da casa, preservando o tipo detectado."""
    nome = sep.tipo.nome_extenso
    return f"{nome} Nº          , DE {meta.ano}"


def _render_articulado(doc, blocks: List[Block], tipo: TipoProposicao):
    # texto normativo proposto: sem linkificação de citações (o articulado da
    # minuta/substitutivo vai a protocolo limpo, sem links azuis no meio)
    for b in blocks:
        if b.type == "paragraph":
            t = plain(b.rich).strip()
            if not t:
                continue
            if _PREAMBULO_RE.search(t) and len(t) <= 120:
                _p(doc, S.PREAMBULO, text=tipo.preambulo)
            elif _tem_pontilhado(b.rich):
                _p_pontilhado(doc, b.rich, S.CORPO)
            else:
                _p(doc, S.CORPO, rich=b.rich, linkificar=False)
        elif b.type == "quote":
            _add_quote_as_transcricao(doc, b.rich)
        elif b.type in ("heading_1", "heading_2", "heading_3", "heading_4"):
            # eventual subtítulo dentro do articulado -> corpo
            _p(doc, S.CORPO, rich=b.rich, linkificar=False)
        elif b.type == "bulleted_list_item" or b.type == "numbered_list_item":
            _p(doc, S.CORPO, rich=b.rich, linkificar=False)
        elif b.type == "table":
            _add_table(doc, b)
        elif b.type == "equation":
            _add_equation(doc, b)


def _render_justificativa(doc, blocks: List[Block], proibidos=frozenset()):
    for b in blocks:
        if b.type == "quote":
            _add_quote_as_transcricao(doc, b.rich)
        elif b.type in ("heading_1", "heading_2", "heading_3", "heading_4"):
            if plain(b.rich).strip():
                _p(doc, S.CORPO, rich=b.rich, proibidos=proibidos)
        elif b.type in ("bulleted_list_item", "numbered_list_item"):
            _add_bullet(doc, b.rich, proibidos)
        elif b.type == "table":
            _add_table(doc, b, proibidos)
        elif b.type == "equation":
            _add_equation(doc, b)
        elif plain(b.rich).strip():
            _p(doc, S.CORPO, rich=b.rich, proibidos=proibidos)


def build_proposicao(sep: PaginaSeparada, meta: MetaDocumento) -> Document:
    doc = Document(str(config.TEMPLATE_PROPOSICAO))
    tipo = sep.tipo

    _p(doc, S.EPIGRAFE, text=_epigrafe_text(sep, meta))
    autoria = meta.autoria_prop
    if tipo.sigla == "PEC":
        # PEC nunca é individual (exige 1/3 da Câmara): "(Do Sr. NOME e outros)",
        # como no modelo oficial da Câmara
        autoria = autoria[:-1] + " e outros)"
    _p(doc, S.AUTORIA, text=autoria)
    # Ementa é parte do texto oficial da proposição: sai limpa, sem hyperlinks
    # (mesma política do articulado), ainda que o autor tenha linkado no Notion.
    _p(doc, S.EMENTA, rich=sep.ementa or [RichText(plain(sep.ementa))], linkificar=False)

    _render_articulado(doc, sep.articulado_blocks, tipo)

    _p(doc, S.JUSTIFICACAO, text="Justificação")
    # os artigos que a própria minuta CRIA não linkam quando citados aqui
    _render_justificativa(doc, sep.justificativa_blocks,
                          dispositivos_criados(sep.articulado_blocks))

    _p(doc, S.FECHO, text=meta.fecho_prop_txt(tipo.local_fecho))
    _p(doc, S.ASSINATURA, text=meta.assinatura_prop)
    _sisconle_rodape(doc, meta)
    _limpar_hyperlinks_orfaos(doc)
    return doc


# ---------------------------------------------------------------------------
# build parecer de comissão (com substitutivo)
# ---------------------------------------------------------------------------
def _relator_txt(par: ParecerSeparado, meta: MetaDocumento) -> str:
    """Relator na ordem de precedência: callout/assinatura da página (já
    resolvidos no splitter) → campo da GUI → placeholder."""
    if par.relator:
        return par.relator
    nome = meta.deputado_nome.strip()
    if nome:
        return f"{meta.tratamento_deputado} {nome}"
    return "Deputado(a) [RELATOR(A)]"


def _rotulo_relator(relator: str) -> str:
    if "Deputado(a)" in relator:
        return "Relator(a)"
    if "Deputada" in relator:
        return "Relatora"
    return "Relator"


def _sisconle_rodape(doc, meta: MetaDocumento):
    """Nº do trabalho SISCONLE no rodapé REAL da página (w:ftr), em todas as
    páginas de todos os documentos — pedido do usuário (04/07/26), que sobrepõe
    o modelo da casa (lá o número é parágrafo no fim do corpo). Preserva o que
    o template já tem no rodapé (timbre da IT, disclaimer da capa): o número
    entra como parágrafo ADICIONAL. Seção sem footer próprio que herda o da
    anterior (corpo da IT) não é desvinculada — herda o rodapé já numerado.
    Cobre first page (titlePg) e páginas pares (evenAndOddHeaders)."""
    for si, section in enumerate(doc.sections):
        footers = [section.footer]
        if section.different_first_page_header_footer:
            footers.append(section.first_page_footer)
        if doc.settings.odd_and_even_pages_header_footer:
            footers.append(section.even_page_footer)
        for footer in footers:
            if footer.is_linked_to_previous:
                if si > 0:
                    continue  # herda (da seção anterior) o rodapé já numerado
                footer.is_linked_to_previous = False  # cria o footer vazio
            if any(meta.sisconle_txt in p.text for p in footer.paragraphs):
                continue
            p0 = footer.paragraphs[0]
            vazio = len(footer.paragraphs) == 1 and not p0.runs and not p0.text.strip()
            para = p0 if vazio else footer.add_paragraph()
            para.style = doc.styles[_style_name(doc, S.SISCONLE)]
            para.add_run(meta.sisconle_txt)


def _render_parecer_blocks(doc, blocks: List[Block], proibidos=frozenset()):
    """Corpo do Relatório/Voto: sub-headings (II.1, II.2…) em TÍTULO SUBITEM;
    o resto segue o padrão do corpo da IT."""
    for b in blocks:
        if b.type in ("heading_1", "heading_2", "heading_3", "heading_4"):
            if plain(b.rich).strip():
                _p(doc, S.TITULO_SUB, rich=b.rich, proibidos=proibidos)
        elif b.type == "paragraph":
            if plain(b.rich).strip():
                _p(doc, S.CORPO, rich=b.rich, proibidos=proibidos)
        elif b.type in ("bulleted_list_item", "numbered_list_item"):
            _add_bullet(doc, b.rich, proibidos)
        elif b.type == "quote":
            _add_quote_as_transcricao(doc, b.rich)
        elif b.type == "callout":
            if plain(b.rich).strip():
                _p(doc, S.CORPO, rich=b.rich, proibidos=proibidos)
        elif b.type == "table":
            _add_table(doc, b, proibidos)
        elif b.type == "equation":
            _add_equation(doc, b)
        elif b.type == "divider":
            continue
        else:
            if plain(b.rich).strip():
                _p(doc, S.CORPO, rich=b.rich, proibidos=proibidos)


def build_parecer(par: ParecerSeparado, meta: MetaDocumento) -> Document:
    """Parecer de comissão no padrão do modelo da casa: cabeçalho de
    identificação, I-Relatório, II-Voto e, em nova página, o Substitutivo —
    tudo em um único .docx."""
    doc = Document(str(config.TEMPLATE_PARECER))
    relator = _relator_txt(par, meta)
    rotulo_rel = _rotulo_relator(relator)
    rotulo_autor = "Autora" if par.autoria.strip().startswith("Deputada") else "Autor"

    # cabeçalho de identificação (os estilos COMISSÃO/EPÍGRAFE forçam caps)
    _p(doc, S.COMISSAO, text=par.comissao or "[COMISSÃO]")
    _p(doc, S.EPIGRAFE, text=par.proposicao or par.titulo)
    if par.apensados:
        _p(doc, S.APENSO, text=f"Apensados: {par.apensados}")
    # dispositivos criados pelo substitutivo: citados no parecer, sem link
    proibidos = dispositivos_criados(par.sub_articulado_blocks)
    if par.ementa:
        _p(doc, S.EMENTA, rich=par.ementa, proibidos=proibidos)
    else:
        _p(doc, S.EMENTA, text="[EMENTA]")
    _p(doc, S.AUTOR_RELATOR, text=f"{rotulo_autor}: {par.autoria or '[AUTOR(A)]'}")
    _p(doc, S.AUTOR_RELATOR, text=f"{rotulo_rel}: {relator}")

    # I – Relatório (o nº SISCONLE fica no rodapé de página, não no corpo)
    _p(doc, S.RELATORIO_VOTO, text=par.relatorio_heading or "I - Relatório")
    _render_parecer_blocks(doc, par.relatorio_blocks, proibidos)
    doc.add_page_break()

    # II – Voto do Relator
    _p(doc, S.RELATORIO_VOTO, text=par.voto_heading or "II - Voto do Relator")
    _render_parecer_blocks(doc, par.voto_blocks, proibidos)
    _p(doc, S.FECHO, text=meta.fecho_prop_txt(config.LOCAL_FECHO_PARECER))
    # assinatura sem o "(Partido/UF)" que o cabeçalho pode trazer — no modelo
    # da casa o relator assina só com o nome
    _p(doc, S.ASSINATURA, text=re.sub(r"\s*\([^)]*\)\s*$", "", relator).strip())
    _p(doc, S.ASSINATURA, text=rotulo_rel)

    # Substitutivo (nova página, mesmo .docx — como no modelo)
    if par.tem_substitutivo:
        doc.add_page_break()
        _p(doc, S.COMISSAO, text=par.comissao or "[COMISSÃO]")
        _p(doc, S.EPIGRAFE, text=par.sub_epigrafe)
        if par.sub_ementa:
            # ementa do substitutivo = texto oficial, sem hyperlinks
            _p(doc, S.EMENTA, rich=par.sub_ementa, linkificar=False)
        _render_articulado(doc, par.sub_articulado_blocks, par.tipo)
        _p(doc, S.FECHO, text=meta.fecho_prop_txt(config.LOCAL_FECHO_PARECER))
        _p(doc, S.ASSINATURA, text=relator)
        _p(doc, S.ASSINATURA, text=rotulo_rel)

    _sisconle_rodape(doc, meta)
    _limpar_hyperlinks_orfaos(doc)
    return doc
