# -*- coding: utf-8 -*-
"""Gera os templates base (corpo vazio) a partir dos documentos-modelo reais.

Preserva timbre (header/footer), estilos nomeados, tabela-cabeçalho e quebras de
seção; remove apenas o conteúdo do corpo. Execução única (ou quando os modelos
mudarem). Uso:  py -m conle_conversor.build_templates
"""
from __future__ import annotations

import shutil
import unicodedata

from docx import Document
from docx.oxml.ns import qn

from . import config


def _fold(s: str) -> str:
    return unicodedata.normalize("NFKD", s or "").encode("ascii", "ignore").decode("ascii").upper()


def _clear_core_props(doc) -> None:
    cp = doc.core_properties
    for attr in ("author", "last_modified_by", "title", "subject", "comments", "category", "keywords"):
        try:
            setattr(cp, attr, "")
        except Exception:
            pass


def _has_sectpr(p_el) -> bool:
    ppr = p_el.find(qn("w:pPr"))
    return ppr is not None and ppr.find(qn("w:sectPr")) is not None


def build_template_it() -> None:
    """IT: mantém tabela-cabeçalho + espaçadores + quebra de seção; esvazia o corpo."""
    config.TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(config.MODELO_IT, config.TEMPLATE_IT)
    doc = Document(str(config.TEMPLATE_IT))
    body = doc.element.body
    children = list(body.iterchildren())

    keep_until = None
    for i, ch in enumerate(children):
        if ch.tag.endswith("}p") and _has_sectpr(ch):
            keep_until = i
            break
    if keep_until is None:
        raise RuntimeError("Não foi encontrada a quebra de seção (sectPr) após o cabeçalho.")

    for ch in children[keep_until + 1:]:
        if ch.tag.endswith("}sectPr"):
            continue  # sectPr final do body — preservar
        body.remove(ch)

    # Limpa o campo TEOR DA SOLICITAÇÃO da tabela-cabeçalho (consolida em 1 parágrafo)
    if doc.tables:
        vistos = []
        for row in doc.tables[0].rows:
            for cell in row.cells:
                tc = cell._tc
                if any(tc is v for v in vistos):
                    continue
                vistos.append(tc)
                if "TEOR DA SOLICITACAO" in _fold(cell.text):
                    for extra in cell.paragraphs[1:]:
                        extra._p.getparent().remove(extra._p)
                    para = cell.paragraphs[0]
                    for r in list(para.runs):
                        r._r.getparent().remove(r._r)
                    para.add_run("TEOR DA SOLICITAÇÃO: ")

    _clear_core_props(doc)
    doc.save(str(config.TEMPLATE_IT))
    print(f"OK template IT  -> {config.TEMPLATE_IT}")


def build_template_proposicao() -> None:
    """Proposição: remove todo o corpo, preservando apenas o sectPr final e os estilos."""
    config.TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(config.MODELO_PROPOSICAO, config.TEMPLATE_PROPOSICAO)
    doc = Document(str(config.TEMPLATE_PROPOSICAO))
    body = doc.element.body
    for ch in list(body.iterchildren()):
        if ch.tag.endswith("}sectPr"):
            continue
        body.remove(ch)
    _clear_core_props(doc)
    doc.save(str(config.TEMPLATE_PROPOSICAO))
    print(f"OK template PROP -> {config.TEMPLATE_PROPOSICAO}")


def build_template_parecer() -> None:
    """Parecer: remove todo o corpo, preservando apenas o sectPr final e os estilos.

    O modelo do parecer não tem tabela-capa nem quebra de seção pós-cabeçalho
    (a quebra de página antes do substitutivo é um w:br manual dentro do corpo,
    removida junto — o builder a recria)."""
    config.TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(config.MODELO_PARECER, config.TEMPLATE_PARECER)
    doc = Document(str(config.TEMPLATE_PARECER))
    body = doc.element.body
    for ch in list(body.iterchildren()):
        if ch.tag.endswith("}sectPr"):
            continue
        body.remove(ch)
    _clear_core_props(doc)
    doc.save(str(config.TEMPLATE_PARECER))
    print(f"OK template PARECER -> {config.TEMPLATE_PARECER}")


def main() -> None:
    build_template_it()
    build_template_proposicao()
    build_template_parecer()


if __name__ == "__main__":
    main()
