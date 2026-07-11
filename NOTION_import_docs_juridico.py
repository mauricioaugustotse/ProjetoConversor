#!/usr/bin/env python3
"""Importa os acervos "Voto Impresso" e "Totalização" (OneDrive) para o Notion.

Fluxo (aprovado em 07/07/2026):
- Inventaria as duas pastas locais, extraindo ZIPs (processos SEI) para staging;
- Deduplica por SHA-256 dentro de cada destino;
- Em cada página destino cria um database inline "Documentos" com metadados
  (Nome, Tipo, Origem, Data, Ano, Nº SEI/Processo, Ordem, Arquivo, Caminho original);
- Cada documento vira uma linha: arquivo original anexado (File Upload API,
  single_part <=20MB / multi_part acima) + texto extraído em blocos nativos
  (PDF via PyMuPDF; HTML do SEI limpo; TXT; DOCX). PDFs sem camada de texto
  ficam só com o anexo. HTML/TXT não são anexados (viram só texto).
- Idempotente: _import_docs_work/manifest.json guarda o progresso por documento.

Uso:
  python NOTION_import_docs_juridico.py --dry-run
  python NOTION_import_docs_juridico.py --target voto_impresso --limit 3
  python NOTION_import_docs_juridico.py            (rodada completa)
"""
from __future__ import annotations

import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent
for _p in (str(PROJECT_ROOT), str(PROJECT_ROOT / "conversores")):
    if _p not in sys.path:
        sys.path.insert(0, _p)

import NOTION_import_codigo_eleitoral_2026_csvs as notion_base  # noqa: E402

import argparse
import csv
import hashlib
import html as html_lib
import json
import logging
import math
import mimetypes
import re
import time
import unicodedata
import zipfile
from datetime import datetime
from typing import Any, Dict, List, Optional, Sequence, Tuple

import requests

LOGGER = logging.getLogger("import_docs_juridico")

BASE_DIR = Path(r"C:\Users\mauri\OneDrive\Documentos\05 - Jurídico e trabalho")
WORK_DIR = PROJECT_ROOT / "_import_docs_work"
STAGING_DIR = WORK_DIR / "staging"
MANIFEST_PATH = WORK_DIR / "manifest.json"
INVENTORY_CSV = WORK_DIR / "inventario.csv"
LOG_PATH = WORK_DIR / "import.log"

NOTION_VERSION = notion_base.DEFAULT_NOTION_VERSION
NOTION_BASE_URL = notion_base.NOTION_BASE_URL

SINGLE_PART_MAX = 20 * 1024 * 1024
MULTI_PART_SIZE = 10 * 1024 * 1024
UPLOAD_TTL_S = 50 * 60  # file uploads não anexados expiram em ~1h
MAX_BLOCKS_PER_DOC = 1500
APPEND_BATCH = 100
SCANNED_CHARS_PER_PAGE = 40

TARGETS: List[Dict[str, str]] = [
    {
        "key": "voto_impresso",
        "titulo": "Voto Impresso",
        "icone": "🗳️",
        "pasta": str(BASE_DIR / "Voto Impresso"),
        "page_id": "39672195-5c64-8072-ad36-d5c265604208",
    },
    {
        "key": "totalizacao",
        "titulo": "Totalização",
        "icone": "🧮",
        "pasta": str(BASE_DIR / "Totalização"),
        "page_id": "39672195-5c64-8025-83ac-ed978b8d6fd5",
    },
]

TEXT_EXTENSIONS = {".txt", ".htm", ".html"}
ATTACH_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".pptx", ".odt", ".rtf", ".png", ".jpg", ".jpeg"}

TIPO_CANONICO = [
    ("termo de concilia", "Termo"),
    ("termo de abertura", "TAP"),
    ("termo de encerramento", "Termo"),
    ("tap", "TAP"),
    ("nota tecnica", "Nota técnica"),
    ("nota técnica", "Nota técnica"),
    ("estudos preliminares", "Estudo"),
    ("estudo preliminar", "Estudo"),
    ("estudo", "Estudo"),
    ("informacao", "Informação"),
    ("informação", "Informação"),
    ("informativo", "Informativo"),
    ("despacho", "Despacho"),
    ("oficio", "Ofício"),
    ("ofício", "Ofício"),
    ("relatorio", "Relatório"),
    ("relatório", "Relatório"),
    ("parecer", "Parecer"),
    ("artigo", "Artigo"),
    ("sentenca", "Sentença"),
    ("sentença", "Sentença"),
    ("memorando", "Memorando"),
    ("e_mail", "E-mail"),
    ("e-mail", "E-mail"),
    ("e mail", "E-mail"),
    ("email", "E-mail"),
    ("minuta", "Minuta"),
    ("planilha", "Planilha"),
    ("carta", "Carta"),
    ("mensagem", "Mensagem"),
    ("requerimento", "Requerimento"),
    ("contrato", "Contrato"),
    ("edital", "Edital"),
    ("ata", "Ata"),
    ("portaria", "Portaria"),
    ("certidao", "Certidão"),
    ("peticao", "Petição"),
    ("procuracao", "Procuração"),
    ("extrato", "Extrato"),
    ("comunicado", "Comunicado"),
    ("anexo", "Anexo"),
    ("plp", "PLP"),
    ("pgp", "Plano"),
    ("plano", "Plano"),
    ("texto", "Texto"),
    ("acao", "Peça judicial"),
    ("ação", "Peça judicial"),
]

RE_SEI_ORDER = re.compile(r"^\[(\d+)\]-(\d+)_(.+)$")
RE_DATE_FULL = re.compile(r"(\d{2})-(\d{2})-(\d{4})")
RE_YEAR = re.compile(r"(?:^|[\s\-(])((?:19|20)\d{2})(?:[\s\-).]|$)")
RE_PROC_SEI = re.compile(r"\d{4}\.\d{2}\.\d{6,9}[-_ ]?\d")
RE_PROC_CNJ = re.compile(r"\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}")
RE_PROC_NUP = re.compile(r"\d{5}\.\d{6}[ /]\d{4}-\d{2}")
RE_SEI_SIGN_DATE = re.compile(
    r"assinado eletronicamente (?:por.{0,300}?)?em (\d{2})/(\d{2})/(\d{4})", re.I | re.S
)
# rodapés de assinatura digital em PDFs (MPF, SEI, PJe): removidos do texto extraído
RE_PDF_SIGN_FOOTER = re.compile(
    r"documento assinado (?:via token )?digitalmente"
    r"|para verificar a assinatura acesse"
    r"|validacaodocumento"
    r"|autenticidade do documento pode ser conferida",
    re.I,
)
PT_ACCENTS = set("áàâãéêíóôõúüçÁÀÂÃÉÊÍÓÔÕÚÜÇ")


# ---------------------------------------------------------------------------
# Utilitários locais
# ---------------------------------------------------------------------------

def norm_ws(text: Any) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def split_rich_text(text: Any) -> List[Dict[str, Any]]:
    """Texto integral em text objects de <=2000 chars, SEM truncar.

    (O chunk_text do módulo base corta em max_chars e insere o marcador
    "[TRUNCADO NO IMPORTADOR NOTION]" — não usar para conteúdo.)
    """
    clean = str(text or "").replace("\x00", "")
    return [notion_base.text_object(clean[i:i + 2000]) for i in range(0, len(clean), 2000)]


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def load_manifest() -> Dict[str, Any]:
    if MANIFEST_PATH.exists():
        return json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    return {"targets": {}}


def save_manifest(manifest: Dict[str, Any]) -> None:
    MANIFEST_PATH.parent.mkdir(parents=True, exist_ok=True)
    tmp = MANIFEST_PATH.with_suffix(".json.tmp")
    tmp.write_text(json.dumps(manifest, ensure_ascii=False, indent=1), encoding="utf-8")
    tmp.replace(MANIFEST_PATH)


def fix_zip_member_name(info: zipfile.ZipInfo) -> str:
    """Corrige mojibake de ZIPs sem flag UTF-8 (zipfile decodifica como cp437).

    Reverte para os bytes originais e escolhe, entre cp850/cp1252/utf-8, a
    decodificação com menos símbolos espúrios e mais acentos pt-BR.
    """
    name = info.filename
    if info.flag_bits & 0x800:
        return name
    try:
        raw = name.encode("cp437")
    except UnicodeEncodeError:
        return name
    candidates = [name]
    for enc in ("cp850", "cp1252", "utf-8"):
        try:
            candidates.append(raw.decode(enc))
        except UnicodeDecodeError:
            continue

    def score(text: str) -> Tuple[int, int]:
        bad = sum(
            1 for c in text
            if 0x2500 <= ord(c) <= 0x25FF or 0x0370 <= ord(c) <= 0x03FF or c in "‡†ƒ…•¤¦§¨�"
        )
        good = sum(1 for c in text if c in PT_ACCENTS)
        return (bad, -good)

    return min(candidates, key=score)


def safe_component(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]', "_", name)
    return norm_ws(cleaned) or "arquivo"


# ---------------------------------------------------------------------------
# Fase 1 — inventário e staging
# ---------------------------------------------------------------------------

def zip_origin_label(zip_path: Path) -> Optional[str]:
    stem = zip_path.stem
    match = re.match(r"SEI_(\d{4})\.(\d{2})\.(\d{6,9})_(\d)(?:\s*-\s*(.+))?$", stem)
    if match:
        base = f"SEI {match.group(1)}.{match.group(2)}.{match.group(3)}-{match.group(4)}"
        extra = norm_ws(match.group(5) or "")
        return f"{base} ({extra})" if extra else base
    if stem.lower().startswith("acao_popular"):
        return "Ação Popular — Resposta STI"
    if stem.lower().startswith("relatoriospf"):
        return None  # zip vazio/redundante (conteúdo já existe extraído na subpasta)
    return stem.replace("_", " ")


def zip_process_number(zip_path: Path) -> str:
    match = re.match(r"SEI_(\d{4})\.(\d{2})\.(\d{6,9})_(\d)", zip_path.stem)
    if match:
        return f"SEI {match.group(1)}.{match.group(2)}.{match.group(3)}-{match.group(4)}"
    return ""


def extract_zip(zip_path: Path, dest_root: Path) -> List[Tuple[Path, str]]:
    """Extrai o zip para staging; retorna [(caminho_extraido, nome_corrigido)]."""
    out: List[Tuple[Path, str]] = []
    dest_root.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(zip_path) as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            fixed = fix_zip_member_name(info)
            base_name = safe_component(Path(fixed.replace("\\", "/")).name)
            target = dest_root / base_name
            if not target.exists() or target.stat().st_size != info.file_size:
                with zf.open(info) as src, open(target, "wb") as dst:
                    while True:
                        chunk = src.read(1024 * 1024)
                        if not chunk:
                            break
                        dst.write(chunk)
            out.append((target, base_name))
    return out


def detect_tipo(raw: str) -> str:
    low = raw.lower().replace("_", " ").replace("-", " ")
    low_ascii = unicodedata.normalize("NFKD", low).encode("ascii", "ignore").decode()
    for prefix, canon in TIPO_CANONICO:
        prefix_ascii = unicodedata.normalize("NFKD", prefix.replace("-", " ")).encode("ascii", "ignore").decode()
        if low_ascii.startswith(prefix_ascii):
            return canon
    first = norm_ws(raw.replace("_", " ").replace("-", " ")).split(" ")[0]
    if first.isalpha() and 3 <= len(first) <= 25:
        return first[:1].upper() + first[1:].lower()
    return "Documento"


def parse_dates(raw: str) -> Tuple[str, Optional[int]]:
    date_iso = ""
    match = RE_DATE_FULL.search(raw)
    if match:
        day, month, year = match.groups()
        try:
            date_iso = datetime(int(year), int(month), int(day)).strftime("%Y-%m-%d")
        except ValueError:
            date_iso = ""
    year_val: Optional[int] = None
    if date_iso:
        year_val = int(date_iso[:4])
    else:
        years = [int(y) for y in RE_YEAR.findall(raw)]
        if years:
            year_val = max(years)
    return date_iso, year_val


def parse_processo(raw: str) -> str:
    for regex in (RE_PROC_CNJ, RE_PROC_SEI, RE_PROC_NUP):
        match = regex.search(raw)
        if match:
            value = match.group(0)
            if regex is RE_PROC_SEI:
                value = re.sub(r"[-_ ]?(\d)$", r"-\1", value)
            return value
    return ""


def doc_from_sei_member(display_name: str) -> Tuple[str, str, Optional[int], str]:
    """[NN]-1234567_Despacho.html -> (titulo, tipo, ordem, num_doc)."""
    stem = Path(display_name).stem
    match = RE_SEI_ORDER.match(stem)
    if not match:
        titulo = norm_ws(stem.replace("_", " "))
        if " " not in titulo and titulo.count("-") >= 2:
            titulo = norm_ws(titulo.replace("-", " "))
        return titulo, detect_tipo(titulo), None, ""
    ordem = int(match.group(1))
    num_doc = match.group(2)
    resto = norm_ws(match.group(3).replace("_", " "))
    if resto.lower() == "e mail":
        resto = "E-mail"
    tipo = detect_tipo(resto)
    titulo = f"{resto} {num_doc}" if resto else f"Documento {num_doc}"
    return titulo, tipo, ordem, num_doc


def inventory_target(target: Dict[str, str]) -> List[Dict[str, Any]]:
    root = Path(target["pasta"])
    docs: List[Dict[str, Any]] = []

    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        rel = path.relative_to(root)
        ext = path.suffix.lower()
        if ext == ".zip":
            origin = zip_origin_label(path)
            if origin is None:
                LOGGER.info("Ignorando zip vazio/redundante: %s", path.name)
                continue
            staging = STAGING_DIR / target["key"] / safe_component(path.stem)
            processo_zip = zip_process_number(path)
            for extracted, display in extract_zip(path, staging):
                inner_ext = extracted.suffix.lower()
                titulo, tipo, ordem, num_doc = doc_from_sei_member(display)
                date_iso, year_val = parse_dates(display)
                docs.append({
                    "arquivo": str(extracted),
                    "nome_exibicao": display,
                    "titulo": titulo,
                    "tipo": tipo,
                    "origem": origin,
                    "ordem": ordem,
                    "processo": processo_zip or parse_processo(display),
                    "num_doc_sei": num_doc,
                    "data": date_iso,
                    "ano": year_val,
                    "caminho_original": f"{path}::{display}",
                    "tamanho": extracted.stat().st_size,
                    "ext": inner_ext,
                })
            continue

        parent_rel = rel.parent
        origem = "Pasta raiz" if str(parent_rel) == "." else str(parent_rel).replace("\\", " / ")
        stem = path.stem
        if stem.lower().endswith(".pdf"):  # nome com extensão dupla ("....pdf.pdf")
            stem = stem[:-4]
        titulo = norm_ws(stem.replace("_", " "))
        date_iso, year_val = parse_dates(path.name)
        docs.append({
            "arquivo": str(path),
            "nome_exibicao": path.name,
            "titulo": titulo,
            "tipo": detect_tipo(titulo),
            "origem": origem,
            "ordem": None,
            "processo": parse_processo(path.name),
            "num_doc_sei": "",
            "data": date_iso,
            "ano": year_val,
            "caminho_original": str(path),
            "tamanho": path.stat().st_size,
            "ext": ext,
        })

    # dedup por hash dentro do destino
    by_hash: Dict[str, Dict[str, Any]] = {}
    duplicates: List[Tuple[str, str]] = []
    for doc in docs:
        digest = sha256_file(Path(doc["arquivo"]))
        doc["sha256"] = digest
        if digest in by_hash:
            kept = by_hash[digest]
            kept.setdefault("origens_extras", []).append(doc["caminho_original"])
            duplicates.append((doc["caminho_original"], kept["caminho_original"]))
        else:
            by_hash[digest] = doc
    unique = list(by_hash.values())
    for dup, kept in duplicates:
        LOGGER.info("Duplicata ignorada: %s (== %s)", dup, kept)

    def sort_key(doc: Dict[str, Any]) -> Tuple:
        return (doc["origem"] != "Pasta raiz", doc["origem"], doc["ordem"] if doc["ordem"] is not None else 10**6, doc["titulo"].lower())

    unique.sort(key=sort_key)
    for doc in unique:
        doc["doc_id"] = doc["sha256"][:16]
    return unique


# ---------------------------------------------------------------------------
# Extração de texto
# ---------------------------------------------------------------------------

def _norm_pdf_line(line: str) -> str:
    """Normaliza linha p/ detectar recorrência: dígitos e códigos hex viram '#'."""
    text = norm_ws(line)
    text = re.sub(r"(?=[0-9A-Fa-f.]*\d)[0-9A-Fa-f]{4,}", "#", text)
    return re.sub(r"\d+", "#", text)


def _recurrent_page_lines(pages: Sequence[str]) -> set:
    """Linhas de cabeçalho/rodapé: recorrentes entre páginas (normalizadas)."""
    if len(pages) < 4:
        return set()
    from collections import Counter

    freq: Counter = Counter()
    for page in pages:
        lines = {_norm_pdf_line(l) for l in page.split("\n") if norm_ws(l)}
        freq.update(l for l in lines if len(l) <= 160)
    threshold = max(3, math.ceil(0.35 * len(pages)))
    return {line for line, count in freq.items() if count >= threshold}


def paragraphs_from_pdf(path: Path) -> Tuple[List[str], bool]:
    import fitz  # PyMuPDF

    with fitz.open(str(path)) as pdf:
        pages = [page.get_text("text") for page in pdf]
    total_chars = sum(len(norm_ws(p)) for p in pages)
    if not pages or total_chars / max(1, len(pages)) < SCANNED_CHARS_PER_PAGE:
        return [], True
    recurrent = _recurrent_page_lines(pages)
    cleaned: List[str] = []
    for page in pages:
        kept = [
            line for line in page.split("\n")
            if not (norm_ws(line) and (_norm_pdf_line(line) in recurrent or RE_PDF_SIGN_FOOTER.search(line)))
        ]
        cleaned.append("\n".join(kept))
    text = "\n".join(cleaned)
    text = re.sub(r"-[ \t]*\n\s*(?=[a-zà-úç])", "", text)  # des-hifenização (mesmo entre páginas)
    paragraphs: List[str] = []
    for block in re.split(r"\n\s*\n", text):
        joined = norm_ws(block.replace("\n", " "))
        if joined:
            paragraphs.append(joined)
    # emenda parágrafos partidos na virada de página: anterior sem pontuação final
    # e continuação começando em minúscula
    merged: List[str] = []
    for para in paragraphs:
        if merged and merged[-1][-1] == "-" and para[:1].islower():
            merged[-1] = merged[-1][:-1] + para
        elif merged and merged[-1][-1] not in ".;:!?\"'»)]}" and para[:1].islower():
            merged[-1] = merged[-1] + " " + para
        else:
            merged.append(para)
    return merged, False


def paragraphs_from_html(path: Path) -> List[str]:
    raw = path.read_bytes()
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1", "replace")
    text = re.sub(r"<(style|script|head)[^>]*>.*?</\1>", " ", text, flags=re.S | re.I)
    text = re.sub(r"<img[^>]*>", " ", text, flags=re.I)
    text = re.sub(r"<(br|/p|/div|/tr|/li|/h[1-6]|/table)[^>]*>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = html_lib.unescape(text)
    text = text.replace("\xa0", " ")
    paragraphs = []
    for line in text.split("\n"):
        cleaned = norm_ws(line)
        if cleaned:
            paragraphs.append(cleaned)
    return paragraphs


def paragraphs_from_txt(path: Path) -> List[str]:
    raw = path.read_bytes()
    for enc in ("utf-8", "cp1252", "latin-1"):
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    paragraphs = []
    for block in re.split(r"\n\s*\n", text):
        joined = norm_ws(block.replace("\n", " "))
        if joined:
            paragraphs.append(joined)
    return paragraphs


def paragraphs_from_docx(path: Path) -> List[str]:
    import docx

    document = docx.Document(str(path))
    paragraphs = [norm_ws(p.text) for p in document.paragraphs if norm_ws(p.text)]
    for table in document.tables:
        for row in table.rows:
            cells = [norm_ws(c.text) for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                paragraphs.append(line)
    return paragraphs


def extract_paragraphs(doc: Dict[str, Any]) -> Tuple[List[str], bool]:
    """Retorna (parágrafos, escaneado_sem_texto)."""
    path = Path(doc["arquivo"])
    ext = doc["ext"]
    try:
        if ext == ".pdf":
            return paragraphs_from_pdf(path)
        if ext in {".htm", ".html"}:
            return paragraphs_from_html(path), False
        if ext == ".txt":
            return paragraphs_from_txt(path), False
        if ext == ".docx":
            return paragraphs_from_docx(path), False
    except Exception as exc:
        LOGGER.warning("Falha ao extrair texto de %s: %s", path.name, exc)
        return [], False
    return [], False


def paragraphs_to_blocks(paragraphs: Sequence[str]) -> Tuple[List[Dict[str, Any]], bool]:
    blocks: List[Dict[str, Any]] = []
    truncated = False
    for para in paragraphs:
        rich = split_rich_text(para)
        for start in range(0, len(rich), 25):
            blocks.append({
                "object": "block",
                "type": "paragraph",
                "paragraph": {"rich_text": rich[start:start + 25]},
            })
            if len(blocks) >= MAX_BLOCKS_PER_DOC:
                truncated = True
                break
        if truncated:
            break
    if truncated:
        blocks.append({
            "object": "block",
            "type": "callout",
            "callout": {
                "rich_text": [notion_base.text_object("Texto truncado por tamanho — íntegra no arquivo anexado (propriedade Arquivo).")],
                "icon": {"type": "emoji", "emoji": "✂️"},
                "color": "yellow_background",
            },
        })
    return blocks, truncated


# ---------------------------------------------------------------------------
# Upload de arquivos (File Upload API)
# ---------------------------------------------------------------------------

def upload_headers(client: "notion_base.NotionClient") -> Dict[str, str]:
    return {
        "Authorization": client.headers["Authorization"],
        "Notion-Version": client.headers["Notion-Version"],
    }


def send_upload_part(client: "notion_base.NotionClient", upload_id: str, data: bytes,
                     filename: str, content_type: str, part_number: Optional[int] = None) -> None:
    url = f"{NOTION_BASE_URL}/v1/file_uploads/{upload_id}/send"
    form: Dict[str, str] = {}
    if part_number is not None:
        form["part_number"] = str(part_number)
    last_error = ""
    for attempt in range(1, 7):
        client.pace()
        try:
            response = requests.post(
                url,
                headers=upload_headers(client),
                files={"file": (filename, data, content_type)},
                data=form or None,
                timeout=300,
            )
        except requests.RequestException as exc:
            last_error = str(exc)
            time.sleep(min(30.0, 2.0 * attempt))
            continue
        if 200 <= response.status_code < 300:
            return
        last_error = f"{response.status_code}: {response.text[:300]}"
        if response.status_code == 429 or response.status_code >= 500:
            time.sleep(min(30.0, 2.0 * attempt))
            continue
        break
    raise RuntimeError(f"Falha no envio do arquivo ({filename}): {last_error}")


def upload_file(client: "notion_base.NotionClient", path: Path, display_name: str) -> str:
    size = path.stat().st_size
    content_type = mimetypes.guess_type(display_name)[0] or "application/octet-stream"
    filename = safe_component(display_name)[:900]
    if size <= SINGLE_PART_MAX:
        created = client.request("POST", "/v1/file_uploads", json_body={
            "mode": "single_part",
            "filename": filename,
            "content_type": content_type,
        })
        upload_id = str(created.get("id", ""))
        send_upload_part(client, upload_id, path.read_bytes(), filename, content_type)
        return upload_id

    parts = math.ceil(size / MULTI_PART_SIZE)
    created = client.request("POST", "/v1/file_uploads", json_body={
        "mode": "multi_part",
        "number_of_parts": parts,
        "filename": filename,
        "content_type": content_type,
    })
    upload_id = str(created.get("id", ""))
    with open(path, "rb") as handle:
        for part_number in range(1, parts + 1):
            data = handle.read(MULTI_PART_SIZE)
            send_upload_part(client, upload_id, data, filename, content_type, part_number=part_number)
            LOGGER.info("  multi-part %d/%d enviado (%s)", part_number, parts, display_name)
    client.request("POST", f"/v1/file_uploads/{upload_id}/complete")
    return upload_id


# ---------------------------------------------------------------------------
# Estrutura no Notion
# ---------------------------------------------------------------------------

def database_schema(origens: Sequence[str]) -> Dict[str, Any]:
    return {
        "Nome": {"title": {}},
        "Tipo": {"select": {"options": []}},
        "Origem": {"select": {"options": [{"name": o} for o in origens]}},
        "Data": {"date": {}},
        "Ano": {"number": {"format": "number"}},
        "Nº SEI/Processo": {"rich_text": {}},
        "Ordem": {"number": {"format": "number"}},
        "Arquivo": {"files": {}},
        "Caminho original": {"rich_text": {}},
    }


def ensure_target_setup(client: "notion_base.NotionClient", target: Dict[str, str],
                        state: Dict[str, Any], origens: Sequence[str], total_docs: int) -> None:
    page_id = target["page_id"]
    if not state.get("renamed"):
        client.request("PATCH", f"/v1/pages/{page_id}", json_body={
            "properties": {"title": {"title": [notion_base.text_object(target["titulo"])]}},
            "icon": {"type": "emoji", "emoji": target["icone"]},
        })
        state["renamed"] = True
        LOGGER.info("Página %s renomeada para '%s'", page_id, target["titulo"])

    if not state.get("callout_id"):
        texto = (
            f"Acervo importado do OneDrive ({Path(target['pasta']).name}) em "
            f"{datetime.now().strftime('%d/%m/%Y')} — {total_docs} documentos únicos, "
            "incluindo o conteúdo dos ZIPs de processos SEI. Arquivo original na propriedade "
            "\"Arquivo\" de cada item; texto pesquisável no corpo das páginas."
        )
        payload = client.request("PATCH", f"/v1/blocks/{page_id}/children", json_body={
            "children": [{
                "object": "block",
                "type": "callout",
                "callout": {
                    "rich_text": [notion_base.text_object(texto)],
                    "icon": {"type": "emoji", "emoji": "📥"},
                    "color": "gray_background",
                },
            }],
        })
        results = payload.get("results") or [{}]
        state["callout_id"] = results[0].get("id", "")

    if not state.get("data_source_id"):
        title_rich = [notion_base.text_object(f"Documentos — {target['titulo']}")]
        payload = client.request("POST", "/v1/databases", json_body={
            "parent": {"type": "page_id", "page_id": page_id},
            "title": title_rich,
            "is_inline": True,
            "initial_data_source": {
                "title": title_rich,
                "properties": database_schema(origens),
            },
        })
        database_id = notion_base.normalize_notion_id(str(payload.get("id", "")))
        data_source_id = notion_base.extract_data_source_id(client, database_id, payload)
        state["database_id"] = database_id
        state["data_source_id"] = data_source_id
        LOGGER.info("Database criado: %s (data_source %s)", database_id, data_source_id)


def doc_properties(doc: Dict[str, Any], attach_upload_id: str = "") -> Dict[str, Any]:
    props: Dict[str, Any] = {
        "Nome": {"title": split_rich_text(doc["titulo"])},
        "Tipo": {"select": {"name": doc["tipo"]}},
        "Origem": {"select": {"name": doc["origem"][:100]}},
        "Nº SEI/Processo": {"rich_text": split_rich_text(doc.get("processo", ""))},
        "Caminho original": {"rich_text": split_rich_text(doc["caminho_original"])},
    }
    if doc.get("data"):
        props["Data"] = {"date": {"start": doc["data"]}}
    if doc.get("ano") is not None:
        props["Ano"] = {"number": doc["ano"]}
    if doc.get("ordem") is not None:
        props["Ordem"] = {"number": doc["ordem"]}
    if attach_upload_id:
        props["Arquivo"] = {"files": [{
            "type": "file_upload",
            "file_upload": {"id": attach_upload_id},
            "name": safe_component(doc["nome_exibicao"])[:100],
        }]}
    return props


def import_document(client: "notion_base.NotionClient", doc: Dict[str, Any],
                    doc_state: Dict[str, Any], data_source_id: str,
                    manifest: Dict[str, Any]) -> None:
    if doc_state.get("status") == "done":
        return
    path = Path(doc["arquivo"])
    should_attach = doc["ext"] in ATTACH_EXTENSIONS

    # 1) upload do arquivo (uploads não anexados expiram; refaz se envelheceu)
    upload_id = doc_state.get("file_upload_id", "")
    if should_attach and not doc_state.get("page_id"):
        age = time.time() - float(doc_state.get("upload_ts", 0))
        if not upload_id or age > UPLOAD_TTL_S:
            LOGGER.info("Upload: %s (%.1f MB)", doc["nome_exibicao"], doc["tamanho"] / 1e6)
            upload_id = upload_file(client, path, doc["nome_exibicao"])
            doc_state["file_upload_id"] = upload_id
            doc_state["upload_ts"] = time.time()
            save_manifest(manifest)

    # 2) texto extraído -> blocos
    paragraphs, scanned = extract_paragraphs(doc)
    if not doc.get("data") and paragraphs:
        match = RE_SEI_SIGN_DATE.search("\n".join(paragraphs))
        if match:
            day, month, year = (int(g) for g in match.groups())
            if 2000 <= year <= 2027:
                try:
                    doc["data"] = datetime(year, month, day).strftime("%Y-%m-%d")
                    doc["ano"] = year
                except ValueError:
                    pass
    blocks, _truncated = paragraphs_to_blocks(paragraphs)
    if scanned:
        blocks = [{
            "object": "block",
            "type": "callout",
            "callout": {
                "rich_text": [notion_base.text_object("PDF sem camada de texto (digitalizado) — conteúdo disponível no arquivo anexado.")],
                "icon": {"type": "emoji", "emoji": "🖨️"},
                "color": "gray_background",
            },
        }]

    # embed do arquivo no topo do corpo (mesmo file_upload da property Arquivo)
    if should_attach and upload_id:
        if doc["ext"] == ".pdf":
            embed_type = "pdf"
        elif doc["ext"] in {".png", ".jpg", ".jpeg"}:
            embed_type = "image"
        else:
            embed_type = "file"
        blocks.insert(0, {
            "object": "block",
            "type": embed_type,
            embed_type: {"type": "file_upload", "file_upload": {"id": upload_id}},
        })

    # 3) cria a página com properties + primeiro lote de children
    if not doc_state.get("page_id"):
        first_batch = blocks[:APPEND_BATCH]
        body: Dict[str, Any] = {
            "parent": {"type": "data_source_id", "data_source_id": data_source_id},
            "properties": doc_properties(doc, attach_upload_id=upload_id if should_attach else ""),
        }
        if first_batch:
            body["children"] = first_batch
        payload = client.request("POST", "/v1/pages", json_body=body)
        doc_state["page_id"] = notion_base.normalize_notion_id(str(payload.get("id", "")))
        doc_state["blocks_total"] = len(blocks)
        doc_state["blocks_appended"] = len(first_batch)
        save_manifest(manifest)

    # 4) lotes restantes
    appended = int(doc_state.get("blocks_appended", 0))
    while appended < len(blocks):
        batch = blocks[appended:appended + APPEND_BATCH]
        client.request("PATCH", f"/v1/blocks/{doc_state['page_id']}/children", json_body={"children": batch})
        appended += len(batch)
        doc_state["blocks_appended"] = appended
        save_manifest(manifest)

    doc_state["status"] = "done"
    doc_state["scanned"] = scanned
    save_manifest(manifest)


# ---------------------------------------------------------------------------
# Relatórios
# ---------------------------------------------------------------------------

def write_inventory_csv(all_docs: Dict[str, List[Dict[str, Any]]]) -> None:
    INVENTORY_CSV.parent.mkdir(parents=True, exist_ok=True)
    with open(INVENTORY_CSV, "w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle, delimiter=";")
        writer.writerow(["destino", "titulo", "tipo", "origem", "ordem", "data", "ano", "processo", "tamanho_mb", "ext", "caminho_original", "origens_extras"])
        for key, docs in all_docs.items():
            for doc in docs:
                writer.writerow([
                    key, doc["titulo"], doc["tipo"], doc["origem"],
                    doc["ordem"] if doc["ordem"] is not None else "",
                    doc["data"], doc["ano"] or "", doc["processo"],
                    f"{doc['tamanho'] / 1e6:.2f}", doc["ext"], doc["caminho_original"],
                    " || ".join(doc.get("origens_extras", [])),
                ])


def print_summary(all_docs: Dict[str, List[Dict[str, Any]]]) -> None:
    for key, docs in all_docs.items():
        print(f"\n=== {key}: {len(docs)} documentos únicos ===")
        by_origin: Dict[str, int] = {}
        by_tipo: Dict[str, int] = {}
        dups = 0
        for doc in docs:
            by_origin[doc["origem"]] = by_origin.get(doc["origem"], 0) + 1
            by_tipo[doc["tipo"]] = by_tipo.get(doc["tipo"], 0) + 1
            dups += len(doc.get("origens_extras", []))
        for origin, count in sorted(by_origin.items()):
            print(f"  origem: {origin:<55} {count:>4}")
        tipos = ", ".join(f"{t}={c}" for t, c in sorted(by_tipo.items(), key=lambda x: -x[1]))
        print(f"  tipos: {tipos}")
        if dups:
            print(f"  duplicatas puladas (mesmo conteúdo): {dups}")


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------

def main() -> int:
    parser = argparse.ArgumentParser(description="Importa Voto Impresso e Totalização para o Notion")
    parser.add_argument("--dry-run", action="store_true", help="só inventário (não toca o Notion)")
    parser.add_argument("--target", choices=["voto_impresso", "totalizacao", "all"], default="all")
    parser.add_argument("--limit", type=int, default=0, help="importa só os N primeiros documentos por destino")
    parser.add_argument("--rate-rps", type=float, default=2.5)
    args = parser.parse_args()

    WORK_DIR.mkdir(parents=True, exist_ok=True)
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s %(levelname)s %(message)s",
        handlers=[logging.StreamHandler(), logging.FileHandler(LOG_PATH, encoding="utf-8")],
    )

    targets = [t for t in TARGETS if args.target in ("all", t["key"])]

    all_docs: Dict[str, List[Dict[str, Any]]] = {}
    for target in targets:
        LOGGER.info("Inventariando %s ...", target["pasta"])
        all_docs[target["key"]] = inventory_target(target)

    write_inventory_csv(all_docs)
    print_summary(all_docs)
    print(f"\nInventário salvo em {INVENTORY_CSV}")
    if args.dry_run:
        return 0

    token = (PROJECT_ROOT / "Chave_Notion.txt").read_text(encoding="utf-8").strip()
    client = notion_base.NotionClient(
        token,
        notion_version=NOTION_VERSION,
        timeout_s=90,
        max_retries=6,
        rate_rps=args.rate_rps,
    )

    manifest = load_manifest()
    exit_code = 0
    for target in targets:
        docs = all_docs[target["key"]]
        if args.limit > 0:
            docs = docs[:args.limit]
        state = manifest["targets"].setdefault(target["key"], {"docs": {}})
        origens = sorted({d["origem"] for d in all_docs[target["key"]]})
        ensure_target_setup(client, target, state, origens, total_docs=len(all_docs[target["key"]]))
        save_manifest(manifest)

        failures: List[str] = []
        for index, doc in enumerate(docs, start=1):
            doc_state = state["docs"].setdefault(doc["doc_id"], {})
            if doc_state.get("status") == "done":
                continue
            LOGGER.info("[%s %d/%d] %s", target["key"], index, len(docs), doc["titulo"])
            try:
                import_document(client, doc, doc_state, state["data_source_id"], manifest)
            except Exception as exc:
                doc_state["erro"] = str(exc)[:500]
                save_manifest(manifest)
                failures.append(f"{doc['titulo']}: {exc}")
                LOGGER.error("FALHA em %s: %s", doc["titulo"], exc)

        done = sum(1 for d in state["docs"].values() if d.get("status") == "done")
        scanned = sum(1 for d in state["docs"].values() if d.get("scanned"))
        print(f"\n=== {target['key']}: {done} importados | {scanned} sem camada de texto | {len(failures)} falhas ===")
        for failure in failures:
            print(f"  FALHA: {failure}")
        if failures:
            exit_code = 1

    return exit_code


if __name__ == "__main__":
    raise SystemExit(main())
