# -*- coding: utf-8 -*-
"""tsje_07: monta os .docx anuais de transcricoes fieis (formato do ensaio:
blocos ano / mes / ID / TITULO / texto), com pagina de rosto explicando o
metodo, e empacota em D:\\Atas TSJE - Transc Claude - 1932-1937.zip.

Inclui so atas com .md pronto (status transcrita/revisada/final). Atas
pendentes/problema viram lista de pendencias no fim do docx do ano.
PDF por ano via docx2pdf (Word) quando disponivel.

Uso: python tsje_07_gerar_docx.py [--ano 1932] [--zip]
"""
import argparse
import csv
import io
import os
import re
import time
import zipfile
from collections import defaultdict

from docx import Document
from docx.shared import Pt, Cm

RAIZ = r'D:\TSJE_TRANSCRICOES'
MANIFEST = os.path.join(RAIZ, 'manifest.csv')
ZIP_FINAL = r'D:\Atas TSJE - Transc Claude - 1932-1937.zip'

MES_NOME = {1: 'Janeiro', 2: 'Fevereiro', 3: 'Março', 4: 'Abril', 5: 'Maio',
            6: 'Junho', 7: 'Julho', 8: 'Agosto', 9: 'Setembro', 10: 'Outubro',
            11: 'Novembro', 12: 'Dezembro'}


def front_matter(path):
    with io.open(path, encoding='utf-8') as f:
        texto = f.read()
    m = re.match(r'---\n(.*?)\n---\n?(.*)', texto, re.S)
    fm = {}
    if not m:
        return fm, texto.strip()
    for linha in m.group(1).split('\n'):
        if ':' in linha:
            k, v = linha.split(':', 1)
            fm[k.strip()] = v.strip()
    return fm, m.group(2).strip()


def fmt_id(arquivo):
    m = re.search(r'\[(\d+)\]', arquivo)
    if not m:
        return ''
    s = m.group(1)
    return f'ID {s[:-3]}.{s[-3:]}' if len(s) > 3 else f'ID {s}'


def rosto(doc, ano, feitas, total, fid_media):
    doc.add_heading(f'Atas do Tribunal Superior de Justiça Eleitoral — {ano}', 0)
    doc.add_paragraph(
        'Transcrição fiel das atas das sessões do TSJE publicadas no '
        'Boletim Eleitoral, feita a partir das imagens das páginas '
        'digitalizadas (banco D:\\TSJE_ATAS).')
    p = doc.add_paragraph()
    p.add_run('Método: ').bold = True
    p.add_run(
        'transcrição frase a frase por Claude Fable 5 (Anthropic) lendo as '
        'imagens dos Boletins; ortografia modernizada, nomes próprios na '
        'grafia original; revisão por gpt-5.6-terra (OpenAI) nos casos '
        'difíceis e em amostra de controle. O texto NÃO é paráfrase: '
        'reproduz o conteúdo integral impresso.')
    p = doc.add_paragraph()
    p.add_run('Convenções: ').bold = True
    p.add_run('[ilegível] = trecho impossível de ler na digitalização; '
              'palavra[?] = leitura incerta.')
    p = doc.add_paragraph()
    p.add_run('Situação deste volume: ').bold = True
    fid = f'; fidelidade média (amostra revisada): {fid_media}' if fid_media else ''
    p.add_run(f'{feitas} de {total} atas transcritas'
              f' ({time.strftime("%d/%m/%Y")}){fid}.')
    doc.add_page_break()


def lacunas_do_ano(ano):
    """Registros pseudo-manifest para os lacuna-*.md (atas recuperadas que o
    indice nao detectou, logo ausentes do manifest). So as do TSJE com corpo."""
    regs = []
    pasta = os.path.join(RAIZ, ano)
    if not os.path.isdir(pasta):
        return regs
    for nome in sorted(os.listdir(pasta)):
        if not nome.startswith('lacuna-') or not nome.endswith('.md'):
            continue
        fm, corpo = front_matter(os.path.join(pasta, nome))
        if not corpo or fm.get('problema') or fm.get('tribunal') != 'superior':
            continue
        d = fm.get('data_sessao', '')
        regs.append({
            'ano': ano, 'status': 'recuperada',
            'transcricao': os.path.join(ano, nome),
            'titulo': fm.get('titulo', ''), 'data_sessao': d or '?',
            'be_data': d if len(d) == 10 else f'{ano}-01-01',
            'arquivo': fm.get('arquivo', ''), 'pag_ini': '1',
            'fidelidade': '', 'recuperada': True})
    return regs


def gerar_ano(ano, regs):
    do_ano = [r for r in regs if r['ano'] == ano]
    prontas = [r for r in do_ano if r['status'] in
               ('transcrita', 'revisada', 'final')
               and os.path.exists(os.path.join(RAIZ, r['transcricao']))]
    prontas += lacunas_do_ano(ano)  # inclui atas recuperadas fora do indice
    prontas.sort(key=lambda r: (r['data_sessao'].replace('?', '9'),
                                r['be_data']))
    if not prontas:
        return None, 0, len(do_ano)

    fids = [float(r['fidelidade']) for r in prontas
            if r['fidelidade'].replace('.', '').isdigit()]
    fid_media = f'{sum(fids)/len(fids):.1f}/10' if fids else ''

    doc = Document()
    est = doc.styles['Normal']
    est.font.name = 'Calibri'
    est.font.size = Pt(11)
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Cm(2.2)

    rosto(doc, ano, len(prontas), len(do_ano), fid_media)

    for r in prontas:
        fm, corpo = front_matter(os.path.join(RAIZ, r['transcricao']))
        mes = MES_NOME.get(int(r['be_data'][5:7]), '')
        doc.add_paragraph(ano)
        doc.add_paragraph(mes)
        doc.add_paragraph(fmt_id(r['arquivo']))
        t = doc.add_paragraph()
        t.add_run(fm.get('titulo', r['titulo'])).bold = True
        if fm.get('presidencia'):
            pr = doc.add_paragraph()
            pr.add_run(fm['presidencia']).italic = True
        for par in corpo.split('\n\n'):
            par = par.strip()
            if par:
                doc.add_paragraph(par)
        fonte = doc.add_paragraph()
        run = fonte.add_run(
            f'[Fonte: Boletim Eleitoral {fm.get("boletim", "")} — '
            f'{r["arquivo"]}, pág. {fm.get("paginas", r["pag_ini"])}]')
        run.italic = True
        run.font.size = Pt(9)
        doc.add_paragraph()

    pend = [r for r in do_ano if r not in prontas
            and r['status'] not in ('duplicada', 'regional')]
    if pend:
        doc.add_page_break()
        doc.add_heading('Pendências deste volume', 1)
        doc.add_paragraph(
            'Atas detectadas no banco mas ainda sem transcrição concluída '
            '(status entre parênteses):')
        for r in pend:
            doc.add_paragraph(f'{r["titulo"]} — {r["arquivo"]}, pág. '
                              f'{r["pag_ini"]} ({r["status"]})',
                              style='List Bullet')

    saida = os.path.join(RAIZ, f'{ano}.docx')
    doc.save(saida)
    return saida, len(prontas), len(do_ano)


def tentar_pdf(docx_path):
    try:
        from docx2pdf import convert
        convert(docx_path, docx_path[:-5] + '.pdf')
        return docx_path[:-5] + '.pdf'
    except Exception:
        return None


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--ano')
    ap.add_argument('--zip', action='store_true')
    ap.add_argument('--pdf', action='store_true')
    args = ap.parse_args()

    with io.open(MANIFEST, encoding='utf-8-sig') as f:
        regs = list(csv.DictReader(f))

    anos = [args.ano] if args.ano else sorted({r['ano'] for r in regs})
    gerados = []
    for ano in anos:
        saida, feitas, total = gerar_ano(ano, regs)
        if saida:
            print(f'{ano}: {feitas}/{total} atas -> {saida}')
            gerados.append(saida)
            if args.pdf:
                pdf = tentar_pdf(saida)
                if pdf:
                    gerados.append(pdf)
                    print(f'  PDF: {pdf}')
        else:
            print(f'{ano}: nenhuma ata transcrita ainda')

    if args.zip:
        # recria o zip com TODOS os volumes ja existentes na pasta
        with zipfile.ZipFile(ZIP_FINAL, 'w', zipfile.ZIP_DEFLATED) as zf:
            for nome in sorted(os.listdir(RAIZ)):
                if re.fullmatch(r'19\d{2}\.(docx|pdf)', nome):
                    zf.write(os.path.join(RAIZ, nome), nome)
        print(f'zip: {ZIP_FINAL}')


if __name__ == '__main__':
    main()
