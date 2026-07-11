# -*- coding: utf-8 -*-
"""Testes OFFLINE dos documentos anexos do Gerador (conle_gerador.anexos + integração).

Nenhum teste toca rede: as chamadas de IA (condensação/visão) são monkeypatchadas.
"""
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from conle_gerador import anexos as anx
from conle_gerador import gerador, llm, prompts


# ------------------------------------------------------------------ helpers
def _silencio(_msg):
    pass


@pytest.fixture(autouse=True)
def _sem_rede(monkeypatch):
    """Garante que nenhum teste chame a OpenAI por engano."""
    def boom(*_a, **_k):
        raise AssertionError("chamada de IA inesperada no teste")
    monkeypatch.setattr(llm, "chat", boom)
    monkeypatch.setattr(llm, "chat_visao", boom)
    yield


# ------------------------------------------------------------------ formatos de texto
def test_txt_utf8(tmp_path):
    p = tmp_path / "oficio.txt"
    p.write_text("Solicito estudo sobre o art. 14, §9º, da CF.", encoding="utf-8")
    docs = anx.processar([str(p)], log=_silencio)
    assert docs[0].texto.startswith("Solicito estudo")
    assert not docs[0].aviso


def test_txt_cp1252(tmp_path):
    p = tmp_path / "oficio_antigo.txt"
    p.write_bytes("Ofício nº 12 — inelegibilidade e filiação.".encode("cp1252"))
    docs = anx.processar([str(p)], log=_silencio)
    assert "Ofício nº 12" in docs[0].texto
    assert "filiação" in docs[0].texto


def test_extensao_desconhecida_cai_em_texto(tmp_path):
    p = tmp_path / "demanda.dat"
    p.write_text("Conteúdo textual em formato exótico.", encoding="utf-8")
    docs = anx.processar([str(p)], log=_silencio)
    assert "formato exótico" in docs[0].texto


def test_html(tmp_path):
    p = tmp_path / "pagina.html"
    p.write_text("<html><style>x{}</style><body><p>Par&aacute;grafo um.</p>"
                 "<p>Dois</p></body></html>", encoding="utf-8")
    docs = anx.processar([str(p)], log=_silencio)
    assert "Parágrafo um." in docs[0].texto
    assert "<p>" not in docs[0].texto


def test_eml(tmp_path):
    p = tmp_path / "mensagem.eml"
    p.write_bytes(b"From: gabinete@camara.leg.br\r\nTo: conle@camara.leg.br\r\n"
                  b"Subject: Pedido de estudo\r\nMIME-Version: 1.0\r\n"
                  b"Content-Type: text/plain; charset=utf-8\r\n\r\n"
                  b"Segue a demanda para estudo da paridade.\r\n")
    docs = anx.processar([str(p)], log=_silencio)
    assert "Subject: Pedido de estudo" in docs[0].texto
    assert "paridade" in docs[0].texto


# ------------------------------------------------------------------ office
def test_docx_com_tabela(tmp_path):
    import docx

    d = docx.Document()
    d.add_paragraph("Demanda: alterar a LC nº 64/1990.")
    tb = d.add_table(rows=1, cols=2)
    tb.rows[0].cells[0].text = "Dispositivo"
    tb.rows[0].cells[1].text = "art. 22"
    p = tmp_path / "demanda.docx"
    d.save(str(p))
    docs = anx.processar([str(p)], log=_silencio)
    assert "LC nº 64/1990" in docs[0].texto
    assert "Dispositivo | art. 22" in docs[0].texto


def test_xlsx(tmp_path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Dados"
    ws.append(["Município", "Vereadores"])
    ws.append(["Alfa", 9])
    p = tmp_path / "dados.xlsx"
    wb.save(str(p))
    docs = anx.processar([str(p)], log=_silencio)
    assert "[Planilha: Dados]" in docs[0].texto
    assert "Alfa\t9" in docs[0].texto


# ------------------------------------------------------------------ pdf e imagem
def test_pdf_com_texto(tmp_path):
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Estudo sobre reeleicao de prefeitos e o art. 14 da CF. " * 8)
    p = tmp_path / "estudo.pdf"
    doc.save(str(p))
    doc.close()
    docs = anx.processar([str(p)], log=_silencio)
    assert "reeleicao de prefeitos" in docs[0].texto
    assert docs[0].origem.startswith("pdf")


def test_pdf_escaneado_usa_ocr_ia(tmp_path, monkeypatch):
    import fitz

    doc = fitz.open()
    doc.new_page()  # página em branco: sem camada de texto
    p = tmp_path / "escaneado.pdf"
    doc.save(str(p))
    doc.close()
    monkeypatch.setattr(llm, "chat_visao",
                        lambda *_a, **_k: "OFÍCIO Nº 5 — solicita minuta de PEC sobre paridade.")
    docs = anx.processar([str(p)], log=_silencio)
    assert "OFÍCIO Nº 5" in docs[0].texto
    assert "OCR IA" in docs[0].origem


def test_imagem_usa_ocr_ia(tmp_path, monkeypatch):
    import fitz

    # gera um PNG minúsculo válido via PyMuPDF
    doc = fitz.open()
    page = doc.new_page(width=80, height=40)
    png = page.get_pixmap().tobytes("png")
    doc.close()
    p = tmp_path / "foto_oficio.png"
    p.write_bytes(png)
    recebidas = {}

    def fake_visao(_sys, _user, imgs, **_k):
        recebidas["n"] = len(imgs)
        return "Texto transcrito da foto."
    monkeypatch.setattr(llm, "chat_visao", fake_visao)
    docs = anx.processar([str(p)], log=_silencio)
    assert docs[0].texto == "Texto transcrito da foto."
    assert recebidas["n"] == 1


# ------------------------------------------------------------------ condensação e tetos
def test_condensacao_de_documento_longo(tmp_path, monkeypatch):
    p = tmp_path / "relatorio.txt"
    p.write_text("Parágrafo relevante sobre quociente eleitoral. " * 900, encoding="utf-8")

    def fake_chat(system, user, **_k):
        assert "Condensa" in system or "condensa" in system
        assert "relatorio.txt" in user
        return "RESUMO FIEL: quociente eleitoral, art. 106 do Código Eleitoral."
    monkeypatch.setattr(llm, "chat", fake_chat)
    docs = anx.processar([str(p)], log=_silencio)
    assert docs[0].texto.startswith("RESUMO FIEL")
    assert "condensado por IA" in docs[0].origem
    assert docs[0].n_chars_original > anx.TETO_DOC


def test_condensacao_falha_trunca(tmp_path, monkeypatch):
    p = tmp_path / "grande.txt"
    p.write_text("x" * (anx.TETO_DOC + 5000), encoding="utf-8")

    def fake_chat(*_a, **_k):
        raise RuntimeError("api fora do ar")
    monkeypatch.setattr(llm, "chat", fake_chat)
    docs = anx.processar([str(p)], log=_silencio)
    assert "truncado" in docs[0].origem
    assert "[… documento truncado por tamanho …]" in docs[0].texto


def test_teto_global(tmp_path):
    paths = []
    for i in range(4):
        p = tmp_path / f"doc{i}.txt"
        p.write_text(f"doc{i} " + "y" * 20_000, encoding="utf-8")
        paths.append(str(p))
    docs = anx.processar(paths, log=_silencio)
    assert sum(len(a.texto) for a in docs) <= anx.TETO_TOTAL + 200  # marcadores de truncagem
    assert all(a.texto for a in docs)  # todos continuam representados


# ------------------------------------------------------------------ falhas não fatais
def test_arquivo_inexistente():
    docs = anx.processar([r"C:\nao\existe\x.pdf"], log=_silencio)
    assert docs[0].texto == ""
    assert "não encontrado" in docs[0].aviso


def test_msg_outlook_nao_suportado(tmp_path):
    p = tmp_path / "mensagem.msg"
    p.write_text("conteudo qualquer", encoding="utf-8")
    docs = anx.processar([str(p)], log=_silencio)
    assert docs[0].texto == ""
    assert ".msg" in docs[0].aviso


def test_falha_de_um_nao_derruba_os_demais(tmp_path):
    ok = tmp_path / "ok.txt"
    ok.write_text("texto válido", encoding="utf-8")
    docs = anx.processar([r"C:\nao\existe.bin", str(ok)], log=_silencio)
    assert docs[0].aviso and docs[1].texto == "texto válido"


# ------------------------------------------------------------------ formatação p/ prompts
def _docs_stub():
    return [anx.Anexo(nome="oficio.pdf", texto="Pedido do deputado.", origem="pdf"),
            anx.Anexo(nome="vazio.png", texto="", origem="imagem", aviso="nenhum texto"),
            anx.Anexo(nome="estudo.docx", texto="A" * 5000, origem="docx")]


def test_formatar_contexto_pula_vazios():
    txt = anx.formatar_contexto(_docs_stub())
    assert "DOCUMENTO ENCAMINHADO Nº 1: “oficio.pdf”" in txt
    assert "vazio.png" not in txt
    assert "DOCUMENTO ENCAMINHADO Nº 2: “estudo.docx”" in txt


def test_resumo_para_analise_corta():
    txt = anx.resumo_para_analise(_docs_stub())
    assert "oficio.pdf" in txt
    assert "íntegra disponível" in txt          # estudo.docx (5000 > ANALISE_CHARS)
    assert len(txt) < 3 * anx.ANALISE_CHARS


def test_contexto_txt_do_gerador_inclui_anexos():
    ctx = {"rag": [], "camara": [], "web": [], "anexos": _docs_stub()}
    txt = gerador._contexto_txt(ctx)
    assert txt.index("DOCUMENTOS ENCAMINHADOS PELO PARLAMENTAR") < txt.index("BASES INTERNAS (RAG)")
    assert prompts.NOTA_ANEXOS[:40] in txt
    assert "Pedido do deputado." in txt


def test_contexto_txt_sem_anexos_intacto():
    txt = gerador._contexto_txt({"rag": [], "camara": [], "web": []})
    assert "DOCUMENTOS ENCAMINHADOS" not in txt
    assert txt.startswith("=== TRECHOS DAS BASES INTERNAS (RAG) ===")


def test_analisar_demanda_recebe_docs(monkeypatch):
    capturado = {}

    def fake_chat(system, user, **_k):
        capturado["system"], capturado["user"] = system, user
        return {"tema": "paridade", "tipo_sigla": "PEC"}
    monkeypatch.setattr(llm, "chat", fake_chat)
    analise = gerador.analisar_demanda("Quero uma PEC.", docs=_docs_stub(), log=_silencio)
    assert analise["tema"] == "paridade"
    assert "DOCUMENTOS ENCAMINHADOS PELO PARLAMENTAR" in capturado["user"]
    assert "Pedido do deputado." in capturado["user"]
    assert "DOCUMENTOS ENCAMINHADOS" in capturado["system"]  # SYS_ANALISE ciente dos anexos


def test_analisar_demanda_sem_docs_nao_muda_prompt(monkeypatch):
    capturado = {}

    def fake_chat(system, user, **_k):
        capturado["user"] = user
        return {"tema": "x", "tipo_sigla": "PL"}
    monkeypatch.setattr(llm, "chat", fake_chat)
    gerador.analisar_demanda("Demanda simples.", log=_silencio)
    assert capturado["user"] == "DEMANDA:\nDemanda simples."


# ------------------------------------------------------------------ integração gerar()
def test_gerar_fim_a_fim_com_anexo(tmp_path, monkeypatch):
    """gerar() com anexo: o texto do documento tem de chegar à ANÁLISE e à REDAÇÃO
    (IT e minuta). Todas as fronteiras (Notion, IA, fontes) são stubadas."""
    from conle_gerador import notion_writer as nw

    p = tmp_path / "oficio_do_gabinete.txt"
    p.write_text("O gabinete solicita alterar o art. 46 da Lei nº 9.096/1995.", encoding="utf-8")

    chamadas = []

    def fake_chat(system, user, **_k):
        chamadas.append((system, user))
        if "analisar a DEMANDA" in system:
            return {"tema": "fidelidade partidária", "tipo_sigla": "PL",
                    "objeto": "Estudo", "palavras_chave_camara": [],
                    "consultas_rag": [], "consultas_web": []}
        if "CORPO ANALÍTICO" in system:
            return {"introducao": ["par 1"], "conclusao": ["fim"]}
        return {"ementa": "Altera a lei.", "articulado": [], "justificativa": ["j"]}

    monkeypatch.setattr(llm, "chat", fake_chat)
    monkeypatch.setattr(nw, "page_id_from_url", lambda _u: "pid-teste")
    monkeypatch.setattr(nw, "verificar_acesso_pagina", lambda _p: None)
    monkeypatch.setattr(nw, "escrever_pagina", lambda _p, blocos, progress=None: len(blocos))
    monkeypatch.setattr(gerador, "_set_titulo", lambda *_a, **_k: None)
    monkeypatch.setattr(gerador, "_salvar_resultado", lambda *_a, **_k: None)

    res = gerador.gerar("Quero um PL sobre fidelidade partidária.", "https://notion.so/x",
                        anexos=[str(p)], usar_rag=False, usar_camara=False, usar_web=False,
                        progress=_silencio)
    assert res.n_blocos > 0
    analise_user = chamadas[0][1]
    assert "oficio_do_gabinete.txt" in analise_user       # análise viu o anexo
    assert "art. 46 da Lei nº 9.096/1995" in analise_user
    redacoes = [u for s, u in chamadas[1:]]
    assert all("DOCUMENTOS ENCAMINHADOS PELO PARLAMENTAR" in u for u in redacoes)  # IT e minuta
    assert all("art. 46 da Lei nº 9.096/1995" in u for u in redacoes)


def test_gerar_anexo_ruim_gera_aviso(tmp_path, monkeypatch):
    """Anexo sem texto aproveitável não aborta a geração — vira aviso no Resultado."""
    from conle_gerador import notion_writer as nw

    def fake_chat(system, user, **_k):
        if "analisar a DEMANDA" in system:
            return {"tema": "t", "tipo_sigla": "PL", "objeto": "o",
                    "palavras_chave_camara": [], "consultas_rag": [], "consultas_web": []}
        return {"introducao": ["x"], "conclusao": ["y"], "ementa": "E.",
                "articulado": [], "justificativa": ["j"]}

    monkeypatch.setattr(llm, "chat", fake_chat)
    monkeypatch.setattr(nw, "page_id_from_url", lambda _u: "pid")
    monkeypatch.setattr(nw, "verificar_acesso_pagina", lambda _p: None)
    monkeypatch.setattr(nw, "escrever_pagina", lambda _p, blocos, progress=None: len(blocos))
    monkeypatch.setattr(gerador, "_set_titulo", lambda *_a, **_k: None)
    monkeypatch.setattr(gerador, "_salvar_resultado", lambda *_a, **_k: None)

    res = gerador.gerar("Demanda.", "https://notion.so/x",
                        anexos=[str(tmp_path / "nao_existe.pdf")],
                        usar_rag=False, usar_camara=False, usar_web=False, progress=_silencio)
    assert any("nao_existe.pdf" in a for a in res.avisos)
    assert any("Nenhum anexo rendeu texto" in a for a in res.avisos)
