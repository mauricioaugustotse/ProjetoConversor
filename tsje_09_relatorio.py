# -*- coding: utf-8 -*-
"""tsje_09: relatorio final didatico do projeto de transcricao das atas do TSJE.

Le o manifest, os .md de lacuna, o diagnostico de completude, o registro de
PDFs mal rotulados e o confronto com o ensaio; produz um docx explicativo em
D:\\TSJE_TRANSCRICOES\\_RELATORIO - Atas TSJE 1932-1937.docx com:
- o que foi feito (numeros por ano: transcritas do indice + recuperadas);
- o que o indice tinha de ruido (TRE, duplicatas, falsos positivos);
- o que FALTA e por que (lacunas recuperaveis ja feitas x boletins ausentes
  do acervo, listados nominalmente);
- achados sobre o acervo (PDFs mal rotulados);
- como o material esta organizado e como conferir.

Uso: python tsje_09_relatorio.py
"""
import csv
import io
import os
import re
import time
from collections import defaultdict

from docx import Document
from docx.shared import Pt, Cm

RAIZ = r'D:\TSJE_TRANSCRICOES'
BASE = os.path.dirname(os.path.abspath(__file__))
MANIFEST = os.path.join(RAIZ, 'manifest.csv')
DIAG = os.path.join(RAIZ, '_diagnostico_completude.csv')
MALROT = os.path.join(RAIZ, '_pdfs_mal_rotulados.csv')
SAIDA = os.path.join(RAIZ, '_RELATORIO - Atas TSJE 1932-1937.docx')

ANOS = ['1932', '1933', '1934', '1935', '1936', '1937']


def ler_csv(path):
    if not os.path.exists(path):
        return []
    with io.open(path, encoding='utf-8-sig') as f:
        return list(csv.DictReader(f))


def contar_lacunas():
    """ano -> nº de atas do TSJE recuperadas (lacuna-*.md com corpo).
    Retorna tambem a contagem de so_resenha confirmadas por leitura de
    imagem e de duplicatas/falsos achados auditados."""
    por_ano = defaultdict(int)
    resenha = dup = 0
    for ano in ANOS:
        pasta = os.path.join(RAIZ, ano)
        if not os.path.isdir(pasta):
            continue
        for nome in os.listdir(pasta):
            if not (nome.startswith(('lacuna-', 'rev-'))
                    and nome.endswith('.md')):
                continue
            with io.open(os.path.join(pasta, nome), encoding='utf-8') as f:
                t = f.read()
            m = re.search(r'---\n(.*?)\n---\n?(.*)', t, re.S)
            if not m or 'tribunal: superior' not in m.group(1):
                continue
            if 'problema: so_resenha' in m.group(1):
                resenha += 1
            elif 'problema: duplicata' in m.group(1):
                dup += 1
            elif m.group(2).strip() and 'problema:' not in m.group(1):
                por_ano[ano] += 1
    return por_ano, resenha, dup


def tabela(doc, cab, linhas, sz=10):
    t = doc.add_table(rows=1, cols=len(cab))
    t.style = 'Table Grid'
    for j, x in enumerate(cab):
        r = t.rows[0].cells[j].paragraphs[0].add_run(str(x))
        r.bold = True
        r.font.size = Pt(sz)
    for ln in linhas:
        cells = t.add_row().cells
        for j, x in enumerate(ln):
            cells[j].paragraphs[0].add_run(str(x)).font.size = Pt(sz)
    return t


def main():
    manifest = ler_csv(MANIFEST)
    diag = ler_csv(DIAG)
    malrot = ler_csv(MALROT)
    lac, n_resenha, n_dup = contar_lacunas()

    por_ano = defaultdict(lambda: defaultdict(int))
    for m in manifest:
        por_ano[m['ano']][m['status']] += 1

    doc = Document()
    est = doc.styles['Normal']
    est.font.name = 'Calibri'
    est.font.size = Pt(11)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2)

    doc.add_heading('Transcrição das Atas do Tribunal Superior de Justiça '
                    'Eleitoral (1932–1937)', 0)
    doc.add_heading('Relatório do trabalho — o que foi feito e o que falta', 2)
    doc.add_paragraph(f'Gerado em {time.strftime("%d/%m/%Y %H:%M")}.')

    # ---- 1. em uma frase
    doc.add_heading('1. Em resumo', 1)
    trans = sum(por_ano[a]['transcrita'] + por_ano[a]['revisada']
                + por_ano[a]['final'] for a in ANOS)
    rec = sum(lac.values())
    ausentes = sum(1 for d in diag if d['situacao'] == 'BOLETIM_AUSENTE_NO_ACERVO')
    falsos = sum(1 for d in diag if d['situacao'] == 'RECUPERAVEL')
    faltam = ausentes + falsos
    doc.add_paragraph(
        f'Foram transcritas fielmente {trans + rec} atas do TSJE do período '
        f'1932–1937 — {trans} a partir do índice do banco e {rec} recuperadas '
        f'na auditoria (atas que existem nos boletins, mas que a detecção '
        f'automática do índice não havia enxergado). O texto é transcrição '
        f'fiel do impresso (não paráfrase): ortografia modernizada, mas '
        f'conteúdo, nomes e votos preservados.')
    doc.add_paragraph(
        f'Em 16/07/2026 o acervo local foi completado com o download de '
        f'todos os Boletins Eleitorais 1932–1937 do arquivo público do TSE '
        f'(AtoM, coleção "Primeira fase da Justiça Eleitoral"): 366 boletins '
        f'novos, elevando o banco a 804 PDFs — essencialmente TODA a coleção '
        f'digitalizada conhecida. Com isso, o diagnóstico de completude '
        f'tornou-se definitivo: {faltam} sessões do TSJE NÃO têm a ata '
        f'formal em nenhum boletim preservado ({ausentes} comprovadas por '
        f'busca no acervo completo + {falsos} cujo único "achado" é ata '
        f'homônima de outro ano, auditada e descartada). Dessas, {n_resenha} '
        f'foram conferidas na imagem: o boletim traz apenas a RESENHA de '
        f'julgamentos ("O Tribunal em sua Nª sessão resolveu…"), não a ata '
        f'formal. A seção 5 detalha e lista os casos.')
    doc.add_paragraph(
        'O material está em D:\\TSJE_TRANSCRICOES: um arquivo .docx por ano '
        '(1932 a 1937) com as atas em ordem cronológica, além deste relatório '
        'e do índice-mestre (D:\\Indice - Atas TSJE 1932-1937.docx).')

    # ---- 2. por ano
    doc.add_heading('2. Quanto foi transcrito, ano a ano', 1)
    linhas = []
    for a in ANOS:
        p = por_ano[a]
        feitas = p['transcrita'] + p['revisada'] + p['final']
        linhas.append([a, feitas, lac.get(a, 0), feitas + lac.get(a, 0),
                       p['regional'], p['duplicada'], p['problema']])
    tot = lambda i: sum(int(l[i]) for l in linhas)
    linhas.append(['Total', tot(1), tot(2), tot(3), tot(4), tot(5), tot(6)])
    tabela(doc, ['Ano', 'Do índice', 'Recuperadas', 'Total TSJE',
                 'Atas de TRE', 'Duplicatas', 'Falsos pos.'], linhas)
    doc.add_paragraph(
        'As três últimas colunas são o "ruído" que a auditoria separou das '
        'atas do Superior: sessões que na verdade eram de Tribunais Regionais '
        '(11 estados diferentes), entradas duplicadas do índice (inclusive '
        'boletins escaneados duas vezes no acervo) e falsos positivos '
        '(capas, avisos de julgamento e votos avulsos que não são atas).',
        style='List Bullet')

    # ---- 3. como foi feito / qualidade
    doc.add_heading('3. Como foi feito e como sabemos que está fiel', 1)
    for t in [
        'Cada ata foi transcrita a partir da IMAGEM da página do boletim '
        '(não do OCR, que é ruidoso nos jornais dos anos 1930). Dígitos e '
        'nomes duvidosos foram conferidos com ampliação e cruzados com a '
        'sequência de sessões e o dia da semana.',
        'Um revisor independente de IA (gpt-5.6-terra) reconferiu uma amostra '
        'de cada ano contra as imagens. Onde ele apontou divergência real, foi '
        'corrigido; onde ele próprio errou (3 casos, confundindo a ata com a '
        'vizinha da mesma página), o parecer foi rejeitado com justificativa '
        'registrada no arquivo. Regra do projeto: nenhuma correção sem '
        'conferência na imagem.',
        'Convenções de fidelidade: ortografia modernizada ("secção"→"seção"), '
        'mas nomes próprios e topônimos na grafia de época (Affonso Penna '
        'Júnior, Minas Geraes, Matto Grosso); erros tipográficos do próprio '
        'jornal preservados e assinalados; [ilegível] e palavra[?] para '
        'trechos que a digitalização não permite ler com certeza.']:
        doc.add_paragraph(t, style='List Bullet')

    # ---- 4. lacunas recuperadas
    doc.add_heading('4. O que a auditoria recuperou', 1)
    doc.add_paragraph(
        f'Reconstruindo a numeração das sessões de cada ano, verificou-se que '
        f'o índice automático havia perdido dezenas de atas (o OCR não '
        f'reconheceu o cabeçalho). {rec} dessas foram recuperadas lendo as '
        f'imagens e estão incluídas nos volumes anuais. Isso é a diferença '
        f'entre "transcrevi o que o índice deu" e "transcrevi o que existe '
        f'no acervo".')

    # ---- 5. o que falta (ausentes no acervo)
    doc.add_heading('5. O que falta — e por quê', 1)
    doc.add_paragraph(
        f'Nenhuma das faltas é erro de transcrição — a busca varreu o texto '
        f'de TODOS os {804} boletins do acervo completado (zips do Drive + '
        f'coleção AtoM/TSE) e a ata formal dessas sessões não está em '
        f'nenhum deles. O que se sabe sobre elas:')
    doc.add_paragraph(
        f'{n_resenha} sessões foram conferidas por leitura da imagem: o '
        f'boletim publica apenas a resenha de julgamentos da sessão, e a '
        f'seção "ACTAS" (que nos boletins de 1936–1937 sai em números '
        f'POSTERIORES, com defasagem de várias sessões) não traz a ata '
        f'formal em nenhum boletim preservado. O registro de cada conferência '
        f'está nos arquivos lacuna-*.md com a marca "so_resenha".',
        style='List Bullet')
    doc.add_paragraph(
        f'{falsos} sessões têm um falso "achado" documentado: o buraco na '
        f'numeração de um ano casa com ata homônima de OUTRO ano (boletins '
        f'mal-arquivados entre anos ou atas de 1934 publicadas com atraso em '
        f'1935). Todas auditadas e registradas como duplicata.',
        style='List Bullet')
    doc.add_paragraph(
        f'As demais constam apenas pela numeração impressa: nem resenha '
        f'localizável. Listadas nominalmente abaixo, para eventual busca em '
        f'fontes físicas (hemerotecas, Diário da Justiça da época).',
        style='List Bullet')
    aus = [d for d in diag if d['situacao'] in
           ('BOLETIM_AUSENTE_NO_ACERVO', 'RECUPERAVEL')]
    por_ano_aus = defaultdict(list)
    for d in aus:
        tp = 'extra.' if d['tipo'] == 'extraordinaria' else 'ord.'
        marca = '*' if d['situacao'] == 'RECUPERAVEL' else ''
        por_ano_aus[d['ano']].append(f"{d['num']}ª {tp}{marca}")
    linhas = [[a, ', '.join(por_ano_aus[a])] for a in ANOS if por_ano_aus[a]]
    tabela(doc, ['Ano', 'Sessões do TSJE sem ata formal no acervo'], linhas, 9)
    doc.add_paragraph(
        '* = casos com falso "achado" auditado (ata homônima de outro ano).',
        style='List Bullet')

    # ---- 6. achados sobre o acervo
    if malrot:
        doc.add_heading('6. Achados sobre o acervo (arquivos mal rotulados)', 1)
        doc.add_paragraph(
            'Durante a transcrição, descobriu-se que alguns PDFs do banco '
            'D:\\TSJE_ATAS têm data ou número ERRADOS no próprio nome do '
            'arquivo (o nome não bate com o boletim impresso). A data errada '
            'está no acervo, não só no índice — vale corrigir na origem:')
        linhas = [[m['id_acervo'], m['nome_do_arquivo_diz'],
                   m['cabecalho_impresso_diz']] for m in malrot]
        tabela(doc, ['ID', 'Nome do arquivo diz', 'Impresso diz'], linhas, 9)

    # ---- 7. organizacao
    doc.add_heading('7. Onde está cada coisa', 1)
    for t in [
        'D:\\TSJE_TRANSCRICOES\\1932.docx … 1937.docx — os volumes anuais com '
        'as atas transcritas em ordem cronológica.',
        'D:\\Indice - Atas TSJE 1932-1937.docx — índice-mestre: localiza cada '
        'ata (data, tipo, nº, boletim, arquivo PDF e página) e traz a '
        'auditoria de completude.',
        'D:\\TSJE_TRANSCRICOES\\<ano>\\*.md — as transcrições individuais '
        '(uma por ata), com metadados e notas de fidelidade no cabeçalho.',
        'D:\\TSJE_TRANSCRICOES\\_diagnostico_completude.csv e '
        '_pdfs_mal_rotulados.csv — as listas técnicas que embasam as seções '
        '5 e 6.']:
        doc.add_paragraph(t, style='List Bullet')

    doc.save(SAIDA)
    print(f'relatorio -> {SAIDA}')
    print(f'  transcritas do indice: {trans} | recuperadas: {rec} | '
          f'ausentes no acervo: {ausentes}')


if __name__ == '__main__':
    main()
