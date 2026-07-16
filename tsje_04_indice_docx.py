# -*- coding: utf-8 -*-
"""tsje_04: gera o indice-mestre .docx das atas do TSJE em D:\\.

Le tsje_indice_atas.csv + tsje_plano.csv e monta:
- resumo geral e tabela por ano;
- uma secao por ano com a tabela das atas do TSJE (data da sessao, tipo, n.,
  boletim, data do BE, arquivo no banco D:\\TSJE_ATAS e pagina do PDF);
- auditoria de completude: lacunas de numeracao ordinal (com deteccao de
  reinicio de serie), meses sem ata, documentos sem OCR, atas com data/numero
  nao lidos, BEs excluidos e numeros de boletim ausentes por ano.

Uso: python tsje_04_indice_docx.py
Saida: D:\\Indice - Atas TSJE 1932-1937.docx
"""
import csv
import io
import os
import re
import time
from collections import defaultdict

from docx import Document
from docx.shared import Pt, Cm

BASE = os.path.dirname(os.path.abspath(__file__))
INDICE = os.path.join(BASE, 'tsje_indice_atas.csv')
PLANO = os.path.join(BASE, 'tsje_plano.csv')
SAIDA = r'D:\Indice - Atas TSJE 1932-1937.docx'
MANIFEST = r'D:\TSJE_TRANSCRICOES\manifest.csv'
CONFRONTO = r'D:\TSJE_TRANSCRICOES\_ensaio_confronto.csv'

TIPO_FMT = {'ordinaria': 'Ordinária', 'extraordinaria': 'Extraordinária', '': '—'}
STATUS_FMT = {'pendente': '—', 'transcrita': 'transcrita',
              'revisada': 'revisada', 'final': 'final',
              'problema': 'PROBLEMA', 'duplicada': 'duplicada (índice)',
              'regional': 'ata de TRE (reclassificada)'}


def ata_id_de(a):
    """Mesma formula de id do tsje_05_preparar (join com o manifest)."""
    mid = re.search(r'\[([0-9a-f]+)\]', a['arquivo_destino'])
    ident = (mid.group(1) if mid
             else re.sub(r'\W', '', a['arquivo_destino'])[:12])
    tipo1 = {'ordinaria': 'o', 'extraordinaria': 'e', '': 'x'}[a['tipo']]
    num = a['num_ordinal'].lstrip('0') or '0'
    dt = a['data_sessao'] if a['data_sessao'] != '?' else 'semdata'
    return f'{dt}-{tipo1}{num}-{ident}-p{int(a["pagina"] or 1)}'


def ler_csv(path):
    with io.open(path, encoding='utf-8-sig') as f:
        return list(csv.DictReader(f))


def fmt_data(iso):
    if len(iso) == 10:
        a, m, d = iso.split('-')
        return f'{d}/{m}/{a}'
    return '?'


def ano_da_ata(a):
    if a['data_sessao'][:4].isdigit():
        return a['data_sessao'][:4]
    return a['data_boletim'][:4]


def tabela(doc, cabecalho, linhas, tamanho=9):
    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.style = 'Table Grid'
    for j, txt in enumerate(cabecalho):
        cel = t.rows[0].cells[j]
        run = cel.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(tamanho)
    for linha in linhas:
        cells = t.add_row().cells
        for j, txt in enumerate(linha):
            run = cells[j].paragraphs[0].add_run(str(txt))
            run.font.size = Pt(tamanho)
    return t


def compacta(nums):
    """[1,2,3,7,9,10] -> '1–3, 7, 9–10'"""
    if not nums:
        return '—'
    partes = []
    ini = prev = nums[0]
    for x in nums[1:]:
        if x == prev + 1:
            prev = x
            continue
        partes.append(f'{ini}–{prev}' if prev > ini else str(ini))
        ini = prev = x
    partes.append(f'{ini}–{prev}' if prev > ini else str(ini))
    return ', '.join(partes)


def analisar_numeracao(atas_tipo):
    """Por ano: maior subsequencia nao-decrescente (LIS) = numeracao aceita;
    numeros fora dela = provaveis erros de OCR. A numeracao do TSJE reinicia
    a cada ano."""
    por_ano = defaultdict(list)
    for a in atas_tipo:
        por_ano[a['data_sessao'][:4]].append(a)
    resultados = []
    for ano in sorted(por_ano):
        grupo = sorted(por_ano[ano], key=lambda a: a['data_sessao'])
        nums = [int(a['num_ordinal']) for a in grupo]
        n = len(nums)
        comp = [1] * n
        pai = [-1] * n
        for i in range(n):
            for j in range(i):
                if nums[j] <= nums[i] and comp[j] + 1 > comp[i]:
                    comp[i], pai[i] = comp[j] + 1, j
        k = max(range(n), key=lambda i: comp[i])
        aceitos_idx = set()
        while k != -1:
            aceitos_idx.add(k)
            k = pai[k]
        aceitos = [grupo[i] for i in sorted(aceitos_idx)]
        suspeitos = [grupo[i] for i in range(n) if i not in aceitos_idx]
        vals = sorted({int(a['num_ordinal']) for a in aceitos})
        faltam = sorted(set(range(vals[0], vals[-1] + 1)) - set(vals))
        resultados.append((ano, aceitos, suspeitos, faltam))
    return resultados


def meses_periodo():
    meses = []
    a, m = 1932, 7
    while (a, m) <= (1937, 11):
        meses.append(f'{a:04d}-{m:02d}')
        m += 1
        if m == 13:
            a, m = a + 1, 1
    return meses


def main():
    atas = ler_csv(INDICE)
    plano = ler_csv(PLANO)
    sup = [a for a in atas if a['tribunal'] == 'superior']
    reg = [a for a in atas if a['tribunal'] == 'regional']
    incluidos = [r for r in plano if r['acao'] == 'incluir']

    manifest = {}
    if os.path.exists(MANIFEST):
        for m in ler_csv(MANIFEST):
            manifest[m['ata_id']] = m
    for a in sup:
        a['_m'] = manifest.get(ata_id_de(a), {})

    doc = Document()
    est = doc.styles['Normal']
    est.font.name = 'Calibri'
    est.font.size = Pt(10)
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Cm(1.8)

    doc.add_heading('Índice das Atas do Tribunal Superior de Justiça '
                    'Eleitoral (1932–1937)', 0)
    doc.add_paragraph(f'Gerado em {time.strftime("%d/%m/%Y %H:%M")}.')

    doc.add_heading('Como ler este índice', 1)
    doc.add_paragraph(
        'Este documento acompanha dois produtos: (1) o BANCO DE BOLETINS em '
        'D:\\TSJE_ATAS — os PDFs originais digitalizados dos Boletins '
        'Eleitorais que contêm atas do TSJE, organizados por ano e datados; '
        'e (2) as TRANSCRIÇÕES FIÉIS em D:\\TSJE_TRANSCRICOES (volumes '
        'anuais .docx no zip "Atas TSJE - Transc Claude").')
    doc.add_paragraph(
        'Cada linha das tabelas anuais é UMA ata: a data em que a sessão '
        'aconteceu, o tipo (ordinária/extraordinária), o número ordinal da '
        'sessão, o Boletim que a publicou (nº e data), o arquivo PDF do '
        'banco onde ela está e a PÁGINA dentro desse PDF. A coluna '
        '"Transcrição" mostra o andamento: "—" = ainda não transcrita; '
        '"transcrita" = texto fiel pronto; "revisada" = já conferida pelo '
        'revisor de IA (gpt-5.6-terra); "final" = revisão aplicada; '
        '"PROBLEMA" = precisa de atenção humana.', style='List Bullet')
    doc.add_paragraph(
        'Onde a data ou o número aparecem como "—" ou "?", o OCR da '
        'digitalização original não permitiu leitura automática — a seção '
        'de auditoria lista esses casos com arquivo e página para '
        'conferência manual (e a transcrição, feita da IMAGEM, resolve a '
        'maioria deles).', style='List Bullet')
    doc.add_paragraph(
        'A seção "Auditoria de completude" responde: falta alguma ata? '
        'Falta algum boletim? Algum documento sem OCR? A seção '
        '"Transcrições" responde: o que já foi transcrito e o que falta.',
        style='List Bullet')

    # ------------------------------------------------------------- resumo
    doc.add_heading('Resumo geral', 1)
    doc.add_paragraph(
        f'Boletins no banco: {len(incluidos)}   |   Atas do TSJE indexadas: '
        f'{len(sup)}   |   Atas de Tribunais Regionais nos mesmos boletins '
        f'(não indexadas): {len(reg)}')
    por_ano_be = defaultdict(int)
    for r in incluidos:
        por_ano_be[r['data_final'][:4]] += 1
    por_ano_ata = defaultdict(lambda: [0, 0, 0])
    for a in sup:
        st = por_ano_ata[ano_da_ata(a)]
        st[0] += 1
        st[1] += a['data_sessao'] != '?'
        st[2] += a['num_ordinal'] != ''
    linhas = []
    for ano in sorted(set(por_ano_be) | set(por_ano_ata)):
        t = por_ano_ata[ano]
        linhas.append([ano, por_ano_be.get(ano, 0), t[0], t[1], t[2]])
    linhas.append(['Total', len(incluidos), len(sup),
                   sum(v[1] for v in por_ano_ata.values()),
                   sum(v[2] for v in por_ano_ata.values())])
    tabela(doc, ['Ano', 'Boletins no banco', 'Atas TSJE', 'Atas com data',
                 'Atas com nº'], linhas, 10)

    # ------------------------------------------------- tabelas por ano
    por_ano = defaultdict(list)
    for a in sup:
        por_ano[ano_da_ata(a)].append(a)
    for ano in sorted(por_ano):
        grupo = sorted(por_ano[ano],
                       key=lambda a: (a['data_sessao'], a['data_boletim'],
                                      int(a['pagina'] or 1)))
        doc.add_heading(f'{ano} — {len(grupo)} atas', 1)
        linhas = []
        for a in grupo:
            num = f'{a["num_ordinal"]}ª' if a['num_ordinal'] else '—'
            nbe = f'n. {int(a["n_boletim_be"])}' if a['n_boletim_be'].isdigit() else '—'
            st = STATUS_FMT.get(a['_m'].get('status', ''), '—')
            linhas.append([fmt_data(a['data_sessao']), TIPO_FMT[a['tipo']],
                           num, nbe, fmt_data(a['data_boletim']),
                           a['arquivo_destino'], a['pagina'], st])
        tabela(doc, ['Data da sessão', 'Tipo', 'Nº', 'Boletim', 'Data do BE',
                     'Arquivo em D:\\TSJE_ATAS\\' + ano, 'Pág.',
                     'Transcrição'], linhas)

    # ------------------------------------------------------- auditoria
    doc.add_heading('Auditoria de completude', 1)

    doc.add_heading('a) Numeração ordinal das sessões — lacunas por ano', 2)
    doc.add_paragraph(
        'A numeração do TSJE reinicia a cada ano. Para cada ano, a maior '
        'sequência não-decrescente em ordem cronológica é tomada como '
        'numeração aceita; números fora dela são prováveis erros de leitura '
        'do OCR. Números faltantes podem indicar ata não indexada (sem '
        'número legível), boletim ausente do acervo ou sessão não publicada.')
    for tipo in ('ordinaria', 'extraordinaria'):
        com_num = [a for a in sup if a['tipo'] == tipo
                   and a['num_ordinal'].isdigit() and a['data_sessao'] != '?']
        doc.add_paragraph().add_run(f'Sessões {TIPO_FMT[tipo].lower()}s '
                                    f'({len(com_num)} com número e data):').bold = True
        for ano, aceitos, suspeitos, faltam in analisar_numeracao(com_num):
            vals = [int(a['num_ordinal']) for a in aceitos]
            txt = (f'{ano}: {min(vals)}ª–{max(vals)}ª ({len(aceitos)} atas); '
                   f'faltam: {compacta(faltam)}')
            if suspeitos:
                exemplos = '; '.join(
                    f'{a["num_ordinal"]}ª em {fmt_data(a["data_sessao"])}'
                    for a in suspeitos[:6])
                extra = f' (+{len(suspeitos)-6})' if len(suspeitos) > 6 else ''
                txt += f'. Nºs suspeitos de OCR: {exemplos}{extra}'
            doc.add_paragraph(txt, style='List Bullet')
    sem_num = sum(1 for a in sup if not a['num_ordinal'])
    sem_tipo = sum(1 for a in sup if not a['tipo'])
    doc.add_paragraph(
        f'Atas sem número ordinal legível: {sem_num} (aparecem nas tabelas '
        f'anuais com "—"). Atas sem tipo identificado: {sem_tipo}.')

    doc.add_heading('b) Meses do período sem nenhuma ata do TSJE', 2)
    com_ata = {a['data_sessao'][:7] for a in sup if a['data_sessao'] != '?'}
    vazios = [m for m in meses_periodo() if m not in com_ata]
    if vazios:
        doc.add_paragraph(', '.join(f'{m[5:]}/{m[:4]}' for m in vazios))
    else:
        doc.add_paragraph('Nenhum — todos os meses de jul/1932 a nov/1937 têm '
                          'ao menos uma ata.')

    doc.add_heading('c) Documentos sem OCR utilizável e campos ilegíveis', 2)
    sem_ocr = [r for r in plano
               if 'sem_texto' in r['flags'] or 'erro_pdf' in r['flags']]
    if sem_ocr:
        tabela(doc, ['Arquivo no acervo', 'Situação'],
               [[r['relpath'], r['flags']] for r in sem_ocr])
    else:
        doc.add_paragraph('Nenhum PDF sem camada de OCR — todos os 815 '
                          'documentos do acervo foram indexáveis.')
    sem_data = [a for a in sup if a['data_sessao'] == '?']
    if sem_data:
        doc.add_paragraph(f'{len(sem_data)} atas com data não lida pelo OCR '
                          f'(localizáveis pelo arquivo/página):')
        tabela(doc, ['Tipo', 'Nº', 'Arquivo', 'Pág.'],
               [[TIPO_FMT[a['tipo']], a['num_ordinal'] or '—',
                 a['arquivo_destino'], a['pagina']] for a in sem_data])
    else:
        doc.add_paragraph('Todas as atas indexadas têm data de sessão lida.')

    doc.add_heading('d) Cobertura dos boletins por ano', 2)
    doc.add_paragraph(
        'Numeração impressa dos BEs vista em TODO o acervo baixado '
        '(incluídos + excluídos). Números ausentes da sequência indicam '
        'boletim que não veio do Drive (ou cujo número o OCR não leu).')
    por_ano_nums = defaultdict(set)
    excl_ano = defaultdict(int)
    for r in plano:
        ano = (r['data_final'][:4] if len(r['data_final']) >= 4
               and r['data_final'][:4].isdigit() else r['ano_zip'])
        if r['n_boletim'].isdigit() and int(r['n_boletim']) <= 260:
            por_ano_nums[ano].add(int(r['n_boletim']))
        if r['acao'] == 'excluir':
            excl_ano[ano] += 1
    linhas = []
    for ano in sorted(por_ano_nums):
        nums = por_ano_nums[ano]
        faltam = sorted(set(range(min(nums), max(nums) + 1)) - nums)
        linhas.append([ano, f'{min(nums)}–{max(nums)}', len(nums),
                       excl_ano.get(ano, 0), compacta(faltam)])
    tabela(doc, ['Ano', 'Faixa de nºs vista', 'BEs com nº lido',
                 'BEs excluídos (sem ata TSJE)', 'Nºs de BE ausentes'],
           linhas)

    # ------------------------------------------- transcricoes: feito x falta
    if manifest:
        doc.add_heading('Transcrições — o que está feito e o que falta', 1)
        doc.add_paragraph(
            'Fluxo de trabalho: as páginas dos Boletins são convertidas em '
            'imagens; o Claude (Fable 5) transcreve cada ata FIELMENTE a '
            'partir da imagem (ortografia modernizada, nomes originais); o '
            'gpt-5.6-terra revisa os casos difíceis e uma amostra de '
            'controle, comparando a transcrição com a imagem; as correções '
            'são aplicadas e o volume anual .docx é gerado. O estado vive '
            'em D:\\TSJE_TRANSCRICOES\\manifest.csv e é retomável a '
            'qualquer momento.')
        doc.add_paragraph(
            'A tabela abaixo separa o que é ALVO (atas do próprio TSJE que '
            'entram nos volumes de transcrição) do que foi filtrado durante '
            'o trabalho: atas que na verdade eram de Tribunais Regionais '
            '(o cabeçalho da seção regional foi partido pelo OCR e o índice '
            'as havia contado como do Superior) e entradas duplicadas do '
            'índice (a mesma sessão detectada duas vezes). O "% do alvo" '
            'mede o avanço só sobre as atas do TSJE ainda a fazer.')
        por_ano_st = defaultdict(lambda: defaultdict(int))
        for m in manifest.values():
            por_ano_st[m['ano']][m['status']] += 1
        linhas = []
        for ano in sorted(por_ano_st):
            st = por_ano_st[ano]
            feitas = st['transcrita'] + st['revisada'] + st['final']
            alvo = feitas + st['pendente'] + st['problema']
            pct = f'{100*feitas/alvo:.0f}%' if alvo else '—'
            linhas.append([ano, alvo, feitas, st['revisada'] + st['final'],
                           st['pendente'], st['problema'],
                           st['regional'], st['duplicada'], pct])
        col = lambda i: sum(int(l[i]) for l in linhas)
        linhas.append(['Total', col(1), col(2), col(3), col(4), col(5),
                       col(6), col(7),
                       f'{100*col(2)/col(1):.0f}%' if col(1) else '—'])
        tabela(doc, ['Ano', 'Alvo (TSJE)', 'Transcritas', 'Revisadas (terra)',
                     'A fazer', 'Problemas', 'De TRE', 'Duplic.',
                     '% do alvo'], linhas, 9)
        fids = [float(m['fidelidade']) for m in manifest.values()
                if m.get('fidelidade', '').replace('.', '').isdigit()]
        if fids:
            doc.add_paragraph(
                f'Fidelidade média atribuída pelo revisor gpt-5.6-terra na '
                f'amostra revisada: {sum(fids)/len(fids):.1f}/10 '
                f'({len(fids)} pareceres).')

        if os.path.exists(CONFRONTO):
            conf = ler_csv(CONFRONTO)
            so_e = [c for c in conf if c['sentido'] == 'so_no_ensaio']
            so_b = [c for c in conf if c['sentido'] == 'so_no_banco']
            doc.add_heading('Confronto com o ensaio anterior '
                            '(Transc IA 1932-1937)', 2)
            doc.add_paragraph(
                f'O ensaio anterior (paráfrases) tem 644 atas; o banco atual '
                f'detectou {len(sup)}. Cruzando por data/tipo/número: '
                f'{len(so_e)} atas do ensaio ainda não casaram com o banco '
                f'e {len(so_b)} atas do banco não estavam no ensaio. '
                f'Atenção: parte das não-casadas se deve às atas cuja data '
                f'o OCR não leu (o casamento melhora à medida que as '
                f'transcrições corrigem datas pela imagem). A lista '
                f'completa está em D:\\TSJE_TRANSCRICOES\\'
                f'_ensaio_confronto.csv.')

    doc.save(SAIDA)
    print(f'{len(sup)} atas do TSJE em {len(por_ano)} anos -> {SAIDA}')


if __name__ == '__main__':
    main()
