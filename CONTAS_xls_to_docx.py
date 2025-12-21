# -*- coding: utf-8 -*-

"""
Este script lê dados de uma planilha (Excel ou CSV), processa os valores
com base em regras específicas, ordena os dados e gera um documento 
Word (.docx) com uma tabela formatada na ordem de colunas desejada 
e um resumo dos totais integrado na tabela.

Alterações nesta versão (somente sobre larguras de colunas / estrutura da tabela Word):
- Helpers para forçar tblGrid (grade) e tcW (largura por célula) em twips.
- Criação MAIS SEGURA de w:tblPr (se não existir) e inserção do w:tblGrid na posição correta.
- Fix de robustez: evita manipulação XML que possa corromper o .docx.
"""

import pandas as pd
import numpy as np
import docx
import os
from tkinter import Tk, filedialog
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Helpers para forçar larguras fixas no Word (tblGrid + tcW) ---
def _twips_from_cm(cm_value: float) -> int:
    # 1 cm ≈ 566.929 twips
    return int(round(float(cm_value) * 567))

def _ensure_tblPr(table):
    """Garante a existência de w:tblPr e o retorna."""
    tbl = table._element
    tbl_pr_list = tbl.xpath('./w:tblPr')
    if tbl_pr_list:
        return tbl_pr_list[0]
    # cria e insere no início de CT_Tbl
    tbl_pr = OxmlElement('w:tblPr')
    tbl.insert(0, tbl_pr)
    return tbl_pr

def _set_table_layout_fixed(table):
    """Força w:tblLayout type='fixed' dentro de w:tblPr."""
    tbl_pr = _ensure_tblPr(table)
    # remove tblLayout anterior
    for child in list(tbl_pr):
        if child.tag == qn('w:tblLayout'):
            tbl_pr.remove(child)
    tbl_layout = OxmlElement('w:tblLayout')
    tbl_layout.set(qn('w:type'), 'fixed')
    tbl_pr.append(tbl_layout)

def _set_cell_width(cell, cm_value: float):
    """Define a largura da célula via w:tcW (dxa) e mantém cell.width (Cm) para compatibilidade."""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        # Remove tcW anterior, se houver
        for child in list(tcPr):
            if child.tag == qn('w:tcW'):
                tcPr.remove(child)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:type'), 'dxa')
        tcW.set(qn('w:w'), str(_twips_from_cm(cm_value)))
        tcPr.append(tcW)
        # Redundância benigna
        cell.width = Cm(cm_value)
    except Exception:
        # fallback silencioso
        cell.width = Cm(cm_value)

def _apply_table_grid(table, column_names, widths_map_cm: dict):
    """Define w:tblGrid com as larguras em twips e ajusta w:tblW (largura total)."""
    tbl = table._element
    # Garante tblPr
    tbl_pr = _ensure_tblPr(table)

    # Remove tblGrid anterior, se existir
    for grid in tbl.xpath('./w:tblGrid'):
        tbl.remove(grid)

    grid = OxmlElement('w:tblGrid')
    total_twips = 0
    for name in column_names:
        width_cm = float(widths_map_cm.get(name, 2.0))
        tw = _twips_from_cm(width_cm)
        total_twips += tw
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(tw))
        grid.append(gridCol)

    # Insere tblGrid logo após tblPr para manter ordem canônica
    # Se não houver outros nós, insere na posição 1
    try:
        tbl.insert(1, grid)
    except Exception:
        tbl.append(grid)

    # Garante w:tblW (largura total)
    # remove tblW anterior
    for child in list(tbl_pr):
        if child.tag == qn('w:tblW'):
            tbl_pr.remove(child)
    tbl_w = OxmlElement('w:tblW')
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_w.set(qn('w:w'), str(total_twips))
    tbl_pr.append(tbl_w)

def set_cell_font(cell, text, bold=False, size=14, color=None, align='left'):
    """
    Função auxiliar para formatar o texto da célula, com fonte padrão 14 e alinhamento justificado.
    """
    p = cell.paragraphs[0]
    p.text = text
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else: # Padrão para a esquerda
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    run = p.runs[0]
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def _verificar_dependencia_excel(caminho_arquivo):
    extensao = os.path.splitext(caminho_arquivo)[1].lower()
    if extensao in ('.xlsx', '.xlsm', '.xltx', '.xltm'):
        try:
            import openpyxl  # noqa: F401
        except Exception:
            return "openpyxl"
    elif extensao == '.xls':
        try:
            import xlrd  # noqa: F401
        except Exception:
            return "xlrd"
    return None

def selecionar_arquivo():
    """Abre o seletor de arquivos do Windows para escolher a planilha de entrada."""
    try:
        root = Tk()
        root.withdraw()
        root.attributes('-topmost', True)
        caminho = filedialog.askopenfilename(
            title="Selecione a planilha de entrada",
            filetypes=[
                ("Planilhas Excel", "*.xlsx *.xls"),
                ("Arquivos CSV", "*.csv"),
                ("Todos os arquivos", "*.*"),
            ],
        )
        root.destroy()
        return caminho if caminho else None
    except Exception as e:
        print(f"ERRO ao abrir o seletor de arquivos: {e}")
        return None

def processar_e_gerar_docx(caminho_arquivo, verbose=False):
    """
    Função principal que executa todo o processo de leitura, cálculo e geração do arquivo.
    
    Argumentos:
        caminho_arquivo (str): O caminho para o arquivo de planilha (.xlsx ou .csv).
        verbose (bool): Se True, imprime um log detalhado do processamento no console.
    """
    # 1. Validação do arquivo de entrada
    if not os.path.exists(caminho_arquivo):
        print(f"ERRO: O arquivo '{caminho_arquivo}' não foi encontrado. "
              f"Verifique se ele está na mesma pasta do script.")
        return

    # 2. Leitura da planilha com tratamento de erros
    try:
        if caminho_arquivo.lower().endswith('.csv'):
            try:
                # Tenta detectar separador automaticamente
                df = pd.read_csv(caminho_arquivo, sep=None, engine='python')
            except Exception:
                # Se falhar, usa “;” como padrão
                df = pd.read_csv(caminho_arquivo, sep=';')
        else:
            dependencia_faltando = _verificar_dependencia_excel(caminho_arquivo)
            if dependencia_faltando:
                print(
                    "ERRO FATAL ao tentar ler a planilha: "
                    f"Dependência '{dependencia_faltando}' não instalada. "
                    f"Instale com: pip install {dependencia_faltando}"
                )
                return
            df = pd.read_excel(caminho_arquivo)
        
        # Normaliza colunas esperadas
        colunas_esperadas = ["Data", "Descrição", "Conta", "Categoria", "Tags", "Valor", "%", "Parcela", "Situação"]
        for col in colunas_esperadas:
            if col not in df.columns:
                df[col] = np.nan

        print(f"Planilha lida com sucesso. Encontradas {len(df)} linhas para processar.")

    except Exception as e:
        print(f"ERRO FATAL ao tentar ler a planilha: {e}")
        return

    # 3. Processamento dos dados
    df['Valor'] = pd.to_numeric(df['Valor'], errors='coerce').fillna(0)
    df['Parcela'] = df['Valor'].copy()
    df['%'] = '1'
    logs = []
    
    contas_excluidas_inversao = ["Itaú - C.Corrente", "Banco do Brasil - C.Corrente"]

    for index, row in df.iterrows():
        valor = row['Valor']
        tags = row['Tags']
        conta = row['Conta']
        situacao = row['Situação']
        
        # Regra: 'Paga' pode inverter negativos (exceto contas da lista)
        deve_inverter_sinal_negativo = False

        # Se 'Paga', inverte sinal negativo. Exceção: contas à frente
        if situacao == 'Paga':
            excecao_geral = conta in contas_excluidas_inversao
            if not excecao_geral:
                deve_inverter_sinal_negativo = True

        if tags == 'carol' and situacao == 'Paga':
            novo_valor = abs(valor)
            logs.append(f"[linha {index:04d}] Regra: Paga, carol. Valor {valor:8.2f} -> {novo_valor:8.2f} (invertido para POSITIVO).")
            df.loc[index, 'Parcela'] = novo_valor
        elif tags == 'm&c':
            novo_valor = valor / 2
            df.loc[index, '%'] = '0,5'
            if deve_inverter_sinal_negativo:
                novo_valor = -novo_valor
                logs.append(f"[linha {index:04d}] Regra: Paga, m&c. Valor {valor:8.2f} -> {novo_valor:8.2f} (invertido e dividido).")
            else:
                logs.append(f"[linha {index:04d}] Regra: m&c (sem inversão). Valor {valor:8.2f} -> {novo_valor:8.2f} (dividido).")
            df.loc[index, 'Parcela'] = novo_valor
        elif tags == 'mauricio':
            novo_valor = valor
            if deve_inverter_sinal_negativo and valor < 0:
                novo_valor = -valor
                logs.append(f"[linha {index:04d}] Regra: Paga, mauricio. Valor {valor:8.2f} -> {novo_valor:8.2f} (invertido para POSITIVO).")
            else:
                logs.append(f"[linha {index:04d}] Regra: mauricio (sem inversão). Valor {valor:8.2f} mantido.")
            df.loc[index, 'Parcela'] = novo_valor
        else:
            novo_valor = valor
            if deve_inverter_sinal_negativo and valor < 0:
                novo_valor = -valor
                logs.append(f"[linha {index:04d}] Regra: Paga, outrem. Valor {valor:8.2f} -> {novo_valor:8.2f} (invertido para POSITIVO).")
            else:
                logs.append(f"[linha {index:04d}] Regra: {tags} (sem inversão). Valor {valor:8.2f} mantido.")
            df.loc[index, 'Parcela'] = novo_valor
    
    # 4. Ordenação e limpeza final
    df['Conta'] = df['Conta'].str.replace('Itaú - C.Corrente', 'Itaú', regex=False)
    print('Limpeza da coluna "Conta" realizada.')
    df.sort_values(by=['Tags', 'Valor'], ascending=[True, True], inplace=True)
    colunas_ordenadas = ["Data", "Descrição", "Conta", "Categoria", "Tags", "Valor", "%", "Parcela", "Situação"]
    df_final = df.reindex(columns=colunas_ordenadas)
    print("Dados e colunas ordenados para o arquivo final.")

    # 5. Cálculos dos totais
    soma_valor_original = df['Valor'].sum()
    filtro_parcela_tags = df['Tags'].isin(['carol', 'm&c'])
    soma_parcela_especifica = df.loc[filtro_parcela_tags, 'Parcela'].sum()
    print("Somas dos totais finalizadas com a lógica atualizada (tags: carol, m&c).")

    # 6. Geração do documento Word (.docx)
    doc = docx.Document()
    
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    print("Layout do documento definido para paisagem com margens estreitas.")

    doc.add_heading('Relatório de Despesas Processado', level=1)
    doc.add_paragraph()

    # --- LÓGICA DE LARGURA FIXA DAS COLUNAS (somente este tema foi alterado) ---
    larguras_fixas_cm = {
        "Data": 3.1, "Descrição": 6.2, "Conta": 2.0, "Categoria": 3.5,
        "Tags": 2.3, "Valor": 3.0, "%": 1.5, "Parcela": 2.6, "Situação": 2.4
    }
    print("Larguras fixas das colunas definidas.")
    
    table = doc.add_table(rows=1, cols=len(df_final.columns))
    table.style = 'Table Grid'

    # Impede o autoajuste e força layout fixo
    try:
        table.autofit = False
        table.allow_autofit = False  # alguns ambientes não têm esta prop; try/except evitaria erro
    except Exception:
        table.autofit = False

    # Força w:tblLayout type='fixed' com criação segura de tblPr
    _set_table_layout_fixed(table)

    # Define a grade (tblGrid) com larguras em twips e largura total (tblW)
    _apply_table_grid(table, list(df_final.columns), larguras_fixas_cm)

    # Cabeçalho
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df_final.columns):
        _set_cell_width(hdr_cells[i], larguras_fixas_cm.get(col_name, 2.0))
        set_cell_font(hdr_cells[i], col_name, bold=True, align='center')

    # Linhas
    for index, row in df_final.iterrows():
        row_cells = table.add_row().cells
        
        for i, col_name in enumerate(df_final.columns):
            _set_cell_width(row_cells[i], larguras_fixas_cm.get(col_name, 2.0))
            font_color = None
            alinhamento = 'justify'
            
            if col_name == 'Parcela':
                tags = row['Tags']
                parcela = row['Parcela']
                is_numeric_parcela = isinstance(parcela, (int, float))
                
                if (tags in ['carol', 'm&c']) and is_numeric_parcela and parcela < 0:
                    font_color = RGBColor(255, 0, 0)

            valor_celula = row[col_name]
            if isinstance(valor_celula, (int, float)):
                texto_formatado = f'{valor_celula:,.2f}'.replace(',', 'TEMP').replace('.', ',').replace('TEMP', '.')
                alinhamento = 'right'
            else:
                texto_formatado = str(valor_celula if pd.notna(valor_celula) else '')
            
            set_cell_font(row_cells[i], texto_formatado, color=font_color, align=alinhamento)
    
    print("Tabela de dados criada com larguras fixas e layout forçado.")

    # Totais
    total_cells = table.add_row().cells
    for i, col_name in enumerate(df_final.columns):
        _set_cell_width(total_cells[i], larguras_fixas_cm.get(col_name, 2.0))
    valor_col_index = list(df_final.columns).index('Valor')
    parcela_col_index = list(df_final.columns).index('Parcela')

    soma_valor_str = f'{soma_valor_original:,.2f}'.replace(',', 'TEMP').replace('.', ',').replace('TEMP', '.')
    soma_parcela_str = f'{soma_parcela_especifica:,.2f}'.replace(',', 'TEMP').replace('.', ',').replace('TEMP', '.')

    set_cell_font(total_cells[0], "TOTAIS", bold=True, align='center')
    set_cell_font(total_cells[valor_col_index], soma_valor_str, bold=True, align='right')
    set_cell_font(total_cells[parcela_col_index], soma_parcela_str, bold=True, align='right')
    print("Linha de totais integrada à tabela.")
    
    # 7. Salva o arquivo .docx (salva atômico: primeiro em tmp, depois renomeia)
    nome_saida = os.path.splitext(os.path.basename(caminho_arquivo))[0]
    pasta_saida = os.path.dirname(caminho_arquivo)
    caminho_saida = os.path.join(pasta_saida, f"{nome_saida}.docx")
    caminho_tmp = f"{caminho_saida}.tmp"
    doc.save(caminho_tmp)
    # renomeia garantindo substituição
    try:
        if os.path.exists(caminho_saida):
            os.remove(caminho_saida)
    except Exception:
        pass
    os.replace(caminho_tmp, caminho_saida)
    print(f"Arquivo Word gerado com sucesso: {caminho_saida}")

if __name__ == "__main__":
    nome_do_arquivo_de_entrada = selecionar_arquivo()
    if nome_do_arquivo_de_entrada:
        print(f"\nArquivo selecionado: '{nome_do_arquivo_de_entrada}'. Iniciando processamento...")
        modo_verbose = True
        processar_e_gerar_docx(nome_do_arquivo_de_entrada, verbose=modo_verbose)
    else:
        print("Nenhum arquivo foi selecionado.")
    
    try:
        input("\\nPressione Enter para sair...")
    except Exception:
        pass
