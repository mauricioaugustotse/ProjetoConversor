# -*- coding: utf-8 -*-
"""Documentos encaminhados pelo parlamentar junto com a demanda (anexos da GUI).

Extrai TEXTO de (quase) qualquer formato para sensibilizar a redação da IT/minuta:
- PDF com texto: PyMuPDF; PDF escaneado (sem camada de texto): OCR via visão da OpenAI;
- Word (.docx nativo; .doc antigo via COM do Word, quando pywin32/Word existirem);
- imagens (png/jpg/webp/tiff/bmp…): transcrição via visão da OpenAI;
- planilhas (.xlsx/.xlsm), .rtf, .eml, HTML e qualquer arquivo de texto (encoding detectado).

Documentos longos são CONDENSADOS por IA (modelo econômico) preservando pedidos,
dispositivos, números e datas — o contexto da redação tem teto de tamanho.
"""
from __future__ import annotations

import base64
import html as _html
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, List, Optional

from . import config_gerador as cfg
from . import llm

# ---------------------------------------------------------------- tetos (chars)
TETO_DOC = 24_000        # acima disso o documento é condensado por IA
CONDENSA_ALVO = 9_000    # tamanho-alvo do texto condensado
CONDENSA_MAX_IN = 110_000  # máximo enviado ao condensador (além disso, trunca antes)
TETO_TOTAL = 60_000      # teto da soma de todos os anexos no contexto da redação
ANALISE_CHARS = 3_500    # fatia de cada documento enviada à etapa de ANÁLISE
MAX_PAG_OCR = 12         # páginas de PDF escaneado transcritas por visão

_EXT_IMAGEM = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff", ".jfif"}
_EXT_TEXTO = {".txt", ".md", ".csv", ".tsv", ".json", ".xml", ".log", ".yaml", ".yml", ".ini"}

SYS_CONDENSA = (
    "Você condensa FIELMENTE documentos encaminhados a uma Consultoria Legislativa. Responda "
    "APENAS com o texto condensado (sem preâmbulo nem comentários), em português, preservando: "
    "o(s) pedido(s) e o objeto; TODOS os dispositivos normativos citados (art./§/inciso, lei, "
    "resolução); números, datas, prazos e valores; nomes de órgãos/pessoas relevantes; e as "
    "transcrições essenciais (entre aspas). Não opine, não acrescente nada que não esteja no "
    "documento, não omita pedidos. Alvo: cerca de {alvo} caracteres."
)

SYS_OCR = (
    "Você transcreve documentos para uma Consultoria Legislativa. As imagens são páginas de um "
    "documento (ou fotos/digitalizações). Transcreva FIELMENTE todo o conteúdo textual legível, "
    "em ordem de leitura, preservando títulos, numeração de dispositivos, tabelas (como texto "
    "alinhado por linhas) e assinaturas/carimbos relevantes. Não resuma, não comente, não "
    "invente texto ilegível — marque trechos ilegíveis com [ilegível]. Responda APENAS com a "
    "transcrição."
)


@dataclass
class Anexo:
    nome: str            # nome do arquivo (sem o caminho)
    texto: str           # texto extraído (já condensado, se foi o caso)
    origem: str = ""     # como foi extraído (pdf, docx, OCR IA…)
    aviso: str = ""      # problema não fatal a reportar ao usuário
    n_chars_original: int = 0


# ================================ extratores ================================
def _ler_texto_bruto(path: Path) -> str:
    """Arquivo de texto com encoding desconhecido. Ordem calibrada para documentos
    brasileiros: UTF-8 estrito -> BOM UTF-16 -> cp1252 estrito -> charset_normalizer
    (o detector erra em amostras curtas, por isso vem por último) -> latin-1."""
    raw = path.read_bytes()
    for enc in ("utf-8-sig", "utf-8"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            pass
    if raw[:2] in (b"\xff\xfe", b"\xfe\xff"):
        try:
            return raw.decode("utf-16")
        except UnicodeDecodeError:
            pass
    if b"\x00" not in raw:  # NUL indica binário/UTF-16 sem BOM — não é cp1252
        try:
            return raw.decode("cp1252")
        except UnicodeDecodeError:
            pass
    try:
        from charset_normalizer import from_bytes
        best = from_bytes(raw).best()
        if best is not None:
            return str(best)
    except Exception:  # noqa: BLE001
        pass
    if b"\x00" in raw:
        raise ValueError("arquivo binário sem texto decodificável")
    return raw.decode("latin-1")


def _strip_html(texto: str) -> str:
    texto = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", texto)
    texto = re.sub(r"(?i)<br\s*/?>|</p>|</div>|</tr>|</li>|</h[1-6]>", "\n", texto)
    texto = re.sub(r"<[^>]+>", " ", texto)
    texto = _html.unescape(texto)
    return re.sub(r"[ \t]+", " ", texto)


def _extrair_pdf(path: Path, log: Callable) -> tuple[str, str]:
    import fitz

    doc = fitz.open(str(path))
    if doc.needs_pass:
        raise ValueError("PDF protegido por senha")
    paginas = [p.get_text("text") for p in doc]
    texto = "\n".join(paginas).strip()
    # camada de texto ausente/irrisória => digitalização: transcreve por visão (IA)
    if len(doc) and len(texto) < 60 * len(doc):
        log(f"   “{path.name}”: PDF sem camada de texto (escaneado) — transcrevendo por visão (IA)…")
        n = min(len(doc), MAX_PAG_OCR)
        imgs = [doc[i].get_pixmap(matrix=fitz.Matrix(2, 2)).tobytes("png") for i in range(n)]
        texto_ocr = llm.chat_visao(
            SYS_OCR, f"Transcreva as {n} página(s) do documento “{path.name}”.", imgs)
        aviso = (f"transcritas as {n} primeiras páginas de {len(doc)} (teto de OCR)"
                 if len(doc) > n else "")
        doc.close()
        if len((texto_ocr or "").strip()) > len(texto):
            return texto_ocr.strip(), f"PDF escaneado · OCR IA ({n} pág.)" + (f" — {aviso}" if aviso else "")
        return texto, "pdf (pouco texto extraível)"
    doc.close()
    return texto, f"pdf ({len(paginas)} pág.)"


def _extrair_docx(path: Path) -> str:
    import docx

    d = docx.Document(str(path))
    partes = [p.text for p in d.paragraphs]
    for tb in d.tables:
        for row in tb.rows:
            partes.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(x for x in partes if x is not None)


def _extrair_doc_antigo(path: Path) -> str:
    """Word 97-2003 via COM (requer pywin32 + Word instalado)."""
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()  # o gerador roda em worker thread da GUI
    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        d = word.Documents.Open(str(path), ReadOnly=True, AddToRecentFiles=False)
        try:
            return d.Content.Text
        finally:
            d.Close(False)
    finally:
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def _extrair_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    partes: List[str] = []
    try:
        for ws in wb.worksheets:
            partes.append(f"[Planilha: {ws.title}]")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= 400:
                    partes.append("[… demais linhas omitidas …]")
                    break
                cells = ["" if c is None else str(c) for c in row]
                if any(x.strip() for x in cells):
                    partes.append("\t".join(cells).rstrip())
    finally:
        wb.close()
    return "\n".join(partes)


def _extrair_rtf(path: Path) -> str:
    from striprtf.striprtf import rtf_to_text

    return rtf_to_text(_ler_texto_bruto(path), errors="ignore")


def _extrair_eml(path: Path) -> str:
    import email
    from email import policy

    msg = email.message_from_bytes(path.read_bytes(), policy=policy.default)
    cab = [f"{k}: {msg.get(k, '')}" for k in ("From", "To", "Date", "Subject") if msg.get(k)]
    corpo = ""
    parte = msg.get_body(preferencelist=("plain", "html"))
    if parte is not None:
        corpo = parte.get_content()
        if parte.get_content_type() == "text/html":
            corpo = _strip_html(corpo)
    anexos = [f"[anexo do e-mail: {fn}]" for fn in
              (p.get_filename() for p in msg.iter_attachments()) if fn]
    return "\n".join(cab) + "\n\n" + corpo + ("\n" + "\n".join(anexos) if anexos else "")


def _extrair_imagem(path: Path, log: Callable) -> str:
    """Qualquer imagem: normaliza para PNG via PyMuPDF e transcreve por visão (IA)."""
    import fitz

    log(f"   “{path.name}”: imagem — transcrevendo por visão (IA)…")
    doc = fitz.open(str(path))  # PyMuPDF abre imagens como documento de 1 página
    png = doc[0].get_pixmap().tobytes("png")
    doc.close()
    texto = llm.chat_visao(SYS_OCR, f"Transcreva o documento da imagem “{path.name}”.", [png])
    return (texto or "").strip()


def extrair(path: Path, log: Callable = print) -> tuple[str, str]:
    """Extrai o texto de UM arquivo. Retorna (texto, origem). Levanta exceção se impossível."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extrair_pdf(path, log)
    if ext == ".docx":
        return _extrair_docx(path), "docx"
    if ext == ".doc":
        try:
            return _extrair_doc_antigo(path), "doc (via Word)"
        except ImportError:
            raise ValueError("formato .doc antigo requer o Word/pywin32 — salve como .docx ou PDF")
    if ext in (".xlsx", ".xlsm"):
        return _extrair_xlsx(path), "planilha"
    if ext == ".rtf":
        return _extrair_rtf(path), "rtf"
    if ext == ".eml":
        return _extrair_eml(path), "e-mail"
    if ext in (".html", ".htm"):
        return _strip_html(_ler_texto_bruto(path)), "html"
    if ext in _EXT_IMAGEM:
        return _extrair_imagem(path, log), "imagem · OCR IA"
    # .txt/.md/.csv/... e QUALQUER extensão desconhecida: tenta como texto
    try:
        texto = _ler_texto_bruto(path)
    except Exception:
        raise ValueError(f"formato “{ext or 'sem extensão'}” não suportado — converta para PDF ou texto")
    if ext == ".msg":  # .msg da Outlook é binário; se decodificou algo, é lixo
        raise ValueError("formato .msg do Outlook não suportado — reenvie como .eml ou PDF")
    return texto, ("texto" if ext in _EXT_TEXTO or not ext else f"texto ({ext})")


# ================================ condensação ================================
def _condensar(texto: str, nome: str, log: Callable) -> tuple[str, bool]:
    """Documento acima do teto -> condensação fiel por IA (modelo econômico).
    Retorna (texto_final, condensou). Em falha da IA, trunca com marcador."""
    log(f"   “{nome}”: {len(texto):,} caracteres — condensando por IA (fiel, com dispositivos e pedidos)…")
    try:
        out = llm.chat(SYS_CONDENSA.format(alvo=CONDENSA_ALVO),
                       f"DOCUMENTO “{nome}”:\n{texto[:CONDENSA_MAX_IN]}",
                       model=cfg.MODEL_ANEXOS)
        out = (out or "").strip()
        if out:
            return out, True
    except Exception as exc:  # noqa: BLE001
        log(f"   condensação falhou ({exc}) — usando o início do documento.")
    return texto[:TETO_DOC] + "\n[… documento truncado por tamanho …]", False


def _normalizar(texto: str) -> str:
    texto = texto.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


def processar(paths: List[str], *, log: Callable = print) -> List[Anexo]:
    """Extrai (e condensa, se preciso) cada arquivo. Nunca lança por causa de UM
    arquivo: falhas viram Anexo com texto vazio + aviso (o chamador reporta)."""
    out: List[Anexo] = []
    for p in paths:
        path = Path(p)
        nome = path.name
        if not path.is_file():
            out.append(Anexo(nome=nome, texto="", aviso="arquivo não encontrado"))
            continue
        try:
            texto, origem = extrair(path, log)
        except Exception as exc:  # noqa: BLE001
            out.append(Anexo(nome=nome, texto="", aviso=f"falha na extração: {exc}"))
            log(f"   ✗ “{nome}”: {exc}")
            continue
        texto = _normalizar(texto)
        n0 = len(texto)
        aviso = ""
        if not texto:
            aviso = "nenhum texto extraído (documento vazio?)"
        elif n0 > TETO_DOC:
            texto, ok = _condensar(texto, nome, log)
            origem += " · condensado por IA" if ok else " · truncado"
        out.append(Anexo(nome=nome, texto=texto, origem=origem, aviso=aviso, n_chars_original=n0))
        if texto:
            log(f"   ✓ “{nome}”: {n0:,} caracteres ({origem}).")
    # teto GLOBAL do contexto: trunca os maiores até caber (mantém todos representados).
    # Piso por documento proporcional ao nº de anexos, para o loop SEMPRE convergir.
    marca = "\n[… documento truncado por tamanho …]"
    piso = max(1_000, min(CONDENSA_ALVO, TETO_TOTAL // max(len(out), 1)))
    while True:
        total = sum(len(a.texto) for a in out)
        if total <= TETO_TOTAL:
            break
        maior = max(out, key=lambda a: len(a.texto))
        novo = max(len(maior.texto) - (total - TETO_TOTAL) - len(marca), piso)
        if novo + len(marca) >= len(maior.texto):
            break  # todos já no piso: não há mais o que cortar
        maior.texto = maior.texto[:novo] + marca
        if "truncado (teto global)" not in maior.origem:
            maior.origem += " · truncado (teto global)"
    return out


# ================================ formatação p/ prompts ================================
def formatar_contexto(anexos: List[Anexo]) -> str:
    """Bloco dos documentos para o CONTEXTO da redação (IT e minuta)."""
    docs = [a for a in anexos if a.texto]
    if not docs:
        return ""
    partes = []
    for i, a in enumerate(docs, 1):
        partes.append(f"--- DOCUMENTO ENCAMINHADO Nº {i}: “{a.nome}” ({a.origem}) ---\n{a.texto}")
    return "\n\n".join(partes)


def resumo_para_analise(anexos: List[Anexo]) -> str:
    """Fatia de cada documento para a etapa de ANÁLISE da demanda (controle de tokens)."""
    docs = [a for a in anexos if a.texto]
    if not docs:
        return ""
    partes = []
    for i, a in enumerate(docs, 1):
        corte = a.texto[:ANALISE_CHARS]
        sufixo = "\n[… íntegra disponível na etapa de redação …]" if len(a.texto) > ANALISE_CHARS else ""
        partes.append(f"--- DOCUMENTO ENCAMINHADO Nº {i}: “{a.nome}” ---\n{corte}{sufixo}")
    return "\n\n".join(partes)
