# -*- coding: utf-8 -*-
"""acv_02: cruza D:\\ACERVO contra C:\\HD_Mau e classifica cada arquivo do D.

Níveis: 1) sha1 binário igual a algum do C -> dup_bin;
        2) sha1 binário repetido dentro do D -> dup_interna (1 representante segue);
        3) sha1 do TEXTO normalizado (normalização idêntica ao org2_03) presente
           no org2_cache_fulltext.tsv do acervo C -> dup_texto;
        4) resto -> omisso (com texto) | omisso_sem_texto.
Texto do lado D vai para acv_cache_texto_d.tsv (retomável). Resultado na tabela
`confronto` do acv_confronto.sqlite + resumo por tema.
"""
import collections
import hashlib
import io
import os
import re
import sqlite3
import subprocess
import sys
import time
import unicodedata
from concurrent.futures import ThreadPoolExecutor

BASE = os.path.dirname(os.path.abspath(__file__))
DB = os.path.join(BASE, 'acv_confronto.sqlite')
CACHE_C = os.path.join(BASE, 'org2_cache_fulltext.tsv')
CACHE_D = os.path.join(BASE, 'acv_cache_texto_d.tsv')
ANTIWORD = r'C:\Program Files\Git\mingw64\bin\antiword.exe'
MIN_NORM = 300
EXTS_TEXTO = ('.pdf', '.doc', '.docx', '.rtf', '.txt')


def lp(p):
    return p if p.startswith('\\\\?\\') else '\\\\?\\' + p


def normalizar(texto):
    t = unicodedata.normalize('NFD', texto.lower())
    t = ''.join(c for c in t if not unicodedata.combining(c))
    return re.sub(r'[^0-9a-z]', '', t)


def extrair_completo(path):
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == '.pdf':
            r = subprocess.run(['pdftotext', '-q', '-enc', 'UTF-8', path, '-'],
                               capture_output=True, timeout=120)
            return r.stdout.decode('utf-8', 'ignore') if r.returncode == 0 else None
        if ext == '.docx':
            import docx
            d = docx.Document(path)
            parts = [p.text for p in d.paragraphs]
            for tb in d.tables:
                for row in tb.rows:
                    parts.extend(c.text for c in row.cells)
            return '\n'.join(parts)
        if ext == '.doc':
            r = subprocess.run([ANTIWORD, '-m', 'UTF-8', path],
                               capture_output=True, timeout=120)
            return r.stdout.decode('utf-8', 'ignore') if r.returncode == 0 else None
        if ext == '.rtf':
            from striprtf.striprtf import rtf_to_text
            raw = io.open(path, encoding='cp1252', errors='ignore').read()
            return rtf_to_text(raw, errors='ignore')
        if ext == '.txt':
            return io.open(path, encoding='utf-8', errors='ignore').read()
    except Exception:
        return None
    return None


def hashes_texto_c():
    """Conjunto de sha1 do texto integral presentes no acervo C (cache org2)."""
    hs = set()
    with io.open(CACHE_C, encoding='utf-8') as f:
        for line in f:
            parts = line.rstrip('\n').split('\t')
            if len(parts) == 5 and parts[3] != 'FALHA' and int(parts[4]) >= MIN_NORM:
                hs.add(parts[3])
    return hs


def main():
    con = sqlite3.connect(DB)
    con.execute("""CREATE TABLE IF NOT EXISTS confronto(
        path TEXT PRIMARY KEY, classe TEXT, ref TEXT)""")

    c_sha = set()
    d_rows = []
    for p, lado, size, sha1, erro in con.execute(
            'SELECT path, lado, size, sha1, erro FROM hashes'):
        if erro:
            continue
        if lado == 'C':
            c_sha.add(sha1)
        else:
            d_rows.append((p, size, sha1))
    print(f'C: {len(c_sha):,} hashes únicos | D: {len(d_rows):,} arquivos')

    con.execute('DELETE FROM confronto')
    lote = []
    vistos_d = {}
    candidatos = []
    for p, size, sha1 in sorted(d_rows):
        if sha1 in c_sha:
            lote.append((p, 'dup_bin', ''))
        elif sha1 in vistos_d:
            lote.append((p, 'dup_interna', vistos_d[sha1]))
        else:
            vistos_d[sha1] = p
            candidatos.append(p)
    print(f'dup_bin: {sum(1 for x in lote if x[1]=="dup_bin"):,} | '
          f'dup_interna: {sum(1 for x in lote if x[1]=="dup_interna"):,} | '
          f'candidatos p/ nível textual: {len(candidatos):,}')

    # nível textual
    hs_c = hashes_texto_c()
    print(f'Hashes de texto do acervo C (cache org2): {len(hs_c):,}')
    cache_d = {}
    if os.path.exists(CACHE_D):
        with io.open(CACHE_D, encoding='utf-8') as f:
            for line in f:
                parts = line.rstrip('\n').split('\t')
                if len(parts) == 4:
                    cache_d[parts[0]] = (parts[1], parts[2], int(parts[3]))

    com_texto = [p for p in candidatos
                 if os.path.splitext(p)[1].lower() in EXTS_TEXTO]
    pend = []
    for p in com_texto:
        try:
            st = os.stat(lp(p))
        except OSError:
            continue
        key = f'{st.st_size}|{int(st.st_mtime)}'
        if p not in cache_d or cache_d[p][0] != key:
            pend.append((p, key))
    print(f'Extraindo texto de {len(pend):,} (cache: {len(com_texto)-len(pend):,})...',
          flush=True)

    def worker(item):
        p, key = item
        txt = extrair_completo(p)
        if txt is None:
            return (p, key, 'FALHA', 0)
        norm = normalizar(txt)
        return (p, key, hashlib.sha1(norm.encode()).hexdigest(), len(norm))

    t0 = time.time()
    with ThreadPoolExecutor(max_workers=8) as ex, \
            io.open(CACHE_D, 'a', encoding='utf-8') as cf:
        for i, (p, key, h, ln) in enumerate(ex.map(worker, pend), 1):
            cf.write(f'{p}\t{key}\t{h}\t{ln}\n')
            cf.flush()
            cache_d[p] = (key, h, ln)
            if i % 500 == 0:
                print(f'  {i:,}/{len(pend):,} ({time.time()-t0:.0f}s)', flush=True)

    stats = collections.Counter()
    for p in candidatos:
        ext = os.path.splitext(p)[1].lower()
        if ext in EXTS_TEXTO and p in cache_d:
            _, h, ln = cache_d[p]
            if h != 'FALHA' and ln >= MIN_NORM and h in hs_c:
                lote.append((p, 'dup_texto', ''))
                stats['dup_texto'] += 1
                continue
        classe = 'omisso' if ext in EXTS_TEXTO else 'omisso_sem_texto'
        lote.append((p, classe, ''))
        stats[classe] += 1

    con.executemany('INSERT INTO confronto VALUES(?,?,?)', lote)
    con.commit()

    print('\n== RESULTADO ==')
    por_classe = collections.Counter(x[1] for x in lote)
    for k, v in por_classe.most_common():
        print(f'  {k:16}: {v:,}')
    print('\n== Omissos por tema ==')
    por_tema = collections.defaultdict(collections.Counter)
    for p, classe, _ in lote:
        if classe.startswith('omisso'):
            tema = p.split('\\')[2] if len(p.split('\\')) > 3 else '(raiz)'
            por_tema[tema][classe] += 1
    for tema, c in sorted(por_tema.items(), key=lambda x: -sum(x[1].values())):
        print(f'  {sum(c.values()):>6,}  {tema}  ({dict(c)})')
    con.close()


if __name__ == '__main__':
    main()
