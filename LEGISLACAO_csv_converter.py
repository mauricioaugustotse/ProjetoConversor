# -*- coding: utf-8 -*-
"""
Converte legislação (Markdown, HTML, PDF ou DOCX) em CSV para importação no Notion/RAG.

Fluxo de trabalho:
1. Localiza o arquivo de entrada automaticamente (ou por seleção manual).
2. Lê o conteúdo conforme extensão e normaliza caracteres/espaçamentos.
3. Detecta a hierarquia normativa (artigos, parágrafos, incisos e alíneas).
4. Limpa títulos/textos, remove ruídos e preserva contexto hierárquico.
5. Escreve CSV final com `ID`, `Nome (Artigo)` e `Texto do Artigo`.
"""

import csv
import re
import os
import glob
import argparse
import unicodedata
from html.parser import HTMLParser
from pathlib import Path

from gui_intuitiva import open_file_panel

SUPPORTED_EXTENSIONS = (".md", ".html", ".htm", ".pdf", ".docx")
SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_INPUT_TARGET = "L9605_files"
CANCELLED_BY_USER = object()
ARTICLE_HEADING_PATTERN = (
    r'Arts?\.\s*\d+(?:\s+a\s*\d+)?(?:\.\d+)*(?:[.\s]*(?:\u00ba|\u00aa))?(?:-[A-Z]+)?(?:\.\s*|\s+)?'
)
REGIMENTO_DOCUMENT = "Regimento Interno da Câmara dos Deputados"
CODIGO_ETICA_DOCUMENT = "Código de Ética e Decoro Parlamentar da Câmara dos Deputados"
STRUCTURE_LEVELS = ("documento", "parte", "livro", "titulo", "capitulo", "secao", "subsecao", "anexo")
STRUCTURAL_HEADER_PATTERN = re.compile(
    r'^(PARTE|LIVRO|T[ÍI]TULO|CAP[ÍI]TULO|SE[ÇC][ÃA]O|SUBSE[ÇC][ÃA]O|ANEXO)\b(?:\s+(.*))?$',
    re.IGNORECASE,
)
HEADING_KIND_TO_KEY = {
    "PARTE": "parte",
    "LIVRO": "livro",
    "TITULO": "titulo",
    "CAPITULO": "capitulo",
    "SECAO": "secao",
    "SUBSECAO": "subsecao",
    "ANEXO": "anexo",
}
HEADING_IDENTIFIER_WORDS = {
    "UNICO", "UNICA", "GERAL",
    "PRIMEIRO", "PRIMEIRA", "SEGUNDO", "SEGUNDA", "TERCEIRO", "TERCEIRA",
    "QUARTO", "QUARTA", "QUINTO", "QUINTA", "SEXTO", "SEXTA",
    "SETIMO", "SETIMA", "OITAVO", "OITAVA", "NONO", "NONA",
    "DECIMO", "DECIMA",
}
GENERIC_DOCUMENT_HEADING_PATTERN = re.compile(
    r'^(?:'
    r'CONSTITUICAO'
    r'|EMENDA\s+CONSTITUCIONAL\s+N[ºO.]?'
    r'|LEI(?:\s+COMPLEMENTAR)?\s+N[ºO.]?'
    r'|DECRETO(?:\s+LEI)?\s+N[ºO.]?'
    r'|MEDIDA\s+PROVISORIA\s+N[ºO.]?'
    r')\b',
    re.IGNORECASE,
)


def normalize_ordinals(text):
    """Padroniza indicadores ordinais duplicados e variações inconsistentes."""
    if not text:
        return text
    text = text.replace('\u00b0', '\u00ba')
    text = text.replace('.\u00ba', '\u00ba').replace('.\u00aa', '\u00aa')
    text = re.sub(r'\u00ba{2,}', '\u00ba', text)
    text = re.sub(r'\u00aa{2,}', '\u00aa', text)
    text = re.sub(r'(\d+)\s*[oO\u00ba]\b', lambda m: f"{m.group(1)}\u00ba", text)
    return text


def normalize_decimal_commas(text):
    """Remove espaços indevidos em torno de vírgulas decimais."""
    if not text:
        return text
    return re.sub(r'(?<=\d)\s*,\s*(?=\d)', ',', text)


def normalize_space_characters(text):
    """Normaliza espaços especiais (NBSP, ZWSP, BOM) para evitar falhas de reconhecimento."""
    if not text:
        return text
    text = text.replace('\ufeff', '')
    text = re.sub(r'[\u00a0\u1680\u2000-\u200a\u202f\u205f\u3000]', ' ', text)
    text = re.sub(r'[\u200b\u200c\u200d\u2060]', '', text)
    return text


def normalize_for_matching(text):
    """Remove acentos e padroniza caixa para comparações estruturais."""
    text = normalize_space_characters(text or "")
    normalized = unicodedata.normalize("NFKD", text)
    without_accents = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    return re.sub(r'\s+', ' ', without_accents).strip().upper()


def get_document_heading(text):
    """Detecta cabeçalhos de documentos normativos para compor metadados de RAG."""
    normalized = normalize_for_matching(text)
    if normalized == "CODIGO DE ETICA E DECORO PARLAMENTAR DA CAMARA DOS DEPUTADOS":
        return CODIGO_ETICA_DOCUMENT
    if normalized == "REGIMENTO INTERNO DA CAMARA DOS DEPUTADOS":
        return REGIMENTO_DOCUMENT
    if GENERIC_DOCUMENT_HEADING_PATTERN.match(normalized):
        return clean_title(text)
    return ""


def is_resolution_heading(text):
    """Identifica resoluções intermediárias que não devem virar itens do Regimento."""
    normalized = normalize_for_matching(text)
    return bool(re.match(r'^RESOLUCAO\s+N[ºO.]?\s*\d+', normalized))


def heading_kind_key(kind):
    normalized = normalize_for_matching(kind)
    return HEADING_KIND_TO_KEY.get(normalized, "")


def is_heading_identifier(value):
    normalized = normalize_for_matching(value).rstrip(".")
    if normalized in HEADING_IDENTIFIER_WORDS:
        return True
    return bool(re.match(r'^(?:[IVXLCDM]+(?:-[A-Z]+)?|\d+(?:-[A-Z]+)?|\d+[ºª](?:-[A-Z]+)?)$', normalized))


def parse_structural_header(text):
    """Retorna (nivel, cabecalho, pendente_descricao) para títulos normativos."""
    match = STRUCTURAL_HEADER_PATTERN.match(text)
    if not match:
        return None

    kind, rest = match.groups()
    key = heading_kind_key(kind)
    if not key:
        return None

    kind_display = normalize_space_characters(kind).strip().upper()
    rest = normalize_space_characters(rest or "").strip()
    if not rest:
        return key, kind_display, True

    first, sep, remainder = rest.partition(" ")
    if is_heading_identifier(first):
        base = f"{kind_display} {first}".strip()
        if remainder.strip():
            return key, f"{base} - {remainder.strip()}", False
        return key, base, True

    return key, f"{kind_display} {rest}".strip(), False


def is_heading_description(text):
    """Reconhece descrições curtas que completam TÍTULO/CAPÍTULO/SEÇÃO."""
    text = normalize_space_characters(text or "").strip()
    if not text or text.startswith("("):
        return False
    if len(text) > 180 or text.endswith((".", ";", ":")):
        return False
    if re.match(fr'^(?:{ARTICLE_HEADING_PATTERN})', text):
        return False
    if re.match(r'^(?:\u00a7\s*\d+|Par\u00e1grafo \u00fanico|[IVXLCDM]+\s*[-\u2013\u2014]|[a-z]\)|\d+[\.\)\-])', text):
        return False

    letters = [ch for ch in text if ch.isalpha()]
    if not letters:
        return False
    return True


def heading_description_continues(text):
    """Indica que o cabeçalho provavelmente continua na próxima linha."""
    normalized = normalize_for_matching(text).rstrip()
    if not normalized:
        return False
    if normalized.endswith(","):
        return True
    last_word = normalized.rsplit(" ", 1)[-1]
    return last_word in {"DE", "DO", "DA", "DOS", "DAS", "E"}


def apply_structure_heading(structure_context, key, heading):
    """Atualiza o contexto hierárquico e limpa níveis inferiores."""
    if key not in STRUCTURE_LEVELS:
        return
    structure_context[key] = heading
    start = STRUCTURE_LEVELS.index(key) + 1
    for lower_key in STRUCTURE_LEVELS[start:]:
        structure_context[lower_key] = ""


def append_structure_description(structure_context, key, description):
    if not key or key not in STRUCTURE_LEVELS or not description:
        return
    current = structure_context.get(key, "")
    if not current:
        structure_context[key] = description
        return
    if description in current:
        return
    if " - " in current:
        structure_context[key] = f"{current} {description}"
    else:
        structure_context[key] = f"{current} - {description}"


def build_metadata_lines(structure_context, item_title):
    lines = []
    document = structure_context.get("documento", "")
    if document:
        lines.append(f"Documento: {document}")
    hierarchy = " > ".join(
        structure_context.get(level, "")
        for level in STRUCTURE_LEVELS
        if level != "documento" and structure_context.get(level)
    )
    if hierarchy:
        lines.append(f"Hierarquia: {hierarchy}")
    if item_title:
        lines.append(f"Dispositivo: {item_title}")
    return lines


def content_from_best_start(content):
    """Para páginas consolidadas, inicia no Regimento anexo; nas demais, preserva o preâmbulo."""
    lines = content.splitlines()
    for index, line in enumerate(lines):
        if get_document_heading(line) == REGIMENTO_DOCUMENT:
            return "\n".join(lines[index:])

    match = re.search(r"Art\.\s*1", content, re.IGNORECASE)
    if not match:
        return ""
    return content


article_sequence_pattern = re.compile(
    r'(\barts?\.)'
    r'(\s*)'
    r'(\d[\wºª-]*(?:,\s*\d[\wºª-]*)+)',
    flags=re.IGNORECASE,
)


def fix_article_number_spacing(text):
    """Garante espaço após vírgulas em sequências de artigos (ex.: 'arts. 21, 26')."""
    if not text:
        return text

    def repl(match):
        prefix, spacing, numbers = match.groups()
        spacing = spacing if spacing else ' '
        normalized_numbers = re.sub(r',\s*(?=\d)', ', ', numbers)
        return f"{prefix}{spacing}{normalized_numbers}"

    return article_sequence_pattern.sub(repl, text)


def normalize_article_heading_punctuation(text):
    """Corrige pontuação ausente em cabeçalhos de artigo mal formatados."""
    if not text:
        return text
    text = re.sub(
        r'^(Art\.\s*\d+(?:\.\d+)*(?:\u00ba|\u00aa)?(?:-[A-Z]+)?)\.(?=[A-Z\u00C0-\u024F])',
        r'\1. ',
        text,
    )
    text = re.sub(
        r'^(Art\.\s*\d+(?:\.\d+)*(?:\u00ba|\u00aa)?(?:-[A-Z]+)?)(?=\s+[A-Z\u00C0-\u024F])',
        r'\1.',
        text,
    )
    return text


def clean_title(raw_title):
    """Funcao centralizada para limpar os titulos dos dispositivos."""
    cleaned = normalize_space_characters(raw_title)
    cleaned = cleaned.strip()
    cleaned = re.sub(r'(?i)^Art\.\s*', 'Art. ', cleaned)
    cleaned = normalize_ordinals(cleaned)
    cleaned = cleaned.rstrip('.')
    cleaned = re.sub(r'\u00a7\s+', '\u00a7', cleaned)
    cleaned = re.sub(
        r'(\u00a7)([1-9])(?!\u00ba|\d)',
        lambda m: f"{m.group(1)}{m.group(2)}\u00ba",
        cleaned,
    )
    cleaned = re.sub(
        r'(?i)(Lei)\s+n[\s.\-]*(?:\u00ba|o)?\s*(?=\d)',
        lambda m: f"{m.group(1)} n\u00ba ",
        cleaned,
    )
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned


def is_standalone_noise_line(text):
    """Remove notas isoladas que sobram de trechos HTML riscados/obsoletos."""
    normalized = normalize_for_matching(text).rstrip(".")
    return normalized in {"PREJUDICADA", "PREJUDICADO"}


def collapse_duplicate_text(text):
    """Remove duplicações simples em que o mesmo conteúdo aparece duas vezes em sequência."""
    if not text:
        return text

    leading_len = len(text) - len(text.lstrip())
    trailing_len = len(text) - len(text.rstrip())
    core = text[leading_len:len(text) - trailing_len if trailing_len else None]

    starts_with_newline = core.startswith('\n') if core else False
    core_stripped = core.strip()

    if core_stripped:
        duplicate_match = re.match(r'(?P<block>.+?)\s+(?P=block)$', core_stripped, re.DOTALL)
        if duplicate_match:
            core_stripped = duplicate_match.group('block')

    if starts_with_newline:
        core_stripped = '\n' + core_stripped

    prefix = text[:leading_len] if leading_len else ''
    suffix = text[len(text) - trailing_len:] if trailing_len else ''
    return f"{prefix}{core_stripped}{suffix}"


class ParagraphHTMLExtractor(HTMLParser):
    """Extrai blocos textuais relevantes de HTML normativo sem bibliotecas externas."""

    block_tags = {"p", "li", "h1", "h2", "h3", "h4", "h5", "h6"}
    skip_tags = {"script", "style", "strike", "s", "del"}

    def __init__(self):
        super().__init__()
        self.capture_depth = 0
        self.skip_depth = 0
        self.current_parts = []
        self.lines = []

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in self.skip_tags:
            self.skip_depth += 1
            return
        if self.skip_depth:
            return
        if tag in self.block_tags:
            if self.capture_depth == 0:
                self.current_parts = []
            self.capture_depth += 1
        elif self.capture_depth and tag == "br":
            self.current_parts.append("\n")

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in self.skip_tags and self.skip_depth:
            self.skip_depth -= 1
            return
        if self.skip_depth:
            return
        if tag in self.block_tags and self.capture_depth:
            self.capture_depth -= 1
            if self.capture_depth == 0:
                self._flush_current_parts()

    def handle_data(self, data):
        if self.skip_depth or not self.capture_depth:
            return
        self.current_parts.append(data)

    def _flush_current_parts(self):
        text = normalize_space_characters("".join(self.current_parts)).replace("\r", "")
        candidate = re.sub(r'\s+', ' ', text).strip()
        if candidate and not is_standalone_noise_line(candidate):
            self.lines.append(candidate)
        self.current_parts = []


def detect_html_encoding(raw_bytes):
    """Tenta descobrir a codificação declarada no HTML."""
    meta_match = re.search(br'charset\s*=\s*["\']?([\w.\-]+)', raw_bytes[:8192], re.IGNORECASE)
    if not meta_match:
        return None
    try:
        return meta_match.group(1).decode("ascii", errors="ignore")
    except Exception:
        return None



def save_current_item(rows, title, text_lines, metadata_lines=None):
    """Funcao auxiliar para salvar o item (artigo, paragrafo, etc.) processado.
    Agrupa as linhas de texto, remove anotacoes/links/lixo e adiciona a lista de resultados."""
    cleaned_text = None
    cleaned_body = None
    title_normalized = clean_title(title) if title else title
    if title and text_lines:
        link_pattern = re.compile(r'\[([^\]]+)\]\s*\([^)]*\)')
        asterisk_header_pattern = re.compile(r'(\s*\*\*[^\*]+\*\*\s*)+')
        bold_inline_pattern = re.compile(r'\*\*(.+?)\*\*')

        normalized_lines = []

        for segment in text_lines:
            if segment is None:
                continue
            segment = normalize_space_characters(segment)
            segment = segment.replace('\r', '')

            for raw_line in segment.split('\n'):
                raw_line = normalize_space_characters(raw_line)
                candidate = raw_line.strip()
                if not candidate:
                    continue
                candidate = re.sub(r'^#{1,6}\s*', '', candidate)

                candidate = link_pattern.sub(r'\1', candidate)
                candidate = re.sub(r'\[([^\]]+)\]', r'\1', candidate)
                candidate = bold_inline_pattern.sub(lambda m: m.group(1).strip(), candidate)
                candidate = asterisk_header_pattern.sub(' ', candidate)

                candidate = candidate.replace('_', '')
                candidate = normalize_ordinals(candidate)
                candidate = normalize_article_heading_punctuation(candidate)
                candidate = re.sub(r'(?i)\bArt\.\s*(?=\d)', 'Art. ', candidate)
                candidate = re.sub(r'\u00a7\s+', '\u00a7', candidate)
                candidate = re.sub(
                    r'(\u00a7)([1-9])(?!\u00ba|\d)',
                    lambda m: f"{m.group(1)}{m.group(2)}\u00ba",
                    candidate,
                )
                candidate = re.sub(
                    r'(?i)(Lei)\s+n[\s.\-]*(?:\u00ba|o)?\s*(?=\d)',
                    lambda m: f"{m.group(1)} n\u00ba ",
                    candidate,
                )
                candidate = re.sub(r'\(\s*(?:https?://|\.\.?/|#)[^)]+\)', '', candidate)
                candidate = re.sub(r'\(\s*[\w\u00C0-\u024F./#%-]*\.html?[^)]*\)', '', candidate, flags=re.IGNORECASE)
                candidate = re.sub(r'\s*\(\s*', ' (', candidate)
                candidate = re.sub(r'\s+\)', ')', candidate)
                candidate = re.sub(r'\)\s*(?=\()', ') (', candidate)
                candidate = re.sub(r'^ \(', '(', candidate)
                candidate = re.sub(r'\(\s*\(', '(', candidate)
                candidate = re.sub(r'\)\s*\)', ')', candidate)
                candidate = re.sub(r'\[\s*\]\s*\([^)]*\)', '', candidate)
                candidate = re.sub(r'https?://\S+', '', candidate)
                candidate = re.sub(r'\s*>\s*', ' ', candidate)

                candidate = re.sub(r'\s+([;:,])', r'\1', candidate)
                candidate = re.sub(r'([;:])(?=\S)', r'\1 ', candidate)
                candidate = re.sub(r'(,)(?=[^\s\d])', r'\1 ', candidate)
                candidate = re.sub(r'[ \t]+', ' ', candidate)
                candidate = normalize_decimal_commas(candidate)
                candidate = fix_article_number_spacing(candidate)
                candidate = candidate.replace('\\', '')
                candidate = candidate.rstrip(',[]* ')
                candidate = collapse_duplicate_text(candidate)

                if candidate:
                    normalized_lines.append(candidate)

        if not normalized_lines and title_normalized:
            normalized_lines.append(title_normalized.strip())

        if normalized_lines:
            cleaned_body = "\n".join(normalized_lines)
            metadata = []
            for metadata_line in metadata_lines or []:
                metadata_line = normalize_space_characters(metadata_line).strip()
                if metadata_line:
                    metadata.append(metadata_line)

            if metadata:
                cleaned_text = "\n".join(metadata + ["", cleaned_body])
            else:
                cleaned_text = cleaned_body
            rows.append({"title": title_normalized or title, "text": cleaned_text})
            print(f"[VERBOSE] Item '{title_normalized or title}' extraido e limpo.")
    return [], "", cleaned_body

def select_input_and_output_via_gui(default_output_name="legislacao.csv"):
    """Abre painel GUI completo para selecionar entrada e saida."""
    gui = open_file_panel(
        title="Legislacao para CSV (Notion)",
        subtitle="Selecione arquivo de legislacao e configure a saida.",
        filetypes=[
            ("Arquivos suportados", "*.md *.html *.htm *.pdf *.docx"),
            ("Markdown", "*.md"),
            ("HTML", "*.html *.htm"),
            ("PDF", "*.pdf"),
            ("Word", "*.docx"),
            ("Todos os arquivos", "*.*"),
        ],
        extensions=[".md", ".html", ".htm", ".pdf", ".docx"],
        initial_files=[],
        allow_add_dir=True,
        recursive_dir=True,
        min_files=1,
        output_label="Pasta de saida",
        initial_output=str(SCRIPT_DIR),
        extra_texts=[("output_name", "Nome do CSV de saida", default_output_name)],
    )
    if gui is None:
        return None, None
    if not gui.get("confirmed"):
        return CANCELLED_BY_USER, None

    files = list(gui.get("files") or [])
    if not files:
        return CANCELLED_BY_USER, None
    input_file = files[0]
    output_dir = str(gui.get("output") or "").strip() or str(SCRIPT_DIR)
    output_name = str((gui.get("texts") or {}).get("output_name", "")).strip() or default_output_name
    output_file = force_output_path(str(Path(output_dir) / Path(output_name).name))
    return input_file, output_file


def resolve_input_path(path_value: str) -> str:
    path = Path(path_value).expanduser()
    if not path.is_absolute():
        path = SCRIPT_DIR / path
    return str(path.resolve())


def find_supported_input_files(folder):
    """Lista arquivos suportados em uma pasta, preferindo documentos diretos e HTML."""
    base = Path(folder).expanduser()
    if not base.is_absolute():
        base = SCRIPT_DIR / base
    base = base.resolve()
    if not base.is_dir():
        return []

    extension_rank = {".html": 0, ".htm": 1, ".md": 2, ".pdf": 3, ".docx": 4}
    candidates = [
        path.resolve()
        for path in base.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    return sorted(
        candidates,
        key=lambda path: (
            len(path.relative_to(base).parts),
            extension_rank.get(path.suffix.lower(), 99),
            str(path).lower(),
        ),
    )


def resolve_input_file(path_value: str):
    """Resolve arquivo de entrada; se for pasta, seleciona o primeiro documento suportado."""
    path = Path(path_value).expanduser()
    if not path.is_absolute():
        path = SCRIPT_DIR / path
    path = path.resolve()

    if path.is_dir():
        candidates = find_supported_input_files(path)
        if not candidates:
            print(f"\n[ERRO] Nenhum arquivo suportado foi encontrado na pasta: '{path}'")
            return None
        if len(candidates) > 1:
            print(f"[AVISO] Múltiplos arquivos suportados na pasta. Usando o primeiro: '{candidates[0]}'")
        else:
            print(f"[INFO] Arquivo encontrado na pasta alvo: '{candidates[0]}'")
        return str(candidates[0])

    return str(path)


def force_output_path(path_value: str) -> str:
    raw_value = str(path_value or "legislacao.csv").strip()
    path = Path(raw_value).expanduser()
    if path.exists() and path.is_dir():
        path = path / "legislacao.csv"
    if not path.is_absolute():
        path = SCRIPT_DIR / path
    path.parent.mkdir(parents=True, exist_ok=True)
    return str(path.resolve())

def find_local_input_file():
    """Procura arquivo suportado na pasta do script."""
    print("\n[INFO] Nenhum arquivo selecionado. Procurando por arquivo suportado na pasta do script...")

    default_target = SCRIPT_DIR / DEFAULT_INPUT_TARGET
    if default_target.is_dir():
        target_candidates = find_supported_input_files(default_target)
        if target_candidates:
            print(f"[INFO] Usando pasta alvo padrão '{DEFAULT_INPUT_TARGET}': '{target_candidates[0]}'")
            return str(target_candidates[0])

    candidates = []
    for name in glob.glob(str(SCRIPT_DIR / "*")):
        if not os.path.isfile(name):
            continue
        if os.path.splitext(name)[1].lower() in SUPPORTED_EXTENSIONS:
            candidates.append(name)
    candidates = sorted(set(candidates))

    if not candidates:
        print("\n[ERRO] Nenhum arquivo suportado foi encontrado. Coloque o arquivo de legislação na mesma pasta.")
        return None

    if len(candidates) > 1:
        print(f"[AVISO] Múltiplos arquivos encontrados. Usando o primeiro: '{candidates[0]}'")
    else:
        print(f"[INFO] Arquivo encontrado: '{candidates[0]}'")
    return candidates[0]

def read_markdown_content(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as md_file:
            return md_file.read()
    except Exception as exc:
        print(f"\n[ERRO] Não foi possível ler o arquivo Markdown: {exc}")
        return None


def read_html_content(file_path):
    try:
        raw_bytes = Path(file_path).read_bytes()
        encodings_to_try = []
        detected_encoding = detect_html_encoding(raw_bytes)
        if detected_encoding:
            encodings_to_try.append(detected_encoding)
        encodings_to_try.extend(["utf-8", "latin-1", "cp1252"])

        decoded_html = None
        for encoding in dict.fromkeys(encodings_to_try):
            try:
                decoded_html = raw_bytes.decode(encoding)
                break
            except (LookupError, UnicodeDecodeError):
                continue

        if decoded_html is None:
            decoded_html = raw_bytes.decode("utf-8", errors="replace")

        parser = ParagraphHTMLExtractor()
        parser.feed(decoded_html)
        parser.close()
        return "\n".join(parser.lines)
    except Exception as exc:
        print(f"\n[ERRO] Não foi possível ler o arquivo HTML: {exc}")
        return None

def read_pdf_content(file_path):
    try:
        import fitz
    except Exception:
        print("\n[ERRO] PyMuPDF (pymupdf) não está instalado. Instale com: python -m pip install pymupdf")
        return None

    try:
        with fitz.open(file_path) as doc:
            return "\n".join(page.get_text("text") for page in doc)
    except Exception as exc:
        print(f"\n[ERRO] Não foi possível ler o arquivo PDF: {exc}")
        return None

def read_docx_content(file_path):
    try:
        from docx import Document
    except Exception:
        print("\n[ERRO] python-docx não está instalado. Instale com: python -m pip install python-docx")
        return None

    try:
        doc = Document(file_path)
        lines = []
        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception as exc:
        print(f"\n[ERRO] Não foi possível ler o arquivo DOCX: {exc}")
        return None

def load_input_content(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".md":
        return read_markdown_content(file_path)
    if ext in {".html", ".htm"}:
        return read_html_content(file_path)
    if ext == ".pdf":
        return read_pdf_content(file_path)
    if ext == ".docx":
        return read_docx_content(file_path)
    print("\n[ERRO] Formato não suportado. Use .md, .html, .htm, .pdf ou .docx.")
    return None

def find_and_convert_markdown_to_csv(input_filename="", output_filename="", use_gui=True):
    """
    Função principal que localiza o arquivo, processa a estrutura e gera o .csv.
    """
    print("--- Conversor Estruturado de Leis (MD/HTML/PDF/DOCX) para CSV (Notion) ---")

    input_filename = str(input_filename or "").strip()
    output_filename = str(output_filename or "").strip()
    if not input_filename and use_gui:
        print("\n[INFO] Abrindo painel GUI para seleção do arquivo...")
        default_output_name = "legislacao.csv"
        gui_input, gui_output = select_input_and_output_via_gui(default_output_name=default_output_name)
        if gui_input is CANCELLED_BY_USER:
            print("\n[INFO] Operação cancelada pelo usuário. Nenhum arquivo foi processado.")
            return
        if gui_input:
            input_filename = gui_input
            if gui_output:
                output_filename = gui_output

    if input_filename:
        input_filename = resolve_input_file(input_filename)
        if not input_filename:
            return
        print(f"[INFO] Arquivo selecionado: '{input_filename}'")
        if os.path.splitext(input_filename)[1].lower() not in SUPPORTED_EXTENSIONS:
            print("\n[ERRO] O arquivo selecionado não possui extensão suportada (.md, .html, .htm, .pdf, .docx).")
            return
    else:
        input_filename = find_local_input_file()
        if not input_filename:
            return
        input_filename = resolve_input_path(input_filename)

    if output_filename:
        output_filename = force_output_path(output_filename)
    else:
        base_name = os.path.splitext(os.path.basename(input_filename))[0]
        output_filename = force_output_path(f"{base_name}.csv")

    print(f"\nIniciando a conversão estruturada para '{output_filename}'...")

    # Expressoes Regulares - correcao aplicada aqui
    art_pattern = re.compile(fr'^({ARTICLE_HEADING_PATTERN})(.*)')
    par_pattern = re.compile(
        r'^('
        r'\u00a7\s*\d+(?:\.\d+)?(?:\s*(?:\u00ba|\u00aa))?(?:-[A-Z]+)?\.?'
        r'|Par\u00e1grafo \u00fanico'
        r')'
        r'(?:\s*[-–—:]\s*|\.\s*)?'
        r'(.*)'
    )
    inc_pattern = re.compile(r'^([IVXLCDM]+)\s*[-\u2013\u2014]\s*(.*)')
    ali_pattern = re.compile(r'^([a-z])\)(.*)')
    num_pattern = re.compile(r'^(\d+)\s*[\.\)\-]\s*(.*)')
    header_search_pattern = re.compile(r'\s*(PARTE|LIVRO|T\u00cdTULO|CAP\u00cdTULO|Se\u00e7\u00e3o|Subse\u00e7\u00e3o|ANEXO)\s', re.IGNORECASE)
    inline_art_pattern = re.compile(ARTICLE_HEADING_PATTERN)
    inline_paragraph_split_pattern = re.compile(
        r'(?=(?:\u00a7\s*\d+(?:-[A-Z]+)?[\w.]*\u00ba?|Par\u00e1grafo \u00fanico)(?:\s*[-–—:]\s*|\s+))',
        re.IGNORECASE,
    )
    footer_start_pattern = re.compile(
        r'^[A-Z\u00C0-\u024F][A-Za-z\u00C0-\u024F]+(?:\s+(?:[a-z\u00C0-\u024F]{1,6}|[A-Z\u00C0-\u024F][A-Za-z\u00C0-\u024F]+)){0,6},\s+\d{1,2}\s+de\s+[A-Za-z\u00C0-\u024F]+\s+de\s+\d{4}',
    )

    content = load_input_content(input_filename)
    if not content:
        print("\n[ERRO] Não foi possível carregar o conteúdo do arquivo de entrada.")
        return

    content_from_articles = content_from_best_start(content)
    if not content_from_articles:
        print("\n[ERRO] 'Art. 1.º' não foi encontrado no arquivo.")
        return

    all_rows = []
    current_text_lines = []
    current_title = ""
    current_metadata_lines = []
    context = {"art": "", "par": "", "inc": "", "ali": "", "num": ""}
    hierarchy_text = {"art": "", "par": "", "inc": "", "ali": "", "num": ""}
    current_level = None
    structure_context = {level: "" for level in STRUCTURE_LEVELS}
    pending_structure_key = ""
    skip_until_document_heading = False
    stop_processing = False

    def decorated_title(title):
        document = normalize_for_matching(structure_context.get("documento", ""))
        if "CODIGO DE ETICA" in document and title:
            return f"Código de Ética, {title}"
        return title

    def reset_current_item():
        nonlocal current_title, current_text_lines, current_level, current_metadata_lines
        current_title, current_text_lines = "", []
        current_level = None
        current_metadata_lines = []

    def save_open_item():
        nonlocal current_title, current_text_lines, current_level, current_metadata_lines
        current_text_lines, current_title, cleaned_parent = save_current_item(
            all_rows,
            current_title,
            current_text_lines,
            current_metadata_lines,
        )
        if cleaned_parent and current_level:
            hierarchy_text[current_level] = cleaned_parent
        current_metadata_lines = []
        return cleaned_parent

    def get_parent_text(*levels):
        """Retorna o texto herdado mais específico disponível para os níveis informados."""
        for level in levels:
            parent_text = hierarchy_text.get(level)
            if parent_text:
                return parent_text

        fallback_parts = []
        order = ("art", "par", "inc", "ali", "num")
        for level in order:
            if level in levels and context.get(level):
                fallback_parts.append(context[level])

        return "\n".join(part for part in fallback_parts if part)

    for line in content_from_articles.splitlines():
        line = normalize_space_characters(line)
        original_line = line.strip()
        if not original_line:
            continue

        expanded_line = original_line
        expanded_line = re.sub(r';\s*(?=[IVXLCDM]+\s*[-–—])', ';\n', expanded_line)
        expanded_line = re.sub(r';\s*(?=[a-z]\))', ';\n', expanded_line)
        expanded_line = re.sub(r';\s*(?=\d+[\.\)\-])', ';\n', expanded_line)

        pending_lines = expanded_line.splitlines()

        while pending_lines:
            sub_line = pending_lines.pop(0)
            sub_line = normalize_space_characters(sub_line)
            line_stripped = sub_line.strip()
            if not line_stripped:
                continue

            line_stripped = re.sub(r'^#{1,6}\s*', '', line_stripped)
            line_stripped = re.sub(r'^[_*]+', '', line_stripped)
            line_stripped = re.sub(r'\[\s*\]\s*\([^)]+\)', '', line_stripped)
            line_stripped = re.sub(r'\[([^\]]+)\]', r'\1', line_stripped)
            line_stripped = re.sub(r'\(\s*(?:https?://|\.\.?/|#)[^)]+\)', '', line_stripped)
            line_stripped = re.sub(r'\(\s*[\w\u00C0-\u024F./#%-]*\.html?[^)]*\)', '', line_stripped, flags=re.IGNORECASE)
            line_stripped = re.sub(r'https?://\S+', '', line_stripped)
            line_stripped = normalize_ordinals(line_stripped)
            line_stripped = normalize_article_heading_punctuation(line_stripped)
            line_stripped = re.sub(r'\u00a7\s+', '\u00a7', line_stripped)
            line_stripped = re.sub(
                r'(\u00a7)([1-9])(?!\u00ba|\d)',
                lambda m: f"{m.group(1)}{m.group(2)}\u00ba",
                line_stripped,
            )
            line_stripped = re.sub(r'^>+\s*', '', line_stripped)
            if '>' in line_stripped:
                normalized = re.sub(r'(?:\s*>\s*){2,}', '\n', line_stripped)
                if normalized != line_stripped:
                    split_parts = normalized.splitlines()
                    if split_parts:
                        pending_lines = split_parts + pending_lines
                        continue
                    line_stripped = normalized
                line_stripped = re.sub(r'\s*>\s*', ' ', line_stripped)

            line_stripped = line_stripped.strip()
            document_heading = get_document_heading(line_stripped)
            if document_heading:
                save_open_item()
                apply_structure_heading(structure_context, "documento", document_heading)
                context = {"art": "", "par": "", "inc": "", "ali": "", "num": ""}
                hierarchy_text = {"art": "", "par": "", "inc": "", "ali": "", "num": ""}
                pending_structure_key = ""
                skip_until_document_heading = False
                reset_current_item()
                print(f"[VERBOSE] Contexto de documento: '{document_heading}'")
                continue

            if skip_until_document_heading:
                continue

            if is_resolution_heading(line_stripped) and structure_context.get("documento") == REGIMENTO_DOCUMENT:
                save_open_item()
                skip_until_document_heading = True
                pending_structure_key = ""
                reset_current_item()
                print(f"[VERBOSE] Ignorando resolução intermediária: '{line_stripped}'")
                continue

            if footer_start_pattern.match(line_stripped):
                print(f"[VERBOSE] Ignorando rodapé normativo: '{line_stripped}'")
                stop_processing = True
                break

            structural_header = parse_structural_header(line_stripped)
            if structural_header:
                save_open_item()
                key, heading, waits_description = structural_header
                apply_structure_heading(structure_context, key, heading)
                pending_structure_key = key if waits_description or heading_description_continues(heading) else ""
                reset_current_item()
                print(f"[VERBOSE] Contexto hierárquico: '{heading}'")
                continue

            if pending_structure_key and is_heading_description(line_stripped):
                append_structure_description(structure_context, pending_structure_key, line_stripped)
                print(f"[VERBOSE] Complemento de cabeçalho: '{line_stripped}'")
                pending_structure_key = pending_structure_key if heading_description_continues(line_stripped) else ""
                continue
            elif pending_structure_key:
                pending_structure_key = ""

            if line_stripped:
                search_pos = 1
                while True:
                    split_match = inline_paragraph_split_pattern.search(line_stripped, search_pos)
                    if not split_match:
                        break
                    prefix = line_stripped[:split_match.start()]
                    trimmed_prefix = prefix.rstrip()
                    has_boundary = not trimmed_prefix or trimmed_prefix[-1] in ".;:–—-)]}"
                    if has_boundary:
                        before = prefix.rstrip(' -–—')
                        after = line_stripped[split_match.start():].lstrip(' -–—')
                        if after:
                            pending_lines.insert(0, after)
                        line_stripped = before.strip()
                        search_pos = 1
                        if not line_stripped:
                            break
                    else:
                        search_pos = split_match.start() + 1
                if not line_stripped:
                    continue

            matches = list(inline_art_pattern.finditer(line_stripped))
            if matches and matches[0].start() == 0:
                segments = []
                last_pos = 0
                for i, m in enumerate(matches):
                    if m.start() > last_pos:
                        seg = line_stripped[last_pos:m.start()].strip()
                        if seg:
                            segments.append(seg)
                    next_start = matches[i + 1].start() if i + 1 < len(matches) else len(line_stripped)
                    seg = line_stripped[m.start():next_start].strip()
                    if seg:
                        segments.append(seg)
                    last_pos = next_start
                if len(segments) > 1 or (segments and segments[0] != line_stripped):
                    pending_lines = segments + pending_lines
                    continue
                if segments:
                    line_stripped = segments[0]

            art_match = art_pattern.match(line_stripped)
            par_match = par_pattern.match(line_stripped)
            inc_match = inc_pattern.match(line_stripped)
            ali_match = ali_pattern.match(line_stripped)
            num_match = num_pattern.match(line_stripped)

            if art_match:
                save_open_item()
                art_title = clean_title(art_match.group(1))
                context = {"art": art_title, "par": "", "inc": "", "ali": "", "num": ""}
                current_level = "art"
                current_title = decorated_title(context["art"])
                current_metadata_lines = build_metadata_lines(structure_context, current_title)
                current_text_lines = [line_stripped]
                hierarchy_text["art"] = ""
                hierarchy_text["par"] = ""
                hierarchy_text["inc"] = ""
                hierarchy_text["ali"] = ""
                hierarchy_text["num"] = ""
            
            elif par_match:
                save_open_item()
                par_title = clean_title(par_match.group(1))
                context["par"] = par_title
                context["inc"] = ""
                context["ali"] = ""
                context["num"] = ""
                current_level = "par"

                current_title = decorated_title(f'{context["art"]}, {context["par"]}')
                current_metadata_lines = build_metadata_lines(structure_context, current_title)
                parent_text = get_parent_text("art")
                current_text_lines = []
                if parent_text:
                    current_text_lines.append(parent_text)
                current_text_lines.append(f"\n{line_stripped}")
                hierarchy_text["inc"] = ""
                hierarchy_text["ali"] = ""
                hierarchy_text["num"] = ""

            elif inc_match:
                save_open_item()
                context["inc"] = inc_match.group(1).strip()
                context["ali"] = ""
                context["num"] = ""
                current_level = "inc"
                title_parts = [context["art"]]
                if context["par"]:
                    title_parts.append(context["par"])
                title_parts.append(context["inc"])
                current_title = decorated_title(", ".join(title_parts))
                current_metadata_lines = build_metadata_lines(structure_context, current_title)
                parent_text = get_parent_text("par", "art")
                current_text_lines = []
                if parent_text:
                    current_text_lines.append(parent_text)
                current_text_lines.append(f"\n{line_stripped}")
                hierarchy_text["ali"] = ""
                hierarchy_text["num"] = ""

            elif ali_match:
                save_open_item()
                alinea = ali_match.group(1).strip() + ")"
                context["ali"] = alinea
                context["num"] = ""
                current_level = "ali"
                title_parts = [context["art"]]
                if context["par"]:
                    title_parts.append(context["par"])
                if context["inc"]:
                    title_parts.append(context["inc"])
                title_parts.append(alinea)
                current_title = decorated_title(", ".join(title_parts))
                current_metadata_lines = build_metadata_lines(structure_context, current_title)
                parent_text = get_parent_text("inc", "par", "art")
                current_text_lines = []
                if parent_text:
                    current_text_lines.append(parent_text)
                current_text_lines.append(f"\n{line_stripped}")
                hierarchy_text["num"] = ""

            elif num_match:
                save_open_item()
                numero = num_match.group(1).strip()
                context["num"] = numero
                current_level = "num"
                title_parts = [context["art"]]
                if context["par"]:
                    title_parts.append(context["par"])
                if context["inc"]:
                    title_parts.append(context["inc"])
                if context["ali"]:
                    title_parts.append(context["ali"])
                title_parts.append(numero)
                current_title = decorated_title(", ".join(title_parts))
                current_metadata_lines = build_metadata_lines(structure_context, current_title)
                parent_text = get_parent_text("ali", "inc", "par", "art")
                current_text_lines = []
                if parent_text:
                    current_text_lines.append(parent_text)
                current_text_lines.append(f"\n{line_stripped}")

            else:
                if current_title:
                    header_in_line_match = header_search_pattern.search(line_stripped)
                    if header_in_line_match:
                        header_value = header_in_line_match.group(1)
                        if not (header_value.isupper() or header_value.istitle()):
                            header_in_line_match = None
                    if header_in_line_match:
                        text_before_header = line_stripped[:header_in_line_match.start()].strip()
                        if text_before_header:
                            current_text_lines.append(text_before_header)
                        
                        save_open_item()
                        
                        print(f"[VERBOSE] Ignorando cabeçalho no meio da linha: '{line_stripped[header_in_line_match.start():]}'")
                        reset_current_item()
                    else:
                        if current_text_lines:
                            current_text_lines[-1] = f"{current_text_lines[-1]} {line_stripped}"
                        else:
                            current_text_lines.append(line_stripped)
        if stop_processing:
            break
    save_open_item()

    footer_pattern = re.compile(
        r'(\b[A-Za-z?-?][^,\n\d]+,\s+\d{1,2}\s+de\s+[A-Za-z?-?]+\s+de\s+\d{4}.*$)',
        re.IGNORECASE | re.DOTALL,
    )
    if all_rows:
        last_text = all_rows[-1]['text']
        match_footer = footer_pattern.search(last_text)
        if match_footer and 'lei n' not in match_footer.group(1).lower():
            cleaned_last = last_text[:match_footer.start()].rstrip()
        else:
            cleaned_last = last_text.strip()
        all_rows[-1]['text'] = cleaned_last

        if not stop_processing:
            print(
                "\n[AVISO] O arquivo terminou sem rodapé normativo detectado. "
                "Confira se a fonte de entrada está completa."
            )
    try:
        with open(output_filename, 'w', newline='', encoding='utf-8-sig') as csv_file:
            writer = csv.writer(csv_file, quoting=csv.QUOTE_ALL)
            writer.writerow(["ID", "Nome (Artigo)", "Texto do Artigo"])
            
            for idx, row in enumerate(all_rows, start=1):
                writer.writerow([idx, row['title'], row['text']])
        
        print("\n---------------------------------------------------------")
        if not all_rows:
            print("[AVISO] Nenhum artigo foi extraído. Verifique o formato do arquivo de entrada.")
        else:
            print(f"[SUCESSO] Conversão concluída! {len(all_rows)} itens foram processados.")
            print(f"O arquivo CSV foi salvo como: '{output_filename}'")
        print("---------------------------------------------------------")

    except Exception as e:
        print(f"\n[ERRO] Ocorreu um erro inesperado ao escrever o arquivo CSV: {e}")


def build_arg_parser():
    parser = argparse.ArgumentParser(description="Conversor estruturado de legislacao (MD/HTML/PDF/DOCX) para CSV.")
    parser.add_argument("--input-file", default="", help="Arquivo ou pasta de entrada (.md, .html, .htm, .pdf, .docx).")
    parser.add_argument("--output-csv", default="", help="Arquivo CSV de saida.")
    parser.add_argument("--no-gui", action="store_true", help="Desativa painel GUI e usa apenas CLI.")
    return parser


def main():
    args = build_arg_parser().parse_args()
    find_and_convert_markdown_to_csv(
        input_filename=args.input_file,
        output_filename=args.output_csv,
        use_gui=(not bool(args.no_gui)),
    )


if __name__ == "__main__":
    main()
