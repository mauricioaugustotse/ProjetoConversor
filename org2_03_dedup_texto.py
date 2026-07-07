# -*- coding: utf-8 -*-
"""Org2 etapa 3: dedup por conteudo textual em HD_Mau.

Estagio A (pre-filtro): usa o trecho (1200 chars) de conteudo/extraido_full.tsv para
agrupar suspeitos por hash do texto normalizado.
Estagio B (confirmacao): extrai o texto COMPLETO so dos suspeitos (pdftotext/antiword/
python-docx/striprtf) e confirma identidade pelo hash do texto completo normalizado.
So remove (Lixeira) o que for confirmado; extracao falhou = nao remove.

Uso: python org2_03_dedup_texto.py            -> so gera plano_org2_dedup_texto.csv
     python org2_03_dedup_texto.py --exec     -> executa (Lixeira) + log_org2_dedup_texto.csv
     python org2_03_dedup_texto.py --sufixos  -> pos-dedup: limpa sufixos " (N)"/" [N]" orfaos
"""
import csv, hashlib, io, os, re, subprocess, sys, time, unicodedata
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor

ROOT = r"C:\Users\mauri\HD_Mau"
BASE = os.path.dirname(os.path.abspath(__file__))
TSV = os.path.join(BASE, "conteudo", "extraido_full.tsv")
PLANO = os.path.join(BASE, "plano_org2_dedup_texto.csv")
LOG = os.path.join(BASE, "log_org2_dedup_texto.csv")
LOG_SUF = os.path.join(BASE, "log_org2_renames_sufixo.csv")
CACHE = os.path.join(BASE, "org2_cache_fulltext.tsv")
ANTIWORD = r"C:\Program Files\Git\mingw64\bin\antiword.exe"
PDFTOTEXT = "pdftotext"
MIN_NORM = 300  # texto normalizado minimo para participar do dedup

RE_SUFIXO = re.compile(r"[ _]\(\d+\)$|[ _]\[\d+\]$")
RE_CNJ = re.compile(r"\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}")

def normalizar(texto):
    """lower + sem acentos + so alfanumerico: robusto a formatacao/extracao."""
    t = unicodedata.normalize("NFD", texto.lower())
    t = "".join(c for c in t if not unicodedata.combining(c))
    return re.sub(r"[^0-9a-z]", "", t)

def extrair_completo(path):
    """Texto completo por extensao. Retorna str ou None (falha)."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            r = subprocess.run([PDFTOTEXT, "-q", "-enc", "UTF-8", path, "-"],
                               capture_output=True, timeout=120)
            return r.stdout.decode("utf-8", "ignore") if r.returncode == 0 else None
        if ext == ".docx":
            import docx
            d = docx.Document(path)
            parts = [p.text for p in d.paragraphs]
            for tb in d.tables:
                for row in tb.rows:
                    parts.extend(c.text for c in row.cells)
            return "\n".join(parts)
        if ext == ".doc":
            r = subprocess.run([ANTIWORD, "-m", "UTF-8", path],
                               capture_output=True, timeout=120)
            return r.stdout.decode("utf-8", "ignore") if r.returncode == 0 else None
        if ext == ".rtf":
            from striprtf.striprtf import rtf_to_text
            raw = io.open(path, encoding="cp1252", errors="ignore").read()
            return rtf_to_text(raw, errors="ignore")
        if ext == ".txt":
            return io.open(path, encoding="utf-8", errors="ignore").read()
    except Exception:
        return None
    return None

def parse_tsv():
    """Registros (caminho_win, trecho). Linhas de continuacao sao anexadas ao anterior."""
    regs = []
    with io.open(TSV, encoding="utf-8", errors="replace") as f:
        for line in f:
            line = line.rstrip("\n")
            if line.startswith("/c/Users/mauri/"):
                parts = line.split("\t", 2)
                if len(parts) == 3:
                    win = parts[0].replace("/c/", "C:\\", 1).replace("/", "\\")
                    regs.append([win, parts[2]])
                    continue
            if regs:
                regs[-1][1] += " " + line
    return regs

def indexar_arvore():
    caminhos = {}
    por_nome = defaultdict(list)
    for dirpath, _, filenames in os.walk(ROOT):
        for name in filenames:
            p = os.path.join(dirpath, name)
            caminhos[p.lower()] = p
            por_nome[name.lower()].append(p)
    return caminhos, por_nome

def casar(regs, caminhos, por_nome):
    """registro -> caminho atual (exato, senao basename unico)."""
    nomes_tsv = defaultdict(int)
    for win, _ in regs:
        nomes_tsv[os.path.basename(win).lower()] += 1
    casados = {}
    for win, trecho in regs:
        lw = win.lower()
        if lw in caminhos:
            casados[caminhos[lw]] = trecho
            continue
        nome = os.path.basename(win).lower()
        if nomes_tsv[nome] == 1 and len(por_nome.get(nome, [])) == 1:
            casados[por_nome[nome][0]] = trecho
    return casados

def score(path):
    stem = os.path.splitext(os.path.basename(path))[0]
    ext = os.path.splitext(path)[1].lower()
    fmt = {".docx": 50, ".doc": 40, ".rtf": 30, ".pdf": 20, ".txt": 10}.get(ext, 0)
    try:
        mt = os.path.getmtime(path)
    except OSError:
        mt = 0
    return (0 if RE_SUFIXO.search(stem) else 1,
            1 if RE_CNJ.search(stem) else 0,
            min(len(stem), 200), fmt, mt)

def carregar_cache():
    cache = {}
    if os.path.exists(CACHE):
        with io.open(CACHE, encoding="utf-8") as f:
            for line in f:
                parts = line.rstrip("\n").split("\t")
                if len(parts) == 5:
                    cache[parts[0]] = (parts[1], parts[2], parts[3], int(parts[4]))
    return cache

EXTS_SUPORTADAS = (".pdf", ".doc", ".docx", ".rtf", ".txt")

def main():
    if "--sufixos" in sys.argv:
        return limpar_sufixos()

    caminhos, por_nome = indexar_arvore()
    print(f"  {len(caminhos)} arquivos na arvore atual")

    if "--full" in sys.argv:
        # Varredura completa: texto integral de toda a base, sem pre-filtro do TSV.
        alvos = [p for p in caminhos.values()
                 if os.path.splitext(p)[1].lower() in EXTS_SUPORTADAS]
        suspeitos = {"full": alvos}
        print(f"Modo FULL: {len(alvos)} arquivos com extensao suportada")
    else:
        print("Parse TSV...")
        regs = parse_tsv()
        print(f"  {len(regs)} registros")
        casados = casar(regs, caminhos, por_nome)
        print(f"  {len(casados)} casados registro->arquivo atual")

        # Estagio A: grupos por hash do trecho normalizado
        grupos = defaultdict(list)
        curtos = 0
        for path, trecho in casados.items():
            norm = normalizar(trecho)
            if len(norm) < MIN_NORM:
                curtos += 1
                continue
            grupos[hashlib.md5(norm.encode()).hexdigest()].append(path)
        suspeitos = {h: g for h, g in grupos.items() if len(g) > 1}
        n_susp = sum(len(g) for g in suspeitos.values())
        print(f"Estagio A: {n_susp} suspeitos em {len(suspeitos)} grupos (texto curto ignorado: {curtos})")
        alvos = [p for g in suspeitos.values() for p in g]

    # Estagio B: texto completo (com cache p/ retomada)
    cache = carregar_cache()
    pendentes = []
    for p in alvos:
        st = os.stat(p)
        key = f"{st.st_size}|{int(st.st_mtime)}"
        if p not in cache or cache[p][0] != key:
            pendentes.append((p, key))
    print(f"Estagio B: extraindo texto completo de {len(pendentes)} arquivos (cache: {len(alvos)-len(pendentes)})...")

    lock_rows = []
    def worker(item):
        p, key = item
        txt = extrair_completo(p)
        if txt is None:
            return (p, key, "FALHA", 0)
        norm = normalizar(txt)
        return (p, key, hashlib.sha1(norm.encode()).hexdigest(), len(norm))

    t0 = time.time()
    with ThreadPoolExecutor(max_workers=8) as ex, io.open(CACHE, "a", encoding="utf-8") as cf:
        for i, res in enumerate(ex.map(worker, pendentes), 1):
            p, key, h, ln = res
            cf.write(f"{p}\t{key}\t-\t{h}\t{ln}\n")
            cache[p] = (key, "-", h, ln)
            if i % 500 == 0:
                print(f"  {i}/{len(pendentes)} ({time.time()-t0:.0f}s)")

    # Grupos confirmados: mesmo sha1 do texto completo (e len>=MIN_NORM)
    confirmados = defaultdict(list)
    falhas = 0
    for g in suspeitos.values():
        for p in g:
            key, _, h, ln = cache.get(p, (None, None, "FALHA", 0))
            if h == "FALHA":
                falhas += 1
            elif ln >= MIN_NORM:
                confirmados[(h, ln)].append(p)
    dups = {k: sorted(set(v), key=score, reverse=True) for k, v in confirmados.items() if len(set(v)) > 1}

    plano = []
    for (h, ln), grupo in sorted(dups.items()):
        keeper = grupo[0]
        for perdedor in grupo[1:]:
            plano.append((perdedor, h, ln, keeper))
    with open(PLANO, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["remover", "sha1_texto", "len_norm", "sobrevivente"])
        w.writerows(plano)
    print(f"CONFIRMADOS: {len(plano)} excedentes em {len(dups)} grupos (falhas de extracao: {falhas})")
    print(f"Plano: {PLANO}")

    if "--exec" not in sys.argv:
        print("Rode com --exec para enviar a Lixeira.")
        return

    from send2trash import send2trash
    with open(LOG, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["removido_para_lixeira", "sha1_texto", "len_norm", "sobrevivente"])
        ok = 0
        for perdedor, h, ln, keeper in plano:
            try:
                send2trash(perdedor)
                w.writerow([perdedor, h, ln, keeper])
                ok += 1
            except Exception as e:
                print(f"ERRO ao remover {perdedor}: {e}")
    print(f"Enviados a Lixeira: {ok}. Log: {LOG}")

def limpar_sufixos():
    """Renomeia arquivos com sufixo ' (N)'/' [N]' quando o nome limpo esta livre."""
    renames = []
    for dirpath, _, filenames in os.walk(ROOT):
        for name in filenames:
            stem, ext = os.path.splitext(name)
            m = RE_SUFIXO.search(stem)
            if not m:
                continue
            limpo = RE_SUFIXO.sub("", stem).rstrip() + ext
            destino = os.path.join(dirpath, limpo)
            if not os.path.exists(destino):
                renames.append((os.path.join(dirpath, name), destino))
    with open(LOG_SUF, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["de", "para"])
        ok = 0
        for de, para in renames:
            try:
                os.rename(de, para)
                w.writerow([de, para])
                ok += 1
            except OSError as e:
                print(f"ERRO: {de}: {e}")
    print(f"Sufixos orfaos limpos: {ok}. Log: {LOG_SUF}")

if __name__ == "__main__":
    main()
