# -*- coding: utf-8 -*-
"""tsje_24: relatorio consolidado da REVISAO das lacunas do TSJE (jul/2026).

Le os artefatos do revisor (tsje_revisao_82.csv, tsje_verificacao_achados,
tsje_acordaos_revisao, tsje_visao_cache) e os .md rev-*, e produz um docx
didatico: por que a revisao foi feita, o metodo em camadas, o que foi
recuperado (lista nominal), o que ficou confirmado ausente (com evidencia)
e as causas-raiz das falsas ausencias.

Uso: python tsje_24_relatorio_revisao.py
Saida: D:\\TSJE_TRANSCRICOES\\_REVISAO - Auditoria das lacunas TSJE.docx
"""
import csv
import glob
import io
import json
import os
import re
import sys
import time
from collections import Counter, defaultdict

from docx import Document
from docx.shared import Pt, Cm

sys.stdout.reconfigure(encoding='utf-8')

TRAB = r'D:\TSJE_TRABALHO'
RAIZ = r'D:\TSJE_TRANSCRICOES'
SAIDA = os.path.join(RAIZ, '_REVISAO - Auditoria das lacunas TSJE.docx')


def ler_csv(path):
    if not os.path.exists(path):
        return []
    return list(csv.DictReader(io.open(path, encoding='utf-8-sig')))


def tabela(doc, cab, linhas, sz=9):
    t = doc.add_table(rows=1, cols=len(cab))
    t.style = 'Table Grid'
    for j, x in enumerate(cab):
        r = t.rows[0].cells[j].paragraphs[0].add_run(str(x))
        r.bold = True
        r.font.size = Pt(sz)
    for ln in linhas:
        cells = t.add_row().cells
        for j, x in enumerate(ln):
            cells[j].paragraphs[0].add_run(str(x)).font.size = Pt(sz)


def main():
    revisao = ler_csv(os.path.join(TRAB, 'tsje_revisao_82.csv'))
    acord = ler_csv(os.path.join(TRAB, 'tsje_acordaos_revisao.csv'))

    # visao: paginas lidas por modelo
    visao = Counter()
    if os.path.exists(os.path.join(TRAB, 'tsje_visao_cache.jsonl')):
        with io.open(os.path.join(TRAB, 'tsje_visao_cache.jsonl'),
                     encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    visao[json.loads(line).get('_modelo', '?')] += 1

    # recuperadas: rev-*.md com corpo (+ lacunas de hoje via recuperada_por)
    recuperadas = []
    for p in glob.glob(os.path.join(RAIZ, '19*', 'rev-*.md')):
        t = io.open(p, encoding='utf-8').read()
        fm = t.split('---', 2)[1]
        corpo = t.split('---', 2)[2].strip()
        if 'problema:' in fm or not corpo:
            continue
        tit = re.search(r'titulo: (.+)', fm)
        d = re.search(r'data_sessao: (.+)', fm)
        arq = re.search(r'arquivo: (.+)', fm)
        recuperadas.append((d.group(1) if d else '?',
                            tit.group(1) if tit else '?',
                            (arq.group(1) if arq else '?')[-55:]))
    recuperadas.sort()

    verd = Counter(r['veredito'] for r in revisao)
    aus = [r for r in revisao if r['veredito'] == 'AUSENTE_CONFIRMADA']
    ind = [r for r in revisao if r['veredito'] == 'INDETERMINADA']

    doc = Document()
    est = doc.styles['Normal']
    est.font.name = 'Calibri'
    est.font.size = Pt(11)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2)

    doc.add_heading('Auditoria das lacunas — Atas do TSJE (1932–1937)', 0)
    doc.add_heading('Relatório da revisão de 16-17/07/2026', 2)
    doc.add_paragraph(f'Gerado em {time.strftime("%d/%m/%Y %H:%M")}.')

    doc.add_heading('1. Por que esta revisão', 1)
    doc.add_paragraph(
        'O diagnóstico inicial apontava 82 sessões do TSJE "sem ata formal '
        'em nenhum boletim preservado". O número era suspeito: o AtoM é o '
        'repositório oficial do acervo histórico e uma omissão desse '
        'tamanho seria improvável. A hipótese central (confirmada): partes '
        'do acervo digitalizado não têm camada de OCR, e o detector '
        'automático de cabeçalhos só enxerga texto.')

    doc.add_heading('2. O método — revisão em camadas', 1)
    for t in [
        'Acervo completado: cruzamento da numeração sequencial dos BEs com '
        'o catálogo AtoM e re-download de tudo que era suspeito; a '
        'deduplicação por sha1 separou o que já tínhamos (173) do que '
        'faltava de verdade (47 boletins novos). Acervo final: 1.228 PDFs.',
        f'Leitura por visão computacional de {sum(visao.values())} páginas '
        f'sem OCR útil ("páginas cegas"): '
        + ', '.join(f'{v} por {k.replace("gpt-5.6-", "")}'
                    for k, v in visao.most_common())
        + ' (funil de custo: sol interrompido, luna assumiu a varredura, '
          'terra revisou as páginas que o luna não leu).',
        'Busca dirigida por sessão-alvo: janela de datas interpolada das '
        'sessões vizinhas (EXCLUSIVE, para não capturar a ata da vizinha), '
        'estendida ao ano seguinte (a seção "ACTAS" saía com defasagem), '
        'com detector textual relaxado (tolerante a OCR corrompido) e '
        'busca alternativa pela data por extenso.',
        'Verificação em dupla camada de todo achado: confirmação por IA de '
        'visão na imagem da página (número, tipo, data e ano impressos, '
        'tribunal, ata formal × resenha) e conferência final humana/Claude '
        'antes da transcrição. Nenhum achado virou transcrição sem passar '
        'pelas duas.']:
        doc.add_paragraph(t, style='List Bullet')

    doc.add_heading('3. Resultado sobre as 82 "ausentes"', 1)
    doc.add_paragraph(
        f'{len(recuperadas)} atas formais foram ENCONTRADAS e transcritas '
        f'(estavam em páginas sem OCR, em boletins do ano seguinte, ou com '
        f'cabeçalho corrompido pelo OCR). '
        f'{len(aus)} sessões seguem ausentes COM EVIDÊNCIA (todos os '
        f'boletins da janela examinados por OCR e visão). '
        f'{len(ind)} permanecem indeterminadas. As demais entradas do '
        f'diagnóstico eram falsos positivos auditados (resenhas, sessões '
        f'vizinhas, atas homônimas de outros anos, atas de TRE).')
    if recuperadas:
        doc.add_heading('Atas recuperadas pela revisão', 2)
        tabela(doc, ['Data da sessão', 'Título', 'Fonte (PDF)'],
               recuperadas)
    if aus:
        doc.add_heading('Ausências confirmadas', 2)
        tabela(doc, ['Ano', 'Sessão', 'Janela de publicação examinada',
                     'BEs examinados'],
               [[r['ano'], f"{r['num']}ª {r['tipo']}", r['janela_pub'],
                 r['candidatos']] for r in aus])
    if ind:
        doc.add_heading('Indeterminadas (para retomada futura)', 2)
        tabela(doc, ['Ano', 'Sessão', 'Pendência'],
               [[r['ano'], f"{r['num']}ª {r['tipo']}", r['evidencia']]
                for r in ind])

    doc.add_heading('4. Causas-raiz das falsas "ausências"', 1)
    for t in [
        'Páginas digitalizadas sem camada de OCR — invisíveis para a '
        'detecção textual (ex.: a 51ª sessão de 1936, achada por visão '
        'na página 4 de um BE do AtoM sem OCR).',
        'Fronteira de ano: atas de dezembro publicadas em BEs de janeiro '
        'do ano seguinte (ex.: 34ª de 1932 e 102ª de 1933, ambas em BEs '
        'de janeiro seguintes).',
        'OCR degradado no cabeçalho: o detector estrito exigia o número '
        'contíguo a "SESSÃO"; ruído de digitalização quebrava o padrão.',
        'Atas homônimas: buraco do ano X casando com ata de mesmo número '
        'de outro ano (boletins mal-arquivados; atas de 1934 publicadas '
        'com atraso em 1935) — a verificação de ano impresso barrou todas.']:
        doc.add_paragraph(t, style='List Bullet')

    if acord:
        doc.add_heading('5. Depuração da fila de acórdãos', 1)
        c = Counter((r.get('fila', ''), r.get('resultado', '')[:24])
                    for r in acord)
        tabela(doc, ['Fila', 'Resultado', 'Qtde'],
               [[f, r, n] for (f, r), n in sorted(c.items())])
        doc.add_paragraph(
            'Ambíguos resolvidos por comparação de teor (relator, partes, '
            'UF, matéria); números ilegíveis recuperados por leitura da '
            'imagem; os sem-teor ganharam o texto bruto do OCR '
            '(confiança "bruta" — sem risco de invenção).')

    doc.add_heading('6. Onde está cada coisa', 1)
    for t in [
        'D:\\TSJE_TRABALHO\\tsje_revisao_82.csv — o veredito fundamentado '
        'de cada sessão investigada.',
        'D:\\TSJE_TRABALHO\\tsje_verificacao_achados.csv — as conferências '
        'de imagem de cada achado (aprovado/reprovado e por quê).',
        'D:\\TSJE_TRABALHO\\tsje_visao_paginas.csv — o que há em cada '
        'página sem OCR (catálogo da leitura por visão).',
        'D:\\TSJE_TRANSCRICOES\\<ano>\\rev-*.md — as transcrições '
        'recuperadas pela revisão (integradas aos volumes e ao Notion).',
        'Scripts tsje_20..24 no repo ProjetoConversor — o revisor completo, '
        're-executável.']:
        doc.add_paragraph(t, style='List Bullet')

    doc.save(SAIDA)
    print(f'-> {SAIDA}')
    print(f'recuperadas={len(recuperadas)} ausentes_confirmadas={len(aus)} '
          f'indeterminadas={len(ind)} | visao={dict(visao)}')


if __name__ == '__main__':
    main()
