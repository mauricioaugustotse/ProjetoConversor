# -*- coding: utf-8 -*-
"""tsje_40: gera um PROJETO DE PESQUISA DE DOUTORADO (Faculdade de Direito da USP)
sobre a Justica Eleitoral na Era Vargas, escrito por gpt-5.6-sol, a partir de:
  - o relatorio Notion "Principais controversias juridicas no TSJE (1932-1937)";
  - as 3 bases estruturadas do TSJE (atas / acordaos / processos) como corpus empirico;
  - o rascunho de proposta ja enviado ao Prof. Jose Levi (USP) e o molde do projeto UnB 2023.

Enfoque: empirico-institucional (historia institucional do direito). Saida: .docx ABNT
na pasta do Google Drive + pagina no Notion.

Fluxo (retomavel; cache em D:\\TSJE_TRABALHO\\projeto_doutorado):
  python tsje_40_projeto_doutorado.py --coletar     # baixa relatorio + casos das bases + docx
  python tsje_40_projeto_doutorado.py --redigir      # gpt-5.6-sol escreve as secoes
  python tsje_40_projeto_doutorado.py --docx         # monta o .docx no Drive
  python tsje_40_projeto_doutorado.py --notion       # grava pagina no Notion
  python tsje_40_projeto_doutorado.py --verificar    # confere casos citados x fontes
  python tsje_40_projeto_doutorado.py --tudo         # coletar->redigir->docx->notion->verificar
Flags: --forcar (regera secoes existentes).
"""
from __future__ import annotations

import argparse
import json
import re
import sys
import time
from pathlib import Path

ROOT = Path(__file__).resolve().parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from conle_gerador import llm
from conle_gerador import notion_rag as nr
from conle_gerador import notion_writer as nw
import _doc_lib

# ------------------------------------------------------------------ constantes
WORK = Path(r"D:\TSJE_TRABALHO\projeto_doutorado")
SECDIR = WORK / "secoes"
FONTES = WORK / "fontes.json"

REPORT_ID = "eb2ddc9cd7f74d45b7aa92665129c906"        # relatorio de controversias
PAGINA_MAE = "39f72195-5c64-8126-95d8-da7259a5a312"   # pagina-mae TSJE (parent no Notion)
DB_PROC = "39f72195-5c64-8164-a8ee-e9bca4503779"
DB_ACOR = "39f72195-5c64-81ab-822c-e7eba4b2def2"
DB_ATAS = "39f72195-5c64-81d0-9f29-f3ebb0839b56"

DRIVE_USP = Path(r"G:\Meu Drive\Justiça eleitoral na era Vargas (1930-1937)"
                 r"\01 - Pesquisa própria (Maurício)\Doutorado - USP (correspondência Dr. Levi)")
DOCX_OUT = DRIVE_USP / "Projeto de pesquisa de doutorado - Racionalizacao do poder e a Justica Eleitoral na Era Vargas - USP 2026.docx"
DOCX_RASCUNHO_USP = DRIVE_USP / "Rascunho proposta de tese USP - Dr. Levi - Justiça Eleitoral na Era Vargas.docx"
DOCX_MOLDE_UNB = (Path(r"G:\Meu Drive\Justiça eleitoral na era Vargas (1930-1937)"
                       r"\01 - Pesquisa própria (Maurício)\Doutorado - UnB 2023")
                  / "Projeto de pesquisa doutorado UnB 2023 - Cortes eleitorais e Assis Brasil.docx")

TITULO = ("Racionalização do poder político como fundamento da construção da ordem jurídica "
          "na Era Vargas: a Justiça Eleitoral na década de 1930")
CANDIDATO = "Maurício Augusto Chiaramonte Vieira"
ORIENTADOR = "Prof. Dr. José Levi Mello do Amaral Júnior"

# classes-nucleo (nomes exatos do schema de `processos`)
CLASSES_PROC = ["CONSULTA", "RECURSO ELEITORAL", "REPRESENTAÇÃO", "HABEAS CORPUS",
                "MANDADO DE SEGURANÇA", "PROCESSO", "ATOS E EXPEDIENTES"]
# casos nominais citados no relatorio (best-effort por classe+numero+UF)
CASOS_NOMINAIS = [
    {"classe": "CONSULTA", "numero": "1", "UF": "ES"},
    {"classe": "CONSULTA", "numero": "54", "UF": "CE"},
    {"classe": "CONSULTA", "numero": "272", "UF": "MG"},
    {"classe": "RECURSO ELEITORAL", "numero": "38", "UF": "SE"},
    {"classe": "REPRESENTAÇÃO", "numero": "66", "UF": "MG"},
    {"classe": "HABEAS CORPUS", "numero": "30", "UF": "RN"},
    {"classe": "MANDADO DE SEGURANÇA", "numero": "3", "UF": "PA"},
    {"classe": "RECURSO ELEITORAL", "numero": "465", "UF": "SP"},
]


# ================================================================= util Notion
def _rt(pr, nome):
    return "".join(x.get("plain_text", "") for x in (pr.get(nome, {}) or {}).get("rich_text", []) or [])


def _sel(pr, nome):
    return ((pr.get(nome, {}) or {}).get("select") or {}).get("name", "")


def _dt(pr, nome):
    return ((pr.get(nome, {}) or {}).get("date") or {}).get("start", "")


def _title(pr):
    return "".join(x.get("plain_text", "") for x in (pr.get("Name", {}) or {}).get("title", []) or [])


def _texto_proc(pg):
    pr = pg.get("properties", {})
    return {
        "nome": _title(pr), "classe": _sel(pr, "classe"), "relator": _sel(pr, "relator"),
        "UF": _sel(pr, "UF"), "numero": _rt(pr, "numero"), "data": _dt(pr, "data_da_decisao"),
        "num_sessao": _rt(pr, "num_sessao"), "fonte": _sel(pr, "fonte"),
        "punchline": _rt(pr, "punchline"), "ementa": _rt(pr, "ementa"),
        "decisao": _rt(pr, "decisao"), "votacao": _rt(pr, "votacao"),
    }


def _amostra_classe(classe, n=4):
    body = {"page_size": 25, "sorts": [{"property": "data_da_decisao", "direction": "ascending"}],
            "filter": {"and": [
                {"property": "classe", "select": {"equals": classe}},
                {"property": "punchline", "rich_text": {"is_not_empty": True}},
            ]}}
    d = nr._req("POST", f"/databases/{DB_PROC}/query", body)
    itens = [_texto_proc(pg) for pg in d.get("results", [])]
    # prioriza os que tem ementa+decisao (mais completos)
    itens.sort(key=lambda x: (bool(x["ementa"]), bool(x["decisao"]), len(x["punchline"])), reverse=True)
    return itens[:n]


def _buscar_caso(c):
    filtros = [{"property": "classe", "select": {"equals": c["classe"]}},
               {"property": "numero", "rich_text": {"equals": c["numero"]}}]
    if c.get("UF"):
        filtros.append({"property": "UF", "select": {"equals": c["UF"]}})
    d = nr._req("POST", f"/databases/{DB_PROC}/query", {"page_size": 5, "filter": {"and": filtros}})
    return [_texto_proc(pg) for pg in d.get("results", [])]


# ================================================================= extrair docx
def _extrair_docx(path: Path) -> str:
    try:
        import docx
    except Exception:
        return ""
    if not path.is_file():
        return f"[arquivo nao encontrado: {path.name}]"
    d = docx.Document(str(path))
    return "\n".join(p.text for p in d.paragraphs if p.text.strip())


# ================================================================= 1) COLETA
def coletar():
    WORK.mkdir(parents=True, exist_ok=True)
    fontes = {}
    print("[coletar] relatorio de controversias (Notion)...")
    fontes["relatorio_md"] = _doc_lib.render_md(_doc_lib.dump_blocos(REPORT_ID), [])
    print(f"          relatorio: {len(fontes['relatorio_md'])} chars")

    print("[coletar] amostra de casos reais por classe (base processos)...")
    casos = {"amostra_por_classe": {}, "nominais": []}
    for cl in CLASSES_PROC:
        try:
            am = _amostra_classe(cl, n=4)
            casos["amostra_por_classe"][cl] = am
            print(f"          {cl}: {len(am)} casos")
        except Exception as e:  # noqa: BLE001
            print(f"          {cl}: ERRO {e}")
        time.sleep(0.3)

    print("[coletar] casos nominais citados no relatorio (best-effort)...")
    for c in CASOS_NOMINAIS:
        try:
            achados = _buscar_caso(c)
            if achados:
                casos["nominais"].append({"busca": c, "achados": achados})
                print(f"          {c['classe']} n.{c['numero']}/{c.get('UF','')}: {len(achados)} achado(s)")
            else:
                print(f"          {c['classe']} n.{c['numero']}/{c.get('UF','')}: (nao localizado)")
        except Exception as e:  # noqa: BLE001
            print(f"          {c}: ERRO {e}")
        time.sleep(0.3)
    fontes["casos"] = casos

    print("[coletar] rascunho USP + molde UnB (docx do Drive)...")
    fontes["rascunho_usp"] = _extrair_docx(DOCX_RASCUNHO_USP)
    fontes["molde_unb"] = _extrair_docx(DOCX_MOLDE_UNB)
    print(f"          rascunho USP: {len(fontes['rascunho_usp'])} chars; molde UnB: {len(fontes['molde_unb'])} chars")

    FONTES.write_text(json.dumps(fontes, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[coletar] salvo em {FONTES}")
    return fontes


# ================================================================= 2) REDACAO
SYS = (
    "Voce e um jurista brasileiro, pesquisador experiente, redigindo um PROJETO DE PESQUISA "
    "para o doutorado em Direito, area de DIREITO DO ESTADO, da Faculdade de Direito da "
    "Universidade de Sao Paulo (Largo Sao Francisco). Registro: portugues culto, formal e "
    "IMPESSOAL (nunca use primeira pessoa; prefira 'a pesquisa', 'pretende-se', 'propoe-se', "
    "'sustenta-se'). Densidade de doutorado: argumentativo, com dialogo teorico e precisao "
    "conceitual. Normas ABNT, citacoes no sistema autor-data (SOBRENOME, ano).\n\n"
    "REGRA ABSOLUTA (anti-invencao): voce SO pode afirmar fatos historicos, numeros de "
    "processo, datas, nomes de relatores, ementas, dispositivos e estatisticas que constem "
    "EXPLICITAMENTE no DOSSIE FACTUAL fornecido no prompt. E terminantemente proibido inventar "
    "processos, datas, decisoes, autores ou citacoes. Se um dado pessoal do candidato faltar "
    "(idiomas, Lattes, linha exata), escreva um marcador entre colchetes, ex.: [inserir idiomas]. "
    "Nao invente referencias bibliograficas.\n\n"
    "Formato de saida: MARKDOWN. Use '##' para o titulo da secao e '###' para subitens; "
    "texto em paragrafos corridos justificaveis; listas com '-' ou '1.' quando fizer sentido. "
    "Nao escreva preambulos como 'aqui esta' nem comentarios meta: devolva apenas o conteudo da secao."
)


def _dossie(fontes, *, casos=True, limite_rel=16000):
    partes = ["=== RELATORIO: PRINCIPAIS CONTROVERSIAS JURIDICAS NO TSJE (1932-1937) ===",
              fontes["relatorio_md"][:limite_rel]]
    if casos:
        partes.append("\n=== AMOSTRA DE CASOS REAIS DAS BASES (corpus empirico; use como exemplos concretos) ===")
        cas = fontes.get("casos", {})
        for cl, itens in (cas.get("amostra_por_classe") or {}).items():
            for it in itens:
                partes.append(
                    f"- [{it['classe']}] {it['nome']} | rel. {it['relator']} | {it['UF']} | {it['data']} "
                    f"| {it['fonte']}\n    punchline: {it['punchline'][:300]}"
                    + (f"\n    ementa: {it['ementa'][:300]}" if it['ementa'] else ""))
        for nn in (cas.get("nominais") or []):
            for it in nn["achados"][:1]:
                partes.append(
                    f"- (NOMINAL) [{it['classe']}] {it['nome']} | rel. {it['relator']} | {it['UF']} | {it['data']}"
                    f"\n    punchline: {it['punchline'][:300]}"
                    + (f"\n    decisao: {it['decisao'][:280]}" if it['decisao'] else ""))
    return "\n".join(partes)


def _alinhamento(fontes):
    return ("=== ALINHAMENTO USP (eixo ja negociado com o Prof. Jose Levi; RESPEITAR) ===\n"
            "Titulo da tese: " + TITULO + "\n"
            "Eixo: a 'racionalizacao do poder politico' (controle juridico-institucional da "
            "representacao) como chave para explicar o surgimento da Justica Eleitoral nos anos 1930; "
            "conexao com o FEDERALISMO; leitura de Karl Loewenstein (Constituicao semantica/nominal/"
            "normativa; missao ao Brasil, 1937). Institucionalismo historico (mudanca institucional "
            "gradual, Mahoney & Thelen) como ferramenta analitica.\n\n"
            "--- Rascunho de proposta enviado ao Prof. Levi (estrutura de capitulos e bibliografia "
            "preliminar; use como base da secao de estrutura e da bibliografia) ---\n"
            + (fontes.get("rascunho_usp") or "")[:6000])


BLOCOS = [
    ("01_resumo_delimitacao",
     "Redija, nesta ordem: (a) um RESUMO do projeto (150-230 palavras, um paragrafo) e uma linha "
     "'Palavras-chave: ' com 5-6 termos; (b) a secao '## 1. Delimitacao do tema e problema de "
     "pesquisa'. Delimite o objeto (a atuacao decisoria do Tribunal Superior de Justica Eleitoral, "
     "1932-1937, como locus de racionalizacao juridica do poder politico), o recorte temporal e a "
     "pergunta central de pesquisa. Ancore em 2-3 controversias concretas do dossie."),
    ("02_justificativa_objetivos",
     "Redija: (a) '## 2. Justificativa e relevancia' (relevancia historico-juridica, originalidade "
     "e o diferencial empirico: existencia de um banco de dados estruturado ineditco construido pelo "
     "candidato, com 571 atas, 1.228 acordaos e 6.732 deliberacoes catalogadas); (b) '## 3. Objetivos', "
     "com '### Objetivo geral' (um paragrafo) e '### Objetivos especificos' (lista de 5-7 itens)."),
    ("03_hipoteses_marco",
     "Redija: (a) '## 4. Hipoteses' (uma hipotese central + 2-3 secundarias, em lista); (b) "
     "'## 5. Marco teorico', articulando: racionalizacao do poder politico; a tipologia constitucional "
     "de Loewenstein; institucionalismo historico (mudanca gradual - Mahoney & Thelen); e a tensao "
     "Direito/Politica. Dialogo teorico denso, nao enumerativo."),
    ("04_metodologia",
     "Redija '## 6. Metodologia e fontes'. Detalhe: (a) abordagem empirico-institucional / historia "
     "institucional do direito, combinando analise qualitativa das controversias com tratamento "
     "quantitativo do corpus; (b) as FONTES PRIMARIAS: as atas, os acordaos e as deliberacoes do TSJE "
     "sistematizados no banco de dados proprio (descreva os campos: classe, relator, UF, data, ementa, "
     "decisao, votacao) e as fontes documentais da epoca; (c) os procedimentos (selecao de controversias "
     "emblematicas, leitura dos acordaos, cruzamento com o contexto politico). Use os numeros do dossie."),
    ("05_estrutura_cronograma",
     "Redija: (a) '## 7. Estrutura provisoria da tese', com 5 a 6 capitulos (titulo + 2-3 linhas cada), "
     "evoluindo a estrutura do rascunho USP para incorporar a analise empirica das controversias e o "
     "recorte ate 1937 (extincao pelo Estado Novo); (b) '## 8. Cronograma' de 48 meses, em lista por "
     "semestres (levantamento; tratamento do corpus; analise; redacao; revisao/defesa)."),
    ("06_bibliografia",
     "Monte '## 9. Referencias' em ABNT. Use como base a bibliografia preliminar do rascunho USP e os "
     "autores efetivamente mencionados nas secoes anteriores. Ordene alfabeticamente por sobrenome. "
     "NAO invente obras: se precisar de uma referencia que nao esteja no material, use o marcador "
     "[completar referencia]. Formate cada entrada no padrao ABNT."),
]


def redigir(forcar=False):
    SECDIR.mkdir(parents=True, exist_ok=True)
    if not FONTES.is_file():
        print("[redigir] fontes.json ausente — rode --coletar antes.")
        return
    fontes = json.loads(FONTES.read_text(encoding="utf-8"))
    dossie = _dossie(fontes)
    alinh = _alinhamento(fontes)
    secoes_ja = {}
    for i, (nome, instrucao) in enumerate(BLOCOS, 1):
        out = SECDIR / f"{nome}.md"
        if out.is_file() and not forcar:
            print(f"[redigir] {nome}: ja existe (use --forcar p/ regerar)")
            secoes_ja[nome] = out.read_text(encoding="utf-8")
            continue
        contexto_prev = ""
        if secoes_ja:
            contexto_prev = ("\n=== SECOES JA REDIGIDAS (para coerencia e evitar repeticao) ===\n"
                             + "\n\n".join(v[:1200] for v in secoes_ja.values())[:4000])
        user = (f"{dossie}\n\n{alinh}\n{contexto_prev}\n\n"
                f"=== TAREFA ({i}/{len(BLOCOS)}) ===\n{instrucao}")
        print(f"[redigir] {nome}: chamando gpt-5.6-sol...")
        txt = llm.chat(SYS, user, max_output_tokens=9000)
        txt = (txt or "").strip()
        out.write_text(txt, encoding="utf-8")
        secoes_ja[nome] = txt
        print(f"          {len(txt)} chars -> {out.name}")
    print("[redigir] concluido.")


def _montar_md():
    partes = []
    for nome, _ in BLOCOS:
        p = SECDIR / f"{nome}.md"
        if p.is_file():
            partes.append(p.read_text(encoding="utf-8").strip())
    return "\n\n".join(partes)


# ================================================================= 3) DOCX
_MD_LINK = re.compile(r"\[([^\]]+)\]\((https?://[^\s)]+)\)")


def _add_hyperlink(par, texto, url):
    """Insere hyperlink real (azul sublinhado) no paragrafo via XML."""
    import docx.opc.constants
    from docx.oxml.shared import OxmlElement, qn
    r_id = par.part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                              is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    cor = OxmlElement("w:color"); cor.set(qn("w:val"), "0563C1"); rpr.append(cor)
    sub = OxmlElement("w:u"); sub.set(qn("w:val"), "single"); rpr.append(sub)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = texto; run.append(t)
    hl.append(run)
    par._p.append(hl)


def _add_fmt_runs(par, texto):
    for seg in re.split(r"(\*\*.+?\*\*|\*[^*\s][^*]*?\*)", texto):
        if not seg:
            continue
        if seg.startswith("**") and seg.endswith("**"):
            r = par.add_run(seg[2:-2]); r.bold = True
        elif seg.startswith("*") and seg.endswith("*") and len(seg) > 2:
            r = par.add_run(seg[1:-1]); r.italic = True
        else:
            par.add_run(seg)


def _add_runs(par, texto):
    """Adiciona runs tratando [links](url), **negrito** e *italico* inline."""
    pos = 0
    for m in _MD_LINK.finditer(texto):
        _add_fmt_runs(par, texto[pos:m.start()])
        _add_hyperlink(par, m.group(1), m.group(2))
        pos = m.end()
    _add_fmt_runs(par, texto[pos:])


def gerar_docx():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    md = _montar_md()
    if not md.strip():
        print("[docx] nenhuma secao redigida — rode --redigir antes.")
        return

    doc = Document()
    # estilo base ABNT
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    sec = doc.sections[0]
    for lado in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, lado, Cm(3))

    def par(texto="", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False, space_after=10,
            line=1.5, before=0):
        p = doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.line_spacing = line
        pf.space_after = Pt(space_after)
        pf.space_before = Pt(before)
        if texto:
            _add_runs(p, texto)
            for r in p.runs:
                r.font.size = Pt(size)
                if bold:
                    r.bold = True
        return p

    # ---- folha de rosto
    par("UNIVERSIDADE DE SÃO PAULO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    par("FACULDADE DE DIREITO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    par("PROGRAMA DE PÓS-GRADUAÇÃO EM DIREITO — ÁREA DE DIREITO DO ESTADO",
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    for _ in range(4):
        par("", space_after=0)
    par(CANDIDATO.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    for _ in range(3):
        par("", space_after=0)
    par(TITULO, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    for _ in range(2):
        par("", space_after=0)
    par("Projeto de pesquisa apresentado ao Programa de Pós-Graduação em Direito da Faculdade "
        "de Direito da Universidade de São Paulo, como parte do processo seletivo para o curso "
        "de Doutorado, na área de concentração Direito do Estado.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    par(f"Orientador pretendido: {ORIENTADOR}.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    for _ in range(6):
        par("", space_after=0)
    par("São Paulo", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    par("2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_page_break()

    # ---- corpo (parser md simples)
    linhas = md.split("\n")
    i = 0
    while i < len(linhas):
        ln = linhas[i].rstrip()
        if not ln.strip():
            i += 1
            continue
        if ln.startswith("### "):
            par(ln[4:].strip(), bold=True, space_after=6, before=8, line=1.5)
        elif ln.startswith("## "):
            par(ln[3:].strip().upper(), bold=True, size=12, space_after=8, before=14, line=1.5)
        elif ln.startswith("# "):
            par(ln[2:].strip().upper(), bold=True, size=13, space_after=10, before=14, line=1.5)
        elif re.match(r"^\s*[-*]\s+", ln):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, re.sub(r"^\s*[-*]\s+", "", ln))
            p.paragraph_format.line_spacing = 1.5
            for r in p.runs:
                r.font.size = Pt(12)
        elif re.match(r"^\s*\d+\.\s+", ln):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\s*\d+\.\s+", "", ln))
            p.paragraph_format.line_spacing = 1.5
            for r in p.runs:
                r.font.size = Pt(12)
        elif ln.lstrip().startswith("|"):
            # bloco de tabela markdown (| cel | cel |)
            bloco = []
            while i < len(linhas) and linhas[i].lstrip().startswith("|"):
                bloco.append(linhas[i].strip())
                i += 1
            i -= 1
            dados = [[c.strip() for c in l.strip("|").split("|")]
                     for l in bloco if not re.match(r"^\|[\s\-:|]+\|$", l)]
            if dados:
                ncols = max(len(r) for r in dados)
                tab = doc.add_table(rows=len(dados), cols=ncols)
                tab.style = "Table Grid"
                for ri, row in enumerate(dados):
                    for ci in range(ncols):
                        cel = row[ci] if ci < len(row) else ""
                        p = tab.cell(ri, ci).paragraphs[0]
                        p.paragraph_format.line_spacing = 1.0
                        _add_runs(p, cel)
                        for r in p.runs:
                            r.font.size = Pt(10)
                doc.add_paragraph()
        else:
            par(ln, line=1.5, space_after=10)
        i += 1

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_OUT))
    print(f"[docx] salvo em {DOCX_OUT}")


# ================================================================= 4) NOTION
def _md_para_blocos(md):
    blocos = []
    for ln in md.split("\n"):
        s = ln.rstrip()
        if not s.strip():
            continue
        if s.startswith("### "):
            blocos.append(nw.bloco_heading(3, s[4:].strip()))
        elif s.startswith("## "):
            blocos.append(nw.bloco_heading(2, s[3:].strip()))
        elif s.startswith("# "):
            blocos.append(nw.bloco_heading(1, s[2:].strip()))
        elif re.match(r"^\s*[-*]\s+", s):
            blocos.append(nw.bloco_bullet(re.sub(r"^\s*[-*]\s+", "", s)))
        elif re.match(r"^\s*\d+\.\s+", s):
            blocos.append(nw.bloco_bullet(re.sub(r"^\s*\d+\.\s+", "", s)))
        else:
            blocos.append(nw.bloco_paragraph(s))
    return blocos


def gravar_notion():
    md = _montar_md()
    if not md.strip():
        print("[notion] nenhuma secao redigida — rode --redigir antes.")
        return
    titulo_pg = "Projeto de pesquisa (doutorado USP) — Justiça Eleitoral na Era Vargas"
    body = {
        "parent": {"page_id": PAGINA_MAE},
        "properties": {"title": [{"type": "text", "text": {"content": titulo_pg}}]},
    }
    print("[notion] criando pagina filha da pagina-mae TSJE...")
    novo = nr._req("POST", "/pages", body)
    pid = novo.get("id")
    print(f"[notion] pagina criada: {pid}")
    # cabecalho
    blocos = [nw.bloco_callout(f"**{TITULO}**\n\nCandidato: {CANDIDATO} · Orientador pretendido: "
                               f"{ORIENTADOR} · USP — Faculdade de Direito · 2026", emoji="🎓")]
    blocos += _md_para_blocos(md)
    n = nw.escrever_pagina(pid, blocos)
    print(f"[notion] {n} blocos escritos. URL: https://www.notion.so/{(pid or '').replace('-','')}")


# ================================================================= 5) VERIFICAR
def verificar():
    if not FONTES.is_file():
        print("[verificar] fontes.json ausente.")
        return
    fontes = json.loads(FONTES.read_text(encoding="utf-8"))
    md = _montar_md()
    corpus = fontes.get("relatorio_md", "") + json.dumps(fontes.get("casos", {}), ensure_ascii=False)
    # numeros de processo citados no texto: "n. 123" / "no 123" / "nº 123"
    citados = set(re.findall(r"n[.ºo]{1,2}\s*(\d{1,4})", md, flags=re.IGNORECASE))
    faltantes = sorted(x for x in citados if x not in corpus)
    print(f"[verificar] {len(citados)} numeros citados no texto; "
          f"{len(faltantes)} nao encontrados literalmente nas fontes.")
    if faltantes:
        print("            revisar (podem ser invencao OU numero derivado de contexto): "
              + ", ".join(f"n. {x}" for x in faltantes[:40]))
    else:
        print("            OK: todos os numeros citados constam nas fontes.")


# ================================================================= CLI
def main():
    ap = argparse.ArgumentParser(description="Gera projeto de doutorado (USP) via gpt-5.6-sol.")
    ap.add_argument("--coletar", action="store_true")
    ap.add_argument("--redigir", action="store_true")
    ap.add_argument("--docx", action="store_true")
    ap.add_argument("--notion", action="store_true")
    ap.add_argument("--verificar", action="store_true")
    ap.add_argument("--tudo", action="store_true")
    ap.add_argument("--forcar", action="store_true")
    a = ap.parse_args()
    if a.tudo:
        coletar(); redigir(a.forcar); gerar_docx(); gravar_notion(); verificar(); return
    if a.coletar:
        coletar()
    if a.redigir:
        redigir(a.forcar)
    if a.docx:
        gerar_docx()
    if a.notion:
        gravar_notion()
    if a.verificar:
        verificar()
    if not any([a.coletar, a.redigir, a.docx, a.notion, a.verificar]):
        ap.print_help()


if __name__ == "__main__":
    main()
